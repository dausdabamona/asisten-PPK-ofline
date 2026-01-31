"""
Full integration test untuk document generation dengan rincian items
"""
from pathlib import Path
from docx import Document
import sys

# Add app to path
BASE_DIR = Path(__file__).parent
sys.path.insert(0, str(BASE_DIR))

from app.services.dokumen_generator import DokumenGenerator

# Initialize generator
generator = DokumenGenerator()

# Simulate form data with rincian_items (like from UI)
form_data = {
    'nama_kegiatan': 'Pembelian Alat Tulis Kantor',
    'kode_akun': '5.1.2.01.01.01.0001',
    'estimasi_biaya': 1000000,
    'tanggal_dokumen': '2024-01-15',
    'hari_tanggal': '15 Januari 2024',
    'penerima_nama': 'Budi Santoso',
    'penerima_nip': '19800515200012001',
    'penerima_jabatan': 'Staf Administrasi',
    
    # Satker data
    'unit_kerja': 'Dinas Pendidikan Kota Bandung',
    'sumber_dana': 'APBD 2024',
    
    # Penandatangan
    'nama_pengajuan': 'Budi Santoso',
    'nama_verifikator': 'Siti Rahayui',
    'nama_ppk': 'Dr. Ahmad Wijaya',
    'nama_atasan': 'Kepala Dinas Pendidikan',
    'nama_kpa': 'Pejabat Pembuat Komitmen',
}

# Add rincian items
rincian_items = [
    {
        'no': 1,
        'item_no': 1,
        'uraian': 'Kertas A4 putih',
        'nama_barang': 'Kertas A4 putih',
        'spesifikasi': 'Kertas putih 80gsm, 1 rim = 500 lembar',
        'volume': 10,
        'satuan': 'rim',
        'harga_satuan': 50000,
        'jumlah': 500000,
        'total_item': 500000,
        'keterangan': 'Untuk administrasi dan surat-menyurat',
    },
    {
        'no': 2,
        'item_no': 2,
        'uraian': 'Tinta printer hitam',
        'nama_barang': 'Tinta printer hitam',
        'spesifikasi': 'Cartridge tinta printer hitam compatible',
        'volume': 5,
        'satuan': 'buah',
        'harga_satuan': 80000,
        'jumlah': 400000,
        'total_item': 400000,
        'keterangan': 'Untuk printer kantor',
    },
    {
        'no': 3,
        'item_no': 3,
        'uraian': 'Ballpoint',
        'nama_barang': 'Ballpoint',
        'spesifikasi': 'Ballpoint standar, tinta biru',
        'volume': 100,
        'satuan': 'buah',
        'harga_satuan': 2000,
        'jumlah': 200000,
        'total_item': 200000,
        'keterangan': 'Stok kantor',
    },
]

form_data['rincian_items'] = rincian_items
form_data['subtotal'] = 1100000
form_data['ppn'] = 110000
form_data['total'] = 1210000

template_path = BASE_DIR / "templates" / "word" / "lembar_permintaan.docx"
output_path = BASE_DIR / "output" / "dokumen" / "test_full_integration.docx"

print("=" * 70)
print("FULL INTEGRATION TEST - Lembar Permintaan dengan Rincian")
print("=" * 70)

print(f"\nTemplate: {template_path}")
print(f"Output: {output_path}")

if template_path.exists():
    print("\n✓ Template found")
    
    # Create output directory
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Prepare data using the same logic as in dokumen_generator
    # Flatten rincian_items to item_N_* placeholders
    data = generator._flatten_rincian_items(form_data)
    
    print(f"\n✓ Data flattened")
    print(f"  - Total keys in data: {len(data)}")
    print(f"  - Rincian items: {len(data.get('rincian_items', []))}")
    
    # Test merge
    success = generator.merge_word_document(template_path, data, output_path)
    
    if success:
        print("✓ Merge successful")
        
        # Verify result
        doc = Document(str(output_path))
        
        # Collect all text from document
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        doc_text = '\n'.join(all_text)
        
        # Check for unreplaced placeholders
        import re
        placeholders = re.findall(r'\{\{[\w_:]+\}\}', doc_text)
        
        if placeholders:
            print(f"\n⚠ FOUND {len(placeholders)} UNREPLACED PLACEHOLDERS:")
            for ph in set(placeholders):
                print(f"  - {ph}")
        else:
            print("\n✓ NO UNREPLACED PLACEHOLDERS!")
        
        # Verify key values
        print("\n=== VERIFICATION OF KEY VALUES ===")
        checks = [
            ('Dinas Pendidikan Kota Bandung', 'Unit Kerja'),
            ('15 Januari 2024', 'Tanggal'),
            ('Kertas A4 putih', 'Item 1 Nama'),
            ('Tinta printer hitam', 'Item 2 Nama'),
            ('Ballpoint', 'Item 3 Nama'),
            ('Budi Santoso', 'Pengajuan'),
            ('Dr. Ahmad Wijaya', 'PPK'),
            ('Kepala Dinas Pendidikan', 'Atasan'),
            ('1,100,000', 'Subtotal (alternative format)'),
        ]
        
        found = 0
        for value, description in checks:
            # Try alternative formats
            variants = [
                value,
                value.replace(',', '.'),
                value.replace(',', ''),
            ]
            
            found_variant = any(v in doc_text for v in variants)
            
            if found_variant:
                print(f"  ✓ {description}")
                found += 1
            else:
                print(f"  ✗ {description} ('{value}' not found)")
        
        print(f"\n✓ {found} of {len(checks)} values verified")
        print(f"\n✓ Test document generated: {output_path}")
        
    else:
        print("✗ Merge failed")
else:
    print(f"✗ Template not found: {template_path}")
