"""
Script untuk membuat template dokumen Perjalanan Dinas dan Swakelola
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates", "word")
os.makedirs(TEMPLATES_DIR, exist_ok=True)


def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, center=False):
    """Add heading with custom style"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_paragraph(doc, text, bold=False, center=False):
    """Add a paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


# =============================================================================
# PERJALANAN DINAS TEMPLATES
# =============================================================================

def create_surat_tugas():
    """Surat Tugas Perjalanan Dinas"""
    doc = Document()

    # Header
    add_heading(doc, "{{kementerian}}", 1, True)
    add_heading(doc, "{{eselon1}}", 2, True)
    add_heading(doc, "{{satker_nama}}", 2, True)
    add_paragraph(doc, "{{satker_alamat}}, {{satker_kota}} - {{satker_provinsi}}", center=True)

    doc.add_paragraph()
    add_heading(doc, "SURAT TUGAS", 1, True)
    add_paragraph(doc, "Nomor: {{nomor_surat_tugas}}", center=True)

    doc.add_paragraph()

    # Isi
    p = doc.add_paragraph()
    p.add_run("Yang bertanda tangan di bawah ini:").bold = False

    # Tabel pemberi tugas
    table = doc.add_table(rows=4, cols=2)
    data = [
        ("Nama", "{{ppk_nama}}"),
        ("NIP", "{{ppk_nip}}"),
        ("Jabatan", "{{ppk_jabatan}}"),
        ("", "Selaku Pejabat Pembuat Komitmen"),
    ]
    for i, (label, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}" if label else value

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Memberikan tugas kepada:").bold = False

    # Tabel pelaksana
    table = doc.add_table(rows=4, cols=2)
    data = [
        ("Nama", "{{pelaksana_nama}}"),
        ("NIP", "{{pelaksana_nip}}"),
        ("Pangkat/Golongan", "{{pelaksana_pangkat}}"),
        ("Jabatan", "{{pelaksana_jabatan}}"),
    ]
    for i, (label, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Untuk melaksanakan perjalanan dinas dalam rangka:").bold = False

    add_paragraph(doc, "{{maksud_perjalanan}}")

    # Tabel perjalanan
    table = doc.add_table(rows=4, cols=2)
    data = [
        ("Tempat tujuan", "{{kota_tujuan}}, {{provinsi_tujuan}}"),
        ("Tanggal berangkat", "{{tanggal_berangkat}}"),
        ("Tanggal kembali", "{{tanggal_kembali}}"),
        ("Lama perjalanan", "{{lama_perjalanan}} (.............) hari"),
    ]
    for i, (label, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"

    doc.add_paragraph()
    add_paragraph(doc, "Demikian surat tugas ini dibuat untuk dilaksanakan dengan penuh tanggung jawab.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=4, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.rows[0].cells[0].text = "{{satker_kota}}, {{tanggal_surat_tugas}}"
    table.rows[1].cells[0].text = "{{ppk_jabatan}},"
    table.rows[2].cells[0].text = ""
    table.rows[3].cells[0].text = "{{ppk_nama}}\nNIP. {{ppk_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "surat_tugas.docx"))
    print("Created: surat_tugas.docx")


def create_sppd():
    """SPPD - Surat Perintah Perjalanan Dinas"""
    doc = Document()

    # Header
    add_heading(doc, "{{kementerian}}", 1, True)
    add_heading(doc, "{{eselon1}}", 2, True)

    doc.add_paragraph()
    add_heading(doc, "SURAT PERINTAH PERJALANAN DINAS (SPPD)", 1, True)
    add_paragraph(doc, "Nomor: {{nomor_sppd}}", center=True)

    doc.add_paragraph()

    # Tabel SPPD
    table = doc.add_table(rows=16, cols=3)
    table.style = 'Table Grid'

    items = [
        ("1", "Pejabat Pembuat Komitmen", "{{ppk_nama}}\nNIP. {{ppk_nip}}"),
        ("2", "Nama/NIP Pegawai yang melakukan\nperjalanan dinas", "{{pelaksana_nama}}\nNIP. {{pelaksana_nip}}"),
        ("3", "a. Pangkat dan Golongan\nb. Jabatan\nc. Tingkat Biaya Perjalanan Dinas", "a. {{pelaksana_pangkat}}\nb. {{pelaksana_jabatan}}\nc. {{tingkat_biaya}}"),
        ("4", "Maksud Perjalanan Dinas", "{{maksud_perjalanan}}"),
        ("5", "Alat angkutan yang dipergunakan", "{{alat_angkut}}"),
        ("6", "a. Tempat Berangkat\nb. Tempat Tujuan", "a. {{kota_asal}}\nb. {{kota_tujuan}}, {{provinsi_tujuan}}"),
        ("7", "a. Lama perjalanan dinas\nb. Tanggal Berangkat\nc. Tanggal harus kembali/tiba di\n    tempat baru", "a. {{lama_perjalanan}} (.............) hari\nb. {{tanggal_berangkat}}\nc. {{tanggal_kembali}}"),
        ("8", "Pengikut: Nama, Tanggal lahir, Keterangan", "-"),
        ("9", "Pembebanan Anggaran\na. Instansi\nb. Akun", "a. {{satker_nama}}\nb. {{kode_akun}}"),
        ("10", "Keterangan lain-lain", ""),
    ]

    for i, (no, label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = no
        row.cells[1].text = label
        row.cells[2].text = value

    doc.add_paragraph()
    add_paragraph(doc, "Dikeluarkan di : {{satker_kota}}")
    add_paragraph(doc, "Pada tanggal    : {{tanggal_surat_tugas}}")

    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=4, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.rows[0].cells[0].text = "{{ppk_jabatan}},"
    table.rows[1].cells[0].text = ""
    table.rows[2].cells[0].text = ""
    table.rows[3].cells[0].text = "{{ppk_nama}}\nNIP. {{ppk_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "sppd.docx"))
    print("Created: sppd.docx")


def create_kuitansi_uang_muka():
    """Kuitansi Uang Muka Perjalanan Dinas"""
    doc = Document()

    add_heading(doc, "KUITANSI UANG MUKA PERJALANAN DINAS", 1, True)

    doc.add_paragraph()

    # Tabel kuitansi
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    items = [
        ("Nomor Bukti", ": {{nomor_sppd}}"),
        ("Sudah Terima Dari", ": Bendahara Pengeluaran {{satker_nama}}"),
        ("Uang Sebanyak", ": Rp {{uang_muka}}"),
        ("Terbilang", ": {{uang_muka_terbilang}}"),
        ("Untuk Pembayaran", ": Uang Muka Perjalanan Dinas ke {{kota_tujuan}}"),
        ("", "  dalam rangka {{maksud_perjalanan}}"),
        ("", "  sesuai Surat Tugas No. {{nomor_surat_tugas}}"),
    ]

    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Setuju Dibayar:\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"
    table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_surat_tugas}}\nYang Menerima,\n\n\n\n{{pelaksana_nama}}\nNIP. {{pelaksana_nip}}"

    doc.add_paragraph()
    add_paragraph(doc, "Lunas Dibayar:")
    add_paragraph(doc, "Bendahara Pengeluaran,")
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "{{bendahara_nama}}")
    add_paragraph(doc, "NIP. {{bendahara_nip}}")

    doc.save(os.path.join(TEMPLATES_DIR, "kuitansi_uang_muka.docx"))
    print("Created: kuitansi_uang_muka.docx")


def create_rincian_biaya_pd():
    """Rincian Biaya Perjalanan Dinas"""
    doc = Document()

    add_heading(doc, "RINCIAN BIAYA PERJALANAN DINAS", 1, True)

    doc.add_paragraph()

    # Informasi perjalanan
    table = doc.add_table(rows=5, cols=2)
    items = [
        ("Lampiran SPPD Nomor", ": {{nomor_sppd}}"),
        ("Tanggal", ": {{tanggal_surat_tugas}}"),
        ("Nama", ": {{pelaksana_nama}}"),
        ("Tujuan", ": {{kota_tujuan}}, {{provinsi_tujuan}}"),
        ("Lama", ": {{lama_perjalanan}} hari"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()

    # Tabel rincian biaya
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'

    # Header
    hdr = table.rows[0]
    hdr.cells[0].text = "No"
    hdr.cells[1].text = "Uraian"
    hdr.cells[2].text = "Jumlah (Rp)"
    hdr.cells[3].text = "Keterangan"

    # Data rows
    items = [
        ("1", "Biaya Transportasi (Pergi-Pulang)", "{{biaya_transport}}", ""),
        ("2", "Uang Harian ({{lama_perjalanan}} hari)", "{{biaya_uang_harian}}", ""),
        ("3", "Biaya Penginapan", "{{biaya_penginapan}}", ""),
        ("4", "Uang Representasi", "{{biaya_representasi}}", ""),
        ("5", "Biaya Lain-lain", "{{biaya_lain_lain}}", ""),
        ("", "JUMLAH", "{{total_biaya}}", ""),
    ]

    for i, (no, uraian, jumlah, ket) in enumerate(items):
        row = table.rows[i + 1]
        row.cells[0].text = no
        row.cells[1].text = uraian
        row.cells[2].text = jumlah
        row.cells[3].text = ket

    # Total terbilang
    table.rows[7].cells[0].text = ""
    table.rows[7].cells[1].text = "Terbilang: {{total_biaya_terbilang}}"

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Telah dibayar sejumlah\nRp {{total_biaya}}\n\nBendahara Pengeluaran,\n\n\n\n{{bendahara_nama}}\nNIP. {{bendahara_nip}}"
    table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_kembali}}\nYang Menerima,\n\n\n\n{{pelaksana_nama}}\nNIP. {{pelaksana_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "rincian_biaya_pd.docx"))
    print("Created: rincian_biaya_pd.docx")


def create_laporan_perjalanan_dinas():
    """Laporan Perjalanan Dinas"""
    doc = Document()

    add_heading(doc, "LAPORAN PERJALANAN DINAS", 1, True)

    doc.add_paragraph()

    # Informasi dasar
    table = doc.add_table(rows=8, cols=2)
    items = [
        ("Berdasarkan Surat Tugas", ": Nomor {{nomor_surat_tugas}}"),
        ("Tanggal", ": {{tanggal_surat_tugas}}"),
        ("Nama", ": {{pelaksana_nama}}"),
        ("NIP", ": {{pelaksana_nip}}"),
        ("Jabatan", ": {{pelaksana_jabatan}}"),
        ("Tempat Tujuan", ": {{kota_tujuan}}, {{provinsi_tujuan}}"),
        ("Tanggal Berangkat", ": {{tanggal_berangkat}}"),
        ("Tanggal Kembali", ": {{tanggal_kembali}}"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()
    add_heading(doc, "I. MAKSUD DAN TUJUAN", 2)
    add_paragraph(doc, "{{maksud_perjalanan}}")

    add_heading(doc, "II. REALISASI PELAKSANAAN", 2)
    add_paragraph(doc, "[Uraian kegiatan yang dilaksanakan selama perjalanan dinas]")

    add_heading(doc, "III. HASIL YANG DICAPAI", 2)
    add_paragraph(doc, "[Hasil-hasil yang dicapai dari perjalanan dinas]")

    add_heading(doc, "IV. KESIMPULAN DAN SARAN", 2)
    add_paragraph(doc, "[Kesimpulan dan saran untuk tindak lanjut]")

    doc.add_paragraph()
    add_paragraph(doc, "Demikian laporan perjalanan dinas ini dibuat dengan sebenarnya untuk dapat dipergunakan sebagaimana mestinya.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Mengetahui,\n{{atasan_jabatan}}\n\n\n\n{{atasan_nama}}\nNIP. {{atasan_nip}}"
    table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_kembali}}\nPembuat Laporan,\n\n\n\n{{pelaksana_nama}}\nNIP. {{pelaksana_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "laporan_perjalanan_dinas.docx"))
    print("Created: laporan_perjalanan_dinas.docx")


def create_kuitansi_rampung():
    """Kuitansi Rampung Perjalanan Dinas"""
    doc = Document()

    add_heading(doc, "KUITANSI/BUKTI PEMBAYARAN PERJALANAN DINAS", 1, True)
    add_paragraph(doc, "(RAMPUNG)", center=True)

    doc.add_paragraph()

    # Tabel kuitansi
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    items = [
        ("Nomor Bukti", ": {{nomor_sppd}}"),
        ("Sudah Terima Dari", ": Bendahara Pengeluaran {{satker_nama}}"),
        ("Untuk Pembayaran", ": Biaya Perjalanan Dinas ke {{kota_tujuan}} dalam rangka {{maksud_perjalanan}}"),
        ("", "  sesuai Surat Tugas No. {{nomor_surat_tugas}}"),
        ("", "  tanggal {{tanggal_surat_tugas}}"),
    ]

    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()

    # Tabel perhitungan
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0]
    hdr.cells[0].text = "No"
    hdr.cells[1].text = "Uraian"
    hdr.cells[2].text = "Jumlah (Rp)"

    items = [
        ("1", "Total Biaya Perjalanan Dinas", "{{total_biaya}}"),
        ("2", "Uang Muka yang telah diterima", "{{uang_muka}}"),
        ("3", "Sisa kurang/(lebih) bayar", "{{kekurangan_bayar}} / ({{kelebihan_bayar}})"),
        ("", "Terbilang:", "{{sisa_terbilang}}"),
    ]

    for i, (no, uraian, jumlah) in enumerate(items):
        row = table.rows[i + 1]
        row.cells[0].text = no
        row.cells[1].text = uraian
        row.cells[2].text = jumlah

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Setuju Dibayar:\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"
    table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_kembali}}\nYang Menerima,\n\n\n\n{{pelaksana_nama}}\nNIP. {{pelaksana_nip}}"

    doc.add_paragraph()
    add_paragraph(doc, "Lunas Dibayar:")
    add_paragraph(doc, "Bendahara Pengeluaran,")
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "{{bendahara_nama}}")
    add_paragraph(doc, "NIP. {{bendahara_nip}}")

    doc.save(os.path.join(TEMPLATES_DIR, "kuitansi_rampung.docx"))
    print("Created: kuitansi_rampung.docx")


def create_daftar_pengeluaran_riil():
    """Daftar Pengeluaran Riil"""
    doc = Document()

    add_heading(doc, "DAFTAR PENGELUARAN RIIL", 1, True)

    doc.add_paragraph()

    add_paragraph(doc, "Yang bertanda tangan di bawah ini:")

    # Identitas
    table = doc.add_table(rows=4, cols=2)
    items = [
        ("Nama", ": {{pelaksana_nama}}"),
        ("NIP", ": {{pelaksana_nip}}"),
        ("Jabatan", ": {{pelaksana_jabatan}}"),
        ("", ""),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()
    add_paragraph(doc, "Berdasarkan Surat Tugas Nomor: {{nomor_surat_tugas}} tanggal {{tanggal_surat_tugas}}, dengan ini menyatakan dengan sesungguhnya bahwa:")

    doc.add_paragraph()
    add_paragraph(doc, "1. Biaya transportasi/perjalanan dinas dari {{kota_asal}} ke {{kota_tujuan}} pulang pergi yang telah dibayarkan sebesar Rp {{biaya_transport}} adalah sesuai dengan pengeluaran yang sebenarnya;")
    add_paragraph(doc, "2. Biaya penginapan yang dibayarkan sebesar Rp {{biaya_penginapan}} adalah sesuai dengan pengeluaran yang sebenarnya;")
    add_paragraph(doc, "3. Biaya lain-lain (seperti airport tax, retribusi, dll) yang telah dibayarkan sebesar Rp {{biaya_lain_lain}} adalah sesuai dengan pengeluaran yang sebenarnya;")

    doc.add_paragraph()
    add_paragraph(doc, "Demikian pernyataan ini dibuat dengan sebenarnya, dan apabila di kemudian hari terdapat kelebihan atas pembayaran, saya bersedia untuk menyetorkan kelebihan tersebut ke Kas Negara.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Mengetahui/Menyetujui:\n{{ppk_jabatan}}\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"
    table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_kembali}}\nPelaksana SPD,\n\n\n\n{{pelaksana_nama}}\nNIP. {{pelaksana_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "daftar_pengeluaran_riil.docx"))
    print("Created: daftar_pengeluaran_riil.docx")


# =============================================================================
# SWAKELOLA TEMPLATES
# =============================================================================

def create_kak_swakelola():
    """KAK Swakelola"""
    doc = Document()

    add_heading(doc, "KERANGKA ACUAN KERJA (KAK)", 1, True)
    add_heading(doc, "KEGIATAN SWAKELOLA", 2, True)

    doc.add_paragraph()

    add_heading(doc, "I. PENDAHULUAN", 2)
    add_heading(doc, "1.1 Latar Belakang", 3)
    add_paragraph(doc, "{{latar_belakang_swakelola}}")

    add_heading(doc, "II. MAKSUD DAN TUJUAN", 2)
    add_heading(doc, "2.1 Maksud", 3)
    add_paragraph(doc, "{{tujuan_swakelola}}")

    add_heading(doc, "2.2 Tujuan", 3)
    add_paragraph(doc, "{{sasaran_swakelola}}")

    add_heading(doc, "III. RUANG LINGKUP", 2)
    add_paragraph(doc, "{{ruang_lingkup_swakelola}}")

    add_heading(doc, "IV. KELUARAN (OUTPUT)", 2)
    add_paragraph(doc, "{{output_swakelola}}")

    add_heading(doc, "V. PELAKSANAAN KEGIATAN", 2)
    add_heading(doc, "5.1 Tipe Swakelola", 3)
    add_paragraph(doc, "{{tipe_swakelola}}")
    add_heading(doc, "5.2 Penyelenggara", 3)
    add_paragraph(doc, "{{penyelenggara}}")
    add_heading(doc, "5.3 Waktu Pelaksanaan", 3)
    add_paragraph(doc, "{{waktu_pelaksanaan}}")
    add_heading(doc, "5.4 Lokasi Kegiatan", 3)
    add_paragraph(doc, "{{lokasi_kegiatan}}")

    add_heading(doc, "VI. BIAYA", 2)
    add_paragraph(doc, "Total pagu anggaran kegiatan ini adalah sebesar Rp {{pagu_swakelola}} ({{pagu_swakelola_terbilang}}).")
    add_paragraph(doc, "Sumber Dana: {{sumber_dana}}")
    add_paragraph(doc, "Kode Akun: {{kode_akun}}")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=4, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.rows[0].cells[0].text = "{{satker_kota}}, {{tanggal_sk_tim}}"
    table.rows[1].cells[0].text = "{{ppk_jabatan}},"
    table.rows[2].cells[0].text = ""
    table.rows[3].cells[0].text = "{{ppk_nama}}\nNIP. {{ppk_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "kak_swakelola.docx"))
    print("Created: kak_swakelola.docx")


def create_sk_tim_swakelola():
    """SK Tim Pelaksana Swakelola"""
    doc = Document()

    add_heading(doc, "{{kementerian}}", 1, True)
    add_heading(doc, "{{eselon1}}", 2, True)
    add_heading(doc, "{{satker_nama}}", 2, True)

    doc.add_paragraph()

    add_heading(doc, "SURAT KEPUTUSAN", 1, True)
    add_paragraph(doc, "Nomor: {{nomor_sk_tim}}", center=True)

    doc.add_paragraph()
    add_paragraph(doc, "TENTANG", center=True)
    add_heading(doc, "TIM PELAKSANA KEGIATAN SWAKELOLA", 1, True)
    add_paragraph(doc, "{{nama_kegiatan}}", center=True)

    doc.add_paragraph()
    add_paragraph(doc, "{{ppk_jabatan}},", center=True)

    doc.add_paragraph()
    add_heading(doc, "Menimbang:", 2)
    add_paragraph(doc, "a. Bahwa dalam rangka pelaksanaan kegiatan {{nama_kegiatan}} tahun anggaran {{tahun_anggaran}}, perlu dibentuk Tim Pelaksana Swakelola;")
    add_paragraph(doc, "b. Bahwa nama-nama yang tercantum dalam Keputusan ini dipandang cakap dan mampu untuk melaksanakan tugas;")

    add_heading(doc, "Mengingat:", 2)
    add_paragraph(doc, "1. Peraturan Presiden Nomor 16 Tahun 2018 tentang Pengadaan Barang/Jasa Pemerintah sebagaimana telah diubah dengan Peraturan Presiden Nomor 12 Tahun 2021;")
    add_paragraph(doc, "2. Peraturan Lembaga LKPP Nomor 3 Tahun 2021 tentang Pedoman Swakelola;")

    add_heading(doc, "MEMUTUSKAN:", 1)
    add_heading(doc, "Menetapkan:", 2)
    add_paragraph(doc, "PERTAMA: Membentuk Tim Pelaksana Swakelola {{nama_kegiatan}} dengan susunan sebagai berikut:")

    # Tabel Tim
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    hdr = table.rows[0]
    hdr.cells[0].text = "No"
    hdr.cells[1].text = "Nama/NIP"
    hdr.cells[2].text = "Jabatan dalam Tim"
    hdr.cells[3].text = "Jabatan di Instansi"

    items = [
        ("1", "{{ketua_tim_nama}}\nNIP. {{ketua_tim_nip}}", "Ketua", "{{ketua_tim_jabatan}}"),
        ("2", "{{sekretaris_nama}}\nNIP. {{sekretaris_nip}}", "Sekretaris", ""),
        ("3", "{{anggota1_nama}}\nNIP. {{anggota1_nip}}", "Anggota", ""),
        ("4", "{{anggota2_nama}}\nNIP. {{anggota2_nip}}", "Anggota", ""),
        ("5", "{{anggota3_nama}}\nNIP. {{anggota3_nip}}", "Anggota", ""),
    ]

    for i, (no, nama, jabatan_tim, jabatan_inst) in enumerate(items):
        row = table.rows[i + 1]
        row.cells[0].text = no
        row.cells[1].text = nama
        row.cells[2].text = jabatan_tim
        row.cells[3].text = jabatan_inst

    doc.add_paragraph()
    add_paragraph(doc, "KEDUA: Tim Pelaksana Swakelola mempunyai tugas:")
    add_paragraph(doc, "a. Melaksanakan kegiatan {{nama_kegiatan}};")
    add_paragraph(doc, "b. Menyusun laporan pelaksanaan kegiatan;")
    add_paragraph(doc, "c. Bertanggung jawab atas hasil pelaksanaan kegiatan kepada PPK.")

    add_paragraph(doc, "KETIGA: Segala biaya yang timbul akibat pelaksanaan keputusan ini dibebankan pada DIPA {{satker_nama}} tahun anggaran {{tahun_anggaran}}.")

    add_paragraph(doc, "KEEMPAT: Keputusan ini berlaku sejak tanggal ditetapkan.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=4, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.rows[0].cells[0].text = "Ditetapkan di : {{satker_kota}}"
    table.rows[1].cells[0].text = "Pada tanggal   : {{tanggal_sk_tim}}"
    table.rows[2].cells[0].text = "{{ppk_jabatan}},"
    table.rows[3].cells[0].text = "\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "sk_tim_swakelola.docx"))
    print("Created: sk_tim_swakelola.docx")


def create_bap_swakelola():
    """Berita Acara Pembayaran Swakelola"""
    doc = Document()

    add_heading(doc, "BERITA ACARA PEMBAYARAN", 1, True)
    add_paragraph(doc, "Nomor: ............", center=True)

    doc.add_paragraph()

    add_paragraph(doc, "Pada hari ini, .......... tanggal .......... bulan .......... tahun .........., bertempat di {{satker_kota}}, kami yang bertanda tangan di bawah ini:")

    doc.add_paragraph()

    # Pihak Pertama
    table = doc.add_table(rows=4, cols=2)
    items = [
        ("Nama", ": {{ppk_nama}}"),
        ("NIP", ": {{ppk_nip}}"),
        ("Jabatan", ": {{ppk_jabatan}}"),
        ("", "  Selanjutnya disebut PIHAK PERTAMA"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()

    # Pihak Kedua
    table = doc.add_table(rows=4, cols=2)
    items = [
        ("Nama", ": {{ketua_tim_nama}}"),
        ("NIP", ": {{ketua_tim_nip}}"),
        ("Jabatan", ": Ketua Tim Pelaksana Swakelola"),
        ("", "  Selanjutnya disebut PIHAK KEDUA"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()
    add_paragraph(doc, "Menyatakan bahwa PIHAK PERTAMA telah melakukan pembayaran kepada PIHAK KEDUA untuk pelaksanaan kegiatan swakelola:")

    # Detail kegiatan
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    items = [
        ("Nama Kegiatan", ": {{nama_kegiatan}}"),
        ("Pagu Anggaran", ": Rp {{pagu_swakelola}}"),
        ("Nilai Pembayaran", ": Rp {{realisasi_biaya}}"),
        ("Terbilang", ": {{realisasi_biaya_terbilang}}"),
        ("Progres", ": {{progres_persen}}%"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()
    add_paragraph(doc, "Demikian Berita Acara ini dibuat untuk dapat dipergunakan sebagaimana mestinya.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "PIHAK PERTAMA\n{{ppk_jabatan}}\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"
    table.rows[0].cells[1].text = "PIHAK KEDUA\nKetua Tim Swakelola\n\n\n\n{{ketua_tim_nama}}\nNIP. {{ketua_tim_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "bap_swakelola.docx"))
    print("Created: bap_swakelola.docx")


def create_laporan_kemajuan():
    """Laporan Kemajuan Swakelola"""
    doc = Document()

    add_heading(doc, "LAPORAN KEMAJUAN PELAKSANAAN SWAKELOLA", 1, True)

    doc.add_paragraph()

    # Informasi kegiatan
    table = doc.add_table(rows=6, cols=2)
    items = [
        ("Nama Kegiatan", ": {{nama_kegiatan}}"),
        ("Tipe Swakelola", ": {{tipe_swakelola}}"),
        ("Pelaksana", ": {{penyelenggara}}"),
        ("Periode Laporan", ": ..........................."),
        ("Nomor SK Tim", ": {{nomor_sk_tim}}"),
        ("Tanggal SK", ": {{tanggal_sk_tim}}"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()
    add_heading(doc, "I. PENDAHULUAN", 2)
    add_paragraph(doc, "{{latar_belakang_swakelola}}")

    add_heading(doc, "II. TUJUAN KEGIATAN", 2)
    add_paragraph(doc, "{{tujuan_swakelola}}")

    add_heading(doc, "III. PROGRES PELAKSANAAN", 2)

    # Tabel progres
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0]
    hdr.cells[0].text = "No"
    hdr.cells[1].text = "Uraian"
    hdr.cells[2].text = "Keterangan"

    items = [
        ("1", "Progres Fisik", "{{progres_persen}}%"),
        ("2", "Realisasi Anggaran", "Rp {{realisasi_biaya}}"),
        ("3", "Sisa Anggaran", "Rp {{sisa_anggaran}}"),
    ]

    for i, (no, uraian, ket) in enumerate(items):
        row = table.rows[i + 1]
        row.cells[0].text = no
        row.cells[1].text = uraian
        row.cells[2].text = ket

    add_heading(doc, "IV. KEGIATAN YANG TELAH DILAKSANAKAN", 2)
    add_paragraph(doc, "{{keterangan_progres}}")

    add_heading(doc, "V. KENDALA DAN PERMASALAHAN", 2)
    add_paragraph(doc, "[Uraian kendala dan permasalahan yang dihadapi]")

    add_heading(doc, "VI. RENCANA KEGIATAN BERIKUTNYA", 2)
    add_paragraph(doc, "[Uraian rencana kegiatan periode berikutnya]")

    doc.add_paragraph()
    add_paragraph(doc, "Demikian laporan kemajuan ini dibuat dengan sebenarnya.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Mengetahui,\n{{ppk_jabatan}}\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"
    table.rows[0].cells[1].text = "{{satker_kota}}, .....................\nKetua Tim Swakelola,\n\n\n\n{{ketua_tim_nama}}\nNIP. {{ketua_tim_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "laporan_kemajuan.docx"))
    print("Created: laporan_kemajuan.docx")


def create_bast_swakelola():
    """BAST Swakelola"""
    doc = Document()

    add_heading(doc, "BERITA ACARA SERAH TERIMA HASIL PEKERJAAN SWAKELOLA", 1, True)
    add_paragraph(doc, "Nomor: ............", center=True)

    doc.add_paragraph()

    add_paragraph(doc, "Pada hari ini, .......... tanggal .......... bulan .......... tahun .........., bertempat di {{satker_kota}}, kami yang bertanda tangan di bawah ini:")

    doc.add_paragraph()

    # Pihak Pertama
    table = doc.add_table(rows=4, cols=2)
    items = [
        ("Nama", ": {{ppk_nama}}"),
        ("NIP", ": {{ppk_nip}}"),
        ("Jabatan", ": {{ppk_jabatan}}"),
        ("", "  Selanjutnya disebut PIHAK PERTAMA"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()

    # Pihak Kedua
    table = doc.add_table(rows=4, cols=2)
    items = [
        ("Nama", ": {{ketua_tim_nama}}"),
        ("NIP", ": {{ketua_tim_nip}}"),
        ("Jabatan", ": Ketua Tim Pelaksana Swakelola"),
        ("", "  Selanjutnya disebut PIHAK KEDUA"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()
    add_paragraph(doc, "Berdasarkan SK Tim Pelaksana Swakelola Nomor: {{nomor_sk_tim}} tanggal {{tanggal_sk_tim}}, dengan ini PIHAK KEDUA menyerahkan hasil pekerjaan swakelola kepada PIHAK PERTAMA dengan rincian sebagai berikut:")

    # Detail kegiatan
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    items = [
        ("Nama Kegiatan", ": {{nama_kegiatan}}"),
        ("Tipe Swakelola", ": {{tipe_swakelola}}"),
        ("Waktu Pelaksanaan", ": {{waktu_pelaksanaan}}"),
        ("Lokasi", ": {{lokasi_kegiatan}}"),
        ("Output/Keluaran", ": {{output_swakelola}}"),
        ("Pagu Anggaran", ": Rp {{pagu_swakelola}}"),
        ("Realisasi Anggaran", ": Rp {{realisasi_biaya}} ({{realisasi_biaya_terbilang}})"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()
    add_paragraph(doc, "PIHAK PERTAMA telah menerima hasil pekerjaan swakelola tersebut dalam keadaan baik dan lengkap.")

    add_paragraph(doc, "Demikian Berita Acara Serah Terima ini dibuat dengan sebenarnya untuk dapat dipergunakan sebagaimana mestinya.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Yang Menerima,\nPIHAK PERTAMA\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"
    table.rows[0].cells[1].text = "Yang Menyerahkan,\nPIHAK KEDUA\n\n\n\n{{ketua_tim_nama}}\nNIP. {{ketua_tim_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "bast_swakelola.docx"))
    print("Created: bast_swakelola.docx")


def main():
    print("Creating Perjalanan Dinas templates...")
    create_surat_tugas()
    create_sppd()
    create_kuitansi_uang_muka()
    create_rincian_biaya_pd()
    create_laporan_perjalanan_dinas()
    create_kuitansi_rampung()
    create_daftar_pengeluaran_riil()

    print("\nCreating Swakelola templates...")
    create_kak_swakelola()
    create_sk_tim_swakelola()
    create_bap_swakelola()
    create_laporan_kemajuan()
    create_bast_swakelola()

    print("\nAll templates created successfully!")


if __name__ == "__main__":
    main()
