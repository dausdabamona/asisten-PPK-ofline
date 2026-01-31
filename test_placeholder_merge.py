"""
Test script untuk debug placeholder replacement di lembar_permintaan.docx
"""
from pathlib import Path
from docx import Document
import sys

# Add app to path
BASE_DIR = Path(__file__).parent
sys.path.insert(0, str(BASE_DIR))

from app.services.dokumen_generator import DokumenGenerator

# Initialize generator
generator = DokumenGenerator()

# Test data matching template placeholders
test_data = {
    'hari_tanggal': '15 Januari 2024',
    'unit_kerja': 'Dinas Pendidikan Kota',
    'sumber_dana': 'APBD 2024',
    'subtotal': 150000,
    'ppn': 15000,
    'total': 165000,
    'nama_pengajuan': 'Budi Santoso',
    'nama_verifikator': 'Siti Rahayui',
    'nama_ppk': 'Dr. Ahmad Wijaya',
    'nama_atasan': 'Kepala Dinas',
    'nama_kpa': 'Pejabat Pembuat Komitmen',
    'item_1_no': '1',
    'item_1_nama': 'Kertas A4',
    'item_1_spek': 'Kertas putih 80gsm',
    'item_1_volume': 10,
    'item_1_satuan': 'rim',
    'item_1_harga': 'Rp 50.000',
    'item_1_total': 'Rp 500.000',
    'item_1_ket': 'Untuk administrasi',
    # Empty items to test placeholder clearing
    'item_2_no': '',
    'item_2_nama': '',
    'item_2_spek': '',
    'item_2_volume': '',
    'item_2_satuan': '',
    'item_2_harga': '',
    'item_2_total': '',
    'item_2_ket': '',
    'item_3_no': '',
    'item_3_nama': '',
    'item_3_spek': '',
    'item_3_volume': '',
    'item_3_satuan': '',
    'item_3_harga': '',
    'item_3_total': '',
    'item_3_ket': '',
    'item_4_no': '',
    'item_4_nama': '',
    'item_4_spek': '',
    'item_4_volume': '',
    'item_4_satuan': '',
    'item_4_harga': '',
    'item_4_total': '',
    'item_4_ket': '',
    'item_5_no': '',
    'item_5_nama': '',
    'item_5_spek': '',
    'item_5_volume': '',
    'item_5_satuan': '',
    'item_5_harga': '',
    'item_5_total': '',
    'item_5_ket': '',
}

template_path = BASE_DIR / "templates" / "word" / "lembar_permintaan.docx"
output_path = BASE_DIR / "output" / "dokumen" / "test_placeholder.docx"

print(f"Template: {template_path}")
print(f"Output: {output_path}")
print(f"Test data keys: {len(test_data)}")

if template_path.exists():
    print("\n✓ Template found")
    
    # Create output directory
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Test merge
    success = generator.merge_word_document(template_path, test_data, output_path)
    
    if success:
        print("✓ Merge successful, opening result...")
        
        # Check result
        doc = Document(str(output_path))
        
        print("\n=== CHECKING RESULTS ===")
        
        # Check paragraphs
        found_placeholders = []
        for para in doc.paragraphs:
            if '{{' in para.text:
                found_placeholders.append(('paragraph', para.text[:100]))
        
        # Check tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if '{{' in cell.text:
                        found_placeholders.append((f'table_{table_idx}', f"[{row_idx},{cell_idx}]: {cell.text[:80]}"))
        
        if found_placeholders:
            print(f"\n⚠ FOUND {len(found_placeholders)} UNREPLACED PLACEHOLDERS:")
            for location, text in found_placeholders:
                print(f"  {location}: {text}")
        else:
            print("\n✓ NO UNREPLACED PLACEHOLDERS FOUND!")
        
        # Check if values are in document
        print("\n=== CHECKING IF VALUES ARE IN DOCUMENT ===")
        doc_text = '\n'.join([para.text for para in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    doc_text += ' ' + cell.text
        
        checks = [
            ('Dinas Pendidikan Kota', 'unit_kerja'),
            ('15 Januari 2024', 'hari_tanggal'),
            ('Budi Santoso', 'nama_pengajuan'),
            ('Kertas A4', 'item_1_nama'),
            ('Rp 500.000', 'item_1_total'),
        ]
        
        for value, name in checks:
            if value in doc_text:
                print(f"  ✓ {name}: '{value}'")
            else:
                print(f"  ✗ {name}: '{value}' NOT FOUND")
        
        print(f"\n✓ Test document generated: {output_path}")
        
    else:
        print("✗ Merge failed")
else:
    print(f"✗ Template not found: {template_path}")
