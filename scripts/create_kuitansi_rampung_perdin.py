#!/usr/bin/env python3
"""
Script untuk membuat template Kuitansi Rampung Perjalanan Dinas.
Format khusus untuk pertanggungjawaban belanja perjalanan dinas.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_single_spacing(paragraph):
    """Set paragraph to single spacing with no space before/after."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    borders = []
    if top:
        borders.append('top')
    if bottom:
        borders.append('bottom')
    if left:
        borders.append('left')
    if right:
        borders.append('right')

    for border_name in borders:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color='D9D9D9'):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)


def create_kuitansi_rampung_perdin():
    """Create Kuitansi Rampung template for Perjalanan Dinas."""
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ========== HEADER: KUITANSI RAMPUNG PERJALANAN DINAS ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(title)
    title_run = title.add_run("KUITANSI RAMPUNG")
    title_run.bold = True
    title_run.font.size = Pt(14)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(subtitle)
    sub_run = subtitle.add_run("PERTANGGUNGJAWABAN BIAYA PERJALANAN DINAS")
    sub_run.font.size = Pt(11)

    doc.add_paragraph()

    # ========== INFO TABLE ==========
    info_table = doc.add_table(rows=7, cols=3)
    info_table.autofit = False

    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(11)

    info_data = [
        ("Nomor Kuitansi", ":", "{{nomor_kuitansi}}"),
        ("Tanggal", ":", "{{tanggal_dokumen:tanggal_panjang}}"),
        ("Nama Kegiatan", ":", "{{nama_kegiatan}}"),
        ("Tujuan Perjalanan", ":", "{{perdin_tujuan}}"),
        ("Maksud Perjalanan", ":", "{{perdin_maksud}}"),
        ("Tanggal Pelaksanaan", ":", "{{tanggal_kegiatan_mulai:tanggal}} s.d. {{tanggal_kegiatan_selesai:tanggal}}"),
        ("Sumber Dana/MAK", ":", "{{kode_akun}}"),
    ]

    for i, (label, sep, value) in enumerate(info_data):
        row = info_table.rows[i]

        # Label cell
        cell_label = row.cells[0]
        p = cell_label.paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(label)
        run.font.size = Pt(10)

        # Separator cell
        cell_sep = row.cells[1]
        p = cell_sep.paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(sep)
        run.font.size = Pt(10)

        # Value cell
        cell_value = row.cells[2]
        p = cell_value.paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(value)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ========== DATA PELAKSANA ==========
    pelaksana_title = doc.add_paragraph()
    set_single_spacing(pelaksana_title)
    run = pelaksana_title.add_run("DATA PELAKSANA PERJALANAN DINAS:")
    run.bold = True
    run.font.size = Pt(10)

    pelaksana_table = doc.add_table(rows=4, cols=3)
    pelaksana_table.autofit = False

    for row in pelaksana_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(11)

    pelaksana_data = [
        ("Nama", ":", "{{penerima_nama}}"),
        ("NIP", ":", "{{penerima_nip}}"),
        ("Jabatan", ":", "{{penerima_jabatan}}"),
        ("Tingkat Biaya Perdin", ":", "Tingkat {{perdin_tingkat}}"),
    ]

    for i, (label, sep, value) in enumerate(pelaksana_data):
        row = pelaksana_table.rows[i]

        p = row.cells[0].paragraphs[0]
        set_single_spacing(p)
        p.add_run(label).font.size = Pt(10)

        p = row.cells[1].paragraphs[0]
        set_single_spacing(p)
        p.add_run(sep).font.size = Pt(10)

        p = row.cells[2].paragraphs[0]
        set_single_spacing(p)
        p.add_run(value).font.size = Pt(10)

    doc.add_paragraph()

    # ========== RINCIAN BIAYA PERJALANAN DINAS ==========
    rincian_title = doc.add_paragraph()
    set_single_spacing(rincian_title)
    run = rincian_title.add_run("RINCIAN BIAYA PERJALANAN DINAS:")
    run.bold = True
    run.font.size = Pt(10)

    # Rincian table
    rincian_table = doc.add_table(rows=1, cols=5)
    rincian_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["No", "Uraian Biaya", "Keterangan", "Jumlah Hari/Unit", "Jumlah (Rp)"]
    header_widths = [Cm(0.8), Cm(5), Cm(4), Cm(2.5), Cm(3.5)]

    for i, (header, width) in enumerate(zip(headers, header_widths)):
        cell = rincian_table.rows[0].cells[i]
        cell.width = width
        set_cell_borders(cell)
        set_cell_shading(cell, 'E8E8E8')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_single_spacing(p)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)

    # Data rows - komponen biaya perjalanan dinas
    komponen_biaya = [
        ("1", "Uang Harian", "{{perdin_lama_hari}} hari x {{tarif_uang_harian:rupiah}}", "{{perdin_lama_hari}}", "{{biaya_uang_harian:rupiah}}"),
        ("2", "Biaya Transport", "{{perdin_alat_angkut}}", "1", "{{biaya_transport:rupiah}}"),
        ("3", "Biaya Penginapan", "{{perdin_lama_hari_inap}} malam x {{tarif_penginapan:rupiah}}", "{{perdin_lama_hari_inap}}", "{{biaya_penginapan:rupiah}}"),
        ("4", "Uang Representatif", "{{keterangan_representatif}}", "-", "{{biaya_representatif:rupiah}}"),
        ("5", "Biaya Lain-lain", "{{keterangan_lainnya}}", "-", "{{biaya_lainnya:rupiah}}"),
    ]

    for no, uraian, keterangan, qty, jumlah in komponen_biaya:
        row = rincian_table.add_row()
        data = [no, uraian, keterangan, qty, jumlah]
        for i, (value, width) in enumerate(zip(data, header_widths)):
            cell = row.cells[i]
            cell.width = width
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            if i == 0 or i == 3:  # No, Qty - center
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 4:  # Jumlah - right
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_single_spacing(p)
            run = p.add_run(value)
            run.font.size = Pt(9)

    # Total row
    total_row = rincian_table.add_row()
    # Merge first 4 cells for "TOTAL"
    total_label_cell = total_row.cells[0]
    set_cell_borders(total_label_cell)
    set_cell_shading(total_label_cell, 'E8E8E8')
    p = total_label_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_single_spacing(p)
    run = p.add_run("TOTAL BIAYA PERJALANAN DINAS")
    run.bold = True
    run.font.size = Pt(9)

    # Merge cells 0-3
    for i in range(1, 4):
        cell = total_row.cells[i]
        set_cell_borders(cell)
        set_cell_shading(cell, 'E8E8E8')

    total_row.cells[0].merge(total_row.cells[1])
    total_row.cells[0].merge(total_row.cells[2])
    total_row.cells[0].merge(total_row.cells[3])

    # Total value
    total_value_cell = total_row.cells[4]
    set_cell_borders(total_value_cell)
    set_cell_shading(total_value_cell, 'E8E8E8')
    p = total_value_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_single_spacing(p)
    run = p.add_run("{{total_realisasi:rupiah}}")
    run.bold = True
    run.font.size = Pt(10)

    # Terbilang
    terbilang_p = doc.add_paragraph()
    terbilang_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(terbilang_p)
    run = terbilang_p.add_run("({{total_realisasi:terbilang}})")
    run.italic = True
    run.font.size = Pt(9)

    doc.add_paragraph()

    # ========== PERHITUNGAN SELISIH ==========
    selisih_title = doc.add_paragraph()
    set_single_spacing(selisih_title)
    run = selisih_title.add_run("PERHITUNGAN:")
    run.bold = True
    run.font.size = Pt(10)

    # Selisih table
    selisih_table = doc.add_table(rows=3, cols=3)
    selisih_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    selisih_data = [
        ("Uang Muka Diterima", ":", "{{uang_muka:rupiah}}"),
        ("Total Realisasi Biaya", ":", "{{total_realisasi:rupiah}}"),
        ("Selisih", ":", "{{selisih:rupiah}}"),
    ]

    for i, (label, sep, value) in enumerate(selisih_data):
        row = selisih_table.rows[i]
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(4)

        # Label
        p = row.cells[0].paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(label)
        run.font.size = Pt(10)
        if i == 2:
            run.bold = True

        # Separator
        p = row.cells[1].paragraphs[0]
        set_single_spacing(p)
        p.add_run(sep).font.size = Pt(10)

        # Value
        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_single_spacing(p)
        run = p.add_run(value)
        run.font.size = Pt(10)
        if i == 2:
            run.bold = True

    # Keterangan selisih
    ket_p = doc.add_paragraph()
    set_single_spacing(ket_p)
    run = ket_p.add_run("Keterangan: {{keterangan_selisih}}")
    run.font.size = Pt(9)
    run.italic = True

    doc.add_paragraph()

    # ========== PERNYATAAN ==========
    pernyataan_p = doc.add_paragraph()
    set_single_spacing(pernyataan_p)
    run = pernyataan_p.add_run(
        "Demikian kuitansi rampung ini dibuat dengan sebenarnya untuk dipergunakan "
        "sebagai bukti pertanggungjawaban biaya perjalanan dinas."
    )
    run.font.size = Pt(9)

    doc.add_paragraph()

    # ========== SIGNATURES (2 columns) ==========
    sig_table = doc.add_table(rows=5, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in sig_table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(7)

    # Row 0: Headers
    left_header = sig_table.rows[0].cells[0]
    p = left_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("Pelaksana Perjalanan Dinas,")
    run.font.size = Pt(10)

    right_header = sig_table.rows[0].cells[1]
    p = right_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("Mengetahui,")
    run.font.size = Pt(10)

    # Row 1: Jabatan
    left_jabatan = sig_table.rows[1].cells[0]
    p = left_jabatan.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("{{penerima_jabatan}}")
    run.font.size = Pt(9)

    right_jabatan = sig_table.rows[1].cells[1]
    p = right_jabatan.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("Pejabat Pembuat Komitmen")
    run.font.size = Pt(9)

    # Row 2: Empty space for signature
    for cell in sig_table.rows[2].cells:
        p = cell.paragraphs[0]
        set_single_spacing(p)
        p.add_run("\n\n\n")

    # Row 3: Names (underlined)
    left_name = sig_table.rows[3].cells[0]
    p = left_name.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("{{penerima_nama}}")
    run.font.size = Pt(10)
    run.underline = True

    right_name = sig_table.rows[3].cells[1]
    p = right_name.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("{{ppk_nama}}")
    run.font.size = Pt(10)
    run.underline = True

    # Row 4: NIP
    left_nip = sig_table.rows[4].cells[0]
    p = left_nip.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("NIP. {{penerima_nip}}")
    run.font.size = Pt(9)

    right_nip = sig_table.rows[4].cells[1]
    p = right_nip.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("NIP. {{ppk_nip}}")
    run.font.size = Pt(9)

    # Save - replace existing kuitansi_rampung.docx
    output_path = Path(__file__).parent.parent / "templates" / "word" / "kuitansi_rampung.docx"
    doc.save(str(output_path))
    print(f"Template saved to: {output_path}")

    return output_path


if __name__ == "__main__":
    create_kuitansi_rampung_perdin()
