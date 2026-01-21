#!/usr/bin/env python3
"""
PPK DOCUMENT FACTORY v3.0 - COMPLETE TEMPLATE CREATOR
======================================================
Membuat semua template dokumen pengadaan dengan format standar pemerintah
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# Output directory
TEMPLATE_DIR = os.path.dirname(os.path.abspath(__file__))
WORD_DIR = os.path.join(TEMPLATE_DIR, "templates", "word")
EXCEL_DIR = os.path.join(TEMPLATE_DIR, "templates", "excel")

os.makedirs(WORD_DIR, exist_ok=True)
os.makedirs(EXCEL_DIR, exist_ok=True)

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:color="000000"/>'
        r'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def create_doc_with_header():
    """Create document with standard header"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Add header
    header = doc.sections[0].header
    h_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    
    # Left: Logo placeholder
    left_cell = h_table.rows[0].cells[0]
    left_cell.width = Inches(1)
    p = left_cell.paragraphs[0]
    p.add_run("[LOGO]")
    
    # Right: Kop surat
    right_cell = h_table.rows[0].cells[1]
    right_cell.width = Inches(5.5)
    
    p1 = right_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("{{satker_kementerian}}")
    run1.bold = True
    run1.font.size = Pt(12)
    
    p2 = right_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("{{satker_eselon1}}")
    run2.bold = True
    run2.font.size = Pt(12)
    
    p3 = right_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("{{satker_nama}}")
    run3.bold = True
    run3.font.size = Pt(14)
    
    p4 = right_cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("{{satker_alamat}}, {{satker_kota}}")
    run4.font.size = Pt(10)
    
    p5 = right_cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Telepon: {{satker_telepon}} | Email: {{satker_email}}")
    run5.font.size = Pt(10)
    
    # Line under header
    p6 = header.add_paragraph()
    p6.add_run("_" * 85)
    
    return doc

def add_ttd_ppk(doc):
    """Add PPK signature block"""
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    # Right column for signature
    cell = table.rows[0].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("{{satker_kota}}, {{tanggal_dokumen:tanggal_long}}")
    
    cell = table.rows[1].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Pejabat Pembuat Komitmen,")
    
    cell = table.rows[2].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n\n")  # Space for signature
    
    cell = table.rows[3].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ppk_nama}}")
    run.bold = True
    run.underline = True
    
    cell = table.rows[4].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("NIP. {{ppk_nip:nip}}")

def add_ttd_dual(doc, pihak1_label="PIHAK PERTAMA", pihak2_label="PIHAK KEDUA"):
    """Add dual signature block"""
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Pihak 1
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(pihak1_label)
    run.bold = True
    
    cell = table.rows[1].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Pejabat Pembuat Komitmen,")
    
    cell = table.rows[2].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n\n")
    
    cell = table.rows[3].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ppk_nama}}")
    run.bold = True
    run.underline = True
    
    cell = table.rows[4].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("NIP. {{ppk_nip:nip}}")
    
    # Pihak 2
    cell = table.rows[0].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(pihak2_label)
    run.bold = True
    
    cell = table.rows[1].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("{{direktur_jabatan}},")
    
    cell = table.rows[2].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n\n")
    
    cell = table.rows[3].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{direktur_nama}}")
    run.bold = True
    run.underline = True
    
    cell = table.rows[4].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("{{penyedia_nama}}")


# ============================================================================
# WORD TEMPLATES
# ============================================================================

def create_spesifikasi_teknis():
    """Template Spesifikasi Teknis"""
    doc = create_doc_with_header()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SPESIFIKASI TEKNIS")
    run.bold = True
    run.font.size = Pt(14)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("{{nama_paket}}")
    
    doc.add_paragraph()
    
    # Content sections
    sections = [
        ("A.", "LATAR BELAKANG", "{{latar_belakang}}"),
        ("B.", "MAKSUD DAN TUJUAN", "{{maksud_tujuan}}"),
        ("C.", "TARGET/SASARAN", "{{target_sasaran}}"),
        ("D.", "RUANG LINGKUP PEKERJAAN", "{{ruang_lingkup}}"),
        ("E.", "SPESIFIKASI TEKNIS BARANG/JASA", "{{spesifikasi_teknis}}"),
        ("F.", "WAKTU PELAKSANAAN", "Pekerjaan dilaksanakan selama {{jangka_waktu}} ({{jangka_waktu_terbilang}}) hari kalender terhitung sejak tanggal Surat Perintah Mulai Kerja (SPMK) diterbitkan."),
        ("G.", "SUMBER DANA", "Pekerjaan ini dibiayai dari {{sumber_dana}} {{satker_nama}} Tahun Anggaran {{tahun_anggaran}} dengan Kode Akun {{kode_akun}}."),
    ]
    
    for num, title_text, content in sections:
        p = doc.add_paragraph()
        run = p.add_run(f"{num} {title_text}")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run(content)
        doc.add_paragraph()
    
    # Spesifikasi table
    p = doc.add_paragraph()
    run = p.add_run("H. RINCIAN SPESIFIKASI BARANG/JASA")
    run.bold = True
    
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    
    headers = ["No", "Uraian Barang/Jasa", "Spesifikasi", "Satuan", "Volume", "Keterangan"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sample row with placeholders
    sample_data = ["1", "{{item.nama_item}}", "{{item.spesifikasi}}", "{{item.satuan}}", "{{item.volume}}", "{{item.keterangan}}"]
    for i, data in enumerate(sample_data):
        table.rows[1].cells[i].text = data
    
    doc.add_paragraph()
    
    # Syarat Penyedia
    p = doc.add_paragraph()
    run = p.add_run("I. SYARAT-SYARAT PENYEDIA")
    run.bold = True
    
    syarat = [
        "Memiliki Surat Izin Usaha sesuai bidang pekerjaan",
        "Memiliki NPWP dan telah memenuhi kewajiban perpajakan",
        "Tidak dalam pengawasan pengadilan, tidak pailit, atau kegiatan usahanya tidak sedang dihentikan",
        "Memiliki pengalaman dalam pekerjaan sejenis",
        "Menyediakan barang/jasa sesuai spesifikasi yang diminta",
    ]
    
    for i, s in enumerate(syarat, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {s}")
    
    add_ttd_ppk(doc)
    
    filepath = os.path.join(WORD_DIR, "spesifikasi_teknis.docx")
    doc.save(filepath)
    print(f"✅ Created: spesifikasi_teknis.docx")
    return filepath


def create_kak():
    """Template Kerangka Acuan Kerja (KAK)"""
    doc = create_doc_with_header()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KERANGKA ACUAN KERJA (KAK)")
    run.bold = True
    run.font.size = Pt(14)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("{{nama_paket}}")
    
    doc.add_paragraph()
    
    # Info table
    info_table = doc.add_table(rows=6, cols=3)
    info_data = [
        ("Kementerian/Lembaga", ":", "{{satker_kementerian}}"),
        ("Unit Eselon I", ":", "{{satker_eselon1}}"),
        ("Satuan Kerja", ":", "{{satker_nama}}"),
        ("Tahun Anggaran", ":", "{{tahun_anggaran}}"),
        ("Sumber Dana", ":", "{{sumber_dana}}"),
        ("Kode Akun/MAK", ":", "{{kode_akun}}"),
    ]
    
    for i, (label, sep, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = sep
        info_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Content sections
    sections = [
        ("1.", "LATAR BELAKANG", "{{latar_belakang}}"),
        ("2.", "MAKSUD DAN TUJUAN", """
Maksud:
{{maksud_pekerjaan}}

Tujuan:
{{tujuan_pekerjaan}}
"""),
        ("3.", "TARGET/SASARAN", "{{target_sasaran}}"),
        ("4.", "RUANG LINGKUP PEKERJAAN", "{{ruang_lingkup}}"),
        ("5.", "PRODUK/KELUARAN (OUTPUT)", "{{output_pekerjaan}}"),
        ("6.", "WAKTU PELAKSANAAN", "Pekerjaan dilaksanakan selama {{jangka_waktu}} ({{jangka_waktu_terbilang}}) hari kalender."),
        ("7.", "TENAGA AHLI/PELAKSANA", "{{tenaga_pelaksana}}"),
        ("8.", "PERALATAN", "{{peralatan}}"),
        ("9.", "METODE PELAKSANAAN", "{{metode_pelaksanaan}}"),
        ("10.", "LAPORAN", "{{ketentuan_laporan}}"),
    ]
    
    for num, title_text, content in sections:
        p = doc.add_paragraph()
        run = p.add_run(f"{num} {title_text}")
        run.bold = True
        
        p = doc.add_paragraph()
        p.add_run(content)
        doc.add_paragraph()
    
    # Biaya
    p = doc.add_paragraph()
    run = p.add_run("11. BIAYA")
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run("Total biaya yang diperlukan untuk pekerjaan ini adalah sebesar:")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{nilai_pagu:rupiah}}")
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("({{nilai_pagu:terbilang}})")
    
    add_ttd_ppk(doc)
    
    filepath = os.path.join(WORD_DIR, "kak.docx")
    doc.save(filepath)
    print(f"✅ Created: kak.docx")
    return filepath


def create_spk():
    """Template Surat Perintah Kerja (SPK)"""
    doc = create_doc_with_header()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SURAT PERINTAH KERJA (SPK)")
    run.bold = True
    run.font.size = Pt(14)
    
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_spk}}")
    
    doc.add_paragraph()
    
    # Intro
    p = doc.add_paragraph()
    p.add_run("Yang bertanda tangan di bawah ini:")
    
    doc.add_paragraph()
    
    # Pihak Pertama
    p = doc.add_paragraph()
    run = p.add_run("I. PIHAK PERTAMA")
    run.bold = True
    
    pihak1_table = doc.add_table(rows=4, cols=3)
    pihak1_data = [
        ("Nama", ":", "{{ppk_nama}}"),
        ("Jabatan", ":", "Pejabat Pembuat Komitmen"),
        ("Alamat", ":", "{{satker_alamat}}, {{satker_kota}}"),
        ("", "", "Selanjutnya disebut PIHAK PERTAMA"),
    ]
    for i, (label, sep, value) in enumerate(pihak1_data):
        pihak1_table.rows[i].cells[0].text = label
        pihak1_table.rows[i].cells[1].text = sep
        pihak1_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Pihak Kedua
    p = doc.add_paragraph()
    run = p.add_run("II. PIHAK KEDUA")
    run.bold = True
    
    pihak2_table = doc.add_table(rows=5, cols=3)
    pihak2_data = [
        ("Nama Perusahaan", ":", "{{penyedia_nama}}"),
        ("Nama Direktur", ":", "{{direktur_nama}}"),
        ("Jabatan", ":", "{{direktur_jabatan}}"),
        ("Alamat", ":", "{{penyedia_alamat}}, {{penyedia_kota}}"),
        ("", "", "Selanjutnya disebut PIHAK KEDUA"),
    ]
    for i, (label, sep, value) in enumerate(pihak2_data):
        pihak2_table.rows[i].cells[0].text = label
        pihak2_table.rows[i].cells[1].text = sep
        pihak2_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Ketentuan
    p = doc.add_paragraph()
    p.add_run("Kedua belah pihak sepakat untuk mengikatkan diri dalam Surat Perintah Kerja untuk melaksanakan pekerjaan dengan ketentuan sebagai berikut:")
    
    doc.add_paragraph()
    
    # Detail pekerjaan
    detail_table = doc.add_table(rows=9, cols=3)
    detail_table.style = 'Table Grid'
    
    detail_data = [
        ("1", "Nama Pekerjaan", "{{nama_paket}}"),
        ("2", "Lingkup Pekerjaan", "{{ruang_lingkup}}"),
        ("3", "Nilai Pekerjaan", "{{nilai_kontrak:rupiah}}\n({{nilai_kontrak:terbilang}})"),
        ("4", "PPN 11%", "{{nilai_ppn:rupiah}}"),
        ("5", "Nilai Total", "{{nilai_bruto:rupiah}}\n({{nilai_bruto:terbilang}})"),
        ("6", "Sumber Dana", "{{sumber_dana}} TA {{tahun_anggaran}}"),
        ("7", "Kode Akun/MAK", "{{kode_akun}}"),
        ("8", "Waktu Pelaksanaan", "{{jangka_waktu}} hari kalender"),
        ("9", "Masa Berlaku SPK", "{{tanggal_mulai:tanggal_long}} s.d. {{tanggal_selesai:tanggal_long}}"),
    ]
    
    for i, (no, label, value) in enumerate(detail_data):
        detail_table.rows[i].cells[0].text = no
        detail_table.rows[i].cells[1].text = label
        detail_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Ketentuan tambahan
    p = doc.add_paragraph()
    run = p.add_run("KETENTUAN DAN SYARAT-SYARAT:")
    run.bold = True
    
    ketentuan = [
        "PIHAK KEDUA wajib melaksanakan pekerjaan sesuai dengan spesifikasi teknis yang telah ditetapkan.",
        "PIHAK KEDUA wajib menyelesaikan pekerjaan dalam jangka waktu yang telah ditentukan.",
        "Apabila PIHAK KEDUA tidak dapat menyelesaikan pekerjaan sesuai waktu yang ditentukan, maka akan dikenakan denda keterlambatan sebesar 1/1000 (satu per seribu) dari nilai kontrak untuk setiap hari keterlambatan.",
        "Pembayaran dilakukan setelah pekerjaan selesai 100% dan diterima dengan baik oleh PIHAK PERTAMA.",
        "PIHAK KEDUA wajib memberikan garansi atas barang/jasa yang diserahkan sesuai ketentuan yang berlaku.",
        "SPK ini mulai berlaku sejak tanggal ditandatangani oleh kedua belah pihak.",
    ]
    
    for i, k in enumerate(ketentuan, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {k}")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Demikian Surat Perintah Kerja ini dibuat untuk dilaksanakan dengan penuh tanggung jawab.")
    
    add_ttd_dual(doc)
    
    filepath = os.path.join(WORD_DIR, "spk.docx")
    doc.save(filepath)
    print(f"✅ Created: spk.docx")
    return filepath


def create_spmk():
    """Template Surat Perintah Mulai Kerja (SPMK)"""
    doc = create_doc_with_header()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SURAT PERINTAH MULAI KERJA (SPMK)")
    run.bold = True
    run.font.size = Pt(14)
    
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_spmk}}")
    
    doc.add_paragraph()
    
    # Intro
    p = doc.add_paragraph()
    p.add_run("Yang bertanda tangan di bawah ini:")
    
    # PPK Info
    ppk_table = doc.add_table(rows=3, cols=3)
    ppk_data = [
        ("Nama", ":", "{{ppk_nama}}"),
        ("Jabatan", ":", "Pejabat Pembuat Komitmen"),
        ("Alamat", ":", "{{satker_nama}}, {{satker_alamat}}"),
    ]
    for i, (label, sep, value) in enumerate(ppk_data):
        ppk_table.rows[i].cells[0].text = label
        ppk_table.rows[i].cells[1].text = sep
        ppk_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Berdasarkan Surat Perintah Kerja (SPK):")
    
    # SPK Reference
    spk_table = doc.add_table(rows=2, cols=3)
    spk_data = [
        ("Nomor SPK", ":", "{{nomor_spk}}"),
        ("Tanggal SPK", ":", "{{tanggal_spk:tanggal_long}}"),
    ]
    for i, (label, sep, value) in enumerate(spk_data):
        spk_table.rows[i].cells[0].text = label
        spk_table.rows[i].cells[1].text = sep
        spk_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("MEMERINTAHKAN")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Kepada:")
    
    # Penyedia Info
    penyedia_table = doc.add_table(rows=4, cols=3)
    penyedia_data = [
        ("Nama Perusahaan", ":", "{{penyedia_nama}}"),
        ("Direktur", ":", "{{direktur_nama}}"),
        ("Alamat", ":", "{{penyedia_alamat}}, {{penyedia_kota}}"),
        ("NPWP", ":", "{{penyedia_npwp:npwp}}"),
    ]
    for i, (label, sep, value) in enumerate(penyedia_data):
        penyedia_table.rows[i].cells[0].text = label
        penyedia_table.rows[i].cells[1].text = sep
        penyedia_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Untuk segera memulai pelaksanaan pekerjaan:")
    
    # Detail pekerjaan
    detail_table = doc.add_table(rows=5, cols=3)
    detail_table.style = 'Table Grid'
    detail_data = [
        ("1", "Nama Pekerjaan", "{{nama_paket}}"),
        ("2", "Nilai Pekerjaan", "{{nilai_kontrak:rupiah}}"),
        ("3", "Waktu Pelaksanaan", "{{jangka_waktu}} hari kalender"),
        ("4", "Tanggal Mulai", "{{tanggal_mulai:tanggal_long}}"),
        ("5", "Tanggal Selesai", "{{tanggal_selesai:tanggal_long}}"),
    ]
    for i, (no, label, value) in enumerate(detail_data):
        detail_table.rows[i].cells[0].text = no
        detail_table.rows[i].cells[1].text = label
        detail_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Demikian Surat Perintah Mulai Kerja ini diterbitkan untuk dilaksanakan sebagaimana mestinya.")
    
    add_ttd_ppk(doc)
    
    # Penerima
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Diterima tanggal: {{tanggal_terima_spmk:tanggal_long}}")
    
    table = doc.add_table(rows=4, cols=1)
    table.rows[0].cells[0].text = "Yang Menerima,"
    table.rows[1].cells[0].text = "{{penyedia_nama}}"
    table.rows[2].cells[0].text = "\n\n\n"
    table.rows[3].cells[0].text = "{{direktur_nama}}"
    for row in table.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(WORD_DIR, "spmk.docx")
    doc.save(filepath)
    print(f"✅ Created: spmk.docx")
    return filepath


def create_bahp():
    """Template Berita Acara Hasil Pemeriksaan (BAHP)"""
    doc = create_doc_with_header()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BERITA ACARA HASIL PEMERIKSAAN")
    run.bold = True
    run.font.size = Pt(14)
    
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_bahp}}")
    
    doc.add_paragraph()
    
    # Intro
    p = doc.add_paragraph()
    p.add_run(f"Pada hari ini, {{hari_pemeriksaan}} tanggal {{tanggal_bahp:tanggal_long}}, kami yang bertanda tangan di bawah ini:")
    
    doc.add_paragraph()
    
    # Tim Pemeriksa
    p = doc.add_paragraph()
    run = p.add_run("TIM PEMERIKSA HASIL PEKERJAAN:")
    run.bold = True
    
    tim_table = doc.add_table(rows=4, cols=4)
    tim_table.style = 'Table Grid'
    
    tim_headers = ["No", "Nama", "NIP", "Jabatan dalam Tim"]
    for i, h in enumerate(tim_headers):
        tim_table.rows[0].cells[i].text = h
        tim_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    tim_data = [
        ("1", "{{pemeriksa1_nama}}", "{{pemeriksa1_nip:nip}}", "Ketua"),
        ("2", "{{pemeriksa2_nama}}", "{{pemeriksa2_nip:nip}}", "Sekretaris"),
        ("3", "{{pemeriksa3_nama}}", "{{pemeriksa3_nip:nip}}", "Anggota"),
    ]
    for i, (no, nama, nip, jabatan) in enumerate(tim_data, 1):
        tim_table.rows[i].cells[0].text = no
        tim_table.rows[i].cells[1].text = nama
        tim_table.rows[i].cells[2].text = nip
        tim_table.rows[i].cells[3].text = jabatan
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Telah mengadakan pemeriksaan terhadap hasil pekerjaan:")
    
    # Detail pekerjaan
    detail_table = doc.add_table(rows=6, cols=3)
    detail_data = [
        ("Nama Pekerjaan", ":", "{{nama_paket}}"),
        ("Nomor SPK", ":", "{{nomor_spk}}"),
        ("Tanggal SPK", ":", "{{tanggal_spk:tanggal_long}}"),
        ("Nilai SPK", ":", "{{nilai_kontrak:rupiah}}"),
        ("Penyedia", ":", "{{penyedia_nama}}"),
        ("Lokasi", ":", "{{lokasi_pekerjaan}}"),
    ]
    for i, (label, sep, value) in enumerate(detail_data):
        detail_table.rows[i].cells[0].text = label
        detail_table.rows[i].cells[1].text = sep
        detail_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Hasil Pemeriksaan
    p = doc.add_paragraph()
    run = p.add_run("HASIL PEMERIKSAAN:")
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run("{{hasil_pemeriksaan}}")
    
    doc.add_paragraph()
    
    # Rincian Pemeriksaan
    p = doc.add_paragraph()
    run = p.add_run("RINCIAN HASIL PEMERIKSAAN:")
    run.bold = True
    
    rincian_table = doc.add_table(rows=2, cols=6)
    rincian_table.style = 'Table Grid'
    
    rincian_headers = ["No", "Uraian Barang/Jasa", "Satuan", "Volume SPK", "Volume Diterima", "Keterangan"]
    for i, h in enumerate(rincian_headers):
        rincian_table.rows[0].cells[i].text = h
        rincian_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    rincian_data = ["1", "{{item.nama_item}}", "{{item.satuan}}", "{{item.volume}}", "{{item.volume_diterima}}", "{{item.keterangan}}"]
    for i, d in enumerate(rincian_data):
        rincian_table.rows[1].cells[i].text = d
    
    doc.add_paragraph()
    
    # Kesimpulan
    p = doc.add_paragraph()
    run = p.add_run("KESIMPULAN:")
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run("{{kesimpulan_pemeriksaan}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Demikian Berita Acara Hasil Pemeriksaan ini dibuat dengan sebenarnya untuk dipergunakan sebagaimana mestinya.")
    
    # TTD Tim Pemeriksa
    doc.add_paragraph()
    ttd_table = doc.add_table(rows=4, cols=3)
    
    for i, jabatan in enumerate(["Ketua,", "Sekretaris,", "Anggota,"]):
        ttd_table.rows[0].cells[i].text = jabatan
        ttd_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i in range(3):
        ttd_table.rows[1].cells[i].text = "\n\n\n"
    
    nama_fields = ["{{pemeriksa1_nama}}", "{{pemeriksa2_nama}}", "{{pemeriksa3_nama}}"]
    for i, nama in enumerate(nama_fields):
        p = ttd_table.rows[2].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(nama)
        run.bold = True
        run.underline = True
    
    nip_fields = ["NIP. {{pemeriksa1_nip:nip}}", "NIP. {{pemeriksa2_nip:nip}}", "NIP. {{pemeriksa3_nip:nip}}"]
    for i, nip in enumerate(nip_fields):
        ttd_table.rows[3].cells[i].text = nip
        ttd_table.rows[3].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Mengetahui PPK
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Mengetahui,")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_ttd_ppk(doc)
    
    filepath = os.path.join(WORD_DIR, "bahp.docx")
    doc.save(filepath)
    print(f"✅ Created: bahp.docx")
    return filepath


def create_bast():
    """Template Berita Acara Serah Terima (BAST)"""
    doc = create_doc_with_header()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BERITA ACARA SERAH TERIMA HASIL PEKERJAAN")
    run.bold = True
    run.font.size = Pt(14)
    
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_bast}}")
    
    doc.add_paragraph()
    
    # Intro
    p = doc.add_paragraph()
    p.add_run(f"Pada hari ini, {{hari_serah_terima}} tanggal {{tanggal_bast:tanggal_long}}, kami yang bertanda tangan di bawah ini:")
    
    doc.add_paragraph()
    
    # Pihak Pertama
    p = doc.add_paragraph()
    run = p.add_run("PIHAK PERTAMA (Yang Menerima):")
    run.bold = True
    
    pihak1_table = doc.add_table(rows=3, cols=3)
    pihak1_data = [
        ("Nama", ":", "{{ppk_nama}}"),
        ("Jabatan", ":", "Pejabat Pembuat Komitmen"),
        ("Alamat", ":", "{{satker_nama}}, {{satker_alamat}}"),
    ]
    for i, (label, sep, value) in enumerate(pihak1_data):
        pihak1_table.rows[i].cells[0].text = label
        pihak1_table.rows[i].cells[1].text = sep
        pihak1_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Pihak Kedua
    p = doc.add_paragraph()
    run = p.add_run("PIHAK KEDUA (Yang Menyerahkan):")
    run.bold = True
    
    pihak2_table = doc.add_table(rows=4, cols=3)
    pihak2_data = [
        ("Nama Perusahaan", ":", "{{penyedia_nama}}"),
        ("Direktur", ":", "{{direktur_nama}}"),
        ("Alamat", ":", "{{penyedia_alamat}}, {{penyedia_kota}}"),
        ("NPWP", ":", "{{penyedia_npwp:npwp}}"),
    ]
    for i, (label, sep, value) in enumerate(pihak2_data):
        pihak2_table.rows[i].cells[0].text = label
        pihak2_table.rows[i].cells[1].text = sep
        pihak2_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Berdasarkan:")
    
    dasar_table = doc.add_table(rows=3, cols=3)
    dasar_data = [
        ("1. Surat Perintah Kerja", ":", "Nomor {{nomor_spk}} tanggal {{tanggal_spk:tanggal_long}}"),
        ("2. Berita Acara Pemeriksaan", ":", "Nomor {{nomor_bahp}} tanggal {{tanggal_bahp:tanggal_long}}"),
        ("3. Nilai Pekerjaan", ":", "{{nilai_kontrak:rupiah}}"),
    ]
    for i, (label, sep, value) in enumerate(dasar_data):
        dasar_table.rows[i].cells[0].text = label
        dasar_table.rows[i].cells[1].text = sep
        dasar_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Maka dengan ini PIHAK KEDUA menyerahkan hasil pekerjaan kepada PIHAK PERTAMA dan PIHAK PERTAMA menerima penyerahan hasil pekerjaan dari PIHAK KEDUA, berupa:")
    
    # Rincian serah terima
    rincian_table = doc.add_table(rows=2, cols=5)
    rincian_table.style = 'Table Grid'
    
    headers = ["No", "Uraian Barang/Jasa", "Satuan", "Volume", "Kondisi"]
    for i, h in enumerate(headers):
        rincian_table.rows[0].cells[i].text = h
        rincian_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    sample = ["1", "{{item.nama_item}}", "{{item.satuan}}", "{{item.volume}}", "{{item.kondisi}}"]
    for i, s in enumerate(sample):
        rincian_table.rows[1].cells[i].text = s
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Dengan ditandatanganinya Berita Acara Serah Terima ini, maka tanggung jawab atas barang/jasa tersebut beralih dari PIHAK KEDUA kepada PIHAK PERTAMA.")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Demikian Berita Acara Serah Terima ini dibuat dalam rangkap 2 (dua) yang masing-masing mempunyai kekuatan hukum yang sama.")
    
    add_ttd_dual(doc, "YANG MENERIMA", "YANG MENYERAHKAN")
    
    filepath = os.path.join(WORD_DIR, "bast.docx")
    doc.save(filepath)
    print(f"✅ Created: bast.docx")
    return filepath


def create_spp_ls():
    """Template Surat Permintaan Pembayaran Langsung (SPP-LS)"""
    doc = create_doc_with_header()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SURAT PERMINTAAN PEMBAYARAN LANGSUNG")
    run.bold = True
    run.font.size = Pt(14)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("(SPP-LS)")
    
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_spp}}")
    
    doc.add_paragraph()
    
    # Info
    info_table = doc.add_table(rows=4, cols=3)
    info_data = [
        ("Tanggal", ":", "{{tanggal_spp:tanggal_long}}"),
        ("Sifat Pembayaran", ":", "Langsung (LS)"),
        ("Jenis Pembayaran", ":", "Pengadaan Barang/Jasa"),
        ("Tahun Anggaran", ":", "{{tahun_anggaran}}"),
    ]
    for i, (label, sep, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = sep
        info_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Sehubungan dengan pelaksanaan kegiatan:")
    
    # Detail Kegiatan
    kegiatan_table = doc.add_table(rows=3, cols=3)
    kegiatan_data = [
        ("Nama Pekerjaan", ":", "{{nama_paket}}"),
        ("Nomor SPK", ":", "{{nomor_spk}} tanggal {{tanggal_spk:tanggal_long}}"),
        ("Penyedia", ":", "{{penyedia_nama}}"),
    ]
    for i, (label, sep, value) in enumerate(kegiatan_data):
        kegiatan_table.rows[i].cells[0].text = label
        kegiatan_table.rows[i].cells[1].text = sep
        kegiatan_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Bersama ini kami ajukan Surat Permintaan Pembayaran (SPP) sebagai berikut:")
    
    # Rincian Pembayaran
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("RINCIAN PEMBAYARAN:")
    run.bold = True
    
    pembayaran_table = doc.add_table(rows=6, cols=3)
    pembayaran_table.style = 'Table Grid'
    
    pembayaran_data = [
        ("1", "Nilai Pekerjaan (DPP)", "{{nilai_kontrak:rupiah}}"),
        ("2", "PPN 11%", "{{nilai_ppn:rupiah}}"),
        ("3", "Nilai Bruto (1+2)", "{{nilai_bruto:rupiah}}"),
        ("4", "{{jenis_pph}} ({{tarif_pph_persen}}%)", "{{nilai_pph:rupiah}}"),
        ("5", "Potongan Lainnya", "Rp 0"),
        ("", "NILAI BERSIH YANG DIBAYARKAN", "{{nilai_bersih:rupiah}}"),
    ]
    for i, (no, label, value) in enumerate(pembayaran_data):
        pembayaran_table.rows[i].cells[0].text = no
        pembayaran_table.rows[i].cells[1].text = label
        pembayaran_table.rows[i].cells[2].text = value
    
    # Bold last row
    for cell in pembayaran_table.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Terbilang: ")
    run = p.add_run("{{nilai_bersih:terbilang}}")
    run.italic = True
    
    doc.add_paragraph()
    
    # Rekening Tujuan
    p = doc.add_paragraph()
    run = p.add_run("REKENING TUJUAN:")
    run.bold = True
    
    rek_table = doc.add_table(rows=4, cols=3)
    rek_data = [
        ("Nama Bank", ":", "{{penyedia_bank}}"),
        ("Nomor Rekening", ":", "{{penyedia_rekening}}"),
        ("Nama Rekening", ":", "{{penyedia_nama_rekening}}"),
        ("NPWP", ":", "{{penyedia_npwp:npwp}}"),
    ]
    for i, (label, sep, value) in enumerate(rek_data):
        rek_table.rows[i].cells[0].text = label
        rek_table.rows[i].cells[1].text = sep
        rek_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Demikian SPP ini dibuat untuk dapat diproses lebih lanjut.")
    
    # TTD PPK dan PPSPM
    doc.add_paragraph()
    ttd_table = doc.add_table(rows=5, cols=2)
    
    ttd_table.rows[0].cells[0].text = "Mengetahui,"
    ttd_table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_spp:tanggal_long}}"
    
    ttd_table.rows[1].cells[0].text = "Pejabat Penandatangan SPM"
    ttd_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen,"
    
    ttd_table.rows[2].cells[0].text = "\n\n\n"
    ttd_table.rows[2].cells[1].text = "\n\n\n"
    
    p = ttd_table.rows[3].cells[0].paragraphs[0]
    run = p.add_run("{{ppspm_nama}}")
    run.bold = True
    run.underline = True
    p = ttd_table.rows[3].cells[1].paragraphs[0]
    run = p.add_run("{{ppk_nama}}")
    run.bold = True
    run.underline = True
    
    ttd_table.rows[4].cells[0].text = "NIP. {{ppspm_nip:nip}}"
    ttd_table.rows[4].cells[1].text = "NIP. {{ppk_nip:nip}}"
    
    for row in ttd_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(WORD_DIR, "spp_ls.docx")
    doc.save(filepath)
    print(f"✅ Created: spp_ls.docx")
    return filepath


def create_drpp():
    """Template Daftar Rincian Permintaan Pembayaran (DRPP)"""
    doc = Document()
    
    # Set landscape
    for section in doc.sections:
        section.orientation = 1  # Landscape
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DAFTAR RINCIAN PERMINTAAN PEMBAYARAN (DRPP)")
    run.bold = True
    run.font.size = Pt(12)
    
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_drpp}}")
    
    # Info
    info_table = doc.add_table(rows=3, cols=3)
    info_data = [
        ("Satuan Kerja", ":", "{{satker_nama}}"),
        ("Tahun Anggaran", ":", "{{tahun_anggaran}}"),
        ("Tanggal", ":", "{{tanggal_spp:tanggal_long}}"),
    ]
    for i, (label, sep, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = sep
        info_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Main table
    main_table = doc.add_table(rows=3, cols=9)
    main_table.style = 'Table Grid'
    
    headers = ["No", "Nama Penerima", "NPWP", "Nama Bank", "No. Rekening", "Kode Akun", "Uraian", "Nilai Bruto (Rp)", "Potongan Pajak (Rp)", "Nilai Bersih (Rp)"]
    
    # Adjust - we have 9 cols but 10 headers, need to merge for header
    headers = ["No", "Nama Penerima/\nBank/Rekening", "NPWP", "Kode Akun", "Uraian Pekerjaan", "Nilai Bruto", "PPN", "PPh", "Nilai Bersih"]
    
    for i, h in enumerate(headers):
        cell = main_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data row
    data = [
        "1",
        "{{penyedia_nama}}\n{{penyedia_bank}}\n{{penyedia_rekening}}",
        "{{penyedia_npwp:npwp}}",
        "{{kode_akun}}",
        "{{nama_paket}}",
        "{{nilai_kontrak:rupiah}}",
        "{{nilai_ppn:rupiah}}",
        "{{nilai_pph:rupiah}}",
        "{{nilai_bersih:rupiah}}",
    ]
    for i, d in enumerate(data):
        main_table.rows[1].cells[i].text = d
    
    # Total row
    main_table.rows[2].cells[0].text = ""
    main_table.rows[2].cells[4].text = "JUMLAH"
    main_table.rows[2].cells[4].paragraphs[0].runs[0].bold = True
    main_table.rows[2].cells[5].text = "{{nilai_kontrak:rupiah}}"
    main_table.rows[2].cells[6].text = "{{nilai_ppn:rupiah}}"
    main_table.rows[2].cells[7].text = "{{nilai_pph:rupiah}}"
    main_table.rows[2].cells[8].text = "{{nilai_bersih:rupiah}}"
    
    doc.add_paragraph()
    
    # TTD
    ttd_table = doc.add_table(rows=4, cols=2)
    ttd_table.rows[0].cells[1].text = f"{{{{satker_kota}}}}, {{{{tanggal_spp:tanggal_long}}}}"
    ttd_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen,"
    ttd_table.rows[2].cells[1].text = "\n\n\n"
    p = ttd_table.rows[3].cells[1].paragraphs[0]
    run = p.add_run("{{ppk_nama}}")
    run.bold = True
    run.underline = True
    ttd_table.rows[3].cells[1].add_paragraph("NIP. {{ppk_nip:nip}}")
    
    for row in ttd_table.rows:
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(WORD_DIR, "drpp.docx")
    doc.save(filepath)
    print(f"✅ Created: drpp.docx")
    return filepath


def create_kuitansi():
    """Template Kuitansi"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KUITANSI")
    run.bold = True
    run.font.size = Pt(16)
    
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_kuitansi}}")
    
    doc.add_paragraph()
    
    # Content table
    content_table = doc.add_table(rows=6, cols=3)
    content_data = [
        ("Sudah terima dari", ":", "Pejabat Pembuat Komitmen {{satker_nama}}"),
        ("Uang sebesar", ":", "{{nilai_bersih:rupiah}}"),
        ("Terbilang", ":", "{{nilai_bersih:terbilang}}"),
        ("Untuk pembayaran", ":", "{{nama_paket}}"),
        ("Nomor SPK", ":", "{{nomor_spk}} tanggal {{tanggal_spk:tanggal_long}}"),
        ("Kode Akun", ":", "{{kode_akun}}"),
    ]
    for i, (label, sep, value) in enumerate(content_data):
        content_table.rows[i].cells[0].text = label
        content_table.rows[i].cells[1].text = sep
        content_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Rincian
    p = doc.add_paragraph()
    run = p.add_run("Rincian Pembayaran:")
    run.bold = True
    
    rincian_table = doc.add_table(rows=5, cols=3)
    rincian_table.style = 'Table Grid'
    rincian_data = [
        ("Nilai Pekerjaan (DPP)", ":", "{{nilai_kontrak:rupiah}}"),
        ("PPN 11%", ":", "{{nilai_ppn:rupiah}}"),
        ("{{jenis_pph}}", ":", "({{nilai_pph:rupiah}})"),
        ("Potongan Lain", ":", "Rp 0"),
        ("JUMLAH BERSIH", ":", "{{nilai_bersih:rupiah}}"),
    ]
    for i, (label, sep, value) in enumerate(rincian_data):
        rincian_table.rows[i].cells[0].text = label
        rincian_table.rows[i].cells[1].text = sep
        rincian_table.rows[i].cells[2].text = value
    rincian_table.rows[-1].cells[0].paragraphs[0].runs[0].bold = True
    rincian_table.rows[-1].cells[2].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # TTD 3 kolom
    ttd_table = doc.add_table(rows=5, cols=3)
    
    headers = ["Mengetahui,\nPejabat Pembuat Komitmen", "Bendahara Pengeluaran", f"{{{{satker_kota}}}}, {{{{tanggal_kuitansi:tanggal_long}}}}\nYang Menerima,"]
    for i, h in enumerate(headers):
        ttd_table.rows[0].cells[i].text = h
        ttd_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i in range(3):
        ttd_table.rows[1].cells[i].text = "\n\n\n"
    
    names = ["{{ppk_nama}}", "{{bendahara_nama}}", "{{direktur_nama}}"]
    for i, name in enumerate(names):
        p = ttd_table.rows[2].cells[i].paragraphs[0]
        run = p.add_run(name)
        run.bold = True
        run.underline = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nips = ["NIP. {{ppk_nip:nip}}", "NIP. {{bendahara_nip:nip}}", "{{penyedia_nama}}"]
    for i, nip in enumerate(nips):
        ttd_table.rows[3].cells[i].text = nip
        ttd_table.rows[3].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(WORD_DIR, "kuitansi.docx")
    doc.save(filepath)
    print(f"✅ Created: kuitansi.docx")
    return filepath


# ============================================================================
# EXCEL TEMPLATES
# ============================================================================

def create_survey_harga():
    """Template Survey Harga"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Survey Harga"
    
    # Styles
    header_font = Font(bold=True, size=12)
    title_font = Font(bold=True, size=14)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # Title
    ws.merge_cells('A1:J1')
    ws['A1'] = 'SURVEY HARGA PASAR'
    ws['A1'].font = title_font
    ws['A1'].alignment = center
    
    ws.merge_cells('A2:J2')
    ws['A2'] = '{{nama_paket}}'
    ws['A2'].alignment = center
    
    ws['A4'] = 'Tanggal Survey:'
    ws['B4'] = '{{tanggal_survey:tanggal_long}}'
    
    # Headers
    headers = ['No', 'Uraian Barang/Jasa', 'Spesifikasi', 'Satuan', 'Volume',
               'Harga Toko 1\n{{survey1_toko}}', 'Harga Toko 2\n{{survey2_toko}}', 
               'Harga Toko 3\n{{survey3_toko}}', 'Harga Rata-rata', 'Jumlah']
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=6, column=col, value=header)
        cell.font = header_font
        cell.border = border
        cell.alignment = center
    
    # Sample data row
    sample_data = ['1', '{{item.nama_item}}', '{{item.spesifikasi}}', '{{item.satuan}}', 
                   '{{item.volume}}', '{{item.harga_survey1}}', '{{item.harga_survey2}}',
                   '{{item.harga_survey3}}', '{{item.harga_rata}}', '{{item.jumlah}}']
    
    for col, data in enumerate(sample_data, 1):
        cell = ws.cell(row=7, column=col, value=data)
        cell.border = border
        cell.alignment = Alignment(vertical='center', wrap_text=True)
    
    # Total row
    ws.cell(row=8, column=1, value='').border = border
    ws.merge_cells('B8:I8')
    ws['B8'] = 'TOTAL'
    ws['B8'].font = header_font
    ws['B8'].alignment = Alignment(horizontal='right')
    ws['B8'].border = border
    ws['J8'] = '{{total_survey}}'
    ws['J8'].font = header_font
    ws['J8'].border = border
    
    # Column widths
    widths = [5, 30, 25, 10, 10, 15, 15, 15, 15, 18]
    for i, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width
    
    # Surveyor info
    ws['A10'] = 'Lokasi Survey:'
    ws['A11'] = f'1. {{{{survey1_toko}}}} - {{{{survey1_alamat}}}}'
    ws['A12'] = f'2. {{{{survey2_toko}}}} - {{{{survey2_alamat}}}}'
    ws['A13'] = f'3. {{{{survey3_toko}}}} - {{{{survey3_alamat}}}}'
    
    # Signature
    ws['H15'] = '{{satker_kota}}, {{tanggal_survey:tanggal_long}}'
    ws['H16'] = 'Petugas Survey,'
    ws['H19'] = '{{surveyor_nama}}'
    ws['H20'] = 'NIP. {{surveyor_nip:nip}}'
    
    filepath = os.path.join(EXCEL_DIR, "survey_harga.xlsx")
    wb.save(filepath)
    print(f"✅ Created: survey_harga.xlsx")
    return filepath


def create_hps():
    """Template HPS (Harga Perkiraan Sendiri)"""
    wb = Workbook()
    ws = wb.active
    ws.title = "HPS"
    
    # Styles
    header_font = Font(bold=True, size=11)
    title_font = Font(bold=True, size=14)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    money_format = '#,##0'
    
    # Header
    ws.merge_cells('A1:G1')
    ws['A1'] = 'HARGA PERKIRAAN SENDIRI (HPS)'
    ws['A1'].font = title_font
    ws['A1'].alignment = center
    
    ws.merge_cells('A2:G2')
    ws['A2'] = '{{nama_paket}}'
    ws['A2'].alignment = center
    
    # Info
    info_data = [
        ('Satuan Kerja', ':', '{{satker_nama}}'),
        ('Tahun Anggaran', ':', '{{tahun_anggaran}}'),
        ('Sumber Dana', ':', '{{sumber_dana}}'),
        ('Kode Akun', ':', '{{kode_akun}}'),
        ('Pagu Anggaran', ':', '{{nilai_pagu:rupiah}}'),
    ]
    
    for i, (label, sep, value) in enumerate(info_data, 4):
        ws[f'A{i}'] = label
        ws[f'B{i}'] = sep
        ws[f'C{i}'] = value
    
    # Table headers
    table_headers = ['No', 'Uraian Barang/Jasa', 'Spesifikasi', 'Satuan', 'Volume', 'Harga Satuan', 'Jumlah']
    for col, header in enumerate(table_headers, 1):
        cell = ws.cell(row=10, column=col, value=header)
        cell.font = header_font
        cell.border = border
        cell.alignment = center
        cell.fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    
    # Sample data
    sample = ['1', '{{item.nama_item}}', '{{item.spesifikasi}}', '{{item.satuan}}', 
              '{{item.volume}}', '{{item.harga_satuan}}', '{{item.jumlah}}']
    for col, data in enumerate(sample, 1):
        cell = ws.cell(row=11, column=col, value=data)
        cell.border = border
    
    # Subtotal
    ws.merge_cells('A12:E12')
    ws['A12'] = 'SUBTOTAL'
    ws['A12'].font = header_font
    ws['A12'].alignment = Alignment(horizontal='right')
    ws['A12'].border = border
    ws['F12'] = ''
    ws['F12'].border = border
    ws['G12'] = '{{subtotal_hps}}'
    ws['G12'].font = header_font
    ws['G12'].border = border
    
    # PPN
    ws.merge_cells('A13:E13')
    ws['A13'] = 'PPN 11%'
    ws['A13'].font = header_font
    ws['A13'].alignment = Alignment(horizontal='right')
    ws['A13'].border = border
    ws['F13'] = ''
    ws['F13'].border = border
    ws['G13'] = '{{ppn_hps}}'
    ws['G13'].font = header_font
    ws['G13'].border = border
    
    # Total HPS
    ws.merge_cells('A14:E14')
    ws['A14'] = 'TOTAL HPS'
    ws['A14'].font = Font(bold=True, size=12)
    ws['A14'].alignment = Alignment(horizontal='right')
    ws['A14'].border = border
    ws['A14'].fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
    ws['F14'] = ''
    ws['F14'].border = border
    ws['G14'] = '{{nilai_hps}}'
    ws['G14'].font = Font(bold=True, size=12)
    ws['G14'].border = border
    ws['G14'].fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
    
    # Terbilang
    ws.merge_cells('A15:G15')
    ws['A15'] = 'Terbilang: {{nilai_hps:terbilang}}'
    ws['A15'].font = Font(italic=True)
    
    # Column widths
    widths = [5, 35, 25, 10, 10, 18, 20]
    for i, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width
    
    # Signature
    ws['E18'] = '{{satker_kota}}, {{tanggal_hps:tanggal_long}}'
    ws['E19'] = 'Pejabat Pembuat Komitmen,'
    ws['E22'] = '{{ppk_nama}}'
    ws['E23'] = 'NIP. {{ppk_nip:nip}}'
    
    filepath = os.path.join(EXCEL_DIR, "hps.xlsx")
    wb.save(filepath)
    print(f"✅ Created: hps.xlsx")
    return filepath


def create_ssp():
    """Template SSP (Surat Setoran Pajak) - PPN dan PPh"""
    wb = Workbook()
    
    # ====== SHEET 1: SSP PPN ======
    ws_ppn = wb.active
    ws_ppn.title = "SSP PPN"
    
    # Styles
    header_font = Font(bold=True, size=11)
    title_font = Font(bold=True, size=14)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # Title
    ws_ppn.merge_cells('A1:F1')
    ws_ppn['A1'] = 'SURAT SETORAN PAJAK (SSP)'
    ws_ppn['A1'].font = title_font
    ws_ppn['A1'].alignment = center
    
    ws_ppn.merge_cells('A2:F2')
    ws_ppn['A2'] = 'PAJAK PERTAMBAHAN NILAI (PPN)'
    ws_ppn['A2'].font = header_font
    ws_ppn['A2'].alignment = center
    
    # Data
    ppn_data = [
        ('NPWP Pemotong', ':', '{{satker_npwp:npwp}}'),
        ('Nama Pemotong', ':', '{{satker_nama}}'),
        ('Alamat', ':', '{{satker_alamat}}, {{satker_kota}}'),
        ('', '', ''),
        ('NPWP Penyedia', ':', '{{penyedia_npwp:npwp}}'),
        ('Nama Penyedia', ':', '{{penyedia_nama}}'),
        ('Alamat Penyedia', ':', '{{penyedia_alamat}}'),
        ('', '', ''),
        ('Kode Akun Pajak', ':', '411211'),
        ('Kode Jenis Setoran', ':', '910'),
        ('Masa Pajak', ':', '{{bulan_pajak}} {{tahun_anggaran}}'),
        ('Tahun Pajak', ':', '{{tahun_anggaran}}'),
        ('', '', ''),
        ('Uraian', ':', '{{nama_paket}}'),
        ('Nomor SPK', ':', '{{nomor_spk}}'),
        ('Tanggal SPK', ':', '{{tanggal_spk:tanggal_long}}'),
        ('DPP', ':', '{{nilai_kontrak:rupiah}}'),
        ('PPN 11%', ':', '{{nilai_ppn:rupiah}}'),
        ('', '', ''),
        ('Jumlah Setoran', ':', '{{nilai_ppn:rupiah}}'),
        ('Terbilang', ':', '{{nilai_ppn:terbilang}}'),
    ]
    
    for i, (label, sep, value) in enumerate(ppn_data, 4):
        ws_ppn[f'A{i}'] = label
        ws_ppn[f'B{i}'] = sep
        ws_ppn.merge_cells(f'C{i}:F{i}')
        ws_ppn[f'C{i}'] = value
        if label in ['Jumlah Setoran', 'DPP', 'PPN 11%']:
            ws_ppn[f'A{i}'].font = header_font
            ws_ppn[f'C{i}'].font = header_font
    
    # Column widths
    ws_ppn.column_dimensions['A'].width = 20
    ws_ppn.column_dimensions['B'].width = 3
    ws_ppn.column_dimensions['C'].width = 40
    
    # ====== SHEET 2: SSP PPh ======
    ws_pph = wb.create_sheet("SSP PPh")
    
    ws_pph.merge_cells('A1:F1')
    ws_pph['A1'] = 'SURAT SETORAN PAJAK (SSP)'
    ws_pph['A1'].font = title_font
    ws_pph['A1'].alignment = center
    
    ws_pph.merge_cells('A2:F2')
    ws_pph['A2'] = '{{jenis_pph}}'
    ws_pph['A2'].font = header_font
    ws_pph['A2'].alignment = center
    
    # Data
    pph_data = [
        ('NPWP Pemotong', ':', '{{satker_npwp:npwp}}'),
        ('Nama Pemotong', ':', '{{satker_nama}}'),
        ('Alamat', ':', '{{satker_alamat}}, {{satker_kota}}'),
        ('', '', ''),
        ('NPWP Penyedia', ':', '{{penyedia_npwp:npwp}}'),
        ('Nama Penyedia', ':', '{{penyedia_nama}}'),
        ('Alamat Penyedia', ':', '{{penyedia_alamat}}'),
        ('', '', ''),
        ('Kode Akun Pajak', ':', '{{kode_akun_pph}}'),
        ('Kode Jenis Setoran', ':', '{{kode_jenis_setoran_pph}}'),
        ('Masa Pajak', ':', '{{bulan_pajak}} {{tahun_anggaran}}'),
        ('Tahun Pajak', ':', '{{tahun_anggaran}}'),
        ('', '', ''),
        ('Uraian', ':', '{{nama_paket}}'),
        ('Nomor SPK', ':', '{{nomor_spk}}'),
        ('Tanggal SPK', ':', '{{tanggal_spk:tanggal_long}}'),
        ('DPP', ':', '{{nilai_kontrak:rupiah}}'),
        ('Tarif {{jenis_pph}}', ':', '{{tarif_pph_persen}}%'),
        ('{{jenis_pph}}', ':', '{{nilai_pph:rupiah}}'),
        ('', '', ''),
        ('Jumlah Setoran', ':', '{{nilai_pph:rupiah}}'),
        ('Terbilang', ':', '{{nilai_pph:terbilang}}'),
    ]
    
    for i, (label, sep, value) in enumerate(pph_data, 4):
        ws_pph[f'A{i}'] = label
        ws_pph[f'B{i}'] = sep
        ws_pph.merge_cells(f'C{i}:F{i}')
        ws_pph[f'C{i}'] = value
        if label in ['Jumlah Setoran', 'DPP', '{{jenis_pph}}']:
            ws_pph[f'A{i}'].font = header_font
            ws_pph[f'C{i}'].font = header_font
    
    # Column widths
    ws_pph.column_dimensions['A'].width = 20
    ws_pph.column_dimensions['B'].width = 3
    ws_pph.column_dimensions['C'].width = 40
    
    filepath = os.path.join(EXCEL_DIR, "ssp.xlsx")
    wb.save(filepath)
    print(f"✅ Created: ssp.xlsx")
    return filepath


# ============================================================================
# MAIN
# ============================================================================

def create_all_templates():
    """Create all templates"""
    print("=" * 60)
    print("CREATING ALL PROCUREMENT TEMPLATES")
    print("=" * 60)
    
    print("\n📄 Word Templates (.docx):")
    create_spesifikasi_teknis()
    create_kak()
    create_spk()
    create_spmk()
    create_bahp()
    create_bast()
    create_spp_ls()
    create_drpp()
    create_kuitansi()
    
    print("\n📊 Excel Templates (.xlsx):")
    create_survey_harga()
    create_hps()
    create_ssp()
    
    print("\n" + "=" * 60)
    print("ALL TEMPLATES CREATED SUCCESSFULLY!")
    print("=" * 60)
    print(f"\nWord templates: {WORD_DIR}")
    print(f"Excel templates: {EXCEL_DIR}")
    print("\nPlaceholder yang digunakan:")
    print("  {{nama_paket}} - Nama pekerjaan")
    print("  {{nilai_kontrak}} - Nilai tanpa format")
    print("  {{nilai_kontrak:rupiah}} - Format Rp 1.000.000")
    print("  {{nilai_kontrak:terbilang}} - Satu juta rupiah")
    print("  {{tanggal_spk:tanggal_long}} - 15 Januari 2026")
    print("  {{ppk_nip:nip}} - 19720103 199503 1 001")
    print("  {{penyedia_npwp:npwp}} - 01.234.567.8-951.000")


if __name__ == "__main__":
    create_all_templates()
