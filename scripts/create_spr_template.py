"""
Script untuk membuat template SPR dan SPPR sesuai format standar.
- SPPR = Surat Perintah Pendebitan Rekening (menggunakan Kartu Debit)
- SPR = Surat Pendebitan Rekening (penarikan tunai via teller bank)
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def create_sppr_template():
    """
    Create SPPR (Surat Perintah Pendebitan Rekening) template.
    Digunakan untuk pendebitan rekening menggunakan Kartu Debit.
    """
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- HEADER ---
    header_center = doc.add_paragraph()
    header_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_center.add_run("{{kementerian}}")
    run.font.size = Pt(12)

    header_sub = doc.add_paragraph()
    header_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = header_sub.add_run("{{satker_nama}}")
    run2.font.size = Pt(11)

    # Garis pemisah
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = separator.add_run("_" * 70)
    run_sep.font.size = Pt(10)

    doc.add_paragraph()  # Spacing

    # --- JUDUL ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("SURAT PERINTAH PENDEBITAN REKENING (SPPR)")
    run_title.bold = True
    run_title.underline = True
    run_title.font.size = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Tanggal: {{tanggal_dokumen}}  Nomor: {{nomor_dokumen}}")

    doc.add_paragraph()  # Spacing

    # --- ISI SURAT ---
    isi1 = doc.add_paragraph()
    isi1.add_run(
        "Saya yang bertandatangan di bawah ini selaku Kuasa Pengguna Anggaran/Pejabat Pembuat "
        "Komitmen atas nama Kuasa Pengguna Anggaran *) memerintahkan Bendahara "
        "Pengeluaran/Bendahara Pengeluaran Pembantu *) agar melakukan pendebitan rekening "
        "menggunakan "
    )
    run_kartu = isi1.add_run("Kartu Debit")
    run_kartu.bold = True
    isi1.add_run(" sejumlah Rp.{{jumlah:rupiah}}")

    isi2 = doc.add_paragraph()
    isi2.add_run("Terbilang: {{terbilang}}")

    isi3 = doc.add_paragraph()
    isi3.add_run("Atas Dasar: Surat Perintah Bayar (SPBy) Nomor {{nomor_spby}} **)")

    doc.add_paragraph()  # Spacing

    # --- TANDA TANGAN ---
    # Lokasi dan tanggal
    ttd_lokasi = doc.add_paragraph()
    ttd_lokasi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ttd_lokasi.add_run("{{lokasi}}, {{tanggal_dokumen}}")

    # Two column signature using table
    ttd_table = doc.add_table(rows=5, cols=2)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    ttd_table.cell(0, 0).text = "Kuasa Pengguna Anggaran/Pejabat Pembuat\nKomitmen atas nama KPA"
    ttd_table.cell(0, 1).text = "Bendahara Pengeluaran/BPP *)"

    # Empty rows for signature
    for i in range(1, 4):
        ttd_table.cell(i, 0).text = ""
        ttd_table.cell(i, 1).text = ""

    # Name and NIP
    ttd_table.cell(4, 0).text = "{{ppk_nama}}\nNIP. {{ppk_nip}}"
    ttd_table.cell(4, 1).text = "{{bendahara_nama}}\nNIP. {{bendahara_nip}}"

    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing

    # --- CATATAN ---
    note1 = doc.add_paragraph()
    note1.add_run("*) Coret yang tidak perlu")
    note1.runs[0].font.size = Pt(9)

    note2 = doc.add_paragraph()
    note2.add_run("**) Penerbitan SPPR dalam rangka penarikan secara tunai untuk mengisi brankas, nomor SPBy dikosongkan.")
    note2.runs[0].font.size = Pt(9)

    doc.add_paragraph()  # Spacing

    # --- TATA CARA ---
    tata_cara_title = doc.add_paragraph()
    tata_cara_title.add_run("Tata cara pengisian Surat Perintah Pendebitan Rekening:")
    tata_cara_title.runs[0].bold = True
    tata_cara_title.runs[0].font.size = Pt(10)

    tata_cara_items = [
        "(1)  diisi nama kementerian;",
        "(2)  diisi nama satker;",
        "(3)  diisi tanggal, bulan, dan tahun surat;",
        "(4)  diisi nomor surat;",
        "(5)  diisi jumlah pendebitan rekening dengan angka;",
        "(6)  diisi uraian jumlah pendebitan rekening dengan huruf;",
        "(7)  diisi nomor Surat Perintah Bayar;",
        "(8)  diisi tempat penandatanganan;",
        "(9)  diisi tanggal, bulan, dan tahun penandatanganan;",
        "(10) diisi tanda tangan KPA/PPK atas nama KPA;",
        "(11) diisi nama KPA/PPK atas nama KPA;",
        "(12) diisi NIP KPA/PPK atas nama KPA;",
        "(13) diisi tanda tangan Bendahara Pengeluaran/BPP;",
        "(14) diisi nama Bendahara Pengeluaran/BPP; dan",
        "(15) diisi NIP Bendahara Pengeluaran/BPP.",
    ]

    for item in tata_cara_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.runs[0].font.size = Pt(9)
        p.paragraph_format.left_indent = Cm(0.5)

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "word" / "sppr.docx"
    doc.save(str(output_path))
    print(f"SPPR template saved to: {output_path}")
    return output_path


def create_spr_template():
    """
    Create SPR (Surat Pendebitan Rekening) template.
    Digunakan untuk pendebitan tunai melalui teller bank.
    """
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- HEADER (KOP SURAT) ---
    header_center = doc.add_paragraph()
    header_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_center.add_run("KOP SURAT SATKER")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()  # Spacing

    # --- JUDUL ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("SURAT PENDEBITAN REKENING (SPR)")
    run_title.bold = True
    run_title.underline = True
    run_title.font.size = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Tanggal: {{tanggal_dokumen}}  No.: {{nomor_dokumen}}")

    doc.add_paragraph()  # Spacing

    # --- KEPADA ---
    kepada = doc.add_paragraph()
    kepada.add_run("Kepada Yth.")

    pimpinan = doc.add_paragraph()
    pimpinan.add_run("Pimpinan setempat {{nama_bank}}")

    doc.add_paragraph()  # Spacing

    # --- ISI SURAT ---
    isi1 = doc.add_paragraph()
    isi1.add_run(
        "Saya yang bertanda tangan di bawah ini selaku Kuasa Pengguna Anggaran/Pejabat "
        "Pembuat Komitmen atas nama Kuasa Pengguna Anggaran*) memerintahkan Bendahara "
        "Pengeluaran/Bendahara Pengeluaran Pembantu*) agar melakukan pendebitan tunai "
        "melalui teller bank yang Saudara pimpin dengan keterangan sebagai berikut:"
    )
    isi1.paragraph_format.first_line_indent = Cm(1)

    # Detail table
    detail_table = doc.add_table(rows=5, cols=2)
    detail_items = [
        ("Nomor Rekening", "{{nomor_rekening}}"),
        ("Nama Rekening", "{{nama_rekening}}"),
        ("Sejumlah", "Rp {{jumlah:rupiah}}"),
        ("Terbilang", "{{terbilang}}"),
        ("Hari/Tanggal", "{{tanggal_dokumen}}"),
    ]

    for i, (label, value) in enumerate(detail_items):
        detail_table.cell(i, 0).text = label
        detail_table.cell(i, 1).text = ": " + value
        detail_table.cell(i, 0).width = Cm(4)
        detail_table.cell(i, 1).width = Cm(10)

    doc.add_paragraph()  # Spacing

    # --- PENUTUP ---
    penutup1 = doc.add_paragraph()
    penutup1.add_run(
        "Berkenaan dengan hal tersebut, mohon bantuan Saudara untuk membantu kelancaran "
        "transaksi dimaksud."
    )

    penutup2 = doc.add_paragraph()
    penutup2.add_run(
        "Demikian disampaikan, atas bantuan dan kerja sama yang baik diucapkan terima kasih."
    )

    doc.add_paragraph()  # Spacing

    # --- TANDA TANGAN ---
    # Lokasi dan tanggal
    ttd_lokasi = doc.add_paragraph()
    ttd_lokasi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ttd_lokasi.add_run("{{lokasi}}, {{tanggal_dokumen}}")

    # Two column signature using table
    ttd_table = doc.add_table(rows=5, cols=2)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    ttd_table.cell(0, 0).text = "Kuasa Pengguna Anggaran/\nPejabat Pembuat Komitmen atas nama KPA*)"
    ttd_table.cell(0, 1).text = "Bendahara Pengeluaran/BPP*)"

    # Empty rows for signature
    for i in range(1, 4):
        ttd_table.cell(i, 0).text = ""
        ttd_table.cell(i, 1).text = ""

    # Name and NIP
    ttd_table.cell(4, 0).text = "{{ppk_nama}}\nNIP. {{ppk_nip}}"
    ttd_table.cell(4, 1).text = "{{bendahara_nama}}\nNIP. {{bendahara_nip}}"

    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing

    # --- CATATAN ---
    note = doc.add_paragraph()
    note.add_run("*) pilih salah satu")
    note.runs[0].font.size = Pt(9)

    doc.add_paragraph()  # Spacing

    # --- TATA CARA ---
    tata_cara_title = doc.add_paragraph()
    tata_cara_title.add_run("Tata cara pengisian Surat Pendebitan Rekening:")
    tata_cara_title.runs[0].bold = True
    tata_cara_title.runs[0].font.size = Pt(10)

    tata_cara_items = [
        "(1)  diisi tanggal, bulan, dan tahun surat;",
        "(2)  diisi nomor surat;",
        "(3)  diisi nama bank cabang;",
        "(4)  diisi nomor rekening Bendahara Pengeluaran/BPP;",
        "(5)  diisi nama rekening Bendahara Pengeluaran/BPP;",
        "(6)  diisi sejumlah uang yang akan ditarik dengan angka;",
        "(7)  diisi terbilang sejumlah uang yang akan ditarik;",
        "(8)  diisi hari dan tanggal penarikan;",
        "(9)  diisi tempat penandatanganan;",
        "(10) diisi tanggal penandatanganan;",
        "(11) diisi tanda tangan KPA/PPK;",
        "(12) diisi nama KPA/PPK;",
        "(13) diisi NIP KPA/PPK;",
        "(14) diisi tanda tangan Bendahara Pengeluaran/BPP;",
        "(15) diisi nama Bendahara Pengeluaran/BPP;",
        "(16) diisi NIP Bendahara Pengeluaran/BPP.",
    ]

    for item in tata_cara_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.runs[0].font.size = Pt(9)
        p.paragraph_format.left_indent = Cm(0.5)

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "word" / "spr.docx"
    doc.save(str(output_path))
    print(f"SPR template saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_sppr_template()
    create_spr_template()
    print("\nBoth SPR and SPPR templates have been created successfully!")
