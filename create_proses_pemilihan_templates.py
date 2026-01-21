"""
Script untuk membuat template Proses Pemilihan Pengadaan Langsung:
1. Surat Undangan Pengadaan Langsung
2. Berita Acara Hasil Pengadaan Langsung (BAHPL)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = "templates/word"


def set_cell_shading(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_kop_surat(doc):
    """Add standard letter header"""
    kop1 = doc.add_paragraph()
    kop1.add_run("{{kementerian}}").bold = True
    kop1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    kop2 = doc.add_paragraph()
    kop2.add_run("{{eselon1}}").bold = True
    kop2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    kop3 = doc.add_paragraph()
    run = kop3.add_run("{{satker_nama}}")
    run.bold = True
    run.font.size = Pt(14)
    kop3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    alamat = doc.add_paragraph()
    alamat.add_run("{{satker_alamat}}, {{satker_kota}}")
    alamat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    border = doc.add_paragraph()
    border.add_run("═" * 75)
    border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()


# ============================================================================
# SURAT UNDANGAN PENGADAAN LANGSUNG
# ============================================================================

def create_undangan_pl():
    """Create Surat Undangan Pengadaan Langsung template"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Nomor dan Perihal
    info_table = doc.add_table(rows=4, cols=3)
    info_data = [
        ("Nomor", ":", "{{nomor_undangan_pl}}"),
        ("Lampiran", ":", "1 (satu) berkas"),
        ("Hal", ":", "Undangan Pengadaan Langsung"),
    ]
    for i, (label, sep, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = sep
        info_table.rows[i].cells[2].text = value
    
    # Tanggal di kanan
    info_table.rows[3].cells[2].text = "{{satker_kota}}, {{tanggal_undangan_pl_fmt}}"
    
    doc.add_paragraph()
    
    # Kepada Yth
    kepada = doc.add_paragraph()
    kepada.add_run("Kepada Yth.\n")
    kepada.add_run("Direktur {{penyedia_nama}}\n").bold = True
    kepada.add_run("di {{penyedia_alamat}}")
    
    doc.add_paragraph()
    
    # Isi Surat
    p1 = doc.add_paragraph()
    p1.add_run(
        "Sehubungan dengan pelaksanaan pengadaan barang/jasa melalui metode Pengadaan Langsung "
        "pada Tahun Anggaran {{tahun_anggaran}}, dengan ini kami mengundang Saudara untuk "
        "menyampaikan penawaran harga dan dokumen kualifikasi untuk paket pekerjaan:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    p1.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Info Paket
    paket_table = doc.add_table(rows=7, cols=3)
    paket_table.style = 'Table Grid'
    paket_data = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Jenis Pengadaan", "{{jenis_pengadaan}}"),
        ("3.", "Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
        ("4.", "HPS", "{{hps_fmt}}"),
        ("5.", "Sumber Dana", "{{sumber_dana}}"),
        ("6.", "Kode Akun/MAK", "{{kode_akun}}"),
        ("7.", "Jangka Waktu", "{{jangka_waktu}} hari kalender"),
    ]
    for i, (no, label, value) in enumerate(paket_data):
        paket_table.rows[i].cells[0].text = no
        paket_table.rows[i].cells[1].text = label
        paket_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Ketentuan
    p2 = doc.add_paragraph()
    p2.add_run("Adapun ketentuan pengadaan langsung ini adalah sebagai berikut:").bold = True
    
    ketentuan = [
        "Penyampaian dokumen penawaran harga dan dokumen kualifikasi paling lambat tanggal {{batas_penawaran_fmt}} pukul {{batas_penawaran_waktu}} WIT.",
        "Dokumen penawaran disampaikan kepada Pejabat Pengadaan dengan alamat: {{alamat_penyerahan}}.",
        "Negosiasi teknis dan harga akan dilaksanakan pada tanggal {{tanggal_negosiasi_fmt}} pukul {{waktu_negosiasi}} WIT.",
        "Dokumen yang harus dilampirkan meliputi: Surat Penawaran Harga, Daftar Harga/RAB, Dokumen Kualifikasi (izin usaha, NPWP, pakta integritas, dll).",
        "Penawaran harga sudah termasuk seluruh biaya untuk menyelesaikan pekerjaan (termasuk pajak).",
        "Spesifikasi teknis dan dokumen pendukung lainnya terlampir.",
    ]
    
    for i, item in enumerate(ketentuan, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    p3 = doc.add_paragraph()
    p3.add_run(
        "Demikian undangan ini kami sampaikan, atas perhatian dan kerjasamanya diucapkan terima kasih."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD Pejabat Pengadaan
    ttd = doc.add_paragraph()
    ttd.add_run("Pejabat Pengadaan,\n\n\n\n")
    ttd.add_run("{{pp_nama}}\n").bold = True
    ttd.add_run("NIP. {{pp_nip}}")
    ttd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tembusan
    doc.add_paragraph()
    tembusan = doc.add_paragraph()
    tembusan.add_run("Tembusan:").italic = True
    tembusan.add_run("\n1. Kuasa Pengguna Anggaran")
    tembusan.add_run("\n2. Pejabat Pembuat Komitmen")
    tembusan.add_run("\n3. Arsip")
    
    filepath = os.path.join(TEMPLATES_DIR, "undangan_pl.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# BERITA ACARA HASIL PENGADAAN LANGSUNG (BAHPL)
# ============================================================================

def create_bahpl():
    """Create Berita Acara Hasil Pengadaan Langsung template"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA HASIL PENGADAAN LANGSUNG").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bahpl}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_bahpl}} tanggal {{tanggal_bahpl_terbilang}} bulan {{bulan_bahpl}} "
        "tahun {{tahun_bahpl_terbilang}} ({{tanggal_bahpl_fmt}}), yang bertanda tangan di bawah ini:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    p1.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Info Pejabat Pengadaan
    pp_table = doc.add_table(rows=4, cols=3)
    pp_data = [
        ("Nama", ":", "{{pp_nama}}"),
        ("NIP", ":", "{{pp_nip}}"),
        ("Jabatan", ":", "Pejabat Pengadaan"),
        ("Alamat", ":", "{{satker_nama}}, {{satker_alamat}}"),
    ]
    for i, (label, sep, value) in enumerate(pp_data):
        pp_table.rows[i].cells[0].text = label
        pp_table.rows[i].cells[1].text = sep
        pp_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    p2 = doc.add_paragraph()
    p2.add_run(
        "berdasarkan Surat Keputusan Kuasa Pengguna Anggaran Nomor {{nomor_sk_pp}} tanggal "
        "{{tanggal_sk_pp}}, telah melaksanakan Pengadaan Langsung untuk pekerjaan:"
    )
    p2.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Info Paket
    p3 = doc.add_paragraph()
    p3.add_run("A. DATA PAKET PENGADAAN").bold = True
    
    paket_table = doc.add_table(rows=7, cols=3)
    paket_table.style = 'Table Grid'
    paket_data = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Jenis Pengadaan", "{{jenis_pengadaan}}"),
        ("3.", "Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
        ("4.", "Sumber Dana", "{{sumber_dana}}"),
        ("5.", "Kode Akun/MAK", "{{kode_akun}}"),
        ("6.", "HPS", "{{hps_fmt}}"),
        ("7.", "Jangka Waktu", "{{jangka_waktu}} hari kalender"),
    ]
    for i, (no, label, value) in enumerate(paket_data):
        paket_table.rows[i].cells[0].text = no
        paket_table.rows[i].cells[1].text = label
        paket_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Proses Pengadaan
    p4 = doc.add_paragraph()
    p4.add_run("B. PROSES PENGADAAN LANGSUNG").bold = True
    
    proses = [
        f"Berdasarkan Surat Undangan Pengadaan Langsung Nomor {{{{nomor_undangan_pl}}}} tanggal {{{{tanggal_undangan_pl_fmt}}}}, telah diundang 1 (satu) penyedia yaitu {{{{penyedia_nama}}}}.",
        "Penyedia telah menyampaikan dokumen penawaran harga dan dokumen kualifikasi sesuai dengan ketentuan yang ditetapkan.",
        f"Negosiasi teknis dan harga telah dilaksanakan pada tanggal {{{{tanggal_negosiasi_fmt}}}}.",
    ]
    
    for i, item in enumerate(proses, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Hasil Evaluasi Kualifikasi
    p5 = doc.add_paragraph()
    p5.add_run("C. HASIL EVALUASI KUALIFIKASI").bold = True
    
    kualifikasi_table = doc.add_table(rows=8, cols=4)
    kualifikasi_table.style = 'Table Grid'
    
    # Header
    kualifikasi_headers = ["No", "Persyaratan Kualifikasi", "Ada", "Tidak Ada"]
    for i, h in enumerate(kualifikasi_headers):
        kualifikasi_table.rows[0].cells[i].text = h
        for para in kualifikasi_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    kualifikasi_items = [
        "Memiliki izin usaha yang masih berlaku",
        "Memiliki NPWP dan telah memenuhi kewajiban pajak",
        "Mempunyai/menguasai peralatan yang diperlukan",
        "Memiliki kemampuan teknis sesuai kebutuhan",
        "Tidak dalam pengawasan pengadilan/bangkrut",
        "Kebenaran dokumen/tidak dipalsukan",
        "Pakta Integritas",
    ]
    
    for i, item in enumerate(kualifikasi_items, 1):
        kualifikasi_table.rows[i].cells[0].text = str(i)
        kualifikasi_table.rows[i].cells[1].text = item
        kualifikasi_table.rows[i].cells[2].text = "☐"
        kualifikasi_table.rows[i].cells[3].text = "☐"
        
        for j in [0, 2, 3]:
            for para in kualifikasi_table.rows[i].cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    p6 = doc.add_paragraph()
    p6.add_run("Kesimpulan Evaluasi Kualifikasi: ")
    p6.add_run("MEMENUHI / TIDAK MEMENUHI").bold = True
    p6.add_run(" *)")
    
    doc.add_paragraph()
    
    # Hasil Negosiasi
    p7 = doc.add_paragraph()
    p7.add_run("D. HASIL NEGOSIASI TEKNIS DAN HARGA").bold = True
    
    doc.add_paragraph()
    
    p8 = doc.add_paragraph()
    p8.add_run("1. Evaluasi Teknis:").bold = True
    
    teknis = doc.add_paragraph()
    teknis.add_run(
        "Spesifikasi teknis yang ditawarkan telah sesuai dengan spesifikasi yang dipersyaratkan "
        "dalam dokumen pengadaan."
    )
    teknis.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    p9 = doc.add_paragraph()
    p9.add_run("2. Negosiasi Harga:").bold = True
    
    harga_table = doc.add_table(rows=4, cols=3)
    harga_table.style = 'Table Grid'
    harga_data = [
        ("a.", "HPS", "{{hps_fmt}}"),
        ("b.", "Harga Penawaran Awal", "{{harga_penawaran_awal_fmt}}"),
        ("c.", "Harga Hasil Negosiasi", "{{harga_negosiasi_fmt}}"),
        ("d.", "Selisih (a - c)", "{{selisih_negosiasi_fmt}}"),
    ]
    for i, (no, label, value) in enumerate(harga_data):
        harga_table.rows[i].cells[0].text = no
        harga_table.rows[i].cells[1].text = label
        harga_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Kesimpulan
    p10 = doc.add_paragraph()
    p10.add_run("E. KESIMPULAN DAN REKOMENDASI").bold = True
    
    kesimpulan = doc.add_paragraph()
    kesimpulan.add_run(
        "Berdasarkan hasil evaluasi kualifikasi dan negosiasi teknis dan harga tersebut di atas, "
        "dengan ini Pejabat Pengadaan menyatakan:"
    )
    
    doc.add_paragraph()
    
    opsi_table = doc.add_table(rows=2, cols=2)
    opsi_table.rows[0].cells[0].text = "☐"
    opsi_table.rows[0].cells[1].text = "PENGADAAN LANGSUNG BERHASIL"
    opsi_table.rows[1].cells[0].text = "☐"
    opsi_table.rows[1].cells[1].text = "PENGADAAN LANGSUNG GAGAL"
    
    doc.add_paragraph()
    
    # Penyedia yang direkomendasikan
    p11 = doc.add_paragraph()
    p11.add_run("Penyedia yang direkomendasikan untuk melaksanakan pekerjaan:").bold = True
    
    penyedia_table = doc.add_table(rows=6, cols=3)
    penyedia_table.style = 'Table Grid'
    penyedia_data = [
        ("1.", "Nama Perusahaan", "{{penyedia_nama}}"),
        ("2.", "Alamat", "{{penyedia_alamat}}"),
        ("3.", "NPWP", "{{penyedia_npwp}}"),
        ("4.", "Nama Direktur", "{{direktur_nama}}"),
        ("5.", "Nomor Rekening", "{{penyedia_rekening}} ({{penyedia_bank}})"),
        ("6.", "Harga Hasil Negosiasi", "{{harga_negosiasi_fmt}}"),
    ]
    for i, (no, label, value) in enumerate(penyedia_data):
        penyedia_table.rows[i].cells[0].text = no
        penyedia_table.rows[i].cells[1].text = label
        penyedia_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Penutup
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara Hasil Pengadaan Langsung ini dibuat dengan sebenarnya untuk "
        "dipergunakan sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    sign_table = doc.add_table(rows=7, cols=2)
    sign_table.rows[0].cells[0].text = "Penyedia,"
    sign_table.rows[0].cells[1].text = "Pejabat Pengadaan,"
    
    sign_table.rows[1].cells[0].text = "{{penyedia_nama}}"
    
    sign_table.rows[5].cells[0].text = "{{direktur_nama}}"
    sign_table.rows[5].cells[1].text = "{{pp_nama}}"
    
    sign_table.rows[6].cells[0].text = "Direktur"
    sign_table.rows[6].cells[1].text = "NIP. {{pp_nip}}"
    
    # Bold names
    for cell in [sign_table.rows[5].cells[0], sign_table.rows[5].cells[1]]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.underline = True
    
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Mengetahui PPK
    doc.add_paragraph()
    doc.add_paragraph()
    
    mengetahui = doc.add_paragraph()
    mengetahui.add_run("Mengetahui,\n")
    mengetahui.add_run("Pejabat Pembuat Komitmen\n\n\n\n")
    mengetahui.add_run("{{ppk_nama}}\n").bold = True
    mengetahui.add_run("NIP. {{ppk_nip}}")
    mengetahui.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Catatan
    doc.add_paragraph()
    catatan = doc.add_paragraph()
    catatan.add_run("*) Coret yang tidak perlu").italic = True
    catatan.paragraph_format.line_spacing = 1.0
    
    filepath = os.path.join(TEMPLATES_DIR, "bahpl.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    
    print("=" * 60)
    print("MEMBUAT TEMPLATE PROSES PEMILIHAN PENGADAAN LANGSUNG")
    print("=" * 60)
    print()
    
    create_undangan_pl()
    create_bahpl()
    
    print()
    print("=" * 60)
    print("SELESAI!")
    print("=" * 60)
    print()
    print("Template yang dibuat:")
    print("┌─────────────────────────────────────────────────────────┐")
    print("│ 1. undangan_pl.docx  - Surat Undangan Pengadaan Langsung│")
    print("│ 2. bahpl.docx        - BA Hasil Pengadaan Langsung     │")
    print("└─────────────────────────────────────────────────────────┘")
    print()
    print("Dokumen ini ditambahkan ke Timeline Manager:")
    print("  SPESIFIKASI → HPS → KAK → UNDANGAN_PL → BAHPL → SPK → ...")
