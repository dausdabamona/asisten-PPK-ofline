#!/usr/bin/env python3
"""
Create Missing Template Drafts
==============================
Script untuk membuat template draft untuk template yang belum ada file-nya
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os

# Path templates
WORD_TEMPLATES_DIR = Path("templates/word")
os.makedirs(WORD_TEMPLATES_DIR, exist_ok=True)

def add_header(doc, title, subtitle=None):
    """Add standard header to document"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.bold = True
    
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(11)
    
    doc.add_paragraph()

def add_satker_info(doc):
    """Add satker info section"""
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Satker"
    cells[1].text = "{{satker_nama}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Kode"
    cells[1].text = "{{satker_kode}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Alamat"
    cells[1].text = "{{satker_alamat}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Kota"
    cells[1].text = "{{satker_kota}}"
    
    doc.add_paragraph()

def create_hps_word():
    """Create HPS Word template"""
    doc = Document()
    
    add_header(doc, "HARGA PERKIRAAN SENDIRI", "{{satker_nama}}")
    
    add_satker_info(doc)
    
    # Paket Info
    p = doc.add_paragraph("INFORMASI PAKET PENGADAAN")
    p.style = 'Heading 2'
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Nama Paket"
    cells[1].text = "{{nama_paket}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Jenis Pengadaan"
    cells[1].text = "{{jenis_pengadaan}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Nilai HPS"
    cells[1].text = "{{nilai_hps:rupiah}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Terbilang"
    cells[1].text = "{{nilai_hps_terbilang:terbilang}}"
    
    doc.add_paragraph()
    
    # Detail Pekerjaan
    p = doc.add_paragraph("DETAIL PEKERJAAN")
    p.style = 'Heading 2'
    
    doc.add_paragraph("Deskripsi: {{spesifikasi_teknis}}")
    doc.add_paragraph("Volume: {{volume_pekerjaan}}")
    doc.add_paragraph("Lokasi: {{lokasi_pekerjaan}}")
    
    doc.add_paragraph()
    
    # TTD
    p = doc.add_paragraph("PPK")
    p.style = 'Heading 2'
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{ppk_nama}}")
    doc.add_paragraph("NIP: {{ppk_nip}}")
    doc.add_paragraph("Tanggal: {{tanggal_hps:tanggal_long}}")
    
    doc.save(WORD_TEMPLATES_DIR / "hps.docx")
    print("✅ Dibuat: hps.docx")

def create_survey_harga_word():
    """Create Survey Harga Word template"""
    doc = Document()
    
    add_header(doc, "BERITA ACARA SURVEY HARGA", "{{satker_nama}}")
    
    add_satker_info(doc)
    
    # Paket Info
    p = doc.add_paragraph("PAKET YANG DISURVEY")
    p.style = 'Heading 2'
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Nama Paket"
    cells[1].text = "{{nama_paket}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Jenis Pengadaan"
    cells[1].text = "{{jenis_pengadaan}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Lokasi Pekerjaan"
    cells[1].text = "{{lokasi_pekerjaan}}"
    
    doc.add_paragraph()
    
    # Survey Detail
    p = doc.add_paragraph("HASIL SURVEY HARGA")
    p.style = 'Heading 2'
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "No"
    hdr_cells[1].text = "Nama Toko / Supplier"
    hdr_cells[2].text = "Tanggal Survey"
    hdr_cells[3].text = "Harga"
    
    for i in range(1, 4):
        cells = table.rows[i].cells
        cells[0].text = str(i)
        cells[1].text = f"{{{{survey{i}_toko}}}}"
        cells[2].text = f"{{{{survey{i}_tanggal:tanggal_short}}}}"
        cells[3].text = f"{{{{survey{i}_harga:rupiah}}}}"
    
    doc.add_paragraph()
    
    # Kesimpulan
    p = doc.add_paragraph("KESIMPULAN")
    p.style = 'Heading 2'
    doc.add_paragraph("HPS ditetapkan berdasarkan survey harga yang dilakukan pada tiga toko/supplier sebagai berikut:")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    p = doc.add_paragraph("PPK")
    p.style = 'Heading 2'
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{ppk_nama}}")
    doc.add_paragraph("NIP: {{ppk_nip}}")
    doc.add_paragraph("Tanggal: {{tanggal_hps:tanggal_long}}")
    
    doc.save(WORD_TEMPLATES_DIR / "survey_harga.docx")
    print("✅ Dibuat: survey_harga.docx")

def create_spm_ls():
    """Create SPM Langsung template"""
    doc = Document()
    
    add_header(doc, "SURAT PERINTAH MEMBAYAR LANGSUNG", "{{satker_nama}}")
    
    add_satker_info(doc)
    
    doc.add_paragraph()
    
    # Nomor SPM
    p = doc.add_paragraph("NOMOR SPM: {{nomor_spp:text}}")
    
    # Paket Info
    p = doc.add_paragraph("PAKET PENGADAAN")
    p.style = 'Heading 2'
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Nama Paket"
    cells[1].text = "{{nama_paket}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Nilai Kontrak"
    cells[1].text = "{{nilai_kontrak:rupiah}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Terbilang"
    cells[1].text = "{{nilai_kontrak_terbilang:terbilang}}"
    
    doc.add_paragraph()
    
    # Penyedia
    p = doc.add_paragraph("DATA PENYEDIA")
    p.style = 'Heading 2'
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Nama"
    cells[1].text = "{{penyedia_nama}}"
    
    cells = table.rows[1].cells
    cells[0].text = "NPWP"
    cells[1].text = "{{penyedia_npwp:npwp}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Nomor Rekening"
    cells[1].text = "{{penyedia_rekening}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Bank"
    cells[1].text = "{{penyedia_bank}}"
    
    doc.add_paragraph()
    
    # TTD
    p = doc.add_paragraph("TANDA TANGAN")
    p.style = 'Heading 2'
    
    doc.add_paragraph("Bendahara:")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{bendahara_nama}}")
    doc.add_paragraph("NIP: {{bendahara_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "spm_ls.docx")
    print("✅ Dibuat: spm_ls.docx")

def create_surat_permohonan_tup():
    """Create Surat Permohonan TUP template"""
    doc = Document()
    
    add_header(doc, "SURAT PERMOHONAN PEMBUKAAN TANDA TUNAI", "{{satker_nama}}")
    
    doc.add_paragraph("Kepada:")
    doc.add_paragraph("Yth. Bendahara Umum {{satker_kota}}")
    doc.add_paragraph()
    
    doc.add_paragraph("Kami mengajukan permohonan pembukaan Tanda Tunai (TUP) dengan data sebagai berikut:")
    doc.add_paragraph()
    
    # Data TUP
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Satker"
    cells[1].text = "{{satker_nama}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Periode"
    cells[1].text = "{{tahun_anggaran}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Nomor Surat"
    cells[1].text = "{{nomor_surat_tugas}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Tanggal"
    cells[1].text = "{{tanggal_surat_tugas:tanggal_long}}"
    
    cells = table.rows[4].cells
    cells[0].text = "Tujuan"
    cells[1].text = "Operasional kegiatan {{satker_nama}}"
    
    doc.add_paragraph()
    
    # Persetujuan
    p = doc.add_paragraph("PERSETUJUAN")
    p.style = 'Heading 2'
    
    doc.add_paragraph()
    doc.add_paragraph("PPK")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{ppk_nama}}")
    doc.add_paragraph("NIP: {{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "surat_permohonan_tup.docx")
    print("✅ Dibuat: surat_permohonan_tup.docx")

if __name__ == "__main__":
    print("=" * 70)
    print("MEMBUAT TEMPLATE DRAFT YANG HILANG")
    print("=" * 70)
    print()
    
    try:
        create_hps_word()
        create_survey_harga_word()
        create_spm_ls()
        create_surat_permohonan_tup()
        
        print()
        print("=" * 70)
        print("✅ SEMUA TEMPLATE BERHASIL DIBUAT!")
        print("=" * 70)
        print()
        print("Template tersimpan di: templates/word/")
        print("Total: 4 template draft baru")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
