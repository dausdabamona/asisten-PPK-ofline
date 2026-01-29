#!/usr/bin/env python3
"""
Create templates for Perjalanan Dinas (Business Travel) documents:
1. SPPD (Surat Perjalanan Dinas)
2. Rincian Biaya Perjalanan Dinas
3. Checklist Dokumen Perjalanan Dinas
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEMPLATE_DIR = Path(__file__).parent.parent / "templates" / "word"


def create_sppd_template():
    """Create SPPD (Surat Perjalanan Dinas) template."""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # === KOP SURAT ===
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kop.add_run("{{kop_surat}}")
    run.bold = True
    run.font.size = Pt(14)

    # Separator line
    sep = doc.add_paragraph("_" * 70)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === TITLE ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SURAT PERJALANAN DINAS (SPPD)")
    run.bold = True
    run.underline = True
    run.font.size = Pt(14)

    # Nomor
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_dokumen}}")

    doc.add_paragraph()

    # === PEJABAT YANG MEMBERI PERINTAH ===
    p1 = doc.add_paragraph()
    p1.add_run("I. Pejabat yang memberi perintah").bold = True

    # Table for pemberi perintah
    table1 = doc.add_table(rows=3, cols=2)
    table1.style = 'Table Grid'

    pemberi_data = [
        ("Nama", "{{nama_ppk}}"),
        ("NIP", "{{nip_ppk}}"),
        ("Jabatan", "Pejabat Pembuat Komitmen"),
    ]

    for i, (label, value) in enumerate(pemberi_data):
        row = table1.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()

    # === PEGAWAI YANG DIPERINTAH ===
    p2 = doc.add_paragraph()
    p2.add_run("II. Pegawai yang diperintahkan").bold = True

    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'

    pegawai_data = [
        ("Nama", "{{penerima_nama}}"),
        ("NIP", "{{penerima_nip}}"),
        ("Pangkat/Gol", "{{penerima_pangkat}}"),
        ("Jabatan", "{{penerima_jabatan}}"),
    ]

    for i, (label, value) in enumerate(pegawai_data):
        row = table2.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()

    # === PERJALANAN DINAS ===
    p3 = doc.add_paragraph()
    p3.add_run("III. Perjalanan Dinas").bold = True

    table3 = doc.add_table(rows=7, cols=2)
    table3.style = 'Table Grid'

    perjalanan_data = [
        ("a. Maksud Perjalanan", "{{tujuan_kegiatan}}"),
        ("b. Alat Angkut", "{{alat_angkut}}"),
        ("c. Tempat Berangkat", "{{kota}}"),
        ("d. Tempat Tujuan", "{{kota_tujuan}}"),
        ("e. Tanggal Berangkat", "{{tanggal_berangkat:tanggal}}"),
        ("f. Tanggal Kembali", "{{tanggal_kembali:tanggal}}"),
        ("g. Lama Perjalanan", "{{lama_hari}} ({{lama_hari_terbilang}}) hari"),
    ]

    for i, (label, value) in enumerate(perjalanan_data):
        row = table3.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)

    doc.add_paragraph()

    # === PEMBEBANAN ANGGARAN ===
    p4 = doc.add_paragraph()
    p4.add_run("IV. Pembebanan Anggaran").bold = True

    table4 = doc.add_table(rows=3, cols=2)
    table4.style = 'Table Grid'

    anggaran_data = [
        ("a. MAK", "{{mak}}"),
        ("b. Akun", "{{akun}}"),
        ("c. Instansi", "{{satker_nama}}"),
    ]

    for i, (label, value) in enumerate(anggaran_data):
        row = table4.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)

    doc.add_paragraph()
    doc.add_paragraph()

    # === TANDA TANGAN ===
    ttd_table = doc.add_table(rows=5, cols=1)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    cell0 = ttd_table.rows[0].cells[0]
    cell0.text = "{{kota}}, {{tanggal_dokumen:tanggal}}"
    cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell1 = ttd_table.rows[1].cells[0]
    cell1.text = "Pejabat Pembuat Komitmen"
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell2 = ttd_table.rows[2].cells[0]
    cell2.text = "\n\n\n"

    cell3 = ttd_table.rows[3].cells[0]
    p = cell3.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{nama_ppk}}")
    run.bold = True

    cell4 = ttd_table.rows[4].cells[0]
    cell4.text = "NIP. {{nip_ppk}}"
    cell4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === HALAMAN 2: TANDA TANGAN TUJUAN ===
    doc.add_page_break()

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("LEMBAR TANDA TANGAN PERJALANAN DINAS")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Tabel tanda tangan di lokasi tujuan
    tbl_ttd = doc.add_table(rows=5, cols=4)
    tbl_ttd.style = 'Table Grid'

    # Header
    headers = ["No", "Tiba di", "Tanggal/Cap/Tanda Tangan", "Berangkat dari"]
    for i, header in enumerate(headers):
        cell = tbl_ttd.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Row 1-4 kosong untuk diisi
    for row_idx in range(1, 5):
        row = tbl_ttd.rows[row_idx]
        row.cells[0].text = str(row_idx)
        row.cells[1].text = ""
        row.cells[2].text = "\n\n\n"
        row.cells[3].text = ""

    # Save
    output_path = TEMPLATE_DIR / "sppd.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


def create_rincian_biaya_pd_template():
    """Create Rincian Biaya Perjalanan Dinas template."""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # === TITLE ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RINCIAN BIAYA PERJALANAN DINAS")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # === INFO UMUM ===
    info_table = doc.add_table(rows=5, cols=2)

    info_data = [
        ("Nama Pelaksana", "{{penerima_nama}}"),
        ("NIP", "{{penerima_nip}}"),
        ("Jabatan", "{{penerima_jabatan}}"),
        ("Tujuan", "{{kota_tujuan}}"),
        ("Lama Perjalanan", "{{lama_hari}} hari"),
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"

    doc.add_paragraph()

    # === TABEL RINCIAN BIAYA ===
    biaya_table = doc.add_table(rows=10, cols=4)
    biaya_table.style = 'Table Grid'

    # Header
    headers = ["No", "Uraian", "Satuan", "Jumlah (Rp)"]
    for i, header in enumerate(headers):
        cell = biaya_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths
    biaya_table.rows[0].cells[0].width = Inches(0.5)
    biaya_table.rows[0].cells[1].width = Inches(3)
    biaya_table.rows[0].cells[2].width = Inches(1.5)
    biaya_table.rows[0].cells[3].width = Inches(1.5)

    # Rincian biaya
    biaya_items = [
        ("1", "Uang Harian", "{{lama_hari}} hari x Rp {{uang_harian}}", "{{total_uang_harian}}"),
        ("2", "Biaya Transport", "", "{{biaya_transport}}"),
        ("", "- Tiket PP", "{{jenis_tiket}}", "{{harga_tiket}}"),
        ("", "- Transport Lokal", "{{transport_lokal}} x Rp", "{{total_transport_lokal}}"),
        ("3", "Biaya Penginapan", "{{lama_inap}} malam x Rp {{harga_inap}}", "{{total_penginapan}}"),
        ("4", "Representasi", "", "{{representasi}}"),
        ("5", "Uang Saku Pertemuan", "", "{{uang_saku}}"),
        ("", "", "", ""),
        ("", "TOTAL", "", "Rp {{nilai:rupiah}}"),
    ]

    for row_idx, (no, uraian, satuan, jumlah) in enumerate(biaya_items, start=1):
        if row_idx >= len(biaya_table.rows):
            break
        row = biaya_table.rows[row_idx]
        row.cells[0].text = no
        row.cells[1].text = uraian
        row.cells[2].text = satuan
        row.cells[3].text = jumlah

        # Bold the TOTAL row
        if uraian == "TOTAL":
            for cell in row.cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # === TANDA TANGAN ===
    ttd_container = doc.add_table(rows=1, cols=2)

    # Kolom kiri - Pelaksana
    cell_left = ttd_container.rows[0].cells[0]
    p1 = cell_left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("Mengetahui,\n")
    p1.add_run("Pejabat Pembuat Komitmen\n\n\n\n")
    run1 = p1.add_run("{{nama_ppk}}\n")
    run1.bold = True
    p1.add_run("NIP. {{nip_ppk}}")

    # Kolom kanan - PPK
    cell_right = ttd_container.rows[0].cells[1]
    p2 = cell_right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("{{kota}}, {{tanggal_dokumen:tanggal}}\n")
    p2.add_run("Yang Melakukan Perjalanan,\n\n\n\n")
    run2 = p2.add_run("{{penerima_nama}}\n")
    run2.bold = True
    p2.add_run("NIP. {{penerima_nip}}")

    # Save
    output_path = TEMPLATE_DIR / "rincian_biaya_pd.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


def create_checklist_pd_template():
    """Create Checklist Dokumen Perjalanan Dinas template."""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # === KOP SURAT ===
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kop.add_run("{{kop_surat}}")
    run.bold = True
    run.font.size = Pt(14)

    # Separator line
    sep = doc.add_paragraph("_" * 70)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === TITLE ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CHECKLIST KELENGKAPAN DOKUMEN")
    run.bold = True
    run.font.size = Pt(14)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run("PERJALANAN DINAS")
    run2.bold = True
    run2.font.size = Pt(14)

    doc.add_paragraph()

    # === INFO UMUM ===
    info = doc.add_paragraph()
    info.add_run("Nama Pelaksana\t: {{penerima_nama}}\n")
    info.add_run("Tujuan\t\t: {{kota_tujuan}}\n")
    info.add_run("Tanggal\t\t: {{tanggal_berangkat:tanggal}} s.d. {{tanggal_kembali:tanggal}}")

    doc.add_paragraph()

    # === TABEL CHECKLIST ===
    checklist_table = doc.add_table(rows=15, cols=4)
    checklist_table.style = 'Table Grid'

    # Header
    headers = ["No", "Dokumen", "Ada", "Keterangan"]
    for i, header in enumerate(headers):
        cell = checklist_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths
    checklist_table.rows[0].cells[0].width = Inches(0.5)
    checklist_table.rows[0].cells[1].width = Inches(3.5)
    checklist_table.rows[0].cells[2].width = Inches(0.7)
    checklist_table.rows[0].cells[3].width = Inches(1.5)

    # Checklist items
    checklist_items = [
        "Surat Tugas",
        "SPPD (Surat Perjalanan Dinas)",
        "Rincian Biaya Perjalanan Dinas",
        "Kuitansi/Tanda Terima",
        "Tiket Pesawat/Kereta/Bus",
        "Boarding Pass",
        "Bill Hotel/Penginapan",
        "Bukti Pembayaran Transport Lokal",
        "Laporan Hasil Perjalanan Dinas",
        "Daftar Hadir/Bukti Kehadiran",
        "Dokumentasi Foto Kegiatan",
        "SPPD yang sudah ditandatangani",
        "Materai (jika nilai >= Rp 5 juta)",
        "Bukti Potong Pajak (jika ada)",
    ]

    for row_idx, item in enumerate(checklist_items, start=1):
        row = checklist_table.rows[row_idx]
        row.cells[0].text = str(row_idx)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = item
        row.cells[2].text = ""  # Untuk di-checklist
        row.cells[3].text = ""  # Keterangan

    doc.add_paragraph()
    doc.add_paragraph()

    # === TANDA TANGAN ===
    ttd_table = doc.add_table(rows=5, cols=1)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    cell0 = ttd_table.rows[0].cells[0]
    cell0.text = "{{kota}}, {{tanggal_dokumen:tanggal}}"
    cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell1 = ttd_table.rows[1].cells[0]
    cell1.text = "Diperiksa oleh,"
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell2 = ttd_table.rows[2].cells[0]
    cell2.text = "\n\n\n"

    cell3 = ttd_table.rows[3].cells[0]
    p = cell3.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{nama_ppk}}")
    run.bold = True

    cell4 = ttd_table.rows[4].cells[0]
    cell4.text = "Pejabat Pembuat Komitmen"
    cell4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = TEMPLATE_DIR / "checklist_perjalanan_dinas.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Creating Perjalanan Dinas templates...")
    print("-" * 50)

    create_sppd_template()
    create_rincian_biaya_pd_template()
    create_checklist_pd_template()

    print("-" * 50)
    print("Done! All templates created successfully.")
