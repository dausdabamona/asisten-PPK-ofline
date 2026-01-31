"""
Create Lembar Permintaan (Formulir Permintaan Realisasi Kegiatan) Template
Based on example from Politeknik Kelautan dan Perikanan Sorong

Features:
- Generate lembar permintaan document dari template
- Simpan data ke database SQLite
- Track dokumen dan item barang/jasa
"""

import sys
import os
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Tuple

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import database dan services
try:
    from app.models.pencairan_models import PencairanManager
    from app.services.dokumen_generator import DokumenGenerator, get_dokumen_generator
    from app.core.config import DATABASE_PATH
    HAS_DB = True
except ImportError:
    HAS_DB = False
    print("⚠️  Database/services not available - template mode only")


def add_horizontal_line(paragraph):
    """Add horizontal line under paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_background(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_lembar_permintaan_template():
    """Create Lembar Permintaan template."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Set default line spacing to 1
    style = doc.styles['Normal']
    style.font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.0
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('FORMULIR PERMINTAAN REALISASI KEGIATAN')
    title_run.bold = True
    title_run.font.size = Pt(12)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('POLITEKNIK KELAUTAN DAN PERIKANAN SORONG')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Header Information (3 columns)
    header_table = doc.add_table(rows=3, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    # Set column widths
    for row in header_table.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)
    
    # Row 1: Hari/Tanggal
    header_table.rows[0].cells[0].text = 'Hari/Tanggal'
    header_table.rows[0].cells[1].text = ': {{hari_tanggal}}'
    
    # Row 2: Unit Kerja
    header_table.rows[1].cells[0].text = 'Unit Kerja'
    header_table.rows[1].cells[1].text = ': {{unit_kerja}}'
    
    # Row 3: Sumber Dana / DIPA
    header_table.rows[2].cells[0].text = 'Sumber Dana'
    header_table.rows[2].cells[1].text = ': {{sumber_dana}}'
    
    # Remove table borders
    def remove_table_borders(table):
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)
    
    remove_table_borders(header_table)
    
    doc.add_paragraph()  # Spacing
    
    # Main table with items
    table = doc.add_table(rows=7, cols=8)
    table.style = 'Light Grid Accent 1'
    
    # Set column widths - auto fit ke content
    col_widths = [Inches(0.35), Inches(2.0), Inches(1.5), Inches(0.55), 
                  Inches(0.65), Inches(1.0), Inches(1.2), Inches(0.8)]
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = width
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['No', 'Nama Barang', 'Spesifikasi', 'Vol', 'Satuan', 
               'Harga Satuan', 'Total', 'Ket']
    
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        set_cell_background(cell, 'D3D3D3')
        
        # Format header text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.0
    
    # Data rows (rows 1-5 for items)
    for row_idx in range(1, 6):
        row = table.rows[row_idx]
        row.cells[0].text = f'{{{{item_{row_idx}_no}}}}'  # No
        row.cells[1].text = f'{{{{item_{row_idx}_nama}}}}'  # Nama Barang
        row.cells[2].text = f'{{{{item_{row_idx}_spek}}}}'  # Spesifikasi
        row.cells[3].text = f'{{{{item_{row_idx}_volume}}}}'  # Volume
        row.cells[4].text = f'{{{{item_{row_idx}_satuan}}}}'  # Satuan
        row.cells[5].text = f'{{{{item_{row_idx}_harga}}}}'  # Harga Satuan
        row.cells[6].text = f'{{{{item_{row_idx}_total}}}}'  # Total
        row.cells[7].text = f'{{{{item_{row_idx}_ket}}}}'  # Keterangan
        
        # Right align numeric columns
        for cell_idx in [3, 4, 5, 6]:
            row.cells[cell_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # SUB TOTAL row
    subtotal_row = table.rows[6]
    subtotal_row.cells[0].text = ''
    subtotal_row.cells[1].text = ''
    subtotal_row.cells[2].text = ''
    subtotal_row.cells[3].text = ''
    subtotal_row.cells[4].text = ''
    subtotal_row.cells[5].text = ''
    
    # Merge cells for SUB TOTAL label
    subtotal_row.cells[5].text = 'SUB TOTAL'
    subtotal_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtotal_row.cells[6].text = '{{subtotal}}'
    subtotal_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtotal_row.cells[7].text = ''
    
    # PPN row
    ppn_row_idx = len(table.rows)
    ppn_row = table.add_row()
    ppn_row.cells[0].text = ''
    ppn_row.cells[1].text = ''
    ppn_row.cells[2].text = ''
    ppn_row.cells[3].text = ''
    ppn_row.cells[4].text = ''
    ppn_row.cells[5].text = 'PPN 10%'
    ppn_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ppn_row.cells[6].text = '{{ppn}}'
    ppn_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ppn_row.cells[7].text = ''
    
    # TOTAL row
    total_row = table.add_row()
    for cell in total_row.cells:
        set_cell_background(cell, 'D3D3D3')
    
    total_row.cells[0].text = ''
    total_row.cells[1].text = ''
    total_row.cells[2].text = ''
    total_row.cells[3].text = ''
    total_row.cells[4].text = ''
    total_row.cells[5].text = 'TOTAL'
    total_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in total_row.cells[5].paragraphs[0].runs:
        run.font.bold = True
    
    total_row.cells[6].text = '{{total}}'
    total_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in total_row.cells[6].paragraphs[0].runs:
        run.font.bold = True
    total_row.cells[7].text = ''
    
    doc.add_paragraph()  # Spacing
    
    # Signature section - 5 penandatangan (2 baris: 2 atas, 3 bawah)
    sig_heading = doc.add_paragraph()
    sig_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_run = sig_heading.add_run('PENANDATANGAN DOKUMEN')
    sig_run.bold = True
    sig_run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Signature positions and codes
    signatures_row1 = [
        ('Yang Mengajukan', 'nama_pengajuan'),
        ('Verifikator\n(PPSPM)', 'nama_verifikator')
    ]
    
    signatures_row2 = [
        ('PPK\n(Pejabat Pembuat Komitmen)', 'nama_ppk'),
        ('Atasan Langsung\nUnit Terkait', 'nama_atasan'),
        ('KPA\n(Kuasa Pengguna Anggaran)', 'nama_kpa')
    ]
    
    # ROW 1: 2 penandatangan
    sig_table_row1 = doc.add_table(rows=5, cols=2)
    remove_table_borders(sig_table_row1)
    
    # Set column widths for row 1
    for row in sig_table_row1.rows:
        for cell in row.cells:
            cell.width = Inches(2.5)
    
    # Headers row 1
    for col_idx, (sig_title, _) in enumerate(signatures_row1):
        cell = sig_table_row1.rows[0].cells[col_idx]
        cell.text = sig_title
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
            para.paragraph_format.line_spacing = 1.0
    
    # Space for signatures (rows 1-3)
    for row_idx in range(1, 4):
        for col_idx in range(2):
            sig_table_row1.rows[row_idx].cells[col_idx].text = ''
    
    # Names row (row 4)
    for col_idx, (_, name_code) in enumerate(signatures_row1):
        cell = sig_table_row1.rows[4].cells[col_idx]
        cell.text = f'{{{{{name_code}}}}}'
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(8)
    
    doc.add_paragraph()  # Spacing between rows
    
    # ROW 2: 3 penandatangan
    sig_table_row2 = doc.add_table(rows=5, cols=3)
    remove_table_borders(sig_table_row2)
    
    # Set column widths for row 2
    for row in sig_table_row2.rows:
        for cell in row.cells:
            cell.width = Inches(1.8)
    
    # Headers row 2
    for col_idx, (sig_title, _) in enumerate(signatures_row2):
        cell = sig_table_row2.rows[0].cells[col_idx]
        cell.text = sig_title
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
            para.paragraph_format.line_spacing = 1.0
    
    # Space for signatures (rows 1-3)
    for row_idx in range(1, 4):
        for col_idx in range(3):
            sig_table_row2.rows[row_idx].cells[col_idx].text = ''
    
    # Names row (row 4)
    for col_idx, (_, name_code) in enumerate(signatures_row2):
        cell = sig_table_row2.rows[4].cells[col_idx]
        cell.text = f'{{{{{name_code}}}}}'
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(8)
    
    doc.add_paragraph()  # Spacing
    
    # Footer notes
    notes = doc.add_paragraph()
    notes_run = notes.add_run('Catatan Bagian Keuangan :')
    notes_run.bold = True
    notes_run.font.size = Pt(9)
    
    note_items = [
        'Jumlah dana sesuai MAK',
        'Realisasi s/d',
        'Sisa Dana',
        'Catatan : *Curet yang tidak perlu'
    ]
    
    for note_item in note_items:
        note_para = doc.add_paragraph(note_item, style='List Bullet')
        for run in note_para.runs:
            run.font.size = Pt(9)
        note_para.paragraph_format.line_spacing = 1.0
    
    # Save
    doc.save('templates/word/lembar_permintaan.docx')
    print("✓ Lembar Permintaan template created: templates/word/lembar_permintaan.docx")


def create_lembar_permintaan(
    hari_tanggal: str,
    unit_kerja: str,
    sumber_dana: str = '',
    items: List[Dict] = None,
    signatures: Dict[str, str] = None,
    kode_transaksi: str = None,
    created_by: str = 'system'
) -> Tuple[Optional[int], Optional[str], Optional[str]]:
    """
    Generate dan simpan lembar permintaan ke database.
    
    Args:
        hari_tanggal: Tanggal (YYYY-MM-DD atau format lain)
        unit_kerja: Nama unit kerja
        sumber_dana: Sumber dana / DIPA
        items: List of item dict dengan keys: item_no, nama_barang, spesifikasi,
               volume, satuan, harga_satuan, total_item, keterangan
        signatures: Dict dengan keys: nama_pengajuan, nama_verifikator, nama_ppk,
                   nama_atasan, nama_kpa
        kode_transaksi: Optional kode transaksi untuk link ke transaksi_pencairan
        created_by: Username yang membuat
    
    Returns:
        Tuple of (lembar_id, doc_path, error_message)
    """
    try:
        if not HAS_DB:
            return None, None, "Database modules tidak tersedia. Install app modules terlebih dahulu."
        
        # Initialize database
        db_manager = PencairanManager(DATABASE_PATH)
        doc_generator = get_dokumen_generator(db_manager)
        
        # Calculate totals dari items
        subtotal = 0
        ppn = 0
        total = 0
        
        if items:
            for item in items:
                item_total = item.get('total_item', 0) or 0
                subtotal += item_total
            
            ppn = subtotal * 0.10  # PPN 10%
            total = subtotal + ppn
        
        # Prepare lembar data
        lembar_data = {
            'hari_tanggal': hari_tanggal,
            'unit_kerja': unit_kerja,
            'sumber_dana': sumber_dana,
            'kode_transaksi': kode_transaksi,
            'subtotal': subtotal,
            'ppn': ppn,
            'total': total,
            'status': 'draft',
            'created_by': created_by,
        }
        
        # Add signatures
        if signatures:
            lembar_data.update({
                'nama_pengajuan': signatures.get('nama_pengajuan', ''),
                'nama_verifikator': signatures.get('nama_verifikator', ''),
                'nama_ppk': signatures.get('nama_ppk', ''),
                'nama_atasan': signatures.get('nama_atasan', ''),
                'nama_kpa': signatures.get('nama_kpa', ''),
            })
        
        # Generate dokumen dan simpan ke database
        lembar_id, doc_path, error = doc_generator.generate_and_save_lembar_permintaan(
            lembar_data,
            items=items,
            created_by=created_by
        )
        
        if error:
            return None, None, error
        
        print(f"✓ Lembar Permintaan berhasil dibuat:")
        print(f"  ID: {lembar_id}")
        print(f"  Unit Kerja: {unit_kerja}")
        print(f"  Total: Rp {total:,.0f}".replace(",", "."))
        print(f"  File: {doc_path}")
        
        return lembar_id, doc_path, None
        
    except Exception as e:
        error_msg = f"Error create_lembar_permintaan: {str(e)}"
        print(f"❌ {error_msg}")
        return None, None, error_msg


def create_lembar_permintaan_from_transaksi(transaksi_id: int) -> Tuple[Optional[int], Optional[str], Optional[str]]:
    """
    Generate lembar permintaan dari data transaksi_pencairan yang sudah ada.
    
    Args:
        transaksi_id: ID transaksi_pencairan
    
    Returns:
        Tuple of (lembar_id, doc_path, error_message)
    """
    try:
        if not HAS_DB:
            return None, None, "Database modules tidak tersedia."
        
        db_manager = PencairanManager(DATABASE_PATH)
        
        # Get transaksi data
        transaksi = db_manager.get_transaksi(transaksi_id)
        if not transaksi:
            return None, None, f"Transaksi {transaksi_id} tidak ditemukan"
        
        # Get rincian transaksi jika ada
        rincian = db_manager.get_rincian_transaksi(transaksi_id) if hasattr(db_manager, 'get_rincian_transaksi') else []
        
        # Prepare items dari rincian
        items = []
        if rincian:
            for idx, item in enumerate(rincian, 1):
                items.append({
                    'item_no': idx,
                    'nama_barang': item.get('nama_barang', ''),
                    'spesifikasi': item.get('spesifikasi', ''),
                    'volume': item.get('volume', 0),
                    'satuan': item.get('satuan', ''),
                    'harga_satuan': item.get('harga_satuan', 0),
                    'total_item': item.get('total_item', 0),
                    'keterangan': item.get('keterangan', ''),
                })
        
        # Prepare signatures
        signatures = {
            'nama_pengajuan': transaksi.get('penerima_nama', ''),
            'nama_verifikator': '',
            'nama_ppk': '',
            'nama_atasan': '',
            'nama_kpa': '',
        }
        
        # Create lembar permintaan
        return create_lembar_permintaan(
            hari_tanggal=datetime.now().strftime('%Y-%m-%d'),
            unit_kerja=transaksi.get('nama_kegiatan', ''),
            sumber_dana=transaksi.get('nomor_mak', ''),
            items=items,
            signatures=signatures,
            kode_transaksi=transaksi.get('kode_transaksi'),
            created_by='system'
        )
        
    except Exception as e:
        error_msg = f"Error create_lembar_permintaan_from_transaksi: {str(e)}"
        print(f"❌ {error_msg}")
        return None, None, error_msg


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Generate Lembar Permintaan'
    )
    parser.add_argument(
        '--template',
        action='store_true',
        help='Generate template saja (tanpa database)'
    )
    parser.add_argument(
        '--transaksi-id',
        type=int,
        help='Create dari transaksi_pencairan ID'
    )
    parser.add_argument(
        '--unit-kerja',
        type=str,
        help='Unit kerja'
    )
    parser.add_argument(
        '--sumber-dana',
        type=str,
        help='Sumber dana'
    )
    
    args = parser.parse_args()
    
    if args.template:
        # Template mode only
        create_lembar_permintaan_template()
    
    elif args.transaksi_id:
        # Create dari transaksi
        lembar_id, doc_path, error = create_lembar_permintaan_from_transaksi(args.transaksi_id)
        if error:
            print(f"❌ {error}")
            sys.exit(1)
        else:
            print(f"✓ Lembar Permintaan ID {lembar_id} created")
            sys.exit(0)
    
    else:
        # Interactive mode
        if not HAS_DB:
            print("⚠️  Database tidak tersedia, generate template mode")
            create_lembar_permintaan_template()
        else:
            print("=== Generate Lembar Permintaan ===\n")
            
            # Collect input
            from datetime import datetime
            hari_tanggal = input(f"Hari/Tanggal [{datetime.now().strftime('%Y-%m-%d')}]: ").strip() or datetime.now().strftime('%Y-%m-%d')
            unit_kerja = input("Unit Kerja: ").strip()
            sumber_dana = input("Sumber Dana (opsional): ").strip()
            
            # Ask for items
            items = []
            print("\nTambah barang/jasa (ENTER 'done' untuk selesai):")
            item_no = 1
            while True:
                print(f"\n--- Item {item_no} ---")
                nama = input(f"Nama barang [{item_no}. ]: ").strip()
                if nama.lower() == 'done':
                    break
                
                spek = input("Spesifikasi: ").strip()
                vol = float(input("Volume: ").strip() or "0")
                satuan = input("Satuan: ").strip()
                harga = float(input("Harga Satuan: ").strip() or "0")
                total_item = vol * harga
                ket = input("Keterangan: ").strip()
                
                items.append({
                    'item_no': item_no,
                    'nama_barang': nama,
                    'spesifikasi': spek,
                    'volume': vol,
                    'satuan': satuan,
                    'harga_satuan': harga,
                    'total_item': total_item,
                    'keterangan': ket,
                })
                
                item_no += 1
            
            # Create lembar
            lembar_id, doc_path, error = create_lembar_permintaan(
                hari_tanggal=hari_tanggal,
                unit_kerja=unit_kerja,
                sumber_dana=sumber_dana,
                items=items or None,
            )
            
            if error:
                print(f"\n❌ Error: {error}")
                sys.exit(1)
            else:
                print(f"\n✓ Berhasil! Lembar Permintaan ID: {lembar_id}")

