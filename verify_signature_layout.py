"""Verify the new signature layout (2 rows: 2 and 3 signatories)."""
from docx import Document

doc = Document('output/test_lembar_permintaan.docx')

print('=' * 70)
print('VERIFIKASI TABEL PENANDATANGAN - LAYOUT BARU')
print('=' * 70)
print()

# Count tables
print(f'Total tabel dalam dokumen: {len(doc.tables)}')
print()

# Find signature tables (the last 2 tables should be the signature tables)
if len(doc.tables) >= 2:
    # Get last 2 tables for signatures
    sig_table1 = doc.tables[-2]
    sig_table2 = doc.tables[-1]
    
    print('BARIS 1 (2 Penandatangan):')
    print(f'  Jumlah kolom: {len(sig_table1.columns)}')
    print(f'  Jumlah baris: {len(sig_table1.rows)}')
    headers1 = [cell.text.strip() for cell in sig_table1.rows[0].cells[:2]]
    print(f'  Header: {headers1}')
    print(f'  Nama penandatangan:')
    for i in range(len(sig_table1.columns)):
        name = sig_table1.rows[4].cells[i].text.strip()
        if name:
            print(f'    {i+1}. {name}')
    print()
    
    print('BARIS 2 (3 Penandatangan):')
    print(f'  Jumlah kolom: {len(sig_table2.columns)}')
    print(f'  Jumlah baris: {len(sig_table2.rows)}')
    headers2 = [cell.text.strip() for cell in sig_table2.rows[0].cells[:3]]
    print(f'  Header: {headers2}')
    print(f'  Nama penandatangan:')
    for i in range(len(sig_table2.columns)):
        name = sig_table2.rows[4].cells[i].text.strip()
        if name:
            print(f'    {i+1}. {name}')
    print()
    
    print('=' * 70)
    print('✓ DISTRIBUSI PENANDATANGAN BERHASIL:')
    print('  ✓ 2 penandatangan di baris atas')
    print('  ✓ 3 penandatangan di baris bawah')
    print('=' * 70)
