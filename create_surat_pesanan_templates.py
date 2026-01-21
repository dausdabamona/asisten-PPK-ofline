"""
Script untuk membuat template Surat Pesanan dan SPMK per jenis pengadaan
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates", "word")
os.makedirs(TEMPLATES_DIR, exist_ok=True)


def add_heading(doc, text, level=1, center=False):
    """Add heading with custom style"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_paragraph(doc, text, bold=False, center=False):
    """Add a paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def create_surat_pesanan():
    """Surat Pesanan untuk Pengadaan Barang"""
    doc = Document()

    # Header
    add_heading(doc, "{{kementerian}}", 1, True)
    add_heading(doc, "{{eselon1}}", 2, True)
    add_heading(doc, "{{satker_nama}}", 2, True)
    add_paragraph(doc, "{{satker_alamat}}, {{satker_kota}} - {{satker_provinsi}}", center=True)

    doc.add_paragraph()
    add_heading(doc, "SURAT PESANAN", 1, True)
    add_paragraph(doc, "Nomor: {{nomor_spmk}}", center=True)

    doc.add_paragraph()

    # Dasar
    add_paragraph(doc, "Menindaklanjuti:", bold=True)
    add_paragraph(doc, "1. Surat Perintah Kerja/Kontrak Nomor: {{nomor_spk}} tanggal {{tanggal_spk}}")
    add_paragraph(doc, "2. Berita Acara Hasil Pengadaan Langsung Nomor: ............ tanggal ............")

    doc.add_paragraph()
    add_paragraph(doc, "Dengan ini kami memesan barang kepada:", bold=True)

    # Tabel Penyedia
    table = doc.add_table(rows=6, cols=2)
    items = [
        ("Nama Perusahaan", "{{penyedia_nama}}"),
        ("Alamat", "{{penyedia_alamat}}, {{penyedia_kota}}"),
        ("NPWP", "{{penyedia_npwp}}"),
        ("Nama Direktur", "{{direktur_nama}}"),
        ("No. Rekening", "{{penyedia_rekening}}"),
        ("Bank", "{{penyedia_bank}}"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"

    doc.add_paragraph()
    add_paragraph(doc, "Untuk menyediakan barang dengan rincian sebagai berikut:", bold=True)

    # Tabel Barang
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'

    # Header
    hdr = table.rows[0]
    headers = ["No", "Uraian Barang", "Satuan", "Volume", "Harga (Rp)"]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h

    # Sample rows
    for i in range(1, 4):
        table.rows[i].cells[0].text = str(i)

    # Total row
    table.rows[4].cells[0].text = ""
    table.rows[4].cells[1].text = "TOTAL"
    table.rows[4].cells[4].text = "{{nilai_kontrak}}"

    doc.add_paragraph()
    add_paragraph(doc, "Terbilang: {{nilai_kontrak_terbilang}}")

    doc.add_paragraph()
    add_paragraph(doc, "Ketentuan:", bold=True)
    add_paragraph(doc, "1. Barang harus diserahkan paling lambat tanggal {{tanggal_selesai}}")
    add_paragraph(doc, "2. Barang harus sesuai dengan spesifikasi yang telah disepakati")
    add_paragraph(doc, "3. Barang diserahkan ke: {{lokasi_pekerjaan}}")
    add_paragraph(doc, "4. Pembayaran akan dilakukan setelah barang diterima dengan baik")

    doc.add_paragraph()
    add_paragraph(doc, "Demikian Surat Pesanan ini dibuat untuk dilaksanakan.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Yang Menerima Pesanan,\n{{penyedia_nama}}\n\n\n\n{{direktur_nama}}\n{{direktur_jabatan}}"
    table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_spmk}}\n{{ppk_jabatan}},\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "surat_pesanan.docx"))
    print("Created: surat_pesanan.docx")


def create_spmk_jasa_lainnya():
    """SPMK untuk Jasa Lainnya"""
    doc = Document()

    # Header
    add_heading(doc, "{{kementerian}}", 1, True)
    add_heading(doc, "{{eselon1}}", 2, True)
    add_heading(doc, "{{satker_nama}}", 2, True)
    add_paragraph(doc, "{{satker_alamat}}, {{satker_kota}} - {{satker_provinsi}}", center=True)

    doc.add_paragraph()
    add_heading(doc, "SURAT PERINTAH MULAI KERJA (SPMK)", 1, True)
    add_paragraph(doc, "Nomor: {{nomor_spmk}}", center=True)

    doc.add_paragraph()

    add_paragraph(doc, "Yang bertanda tangan di bawah ini:")

    # Tabel PPK
    table = doc.add_table(rows=4, cols=2)
    items = [
        ("Nama", "{{ppk_nama}}"),
        ("NIP", "{{ppk_nip}}"),
        ("Jabatan", "{{ppk_jabatan}}"),
        ("", "Selaku Pejabat Pembuat Komitmen"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}" if label else value

    doc.add_paragraph()
    add_paragraph(doc, "Berdasarkan Surat Perintah Kerja/Kontrak Nomor: {{nomor_spk}} tanggal {{tanggal_spk}}, dengan ini memerintahkan kepada:")

    # Tabel Penyedia
    table = doc.add_table(rows=5, cols=2)
    items = [
        ("Nama Perusahaan", "{{penyedia_nama}}"),
        ("Alamat", "{{penyedia_alamat}}, {{penyedia_kota}}"),
        ("Direktur", "{{direktur_nama}}"),
        ("NPWP", "{{penyedia_npwp}}"),
        ("", "Selanjutnya disebut PENYEDIA"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}" if label else value

    doc.add_paragraph()
    add_paragraph(doc, "Untuk segera memulai pelaksanaan pekerjaan:", bold=True)

    # Tabel pekerjaan
    table = doc.add_table(rows=5, cols=2)
    items = [
        ("Nama Pekerjaan", "{{nama_paket}}"),
        ("Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
        ("Nilai Kontrak", "Rp {{nilai_kontrak}} ({{nilai_kontrak_terbilang}})"),
        ("Jangka Waktu", "{{jangka_waktu}} hari kalender"),
        ("Tanggal Selesai", "{{tanggal_selesai}}"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"

    doc.add_paragraph()
    add_paragraph(doc, "Ketentuan:", bold=True)
    add_paragraph(doc, "1. PENYEDIA wajib memulai pekerjaan terhitung sejak tanggal {{tanggal_mulai}}")
    add_paragraph(doc, "2. PENYEDIA wajib menyelesaikan pekerjaan sesuai dengan Kerangka Acuan Kerja (KAK)")
    add_paragraph(doc, "3. Pekerjaan harus diselesaikan dalam jangka waktu {{jangka_waktu}} hari kalender")
    add_paragraph(doc, "4. Apabila terjadi keterlambatan, akan dikenakan denda sesuai ketentuan kontrak")

    doc.add_paragraph()
    add_paragraph(doc, "Demikian Surat Perintah Mulai Kerja ini dibuat untuk dilaksanakan dengan penuh tanggung jawab.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Yang Menerima Perintah,\n{{penyedia_nama}}\n\n\n\n{{direktur_nama}}\n{{direktur_jabatan}}"
    table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_spmk}}\n{{ppk_jabatan}},\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "spmk_jasa_lainnya.docx"))
    print("Created: spmk_jasa_lainnya.docx")


def create_spmk_konstruksi():
    """SPMK untuk Konstruksi"""
    doc = Document()

    # Header
    add_heading(doc, "{{kementerian}}", 1, True)
    add_heading(doc, "{{eselon1}}", 2, True)
    add_heading(doc, "{{satker_nama}}", 2, True)
    add_paragraph(doc, "{{satker_alamat}}, {{satker_kota}} - {{satker_provinsi}}", center=True)

    doc.add_paragraph()
    add_heading(doc, "SURAT PERINTAH MULAI KERJA (SPMK)", 1, True)
    add_heading(doc, "PEKERJAAN KONSTRUKSI", 2, True)
    add_paragraph(doc, "Nomor: {{nomor_spmk}}", center=True)

    doc.add_paragraph()

    add_paragraph(doc, "Yang bertanda tangan di bawah ini:")

    # Tabel PPK
    table = doc.add_table(rows=4, cols=2)
    items = [
        ("Nama", "{{ppk_nama}}"),
        ("NIP", "{{ppk_nip}}"),
        ("Jabatan", "{{ppk_jabatan}}"),
        ("", "Selaku Pejabat Pembuat Komitmen"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}" if label else value

    doc.add_paragraph()
    add_paragraph(doc, "Berdasarkan Surat Perjanjian/Kontrak Nomor: {{nomor_spk}} tanggal {{tanggal_spk}}, dengan ini memerintahkan kepada:")

    # Tabel Penyedia
    table = doc.add_table(rows=5, cols=2)
    items = [
        ("Nama Perusahaan", "{{penyedia_nama}}"),
        ("Alamat", "{{penyedia_alamat}}, {{penyedia_kota}}"),
        ("Direktur", "{{direktur_nama}}"),
        ("NPWP", "{{penyedia_npwp}}"),
        ("", "Selanjutnya disebut PENYEDIA JASA"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}" if label else value

    doc.add_paragraph()
    add_paragraph(doc, "Untuk segera memulai pelaksanaan pekerjaan konstruksi:", bold=True)

    # Tabel pekerjaan
    table = doc.add_table(rows=6, cols=2)
    items = [
        ("Nama Pekerjaan", "{{nama_paket}}"),
        ("Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
        ("Nilai Kontrak", "Rp {{nilai_kontrak}} ({{nilai_kontrak_terbilang}})"),
        ("Jangka Waktu Pelaksanaan", "{{jangka_waktu}} hari kalender"),
        ("Masa Pemeliharaan", "{{masa_pemeliharaan}} hari kalender"),
        ("Tanggal Selesai", "{{tanggal_selesai}}"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"

    doc.add_paragraph()
    add_paragraph(doc, "Ketentuan:", bold=True)
    add_paragraph(doc, "1. PENYEDIA JASA wajib memulai mobilisasi dan pekerjaan persiapan terhitung sejak tanggal {{tanggal_mulai}}")
    add_paragraph(doc, "2. PENYEDIA JASA wajib melaksanakan pekerjaan sesuai dengan gambar teknis dan spesifikasi")
    add_paragraph(doc, "3. Pekerjaan harus diselesaikan dalam jangka waktu {{jangka_waktu}} hari kalender")
    add_paragraph(doc, "4. PENYEDIA JASA wajib menunjuk Pelaksana Lapangan dan menyerahkan jadwal pelaksanaan (Time Schedule)")
    add_paragraph(doc, "5. PENYEDIA JASA wajib melaporkan progres pekerjaan secara berkala")
    add_paragraph(doc, "6. Apabila terjadi keterlambatan, akan dikenakan denda sesuai ketentuan kontrak")

    doc.add_paragraph()
    add_paragraph(doc, "Demikian Surat Perintah Mulai Kerja ini dibuat untuk dilaksanakan dengan penuh tanggung jawab.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Tanda tangan
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Yang Menerima Perintah,\n{{penyedia_nama}}\n\n\n\n{{direktur_nama}}\n{{direktur_jabatan}}"
    table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_spmk}}\n{{ppk_jabatan}},\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "spmk_konstruksi.docx"))
    print("Created: spmk_konstruksi.docx")


def main():
    print("Creating Surat Pesanan & SPMK templates...")
    create_surat_pesanan()
    create_spmk_jasa_lainnya()
    create_spmk_konstruksi()
    print("\nAll templates created successfully!")


if __name__ == "__main__":
    main()
