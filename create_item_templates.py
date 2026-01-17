#!/usr/bin/env python3
"""
PPK DOCUMENT FACTORY v3.0 - TEMPLATE WITH ITEM LOOPS
=====================================================
Templates dengan dukungan loop untuk item_barang
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# Output directory
TEMPLATE_DIR = os.path.dirname(os.path.abspath(__file__))
WORD_DIR = os.path.join(TEMPLATE_DIR, "templates", "word")
EXCEL_DIR = os.path.join(TEMPLATE_DIR, "templates", "excel")

os.makedirs(WORD_DIR, exist_ok=True)
os.makedirs(EXCEL_DIR, exist_ok=True)


def create_spesifikasi_teknis_with_items():
    """Template Spesifikasi Teknis dengan tabel item"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SPESIFIKASI TEKNIS")
    run.bold = True
    run.font.size = Pt(14)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("{{nama_paket}}")
    
    doc.add_paragraph()
    
    # Latar Belakang
    p = doc.add_paragraph()
    run = p.add_run("A. LATAR BELAKANG")
    run.bold = True
    p = doc.add_paragraph("{{latar_belakang}}")
    doc.add_paragraph()
    
    # Ruang Lingkup
    p = doc.add_paragraph()
    run = p.add_run("B. RUANG LINGKUP PEKERJAAN")
    run.bold = True
    p = doc.add_paragraph("{{ruang_lingkup}}")
    doc.add_paragraph()
    
    # RINCIAN ITEM BARANG
    p = doc.add_paragraph()
    run = p.add_run("C. RINCIAN SPESIFIKASI BARANG/JASA")
    run.bold = True
    
    doc.add_paragraph()
    
    # Create table with item loop
    table = doc.add_table(rows=3, cols=7)
    table.style = 'Table Grid'
    
    # Headers
    headers = ["No", "Uraian", "Spesifikasi", "Satuan", "Volume", "Harga Satuan", "Jumlah"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Template row with item placeholders (will be duplicated for each item)
    item_row = table.rows[1]
    item_row.cells[0].text = "{{no}}"
    item_row.cells[1].text = "{{item.uraian}}"
    item_row.cells[2].text = "{{item.spesifikasi}}"
    item_row.cells[3].text = "{{item.satuan}}"
    item_row.cells[4].text = "{{item.volume}}"
    item_row.cells[5].text = "{{item.harga_satuan}}"
    item_row.cells[6].text = "{{item.total}}"
    
    # Total row
    total_row = table.rows[2]
    total_row.cells[0].merge(total_row.cells[4])
    total_row.cells[0].text = "TOTAL"
    total_row.cells[0].paragraphs[0].runs[0].bold = True
    total_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row.cells[5].text = ""
    total_row.cells[6].text = "{{subtotal_item_fmt}}"
    total_row.cells[6].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Waktu
    p = doc.add_paragraph()
    run = p.add_run("D. WAKTU PELAKSANAAN")
    run.bold = True
    p = doc.add_paragraph()
    p.add_run("Pekerjaan dilaksanakan selama {{jangka_waktu}} hari kalender.")
    
    doc.add_paragraph()
    
    # TTD
    ttd_table = doc.add_table(rows=5, cols=2)
    ttd_table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_dokumen:tanggal_long}}"
    ttd_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen,"
    ttd_table.rows[2].cells[1].text = "\n\n\n"
    p = ttd_table.rows[3].cells[1].paragraphs[0]
    run = p.add_run("{{ppk_nama}}")
    run.bold = True
    run.underline = True
    ttd_table.rows[4].cells[1].text = "NIP. {{ppk_nip:nip}}"
    
    for row in ttd_table.rows:
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(WORD_DIR, "spesifikasi_teknis.docx")
    doc.save(filepath)
    print(f"✅ Created: spesifikasi_teknis.docx (with item loop)")
    return filepath


def create_hps_excel():
    """Template HPS Excel dengan item loop"""
    wb = Workbook()
    ws = wb.active
    ws.title = "HPS"
    
    # Styles
    header_font = Font(bold=True, size=11)
    title_font = Font(bold=True, size=14)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    right = Alignment(horizontal='right', vertical='center')
    
    # Title
    ws.merge_cells('A1:G1')
    ws['A1'] = 'HARGA PERKIRAAN SENDIRI (HPS)'
    ws['A1'].font = title_font
    ws['A1'].alignment = center
    
    ws.merge_cells('A2:G2')
    ws['A2'] = '{{nama_paket}}'
    ws['A2'].alignment = center
    
    # Info
    ws['A4'] = 'Satuan Kerja'
    ws['B4'] = ':'
    ws['C4'] = '{{satker_nama}}'
    
    ws['A5'] = 'Tahun Anggaran'
    ws['B5'] = ':'
    ws['C5'] = '{{tahun_anggaran}}'
    
    ws['A6'] = 'Pagu Anggaran'
    ws['B6'] = ':'
    ws['C6'] = '{{nilai_pagu:rupiah}}'
    
    # Table header
    row = 8
    headers = ['No', 'Uraian Barang/Jasa', 'Spesifikasi', 'Satuan', 'Volume', 'Harga Satuan', 'Jumlah']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=header)
        cell.font = header_font
        cell.border = border
        cell.alignment = center
        cell.fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    
    # Item row template (row 9 will be duplicated)
    row = 9
    item_data = ['{{no}}', '{{item.uraian}}', '{{item.spesifikasi}}', 
                 '{{item.satuan}}', '{{item.volume}}', '{{item.harga_dasar}}', '{{item.total}}']
    for col, data in enumerate(item_data, 1):
        cell = ws.cell(row=row, column=col, value=data)
        cell.border = border
    
    # Subtotal row (will be below items after processing)
    row = 10
    ws.merge_cells(f'A{row}:E{row}')
    ws[f'A{row}'] = 'SUBTOTAL'
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].alignment = right
    ws[f'A{row}'].border = border
    ws[f'F{row}'] = ''
    ws[f'F{row}'].border = border
    ws[f'G{row}'] = '{{subtotal_item}}'
    ws[f'G{row}'].font = header_font
    ws[f'G{row}'].border = border
    
    # PPN
    row = 11
    ws.merge_cells(f'A{row}:E{row}')
    ws[f'A{row}'] = 'PPN 11%'
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].alignment = right
    ws[f'A{row}'].border = border
    ws[f'F{row}'] = ''
    ws[f'F{row}'].border = border
    ws[f'G{row}'] = '{{ppn_item}}'
    ws[f'G{row}'].font = header_font
    ws[f'G{row}'].border = border
    
    # Total HPS
    row = 12
    ws.merge_cells(f'A{row}:E{row}')
    ws[f'A{row}'] = 'TOTAL HPS'
    ws[f'A{row}'].font = Font(bold=True, size=12)
    ws[f'A{row}'].alignment = right
    ws[f'A{row}'].fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
    ws[f'A{row}'].border = border
    ws[f'F{row}'] = ''
    ws[f'F{row}'].border = border
    ws[f'G{row}'] = '{{grand_total_item}}'
    ws[f'G{row}'].font = Font(bold=True, size=12)
    ws[f'G{row}'].fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
    ws[f'G{row}'].border = border
    
    # Terbilang
    row = 13
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = 'Terbilang: {{grand_total_item_terbilang}}'
    ws[f'A{row}'].font = Font(italic=True)
    
    # Column widths
    widths = [5, 35, 25, 10, 10, 18, 20]
    for i, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width
    
    # TTD
    ws['E16'] = '{{satker_kota}}, {{tanggal_hps:tanggal_long}}'
    ws['E17'] = 'Pejabat Pembuat Komitmen,'
    ws['E20'] = '{{ppk_nama}}'
    ws['E20'].font = Font(bold=True, underline='single')
    ws['E21'] = 'NIP. {{ppk_nip:nip}}'
    
    filepath = os.path.join(EXCEL_DIR, "hps.xlsx")
    wb.save(filepath)
    print(f"✅ Created: hps.xlsx (with item loop)")
    return filepath


def create_survey_harga_excel():
    """Template Survey Harga Excel dengan item loop"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Survey Harga"
    
    # Styles
    header_font = Font(bold=True, size=11)
    title_font = Font(bold=True, size=14)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # Title
    ws.merge_cells('A1:J1')
    ws['A1'] = 'SURVEY HARGA PASAR'
    ws['A1'].font = title_font
    ws['A1'].alignment = center
    
    ws.merge_cells('A2:J2')
    ws['A2'] = '{{nama_paket}}'
    ws['A2'].alignment = center
    
    ws['A4'] = 'Tanggal Survey:'
    ws['B4'] = '{{tanggal_survey:tanggal_long}}'
    
    # Table header
    row = 6
    headers = ['No', 'Uraian', 'Spesifikasi', 'Satuan', 'Volume',
               'Toko 1', 'Toko 2', 'Toko 3', 'Rata-rata', 'Jumlah']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=header)
        cell.font = header_font
        cell.border = border
        cell.alignment = center
    
    # Item row template
    row = 7
    item_data = ['{{no}}', '{{item.uraian}}', '{{item.spesifikasi}}', '{{item.satuan}}',
                 '{{item.volume}}', '{{item.harga_survey1}}', '{{item.harga_survey2}}',
                 '{{item.harga_survey3}}', '{{item.harga_rata}}', '{{item.total}}']
    for col, data in enumerate(item_data, 1):
        cell = ws.cell(row=row, column=col, value=data)
        cell.border = border
    
    # Total row
    row = 8
    ws.merge_cells(f'A{row}:I{row}')
    ws[f'A{row}'] = 'TOTAL'
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].alignment = Alignment(horizontal='right')
    ws[f'A{row}'].border = border
    ws[f'J{row}'] = '{{subtotal_item}}'
    ws[f'J{row}'].font = header_font
    ws[f'J{row}'].border = border
    
    # Column widths
    widths = [5, 25, 20, 8, 8, 15, 15, 15, 15, 18]
    for i, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width
    
    # Lokasi Survey
    ws['A10'] = 'Lokasi Survey:'
    ws['A11'] = '1. {{survey1_toko}} - {{survey1_alamat}}'
    ws['A12'] = '2. {{survey2_toko}} - {{survey2_alamat}}'
    ws['A13'] = '3. {{survey3_toko}} - {{survey3_alamat}}'
    
    # TTD
    ws['H15'] = '{{satker_kota}}, {{tanggal_survey:tanggal_long}}'
    ws['H16'] = 'Petugas Survey,'
    ws['H19'] = '{{surveyor_nama}}'
    ws['H19'].font = Font(bold=True, underline='single')
    ws['H20'] = 'NIP. {{surveyor_nip:nip}}'
    
    filepath = os.path.join(EXCEL_DIR, "survey_harga.xlsx")
    wb.save(filepath)
    print(f"✅ Created: survey_harga.xlsx (with item loop)")
    return filepath


def create_drpp_word():
    """Template DRPP dengan item loop"""
    doc = Document()
    
    # Set landscape
    for section in doc.sections:
        section.orientation = 1
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DAFTAR RINCIAN PERMINTAAN PEMBAYARAN (DRPP)")
    run.bold = True
    run.font.size = Pt(12)
    
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_drpp}}")
    
    doc.add_paragraph()
    
    # Info
    info_table = doc.add_table(rows=2, cols=3)
    info_table.rows[0].cells[0].text = "Satuan Kerja"
    info_table.rows[0].cells[1].text = ":"
    info_table.rows[0].cells[2].text = "{{satker_nama}}"
    info_table.rows[1].cells[0].text = "Tahun Anggaran"
    info_table.rows[1].cells[1].text = ":"
    info_table.rows[1].cells[2].text = "{{tahun_anggaran}}"
    
    doc.add_paragraph()
    
    # Main table with items
    table = doc.add_table(rows=4, cols=8)
    table.style = 'Table Grid'
    
    # Headers
    headers = ["No", "Uraian Barang/Jasa", "Satuan", "Volume", "Harga Satuan", "Jumlah", "PPh", "Keterangan"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Item row template
    item_row = table.rows[1]
    item_data = ["{{no}}", "{{item.uraian}}", "{{item.satuan}}", "{{item.volume}}", 
                 "{{item.harga_satuan}}", "{{item.total}}", "", "{{item.keterangan}}"]
    for i, data in enumerate(item_data):
        item_row.cells[i].text = data
    
    # Subtotal
    subtotal_row = table.rows[2]
    subtotal_row.cells[0].merge(subtotal_row.cells[3])
    subtotal_row.cells[0].text = "SUBTOTAL"
    subtotal_row.cells[0].paragraphs[0].runs[0].bold = True
    subtotal_row.cells[4].text = ""
    subtotal_row.cells[5].text = "{{subtotal_item_fmt}}"
    
    # PPN & Total
    total_row = table.rows[3]
    total_row.cells[0].merge(total_row.cells[3])
    total_row.cells[0].text = "PPN 11%"
    total_row.cells[4].text = ""
    total_row.cells[5].text = "{{ppn_item_fmt}}"
    
    doc.add_paragraph()
    
    # Grand Total
    p = doc.add_paragraph()
    p.add_run("TOTAL PEMBAYARAN: ")
    run = p.add_run("{{grand_total_item_fmt}}")
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run("Terbilang: ")
    run = p.add_run("{{grand_total_item_terbilang}}")
    run.italic = True
    
    doc.add_paragraph()
    
    # TTD
    ttd_table = doc.add_table(rows=4, cols=2)
    ttd_table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_spp:tanggal_long}}"
    ttd_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen,"
    ttd_table.rows[2].cells[1].text = "\n\n\n"
    p = ttd_table.rows[3].cells[1].paragraphs[0]
    run = p.add_run("{{ppk_nama}}")
    run.bold = True
    run.underline = True
    
    for row in ttd_table.rows:
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(WORD_DIR, "drpp.docx")
    doc.save(filepath)
    print(f"✅ Created: drpp.docx (with item loop)")
    return filepath


def create_all_item_templates():
    """Create all templates with item loop support"""
    print("=" * 60)
    print("CREATING TEMPLATES WITH ITEM LOOP SUPPORT")
    print("=" * 60)
    
    print("\n📄 Word Templates:")
    create_spesifikasi_teknis_with_items()
    create_drpp_word()
    
    print("\n📊 Excel Templates:")
    create_hps_excel()
    create_survey_harga_excel()
    
    print("\n" + "=" * 60)
    print("DONE! Templates with item loop created.")
    print("=" * 60)
    print("""
Placeholder untuk loop item:
  {{no}}                    - Nomor urut item (auto)
  {{item.uraian}}           - Uraian/nama barang
  {{item.spesifikasi}}      - Spesifikasi teknis
  {{item.satuan}}           - Satuan
  {{item.volume}}           - Volume/jumlah
  {{item.harga_dasar}}      - Harga satuan (angka)
  {{item.harga_satuan}}     - Harga satuan (format rupiah)
  {{item.total}}            - Total item (format rupiah)
  {{item.harga_survey1}}    - Harga survey toko 1
  {{item.harga_survey2}}    - Harga survey toko 2
  {{item.harga_survey3}}    - Harga survey toko 3
  {{item.harga_rata}}       - Harga rata-rata survey
  {{item.keterangan}}       - Keterangan item

Placeholder summary:
  {{subtotal_item}}         - Subtotal (angka)
  {{subtotal_item_fmt}}     - Subtotal (format rupiah)
  {{ppn_item}}              - PPN 11% (angka)
  {{ppn_item_fmt}}          - PPN (format rupiah)
  {{grand_total_item}}      - Grand total (angka)
  {{grand_total_item_fmt}}  - Grand total (format rupiah)
  {{grand_total_item_terbilang}} - Terbilang
""")


if __name__ == "__main__":
    create_all_item_templates()
