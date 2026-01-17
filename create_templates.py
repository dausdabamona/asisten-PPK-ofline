"""
PPK DOCUMENT FACTORY v3.0 - Sample Template Creator
====================================================
Create sample Word and Excel templates with placeholders
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
WORD_DIR = os.path.join(SCRIPT_DIR, "templates", "word")
EXCEL_DIR = os.path.join(SCRIPT_DIR, "templates", "excel")

os.makedirs(WORD_DIR, exist_ok=True)
os.makedirs(EXCEL_DIR, exist_ok=True)


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_kop_surat(doc):
    """Add header with kop surat"""
    header = doc.sections[0].header
    
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run("{{kementerian}}\n")
    run.bold = True
    run.font.size = Pt(11)
    
    run = p.add_run("{{eselon1}}\n")
    run.bold = True
    run.font.size = Pt(11)
    
    run = p.add_run("{{satker_nama}}\n")
    run.bold = True
    run.font.size = Pt(14)
    
    run = p.add_run("{{satker_alamat}}, {{satker_kota}}\n")
    run.font.size = Pt(9)
    
    run = p.add_run("Telepon: {{satker_telepon}} | Email: {{satker_email}}")
    run.font.size = Pt(9)
    
    # Border line
    p2 = header.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)


# ============================================================================
# WORD TEMPLATES
# ============================================================================

def create_spesifikasi_template():
    """Create Spesifikasi Teknis template"""
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SPESIFIKASI TEKNIS\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run("Nomor: {{nomor_spesifikasi}}")
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Content
    p = doc.add_paragraph()
    p.add_run("Nama Pekerjaan\t: ").bold = True
    p.add_run("{{nama_paket}}\n")
    p.add_run("Tahun Anggaran\t: ").bold = True
    p.add_run("{{tahun_anggaran}}\n")
    p.add_run("Sumber Dana\t\t: ").bold = True
    p.add_run("{{sumber_dana}}\n")
    
    doc.add_paragraph()
    
    # Section A
    p = doc.add_paragraph()
    p.add_run("A. LATAR BELAKANG").bold = True
    
    p = doc.add_paragraph()
    p.add_run("{{latar_belakang}}")
    
    doc.add_paragraph()
    
    # Section B
    p = doc.add_paragraph()
    p.add_run("B. MAKSUD DAN TUJUAN").bold = True
    
    p = doc.add_paragraph()
    p.add_run("{{maksud_tujuan}}")
    
    doc.add_paragraph()
    
    # Section C
    p = doc.add_paragraph()
    p.add_run("C. RUANG LINGKUP").bold = True
    
    p = doc.add_paragraph()
    p.add_run("{{ruang_lingkup}}")
    
    doc.add_paragraph()
    
    # Section D
    p = doc.add_paragraph()
    p.add_run("D. SPESIFIKASI TEKNIS").bold = True
    
    p = doc.add_paragraph()
    p.add_run("{{spesifikasi_teknis}}")
    
    doc.add_paragraph()
    
    # Section E
    p = doc.add_paragraph()
    p.add_run("E. VOLUME PEKERJAAN").bold = True
    
    p = doc.add_paragraph()
    p.add_run("{{volume_pekerjaan}}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("{{satker_kota}}, {{tanggal_spesifikasi:tanggal_long}}\n")
    p.add_run("Pejabat Pembuat Komitmen\n\n\n\n\n")
    run = p.add_run("{{ppk_nama}}\n")
    run.bold = True
    run.underline = True
    p.add_run("NIP. {{ppk_nip:nip}}")
    
    doc.save(os.path.join(WORD_DIR, "spesifikasi_teknis.docx"))
    print("✅ Created: spesifikasi_teknis.docx")


def create_kak_template():
    """Create KAK template"""
    doc = Document()
    
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KERANGKA ACUAN KERJA (KAK)\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run("{{nama_paket:upper}}")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Info
    p = doc.add_paragraph()
    p.add_run("Nomor\t\t\t: ").bold = True
    p.add_run("{{nomor_kak}}\n")
    p.add_run("Tahun Anggaran\t: ").bold = True
    p.add_run("{{tahun_anggaran}}\n")
    p.add_run("Sumber Dana\t\t: ").bold = True
    p.add_run("{{sumber_dana}}\n")
    p.add_run("Kode Akun\t\t: ").bold = True
    p.add_run("{{kode_akun}}\n")
    p.add_run("Nilai Pagu\t\t: ").bold = True
    p.add_run("{{nilai_pagu:rupiah}}\n")
    
    doc.add_paragraph()
    
    # Sections (similar to spesifikasi)
    sections = [
        ("I. LATAR BELAKANG", "{{latar_belakang}}"),
        ("II. MAKSUD DAN TUJUAN", "{{maksud_tujuan}}"),
        ("III. RUANG LINGKUP PEKERJAAN", "{{ruang_lingkup}}"),
        ("IV. SPESIFIKASI TEKNIS", "{{spesifikasi_teknis}}"),
        ("V. WAKTU PELAKSANAAN", "{{jangka_waktu}} hari kalender"),
        ("VI. TENAGA AHLI", "Sesuai kebutuhan pekerjaan"),
        ("VII. PELAPORAN", "Penyedia wajib menyampaikan laporan pelaksanaan pekerjaan")
    ]
    
    for title, content in sections:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p = doc.add_paragraph()
        p.add_run(content)
        doc.add_paragraph()
    
    # TTD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("{{satker_kota}}, {{tanggal_kak:tanggal_long}}\n")
    p.add_run("Pejabat Pembuat Komitmen\n\n\n\n\n")
    run = p.add_run("{{ppk_nama}}\n")
    run.bold = True
    run.underline = True
    p.add_run("NIP. {{ppk_nip:nip}}")
    
    doc.save(os.path.join(WORD_DIR, "kak.docx"))
    print("✅ Created: kak.docx")


def create_spk_template():
    """Create SPK template"""
    doc = Document()
    
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SURAT PERINTAH KERJA\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run("Nomor: {{nomor_spk}}")
    
    doc.add_paragraph()
    
    # Opening
    p = doc.add_paragraph()
    p.add_run("Yang bertanda tangan di bawah ini:\n\n")
    
    # PPK
    p.add_run("Nama\t\t\t: ").bold = True
    p.add_run("{{ppk_nama}}\n")
    p.add_run("NIP\t\t\t: ").bold = True
    p.add_run("{{ppk_nip}}\n")
    p.add_run("Jabatan\t\t\t: ").bold = True
    p.add_run("{{ppk_jabatan}}\n")
    p.add_run("Alamat\t\t\t: ").bold = True
    p.add_run("{{satker_alamat}}\n\n")
    
    run = p.add_run("Selanjutnya disebut sebagai ")
    run = p.add_run("PIHAK PERTAMA\n\n")
    run.bold = True
    
    # Penyedia
    p.add_run("Nama\t\t\t: ").bold = True
    p.add_run("{{direktur_nama}}\n")
    p.add_run("Jabatan\t\t\t: ").bold = True
    p.add_run("{{direktur_jabatan}}\n")
    p.add_run("Bertindak untuk dan atas nama:\n")
    p.add_run("Nama Perusahaan\t: ").bold = True
    p.add_run("{{penyedia_nama}}\n")
    p.add_run("Alamat\t\t\t: ").bold = True
    p.add_run("{{penyedia_alamat}}\n")
    p.add_run("NPWP\t\t\t: ").bold = True
    p.add_run("{{penyedia_npwp}}\n\n")
    
    run = p.add_run("Selanjutnya disebut sebagai ")
    run = p.add_run("PIHAK KEDUA\n\n")
    run.bold = True
    
    p.add_run("Dengan ini PIHAK PERTAMA memberikan pekerjaan kepada PIHAK KEDUA ")
    p.add_run("dan PIHAK KEDUA menerima pekerjaan tersebut, yaitu:\n")
    
    doc.add_paragraph()
    
    # Details table
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    
    details = [
        ("Nama Pekerjaan", ":", "{{nama_paket}}"),
        ("Lokasi Pekerjaan", ":", "{{lokasi_pekerjaan}}"),
        ("Sumber Dana", ":", "{{sumber_dana}}"),
        ("Nilai Pekerjaan", ":", "{{nilai_kontrak:rupiah}}"),
        ("PPN 11%", ":", "{{nilai_ppn:rupiah}}"),
        ("Nilai SPK", ":", "{{nilai_bruto:rupiah}}"),
        ("Jangka Waktu", ":", "{{jangka_waktu}} hari kalender"),
        ("Tanggal Mulai", ":", "{{tanggal_mulai:tanggal_long}}"),
        ("Tanggal Selesai", ":", "{{tanggal_selesai:tanggal_long}}"),
    ]
    
    for i, (label, sep, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = sep
        table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Terbilang: ").bold = True
    p.add_run("({{nilai_bruto:terbilang}})")
    
    doc.add_paragraph()
    
    # Ketentuan
    p = doc.add_paragraph()
    p.add_run("KETENTUAN:\n").bold = True
    p.add_run("1. PIHAK KEDUA harus menyelesaikan pekerjaan sesuai dengan spesifikasi teknis.\n")
    p.add_run("2. Penyerahan pekerjaan dilakukan setelah pemeriksaan dan dinyatakan sesuai.\n")
    p.add_run("3. Pembayaran dilakukan setelah pekerjaan diterima dengan baik.\n")
    p.add_run("4. PIHAK KEDUA wajib memberikan jaminan atas hasil pekerjaan.\n")
    
    doc.add_paragraph()
    
    # TTD
    table = doc.add_table(rows=1, cols=2)
    
    cell1 = table.rows[0].cells[0]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("PIHAK KEDUA\n").bold = True
    p1.add_run("{{direktur_jabatan}}\n")
    p1.add_run("{{penyedia_nama}}\n\n\n\n\n")
    run = p1.add_run("{{direktur_nama}}\n")
    run.bold = True
    run.underline = True
    
    cell2 = table.rows[0].cells[1]
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("PIHAK PERTAMA\n").bold = True
    p2.add_run("Pejabat Pembuat Komitmen\n\n\n\n\n\n")
    run = p2.add_run("{{ppk_nama}}\n")
    run.bold = True
    run.underline = True
    p2.add_run("NIP. {{ppk_nip:nip}}")
    
    doc.save(os.path.join(WORD_DIR, "spk.docx"))
    print("✅ Created: spk.docx")


def create_spmk_template():
    """Create SPMK template"""
    doc = Document()
    
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SURAT PERINTAH MULAI KERJA\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run("Nomor: {{nomor_spmk}}")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Menunjuk Surat Perintah Kerja:\n")
    p.add_run("Nomor\t\t: {{nomor_spk}}\n")
    p.add_run("Tanggal\t\t: {{tanggal_spk:tanggal_long}}\n")
    p.add_run("Paket Pekerjaan\t: {{nama_paket}}\n\n")
    
    p.add_run("Dengan ini memerintahkan kepada:\n\n")
    
    p.add_run("Nama Perusahaan\t: ").bold = True
    p.add_run("{{penyedia_nama}}\n")
    p.add_run("Alamat\t\t\t: ").bold = True
    p.add_run("{{penyedia_alamat}}\n\n")
    
    p.add_run("Untuk segera memulai pelaksanaan pekerjaan dengan ketentuan sebagai berikut:\n")
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    details = [
        ("Tanggal Mulai Kerja", ":", "{{tanggal_mulai:tanggal_long}}"),
        ("Tanggal Selesai", ":", "{{tanggal_selesai:tanggal_long}}"),
        ("Jangka Waktu", ":", "{{jangka_waktu}} hari kalender"),
        ("Lokasi Pekerjaan", ":", "{{lokasi_pekerjaan}}"),
    ]
    
    for i, (label, sep, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = sep
        table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("\nDemikian Surat Perintah Mulai Kerja ini diterbitkan untuk dilaksanakan ")
    p.add_run("dengan penuh tanggung jawab.")
    
    doc.add_paragraph()
    
    # TTD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("{{satker_kota}}, {{tanggal_spmk:tanggal_long}}\n")
    p.add_run("Pejabat Pembuat Komitmen\n\n\n\n\n")
    run = p.add_run("{{ppk_nama}}\n")
    run.bold = True
    run.underline = True
    p.add_run("NIP. {{ppk_nip:nip}}")
    
    doc.save(os.path.join(WORD_DIR, "spmk.docx"))
    print("✅ Created: spmk.docx")


def create_bahp_template():
    """Create BAHP template"""
    doc = Document()
    
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BERITA ACARA HASIL PEMERIKSAAN\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run("Nomor: {{nomor_bahp}}")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Pada hari ini, {{tanggal_bahp:tanggal_full}}, kami yang bertanda tangan ")
    p.add_run("di bawah ini:\n\n")
    
    p.add_run("1. ").bold = True
    p.add_run("Nama\t\t: {{pemeriksa1_nama}}\n")
    p.add_run("   NIP\t\t: {{pemeriksa1_nip}}\n")
    p.add_run("   Jabatan\t: {{pemeriksa1_jabatan}}\n\n")
    
    p.add_run("2. ").bold = True
    p.add_run("Nama\t\t: {{pemeriksa2_nama}}\n")
    p.add_run("   NIP\t\t: {{pemeriksa2_nip}}\n")
    p.add_run("   Jabatan\t: {{pemeriksa2_jabatan}}\n\n")
    
    p.add_run("Selaku Pejabat/Panitia Pemeriksa Hasil Pekerjaan,\n\n")
    p.add_run("Telah melakukan pemeriksaan terhadap hasil pekerjaan:\n")
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    details = [
        ("Nama Pekerjaan", ":", "{{nama_paket}}"),
        ("Nomor SPK", ":", "{{nomor_spk}}"),
        ("Tanggal SPK", ":", "{{tanggal_spk:tanggal_long}}"),
        ("Nilai Kontrak", ":", "{{nilai_bruto:rupiah}}"),
        ("Penyedia", ":", "{{penyedia_nama}}"),
    ]
    
    for i, (label, sep, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = sep
        table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("HASIL PEMERIKSAAN:\n").bold = True
    p.add_run("{{hasil_pemeriksaan}}")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("KESIMPULAN:\n").bold = True
    p.add_run("Berdasarkan hasil pemeriksaan tersebut di atas, pekerjaan dinyatakan ")
    p.add_run("{{kesimpulan:upper}} ").bold = True
    p.add_run("dengan spesifikasi teknis dalam kontrak.")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Demikian Berita Acara Hasil Pemeriksaan ini dibuat dengan sebenarnya ")
    p.add_run("untuk dapat dipergunakan sebagaimana mestinya.")
    
    doc.add_paragraph()
    
    # TTD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Pejabat/Panitia Pemeriksa Hasil Pekerjaan\n\n").bold = True
    
    table = doc.add_table(rows=1, cols=2)
    
    for i, num in enumerate(['1', '2']):
        cell = table.rows[0].cells[i]
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(f"{{{{pemeriksa{num}_jabatan}}}}\n\n\n\n\n")
        run = cp.add_run(f"{{{{pemeriksa{num}_nama}}}}\n")
        run.bold = True
        run.underline = True
        cp.add_run(f"NIP. {{{{pemeriksa{num}_nip}}}}")
    
    doc.save(os.path.join(WORD_DIR, "bahp.docx"))
    print("✅ Created: bahp.docx")


def create_bast_template():
    """Create BAST template"""
    doc = Document()
    
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BERITA ACARA SERAH TERIMA HASIL PEKERJAAN\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run("Nomor: {{nomor_bast}}")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Pada hari ini, {{tanggal_bast:tanggal_full}}, kami yang bertanda tangan ")
    p.add_run("di bawah ini:\n\n")
    
    # Pihak 1
    p.add_run("I. ").bold = True
    p.add_run("Nama\t\t: {{direktur_nama}}\n")
    p.add_run("   Jabatan\t: {{direktur_jabatan}}\n")
    p.add_run("   Perusahaan\t: {{penyedia_nama}}\n")
    p.add_run("   Selanjutnya disebut ")
    run = p.add_run("PIHAK PERTAMA ")
    run.bold = True
    p.add_run("(Yang Menyerahkan)\n\n")
    
    # Pihak 2
    p.add_run("II. ").bold = True
    p.add_run("Nama\t\t: {{ppk_nama}}\n")
    p.add_run("   NIP\t\t: {{ppk_nip}}\n")
    p.add_run("   Jabatan\t: {{ppk_jabatan}}\n")
    p.add_run("   Selanjutnya disebut ")
    run = p.add_run("PIHAK KEDUA ")
    run.bold = True
    p.add_run("(Yang Menerima)\n\n")
    
    p.add_run("Berdasarkan:\n")
    p.add_run("1. Surat Perintah Kerja Nomor: {{nomor_spk}} ")
    p.add_run("Tanggal: {{tanggal_spk:tanggal_long}}\n")
    p.add_run("2. Berita Acara Hasil Pemeriksaan Nomor: {{nomor_bahp}}\n\n")
    
    p.add_run("PIHAK PERTAMA menyerahkan kepada PIHAK KEDUA dan PIHAK KEDUA ")
    p.add_run("menerima dari PIHAK PERTAMA, hasil pekerjaan:\n")
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    details = [
        ("Nama Pekerjaan", ":", "{{nama_paket}}"),
        ("Nilai Pekerjaan", ":", "{{nilai_bruto:rupiah}}"),
        ("Volume yang Diserahkan", ":", "{{volume_serah_terima}}"),
        ("Kondisi", ":", "{{kondisi_barang}}"),
    ]
    
    for i, (label, sep, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = sep
        table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Demikian Berita Acara Serah Terima ini dibuat dengan sebenarnya dalam ")
    p.add_run("rangkap 2 (dua) untuk dapat dipergunakan sebagaimana mestinya.")
    
    doc.add_paragraph()
    
    # TTD
    table = doc.add_table(rows=1, cols=2)
    
    cell1 = table.rows[0].cells[0]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("PIHAK PERTAMA\n(Yang Menyerahkan)\n").bold = True
    p1.add_run("{{direktur_jabatan}}\n")
    p1.add_run("{{penyedia_nama}}\n\n\n\n\n")
    run = p1.add_run("{{direktur_nama}}\n")
    run.bold = True
    run.underline = True
    
    cell2 = table.rows[0].cells[1]
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("PIHAK KEDUA\n(Yang Menerima)\n").bold = True
    p2.add_run("Pejabat Pembuat Komitmen\n\n\n\n\n\n")
    run = p2.add_run("{{ppk_nama}}\n")
    run.bold = True
    run.underline = True
    p2.add_run("NIP. {{ppk_nip:nip}}")
    
    doc.save(os.path.join(WORD_DIR, "bast.docx"))
    print("✅ Created: bast.docx")


def create_spp_ls_template():
    """Create SPP-LS template"""
    doc = Document()
    
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SURAT PERMINTAAN PEMBAYARAN LANGSUNG (SPP-LS)\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run("Nomor: {{nomor_spp_ls}}")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Tanggal: {{tanggal_spp:tanggal_long}}\n")
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=14, cols=3)
    table.style = 'Table Grid'
    
    details = [
        ("Tahun Anggaran", ":", "{{tahun_anggaran}}"),
        ("Jenis Pembayaran", ":", "SPP-LS Kontraktual"),
        ("Nama Pekerjaan", ":", "{{nama_paket}}"),
        ("Nomor Kontrak/SPK", ":", "{{nomor_spk}}"),
        ("Tanggal Kontrak", ":", "{{tanggal_spk:tanggal_long}}"),
        ("", "", ""),
        ("Nilai Bruto", ":", "{{nilai_bruto:rupiah}}"),
        ("PPN 11%", ":", "{{nilai_ppn:rupiah}}"),
        ("{{jenis_pph}}", ":", "{{nilai_pph:rupiah}}"),
        ("Nilai Bersih", ":", "{{nilai_bersih:rupiah}}"),
        ("Terbilang", ":", "({{nilai_bersih:terbilang}})"),
        ("", "", ""),
        ("Nama Penerima", ":", "{{penyedia_nama}}"),
        ("No. Rekening", ":", "{{penyedia_rekening}} - {{penyedia_bank}}"),
    ]
    
    for i, (label, sep, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = sep
        table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    table = doc.add_table(rows=1, cols=2)
    
    cell1 = table.rows[0].cells[0]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("Mengetahui/Menyetujui\nPejabat Pembuat Komitmen\n\n\n\n\n")
    run = p1.add_run("{{ppk_nama}}\n")
    run.bold = True
    run.underline = True
    p1.add_run("NIP. {{ppk_nip:nip}}")
    
    cell2 = table.rows[0].cells[1]
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("{{satker_kota}}, {{tanggal_spp:tanggal_long}}\n")
    p2.add_run("Pejabat Penandatangan SPM\n\n\n\n\n")
    run = p2.add_run("{{ppspm_nama}}\n")
    run.bold = True
    run.underline = True
    p2.add_run("NIP. {{ppspm_nip:nip}}")
    
    doc.save(os.path.join(WORD_DIR, "spp_ls.docx"))
    print("✅ Created: spp_ls.docx")


def create_kuitansi_template():
    """Create Kuitansi template"""
    doc = Document()
    
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KUITANSI / BUKTI PEMBAYARAN\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run("Nomor: {{nomor_kuitansi}}")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Sudah terima dari\t: ").bold = True
    p.add_run("Kuasa Pengguna Anggaran\n")
    p.add_run("\t\t\t  {{satker_nama}}\n\n")
    
    p.add_run("Uang Sebesar\t\t: ").bold = True
    p.add_run("{{nilai_bersih:rupiah}}\n")
    p.add_run("\t\t\t  ({{nilai_bersih:terbilang}})\n\n")
    
    p.add_run("Untuk Pembayaran\t: ").bold = True
    p.add_run("{{nama_paket}}\n")
    
    doc.add_paragraph()
    
    # Rincian
    p = doc.add_paragraph()
    p.add_run("RINCIAN:\n").bold = True
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    details = [
        ("Nilai Kontrak", ":", "{{nilai_kontrak:rupiah}}"),
        ("PPN 11%", ":", "{{nilai_ppn:rupiah}}"),
        ("{{jenis_pph}}", ":", "{{nilai_pph:rupiah}}"),
        ("", "", ""),
        ("NILAI BERSIH", ":", "{{nilai_bersih:rupiah}}"),
    ]
    
    for i, (label, sep, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = sep
        table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD 3 kolom
    table = doc.add_table(rows=1, cols=3)
    
    # Bendahara
    cell1 = table.rows[0].cells[0]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("Lunas dibayar\nTgl: .....................\nBendahara Pengeluaran\n\n\n\n\n")
    run = p1.add_run("{{bendahara_nama}}\n")
    run.bold = True
    run.underline = True
    p1.add_run("NIP. {{bendahara_nip:nip}}")
    
    # PPK
    cell2 = table.rows[0].cells[1]
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Setuju dibayar\n{{satker_kota}}, ...............\nPejabat Pembuat Komitmen\n\n\n\n\n")
    run = p2.add_run("{{ppk_nama}}\n")
    run.bold = True
    run.underline = True
    p2.add_run("NIP. {{ppk_nip:nip}}")
    
    # Penerima
    cell3 = table.rows[0].cells[2]
    p3 = cell3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Yang menerima\n{{direktur_jabatan}}\n{{penyedia_nama}}\n\n\n\n\n")
    run = p3.add_run("{{direktur_nama}}\n")
    run.bold = True
    run.underline = True
    
    doc.save(os.path.join(WORD_DIR, "kuitansi.docx"))
    print("✅ Created: kuitansi.docx")


def create_drpp_template():
    """Create DRPP template"""
    doc = Document()
    
    # Landscape
    for section in doc.sections:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DAFTAR RINCIAN PERMINTAAN PEMBAYARAN\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run("Nomor: {{nomor_drpp}}")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Nomor SPP-LS\t: {{nomor_spp_ls}}\n")
    p.add_run("Tanggal\t\t: {{tanggal_spp:tanggal_long}}\n")
    
    doc.add_paragraph()
    
    # Table header
    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    
    headers = ["No", "Kode Akun", "Uraian", "Bukti", "Nilai (Rp)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for para in table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    # Sample row
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "{{kode_akun}}"
    table.rows[1].cells[2].text = "{{nama_paket}}"
    table.rows[1].cells[3].text = "SPK"
    table.rows[1].cells[4].text = "{{nilai_bersih:angka}}"
    
    # Total
    table.rows[2].cells[2].text = "JUMLAH"
    table.rows[2].cells[4].text = "{{nilai_bersih:angka}}"
    for para in table.rows[2].cells[2].paragraphs:
        for run in para.runs:
            run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Terbilang: ").bold = True
    p.add_run("({{nilai_bersih:terbilang}})")
    
    doc.add_paragraph()
    
    # TTD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("{{satker_kota}}, {{tanggal_spp:tanggal_long}}\n")
    p.add_run("Pejabat Pembuat Komitmen\n\n\n\n\n")
    run = p.add_run("{{ppk_nama}}\n")
    run.bold = True
    run.underline = True
    p.add_run("NIP. {{ppk_nip:nip}}")
    
    doc.save(os.path.join(WORD_DIR, "drpp.docx"))
    print("✅ Created: drpp.docx")


# ============================================================================
# EXCEL TEMPLATES
# ============================================================================

def create_hps_template():
    """Create HPS Excel template"""
    wb = Workbook()
    ws = wb.active
    ws.title = "HPS"
    
    # Styles
    header_font = Font(bold=True, size=14)
    bold_font = Font(bold=True)
    currency_format = '#,##0'
    
    header_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Title
    ws.merge_cells('A1:G1')
    ws['A1'] = "HARGA PERKIRAAN SENDIRI (HPS)"
    ws['A1'].font = header_font
    ws['A1'].alignment = Alignment(horizontal='center')
    
    ws.merge_cells('A2:G2')
    ws['A2'] = "{{nama_paket}}"
    ws['A2'].alignment = Alignment(horizontal='center')
    
    # Info
    ws['A4'] = "Nomor HPS"
    ws['B4'] = ":"
    ws['C4'] = "{{nomor_hps}}"
    
    ws['A5'] = "Tahun Anggaran"
    ws['B5'] = ":"
    ws['C5'] = "{{tahun_anggaran}}"
    
    ws['A6'] = "Nilai Pagu"
    ws['B6'] = ":"
    ws['C6'] = "{{nilai_pagu}}"
    
    # Table header
    headers = ["No", "Uraian Barang/Jasa", "Satuan", "Volume", "Harga Satuan", "Jumlah", "Keterangan"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=8, column=col)
        cell.value = header
        cell.font = bold_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center')
    
    # Item placeholder row
    ws.cell(row=9, column=1).value = "{{item.no}}"
    ws.cell(row=9, column=2).value = "{{item.nama_item}}"
    ws.cell(row=9, column=3).value = "{{item.satuan}}"
    ws.cell(row=9, column=4).value = "{{item.volume}}"
    ws.cell(row=9, column=5).value = "{{item.harga_satuan}}"
    ws.cell(row=9, column=6).value = "{{item.jumlah}}"
    ws.cell(row=9, column=7).value = "{{item.keterangan}}"
    
    for col in range(1, 8):
        ws.cell(row=9, column=col).border = thin_border
    
    # Total row
    ws.cell(row=10, column=5).value = "TOTAL HPS"
    ws.cell(row=10, column=5).font = bold_font
    ws.cell(row=10, column=6).value = "{{nilai_hps}}"
    ws.cell(row=10, column=6).font = bold_font
    
    # PPN
    ws.cell(row=11, column=5).value = "PPN 11%"
    ws.cell(row=11, column=6).value = "{{nilai_ppn}}"
    
    # Total + PPN
    ws.cell(row=12, column=5).value = "TOTAL + PPN"
    ws.cell(row=12, column=5).font = bold_font
    ws.cell(row=12, column=6).value = "{{nilai_bruto}}"
    ws.cell(row=12, column=6).font = bold_font
    
    # Terbilang
    ws['A14'] = "Terbilang:"
    ws['A14'].font = bold_font
    ws.merge_cells('B14:G14')
    ws['B14'] = "({{nilai_bruto:terbilang}})"
    
    # TTD
    ws['E17'] = "{{satker_kota}}, {{tanggal_hps:tanggal_long}}"
    ws['E18'] = "Pejabat Pembuat Komitmen"
    ws['E22'] = "{{ppk_nama}}"
    ws['E22'].font = bold_font
    ws['E23'] = "NIP. {{ppk_nip}}"
    
    # Column widths
    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 35
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 10
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 15
    ws.column_dimensions['G'].width = 15
    
    wb.save(os.path.join(EXCEL_DIR, "hps.xlsx"))
    print("✅ Created: hps.xlsx")


def create_survey_harga_template():
    """Create Survey Harga Excel template"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Survey Harga"
    
    header_font = Font(bold=True, size=14)
    bold_font = Font(bold=True)
    header_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Title
    ws.merge_cells('A1:I1')
    ws['A1'] = "DAFTAR SURVEY HARGA PASAR"
    ws['A1'].font = header_font
    ws['A1'].alignment = Alignment(horizontal='center')
    
    ws.merge_cells('A2:I2')
    ws['A2'] = "{{nama_paket}}"
    ws['A2'].alignment = Alignment(horizontal='center')
    
    # Survey sources
    ws['A4'] = "Sumber Harga 1"
    ws['B4'] = ":"
    ws['C4'] = "{{survey1_toko}}"
    ws['D4'] = "Tanggal:"
    ws['E4'] = "{{survey1_tanggal}}"
    
    ws['A5'] = "Sumber Harga 2"
    ws['B5'] = ":"
    ws['C5'] = "{{survey2_toko}}"
    ws['D5'] = "Tanggal:"
    ws['E5'] = "{{survey2_tanggal}}"
    
    ws['A6'] = "Sumber Harga 3"
    ws['B6'] = ":"
    ws['C6'] = "{{survey3_toko}}"
    ws['D6'] = "Tanggal:"
    ws['E6'] = "{{survey3_tanggal}}"
    
    # Table header
    headers = ["No", "Uraian Barang/Jasa", "Satuan", "Volume", 
               "Harga 1", "Harga 2", "Harga 3", "Rata-rata", "Jumlah"]
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=8, column=col)
        cell.value = header
        cell.font = bold_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center')
    
    # Item row
    ws.cell(row=9, column=1).value = "{{item.no}}"
    ws.cell(row=9, column=2).value = "{{item.nama_item}}"
    ws.cell(row=9, column=3).value = "{{item.satuan}}"
    ws.cell(row=9, column=4).value = "{{item.volume}}"
    ws.cell(row=9, column=5).value = "{{item.harga_survey1}}"
    ws.cell(row=9, column=6).value = "{{item.harga_survey2}}"
    ws.cell(row=9, column=7).value = "{{item.harga_survey3}}"
    ws.cell(row=9, column=8).value = "{{item.harga_rata_rata}}"
    ws.cell(row=9, column=9).value = "{{item.jumlah}}"
    
    for col in range(1, 10):
        ws.cell(row=9, column=col).border = thin_border
    
    # Total
    ws.cell(row=10, column=8).value = "TOTAL"
    ws.cell(row=10, column=8).font = bold_font
    ws.cell(row=10, column=9).value = "{{nilai_hps}}"
    
    # Column widths
    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 8
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 12
    ws.column_dimensions['G'].width = 12
    ws.column_dimensions['H'].width = 12
    ws.column_dimensions['I'].width = 15
    
    wb.save(os.path.join(EXCEL_DIR, "survey_harga.xlsx"))
    print("✅ Created: survey_harga.xlsx")


def create_ssp_template():
    """Create SSP Excel template with PPN and PPh sheets"""
    wb = Workbook()
    
    # PPN Sheet
    ws_ppn = wb.active
    ws_ppn.title = "PPN"
    
    ws_ppn.merge_cells('A1:D1')
    ws_ppn['A1'] = "SURAT SETORAN PAJAK - PPN"
    ws_ppn['A1'].font = Font(bold=True, size=14)
    ws_ppn['A1'].alignment = Alignment(horizontal='center')
    
    ws_ppn['A3'] = "NPWP"
    ws_ppn['B3'] = ":"
    ws_ppn['C3'] = "{{penyedia_npwp}}"
    
    ws_ppn['A4'] = "Nama WP"
    ws_ppn['B4'] = ":"
    ws_ppn['C4'] = "{{penyedia_nama}}"
    
    ws_ppn['A5'] = "Alamat"
    ws_ppn['B5'] = ":"
    ws_ppn['C5'] = "{{penyedia_alamat}}"
    
    ws_ppn['A7'] = "Kode Akun Pajak"
    ws_ppn['B7'] = ":"
    ws_ppn['C7'] = "411211"
    
    ws_ppn['A8'] = "Kode Jenis Setoran"
    ws_ppn['B8'] = ":"
    ws_ppn['C8'] = "100"
    
    ws_ppn['A9'] = "Masa Pajak"
    ws_ppn['B9'] = ":"
    ws_ppn['C9'] = "{{tanggal_spp:tanggal_short}}"
    
    ws_ppn['A11'] = "Jumlah Setoran"
    ws_ppn['B11'] = ":"
    ws_ppn['C11'] = "{{nilai_ppn}}"
    ws_ppn['A11'].font = Font(bold=True)
    ws_ppn['C11'].font = Font(bold=True)
    
    ws_ppn['A12'] = "Terbilang"
    ws_ppn['B12'] = ":"
    ws_ppn.merge_cells('C12:D12')
    ws_ppn['C12'] = "({{nilai_ppn:terbilang}})"
    
    ws_ppn.column_dimensions['A'].width = 20
    ws_ppn.column_dimensions['C'].width = 30
    
    # PPh Sheet
    ws_pph = wb.create_sheet("PPh")
    
    ws_pph.merge_cells('A1:D1')
    ws_pph['A1'] = "SURAT SETORAN PAJAK - {{jenis_pph}}"
    ws_pph['A1'].font = Font(bold=True, size=14)
    ws_pph['A1'].alignment = Alignment(horizontal='center')
    
    ws_pph['A3'] = "NPWP"
    ws_pph['B3'] = ":"
    ws_pph['C3'] = "{{penyedia_npwp}}"
    
    ws_pph['A4'] = "Nama WP"
    ws_pph['B4'] = ":"
    ws_pph['C4'] = "{{penyedia_nama}}"
    
    ws_pph['A5'] = "Alamat"
    ws_pph['B5'] = ":"
    ws_pph['C5'] = "{{penyedia_alamat}}"
    
    ws_pph['A7'] = "Kode Akun Pajak"
    ws_pph['B7'] = ":"
    ws_pph['C7'] = "411122"
    
    ws_pph['A8'] = "Kode Jenis Setoran"
    ws_pph['B8'] = ":"
    ws_pph['C8'] = "100"
    
    ws_pph['A9'] = "Masa Pajak"
    ws_pph['B9'] = ":"
    ws_pph['C9'] = "{{tanggal_spp:tanggal_short}}"
    
    ws_pph['A11'] = "Jumlah Setoran"
    ws_pph['B11'] = ":"
    ws_pph['C11'] = "{{nilai_pph}}"
    ws_pph['A11'].font = Font(bold=True)
    ws_pph['C11'].font = Font(bold=True)
    
    ws_pph['A12'] = "Terbilang"
    ws_pph['B12'] = ":"
    ws_pph.merge_cells('C12:D12')
    ws_pph['C12'] = "({{nilai_pph:terbilang}})"
    
    ws_pph.column_dimensions['A'].width = 20
    ws_pph.column_dimensions['C'].width = 30
    
    wb.save(os.path.join(EXCEL_DIR, "ssp.xlsx"))
    print("✅ Created: ssp.xlsx")


# ============================================================================
# MAIN
# ============================================================================

def create_all_templates():
    """Create all sample templates"""
    print("\n" + "=" * 60)
    print("CREATING SAMPLE TEMPLATES")
    print("=" * 60 + "\n")
    
    # Word templates
    print("📄 Word Templates:")
    create_spesifikasi_template()
    create_kak_template()
    create_spk_template()
    create_spmk_template()
    create_bahp_template()
    create_bast_template()
    create_spp_ls_template()
    create_kuitansi_template()
    create_drpp_template()
    
    print("\n📊 Excel Templates:")
    create_hps_template()
    create_survey_harga_template()
    create_ssp_template()
    
    print("\n" + "=" * 60)
    print("ALL TEMPLATES CREATED SUCCESSFULLY!")
    print("=" * 60)
    print(f"\nWord templates: {WORD_DIR}")
    print(f"Excel templates: {EXCEL_DIR}")


if __name__ == "__main__":
    create_all_templates()
