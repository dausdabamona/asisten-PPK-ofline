"""
Script untuk membuat template Nota Dinas Persiapan Pengadaan
Untuk dikirim ke Pejabat Pengadaan dengan lampiran RKAKL dan RUP
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = "templates/word"


def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def create_nota_dinas_pp_template():
    """Create Nota Dinas untuk Pejabat Pengadaan template"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    # ========== KOP SURAT ==========
    # Kementerian
    kop1 = doc.add_paragraph()
    kop1_run = kop1.add_run("{{kementerian}}")
    kop1_run.bold = True
    kop1_run.font.size = Pt(12)
    kop1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Eselon 1
    kop2 = doc.add_paragraph()
    kop2_run = kop2.add_run("{{eselon1}}")
    kop2_run.bold = True
    kop2_run.font.size = Pt(12)
    kop2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Satker
    kop3 = doc.add_paragraph()
    kop3_run = kop3.add_run("{{satker_nama}}")
    kop3_run.bold = True
    kop3_run.font.size = Pt(14)
    kop3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Alamat
    alamat = doc.add_paragraph()
    alamat_run = alamat.add_run("{{satker_alamat}}, {{satker_kota}}")
    alamat_run.font.size = Pt(10)
    alamat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Telepon/Email
    kontak = doc.add_paragraph()
    kontak_run = kontak.add_run("Telepon: {{satker_telepon}} | Email: {{satker_email}}")
    kontak_run.font.size = Pt(10)
    kontak.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Garis pembatas
    border = doc.add_paragraph()
    border.add_run("═" * 75)
    border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ========== HEADER NOTA DINAS ==========
    title = doc.add_paragraph()
    title_run = title.add_run("NOTA DINAS")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Nomor: {{nomor_nota_dinas_pp}}")
    subtitle_run.font.size = Pt(11)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ========== INFO TABLE ==========
    info_table = doc.add_table(rows=5, cols=3)
    
    # Kepada
    info_table.rows[0].cells[0].text = "Kepada"
    info_table.rows[0].cells[1].text = ":"
    info_table.rows[0].cells[2].text = "Yth. Pejabat Pengadaan\n{{satker_nama}}"
    
    # Dari
    info_table.rows[1].cells[0].text = "Dari"
    info_table.rows[1].cells[1].text = ":"
    info_table.rows[1].cells[2].text = "Pejabat Pembuat Komitmen"
    
    # Tanggal
    info_table.rows[2].cells[0].text = "Tanggal"
    info_table.rows[2].cells[1].text = ":"
    info_table.rows[2].cells[2].text = "{{tanggal_nota_dinas_pp_fmt}}"
    
    # Lampiran
    info_table.rows[3].cells[0].text = "Lampiran"
    info_table.rows[3].cells[1].text = ":"
    info_table.rows[3].cells[2].text = "1 (satu) berkas"
    
    # Perihal
    info_table.rows[4].cells[0].text = "Perihal"
    info_table.rows[4].cells[1].text = ":"
    perihal_cell = info_table.rows[4].cells[2]
    perihal_cell.text = "Permintaan Pelaksanaan Pengadaan\n{{nama_paket}}"
    for para in perihal_cell.paragraphs:
        for run in para.runs:
            run.bold = True
    
    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(12)
    
    doc.add_paragraph()
    
    # ========== ISI NOTA DINAS ==========
    # Pembuka
    p1 = doc.add_paragraph()
    p1.add_run(
        "Dalam rangka pelaksanaan kegiatan pengadaan barang/jasa pada "
        "{{satker_nama}} Tahun Anggaran {{tahun_anggaran}}, bersama ini kami sampaikan "
        "permintaan pelaksanaan pengadaan sebagai berikut:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    p1.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # ========== TABEL INFORMASI PAKET ==========
    paket_table = doc.add_table(rows=10, cols=3)
    paket_table.style = 'Table Grid'
    
    paket_info = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Kode Paket", "{{kode_paket}}"),
        ("3.", "Kode RUP", "{{kode_rup}}"),
        ("4.", "Jenis Pengadaan", "{{jenis_pengadaan}}"),
        ("5.", "Metode Pemilihan", "{{metode_pemilihan}}"),
        ("6.", "Sumber Dana", "{{sumber_dana}}"),
        ("7.", "Kode Akun/MAK", "{{kode_akun}}"),
        ("8.", "Pagu Anggaran", "{{nilai_pagu_fmt}}"),
        ("9.", "HPS", "{{nilai_hps_fmt}}"),
        ("10.", "Jangka Waktu Pelaksanaan", "{{jangka_waktu}} hari kalender"),
    ]
    
    for i, (no, label, value) in enumerate(paket_info):
        paket_table.rows[i].cells[0].text = no
        paket_table.rows[i].cells[1].text = label
        paket_table.rows[i].cells[2].text = value
    
    # Set column widths
    for row in paket_table.rows:
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(9)
    
    doc.add_paragraph()
    
    # ========== LAMPIRAN ==========
    lamp_title = doc.add_paragraph()
    lamp_title.add_run("Dokumen yang dilampirkan:").bold = True
    
    lampiran_list = [
        "Salinan RKAKL/POK yang telah disahkan",
        "Salinan RUP yang telah diumumkan",
        "Kerangka Acuan Kerja (KAK)/Spesifikasi Teknis",
        "Harga Perkiraan Sendiri (HPS)",
        "Rancangan Kontrak/SPK",
    ]
    
    for i, lamp in enumerate(lampiran_list, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {lamp}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # ========== PENUTUP ==========
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian nota dinas ini disampaikan untuk dapat ditindaklanjuti dengan "
        "melaksanakan proses pemilihan penyedia sesuai ketentuan yang berlaku."
    )
    penutup.paragraph_format.first_line_indent = Cm(1.25)
    penutup.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========== TANDA TANGAN ==========
    sign_table = doc.add_table(rows=6, cols=2)
    
    # Kolom kanan untuk tanda tangan
    sign_table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_nota_dinas_pp_fmt}}"
    sign_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen,"
    sign_table.rows[2].cells[1].text = ""  # Space for signature
    sign_table.rows[3].cells[1].text = ""
    sign_table.rows[4].cells[1].text = "{{ppk_nama}}"
    sign_table.rows[5].cells[1].text = "NIP. {{ppk_nip}}"
    
    # Bold nama
    for para in sign_table.rows[4].cells[1].paragraphs:
        for run in para.runs:
            run.bold = True
            run.underline = True
    
    # Alignment
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set widths
    for row in sign_table.rows:
        row.cells[0].width = Cm(8)
        row.cells[1].width = Cm(7)
    
    # ========== PAGE BREAK - LAMPIRAN CHECKLIST ==========
    doc.add_page_break()
    
    # Header lampiran
    lamp_header = doc.add_paragraph()
    lamp_header.add_run("LAMPIRAN NOTA DINAS").bold = True
    lamp_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lamp_nomor = doc.add_paragraph()
    lamp_nomor.add_run("Nomor: {{nomor_nota_dinas_pp}}")
    lamp_nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lamp_tanggal = doc.add_paragraph()
    lamp_tanggal.add_run("Tanggal: {{tanggal_nota_dinas_pp_fmt}}")
    lamp_tanggal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    checklist_title = doc.add_paragraph()
    checklist_title.add_run("CHECKLIST KELENGKAPAN DOKUMEN PERSIAPAN PENGADAAN").bold = True
    checklist_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Checklist table
    check_table = doc.add_table(rows=12, cols=4)
    check_table.style = 'Table Grid'
    
    # Header
    headers = ["No", "Dokumen", "Ada", "Tidak Ada"]
    for i, h in enumerate(headers):
        cell = check_table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    checklist_items = [
        "RKAKL/POK yang telah disahkan",
        "RUP yang telah diumumkan di SIRUP",
        "Identifikasi Kebutuhan",
        "Penetapan Spesifikasi Teknis/KAK",
        "Survei Harga Pasar (min. 3 sumber)",
        "Perhitungan HPS",
        "Dokumen HPS yang telah ditetapkan",
        "Rancangan Kontrak/SPK",
        "Dokumen Pemilihan (jika diperlukan)",
        "Pakta Integritas",
        "Surat Penunjukan PPK (SK)",
    ]
    
    for i, item in enumerate(checklist_items, 1):
        row = check_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = item
        row.cells[2].text = "☐"
        row.cells[3].text = "☐"
        
        # Center alignment for No and checkboxes
        for j in [0, 2, 3]:
            for para in row.cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set column widths
    for row in check_table.rows:
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(10)
        row.cells[2].width = Cm(2)
        row.cells[3].width = Cm(2)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Keterangan
    ket = doc.add_paragraph()
    ket.add_run("Keterangan:").bold = True
    
    ket1 = doc.add_paragraph()
    ket1.add_run("Dokumen yang dicentang (☑) pada kolom 'Ada' merupakan dokumen yang dilampirkan.")
    ket1.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Tanda tangan verifikasi
    verif_table = doc.add_table(rows=6, cols=2)
    
    verif_table.rows[0].cells[0].text = "Disiapkan oleh:"
    verif_table.rows[0].cells[1].text = "Diperiksa oleh:"
    
    verif_table.rows[4].cells[0].text = "{{ppk_nama}}"
    verif_table.rows[4].cells[1].text = "{{pejabat_pengadaan_nama}}"
    
    verif_table.rows[5].cells[0].text = "PPK"
    verif_table.rows[5].cells[1].text = "Pejabat Pengadaan"
    
    for row in verif_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    filepath = os.path.join(TEMPLATES_DIR, "nota_dinas_pp.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


def create_surat_permintaan_pengadaan_template():
    """Create Surat Permintaan Pelaksanaan Pengadaan (alternatif format)"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("SURAT PERMINTAAN PELAKSANAAN PENGADAAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_surat_permintaan}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Isi surat
    p1 = doc.add_paragraph()
    p1.add_run("Yang bertanda tangan di bawah ini:")
    
    # Info PPK
    ppk_table = doc.add_table(rows=4, cols=3)
    ppk_info = [
        ("Nama", ":", "{{ppk_nama}}"),
        ("NIP", ":", "{{ppk_nip}}"),
        ("Jabatan", ":", "Pejabat Pembuat Komitmen"),
        ("Satuan Kerja", ":", "{{satker_nama}}"),
    ]
    for i, (label, sep, value) in enumerate(ppk_info):
        ppk_table.rows[i].cells[0].text = label
        ppk_table.rows[i].cells[1].text = sep
        ppk_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    p2 = doc.add_paragraph()
    p2.add_run(
        "Dengan ini menyampaikan permintaan untuk melaksanakan proses pemilihan penyedia "
        "barang/jasa dengan rincian sebagai berikut:"
    )
    p2.paragraph_format.first_line_indent = Cm(1.25)
    
    doc.add_paragraph()
    
    # Info Paket
    paket_table = doc.add_table(rows=9, cols=3)
    paket_table.style = 'Table Grid'
    
    paket_info = [
        ("1.", "Nama Paket Pekerjaan", "{{nama_paket}}"),
        ("2.", "Kode RUP", "{{kode_rup}}"),
        ("3.", "Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
        ("4.", "Sumber Dana", "{{sumber_dana}} - MAK {{kode_akun}}"),
        ("5.", "Nilai Pagu", "{{nilai_pagu_fmt}}"),
        ("6.", "Nilai HPS", "{{nilai_hps_fmt}}"),
        ("7.", "Metode Pemilihan", "{{metode_pemilihan}}"),
        ("8.", "Jenis Kontrak", "{{jenis_kontrak}}"),
        ("9.", "Waktu Pelaksanaan", "{{jangka_waktu}} hari kalender"),
    ]
    
    for i, (no, label, value) in enumerate(paket_info):
        paket_table.rows[i].cells[0].text = no
        paket_table.rows[i].cells[1].text = label
        paket_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Lampiran
    lamp = doc.add_paragraph()
    lamp.add_run("Bersama surat ini dilampirkan:").bold = True
    
    lampiran = [
        "Kerangka Acuan Kerja (KAK) / Spesifikasi Teknis",
        "Harga Perkiraan Sendiri (HPS)",
        "Rancangan Kontrak",
        "Copy RKAKL/DIPA/POK",
        "Copy RUP",
        "Dokumen pendukung lainnya",
    ]
    
    for i, l in enumerate(lampiran, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {l}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # Penutup
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian surat permintaan ini dibuat untuk dapat ditindaklanjuti sesuai "
        "ketentuan peraturan perundang-undangan yang berlaku."
    )
    penutup.paragraph_format.first_line_indent = Cm(1.25)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    sign_table = doc.add_table(rows=5, cols=2)
    sign_table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_hari_ini_fmt}}"
    sign_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen,"
    sign_table.rows[3].cells[1].text = "{{ppk_nama}}"
    sign_table.rows[4].cells[1].text = "NIP. {{ppk_nip}}"
    
    for para in sign_table.rows[3].cells[1].paragraphs:
        for run in para.runs:
            run.bold = True
            run.underline = True
    
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    filepath = os.path.join(TEMPLATES_DIR, "surat_permintaan_pengadaan.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


if __name__ == "__main__":
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    
    print("Creating Nota Dinas templates...")
    print()
    
    create_nota_dinas_pp_template()
    create_surat_permintaan_pengadaan_template()
    
    print()
    print("=" * 50)
    print("Templates created!")
    print()
    print("Placeholders yang digunakan:")
    print("  {{nomor_nota_dinas_pp}}     - Nomor nota dinas")
    print("  {{tanggal_nota_dinas_pp_fmt}} - Tanggal nota dinas")
    print("  {{kode_rup}}                - Kode RUP")
    print("  {{metode_pemilihan}}        - Metode pemilihan penyedia")
    print("  {{pejabat_pengadaan_nama}}  - Nama Pejabat Pengadaan")
