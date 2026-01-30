#!/usr/bin/env python3
"""
Script untuk membuat template Kuitansi Uang Muka yang simpel.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_single_spacing(paragraph):
    """Set paragraph to single spacing with no space before/after."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    """Set cell margins."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(margin_value))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)


def create_kuitansi_um_template():
    """Create simple Kuitansi Uang Muka template."""
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== HEADER: KUITANSI ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(title)
    title_run = title.add_run("KUITANSI")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Subtitle: PEMBAYARAN UANG MUKA
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(subtitle)
    sub_run = subtitle.add_run("PEMBAYARAN UANG MUKA")
    sub_run.font.size = Pt(12)

    doc.add_paragraph()

    # ========== INFO TABLE ==========
    info_table = doc.add_table(rows=5, cols=3)
    info_table.autofit = False

    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(11)

    info_data = [
        ("Nomor Kuitansi", ":", "{{nomor_kuitansi}}"),
        ("Tanggal", ":", "{{tanggal_dokumen:tanggal_panjang}}"),
        ("Nama Kegiatan", ":", "{{nama_kegiatan}}"),
        ("Sumber Dana/MAK", ":", "{{kode_akun}}"),
        ("Unit Kerja", ":", "{{unit_kerja}}"),
    ]

    for i, (label, sep, value) in enumerate(info_data):
        row = info_table.rows[i]

        # Label cell
        cell_label = row.cells[0]
        p = cell_label.paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(label)
        run.font.size = Pt(11)

        # Separator cell
        cell_sep = row.cells[1]
        p = cell_sep.paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(sep)
        run.font.size = Pt(11)

        # Value cell
        cell_value = row.cells[2]
        p = cell_value.paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(value)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # ========== NILAI UANG MUKA ==========
    nilai_p = doc.add_paragraph()
    set_single_spacing(nilai_p)
    nilai_p.add_run("Telah diterima uang sebesar:").font.size = Pt(11)

    # Amount box
    amount_table = doc.add_table(rows=1, cols=1)
    amount_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    amount_cell = amount_table.rows[0].cells[0]
    amount_cell.width = Cm(10)

    # Add border to amount cell
    tc = amount_cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

    amount_p = amount_cell.paragraphs[0]
    amount_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(amount_p)
    amount_run = amount_p.add_run("{{uang_muka:rupiah}}")
    amount_run.bold = True
    amount_run.font.size = Pt(14)

    # Terbilang
    terbilang_p = doc.add_paragraph()
    terbilang_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(terbilang_p)
    terbilang_run = terbilang_p.add_run("({{uang_muka:terbilang}} rupiah)")
    terbilang_run.italic = True
    terbilang_run.font.size = Pt(10)

    doc.add_paragraph()

    # Keterangan persentase
    persen_p = doc.add_paragraph()
    set_single_spacing(persen_p)
    persen_p.add_run("Persentase Uang Muka: {{persentase_um}} dari estimasi biaya {{estimasi_biaya:rupiah}}").font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # ========== SIGNATURES (2 columns) ==========
    sig_table = doc.add_table(rows=5, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in sig_table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(7)

    # Row 0: Headers
    left_header = sig_table.rows[0].cells[0]
    p = left_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("Yang Menyerahkan,")
    run.font.size = Pt(11)

    right_header = sig_table.rows[0].cells[1]
    p = right_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("Yang Menerima,")
    run.font.size = Pt(11)

    # Row 1: Jabatan
    left_jabatan = sig_table.rows[1].cells[0]
    p = left_jabatan.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("Bendahara Pengeluaran")
    run.font.size = Pt(10)

    right_jabatan = sig_table.rows[1].cells[1]
    p = right_jabatan.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("{{penerima_jabatan}}")
    run.font.size = Pt(10)

    # Row 2: Empty space for signature
    for cell in sig_table.rows[2].cells:
        p = cell.paragraphs[0]
        set_single_spacing(p)
        p.add_run("\n\n\n")

    # Row 3: Names (underlined)
    left_name = sig_table.rows[3].cells[0]
    p = left_name.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("{{bendahara_nama}}")
    run.font.size = Pt(11)
    run.underline = True

    right_name = sig_table.rows[3].cells[1]
    p = right_name.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("{{penerima_nama}}")
    run.font.size = Pt(11)
    run.underline = True

    # Row 4: NIP
    left_nip = sig_table.rows[4].cells[0]
    p = left_nip.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("NIP. {{bendahara_nip}}")
    run.font.size = Pt(10)

    right_nip = sig_table.rows[4].cells[1]
    p = right_nip.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("NIP. {{penerima_nip}}")
    run.font.size = Pt(10)

    # ========== FOOTER NOTE ==========
    doc.add_paragraph()
    doc.add_paragraph()

    note_p = doc.add_paragraph()
    set_single_spacing(note_p)
    note_run = note_p.add_run("Catatan: Kuitansi ini sebagai bukti penerimaan uang muka untuk kegiatan tersebut di atas.")
    note_run.font.size = Pt(9)
    note_run.italic = True

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "word" / "kuitansi_uang_muka.docx"
    doc.save(str(output_path))
    print(f"Template saved to: {output_path}")

    return output_path


if __name__ == "__main__":
    create_kuitansi_um_template()
