#!/usr/bin/env python3
"""
Update Kuitansi Uang Muka template to include auto-generated nomor tanda terima.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEMPLATE_DIR = Path(__file__).parent.parent / "templates" / "word"


def create_kuitansi_uang_muka_template():
    """Create Kuitansi Uang Muka template with auto-generated nomor tanda terima."""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # === KOP SURAT ===
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kop.add_run("{{kop_surat}}")
    run.bold = True
    run.font.size = Pt(14)

    # Separator line
    sep = doc.add_paragraph("_" * 70)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === TITLE ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KUITANSI / TANDA TERIMA UANG MUKA")
    run.bold = True
    run.underline = True
    run.font.size = Pt(14)

    # Nomor Tanda Terima (Auto-generated)
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_tanda_terima}}")

    doc.add_paragraph()

    # === SUDAH TERIMA DARI ===
    p1 = doc.add_paragraph()
    p1.add_run("Sudah terima dari\t: ").bold = True
    p1.add_run("Bendahara Pengeluaran {{satker_nama}}")

    # === UANG SEJUMLAH ===
    p2 = doc.add_paragraph()
    p2.add_run("Uang sejumlah\t\t: ").bold = True
    p2.add_run("Rp {{nilai:rupiah}}")

    # === TERBILANG ===
    p3 = doc.add_paragraph()
    p3.add_run("Terbilang\t\t: ").bold = True
    p3.add_run("{{nilai:terbilang}}")

    # === UNTUK KEPERLUAN ===
    p4 = doc.add_paragraph()
    p4.add_run("Untuk keperluan\t: ").bold = True
    p4.add_run("Uang muka kegiatan {{nama_kegiatan}}")

    doc.add_paragraph()

    # === TABEL RINCIAN ===
    rincian_title = doc.add_paragraph()
    rincian_title.add_run("Rincian:").bold = True

    # Table header
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    headers = ["No", "Uraian", "Volume", "Satuan", "Jumlah"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Rincian rows (placeholder)
    rincian_row = table.add_row()
    rincian_row.cells[0].text = "{{rincian_no}}"
    rincian_row.cells[1].text = "{{rincian_uraian}}"
    rincian_row.cells[2].text = "{{rincian_volume}}"
    rincian_row.cells[3].text = "{{rincian_satuan}}"
    rincian_row.cells[4].text = "{{rincian_jumlah}}"

    # Total row
    total_row = table.add_row()
    total_row.cells[0].merge(total_row.cells[3])
    total_row.cells[0].text = "TOTAL"
    total_row.cells[0].paragraphs[0].runs[0].bold = True
    total_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row.cells[4].text = "{{subtotal}}"
    total_row.cells[4].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # === TANDA TANGAN ===
    ttd_container = doc.add_table(rows=1, cols=2)

    # Kolom kiri - Mengetahui
    cell_left = ttd_container.rows[0].cells[0]
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.add_run("Mengetahui,\n")
    p_left.add_run("Pejabat Pembuat Komitmen\n\n\n\n")
    run_left = p_left.add_run("{{nama_ppk}}\n")
    run_left.bold = True
    p_left.add_run("NIP. {{nip_ppk}}")

    # Kolom kanan - Penerima
    cell_right = ttd_container.rows[0].cells[1]
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.add_run("{{kota}}, {{tanggal_dokumen:tanggal}}\n")
    p_right.add_run("Yang Menerima,\n\n\n\n")
    run_right = p_right.add_run("{{penerima_nama}}\n")
    run_right.bold = True
    p_right.add_run("NIP. {{penerima_nip}}")

    # Save
    output_path = TEMPLATE_DIR / "kuitansi_uang_muka.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Updating Kuitansi Uang Muka template...")
    print("-" * 50)
    create_kuitansi_uang_muka_template()
    print("-" * 50)
    print("Done!")
