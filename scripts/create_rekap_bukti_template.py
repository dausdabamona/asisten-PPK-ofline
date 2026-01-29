"""
Script untuk membuat template Rekap Bukti Pengeluaran (REKAP_BKT).
- Format Word (.docx)
- Spasi 1 (single spacing)
- Tabel dengan kolom nomor yang lebar menyesuaikan isi
"""

from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pathlib import Path


def set_cell_border(cell, **kwargs):
    """
    Set cell border.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="{top_val}" w:sz="{top_sz}" w:color="{top_color}"/>'
        r'<w:left w:val="{left_val}" w:sz="{left_sz}" w:color="{left_color}"/>'
        r'<w:bottom w:val="{bottom_val}" w:sz="{bottom_sz}" w:color="{bottom_color}"/>'
        r'<w:right w:val="{right_val}" w:sz="{right_sz}" w:color="{right_color}"/>'
        r'</w:tcBorders>'.format(
            top_val=kwargs.get("top", {}).get("val", "single"),
            top_sz=kwargs.get("top", {}).get("sz", 4),
            top_color=kwargs.get("top", {}).get("color", "000000"),
            left_val=kwargs.get("left", {}).get("val", "single"),
            left_sz=kwargs.get("left", {}).get("sz", 4),
            left_color=kwargs.get("left", {}).get("color", "000000"),
            bottom_val=kwargs.get("bottom", {}).get("val", "single"),
            bottom_sz=kwargs.get("bottom", {}).get("sz", 4),
            bottom_color=kwargs.get("bottom", {}).get("color", "000000"),
            right_val=kwargs.get("right", {}).get("val", "single"),
            right_sz=kwargs.get("right", {}).get("sz", 4),
            right_color=kwargs.get("right", {}).get("color", "000000"),
        )
    )
    tcPr.append(tcBorders)


def set_single_spacing(paragraph):
    """Set paragraph to single spacing."""
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def create_rekap_bukti_template():
    """
    Create Rekap Bukti Pengeluaran template.
    """
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # --- HEADER (KOP SURAT) ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("{{kementerian}}")
    run.font.size = Pt(12)
    run.font.bold = True
    set_single_spacing(header)

    header_sub = doc.add_paragraph()
    header_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = header_sub.add_run("{{satker_nama}}")
    run2.font.size = Pt(11)
    set_single_spacing(header_sub)

    # Garis pemisah
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = separator.add_run("_" * 80)
    run_sep.font.size = Pt(8)
    set_single_spacing(separator)

    # Spacing
    spacing = doc.add_paragraph()
    set_single_spacing(spacing)

    # --- JUDUL ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("REKAP BUKTI PENGELUARAN")
    run_title.bold = True
    run_title.font.size = Pt(14)
    set_single_spacing(title)

    # Nomor dokumen
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_dokumen}}")
    set_single_spacing(nomor)

    # Spacing
    doc.add_paragraph()

    # --- INFO KEGIATAN ---
    info_table = doc.add_table(rows=5, cols=3)
    info_items = [
        ("Nama Kegiatan", ":", "{{nama_kegiatan}}"),
        ("Kode Akun", ":", "{{kode_akun}}"),
        ("Penerima", ":", "{{penerima_nama}}"),
        ("NIP", ":", "{{penerima_nip}}"),
        ("Tanggal", ":", "{{tanggal_dokumen}}"),
    ]

    for i, (label, sep, value) in enumerate(info_items):
        cell0 = info_table.cell(i, 0)
        cell1 = info_table.cell(i, 1)
        cell2 = info_table.cell(i, 2)

        cell0.text = label
        cell1.text = sep
        cell2.text = value

        # Set font size
        for cell in [cell0, cell1, cell2]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
                set_single_spacing(para)

    # Set column widths for info table
    info_table.columns[0].width = Cm(3.5)
    info_table.columns[1].width = Cm(0.5)
    info_table.columns[2].width = Cm(12)

    # Spacing
    doc.add_paragraph()

    # --- TABEL RINCIAN ---
    rincian_title = doc.add_paragraph()
    rincian_title.add_run("Rincian Pengeluaran:")
    rincian_title.runs[0].bold = True
    rincian_title.runs[0].font.size = Pt(11)
    set_single_spacing(rincian_title)

    # Create table with 6 columns
    # No | Uraian | Volume | Satuan | Harga Satuan | Jumlah
    rincian_table = doc.add_table(rows=2, cols=6)
    rincian_table.style = 'Table Grid'
    rincian_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["No", "Uraian", "Volume", "Satuan", "Harga Satuan", "Jumlah"]
    header_row = rincian_table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # Format header
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            set_single_spacing(para)

    # Data row template (will be repeated for each item)
    data_row = rincian_table.rows[1]
    data_values = [
        "{{rincian_no}}",
        "{{rincian_uraian}}",
        "{{rincian_volume}}",
        "{{rincian_satuan}}",
        "{{rincian_harga_satuan}}",
        "{{rincian_jumlah}}"
    ]
    for i, value in enumerate(data_values):
        cell = data_row.cells[i]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in cell.paragraphs:
            # Align numbers to right, text to left
            if i in [0, 2, 4, 5]:  # No, Volume, Harga, Jumlah - right align
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif i == 3:  # Satuan - center
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:  # Uraian - left
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(10)
            set_single_spacing(para)

    # Set column widths - No column narrow, Uraian wide
    rincian_table.columns[0].width = Cm(1)      # No - narrow
    rincian_table.columns[1].width = Cm(6)      # Uraian - wide
    rincian_table.columns[2].width = Cm(1.5)    # Volume
    rincian_table.columns[3].width = Cm(2)      # Satuan
    rincian_table.columns[4].width = Cm(3)      # Harga Satuan
    rincian_table.columns[5].width = Cm(3)      # Jumlah

    # Spacing
    doc.add_paragraph()

    # --- TOTAL SECTION ---
    total_table = doc.add_table(rows=4, cols=2)

    total_items = [
        ("SUBTOTAL", "{{subtotal}}"),
        ("PPN ({{ppn_persen}}%)", "{{ppn_nilai}}"),
        ("PPh ({{pph_persen}}%)", "{{pph_nilai}}"),
        ("GRAND TOTAL", "{{grand_total}}"),
    ]

    for i, (label, value) in enumerate(total_items):
        cell0 = total_table.cell(i, 0)
        cell1 = total_table.cell(i, 1)

        cell0.text = label
        cell1.text = value

        # Format
        for para in cell0.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.size = Pt(11)
                if i == 3:  # GRAND TOTAL bold
                    run.bold = True
            set_single_spacing(para)

        for para in cell1.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.size = Pt(11)
                if i == 3:  # GRAND TOTAL bold
                    run.bold = True
            set_single_spacing(para)

    # Set total table column widths
    total_table.columns[0].width = Cm(13)
    total_table.columns[1].width = Cm(3.5)

    # Spacing
    doc.add_paragraph()

    # --- TERBILANG ---
    terbilang = doc.add_paragraph()
    terbilang.add_run("Terbilang: ")
    terbilang.add_run("{{terbilang}}")
    terbilang.runs[0].bold = True
    terbilang.runs[0].font.size = Pt(11)
    terbilang.runs[1].font.size = Pt(11)
    terbilang.runs[1].italic = True
    set_single_spacing(terbilang)

    # Spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # --- TANDA TANGAN ---
    ttd_intro = doc.add_paragraph()
    ttd_intro.add_run("Demikian rekap bukti pengeluaran ini dibuat dengan sebenar-benarnya.")
    ttd_intro.runs[0].font.size = Pt(11)
    set_single_spacing(ttd_intro)

    # Spacing
    doc.add_paragraph()

    # Lokasi dan tanggal
    ttd_lokasi = doc.add_paragraph()
    ttd_lokasi.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ttd_lokasi.add_run("{{lokasi}}, {{tanggal_dokumen}}")
    ttd_lokasi.runs[0].font.size = Pt(11)
    set_single_spacing(ttd_lokasi)

    # Tanda tangan table (2 columns)
    ttd_table = doc.add_table(rows=6, cols=2)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header - jabatan
    ttd_table.cell(0, 0).text = "Penerima/Pelaksana,"
    ttd_table.cell(0, 1).text = "Mengetahui,\nPejabat Pembuat Komitmen"

    # Empty rows for signature space
    for i in range(1, 5):
        ttd_table.cell(i, 0).text = ""
        ttd_table.cell(i, 1).text = ""

    # Name and NIP
    ttd_table.cell(5, 0).text = "{{penerima_nama}}\nNIP. {{penerima_nip}}"
    ttd_table.cell(5, 1).text = "{{ppk_nama}}\nNIP. {{ppk_nip}}"

    # Format TTD table
    for row in ttd_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(11)
                set_single_spacing(para)

    # Set TTD table column widths
    ttd_table.columns[0].width = Cm(8)
    ttd_table.columns[1].width = Cm(8)

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "word" / "rekap_bukti_pengeluaran.docx"
    doc.save(str(output_path))
    print(f"Rekap Bukti Pengeluaran template saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_rekap_bukti_template()
    print("\nTemplate Rekap Bukti Pengeluaran berhasil dibuat!")
