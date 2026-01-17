"""
Script untuk membuat template Survey Harga dan HPS
dengan format tabel item yang benar
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = "templates/word"

def set_cell_shading(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_survey_harga_template():
    """Create Survey Harga template with item loop"""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("FORM SURVEY HARGA")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Paket info
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run("Nama Paket: ").bold = True
    info.add_run("{{nama_paket}}")
    
    info2 = doc.add_paragraph()
    info2.add_run("Kode Paket: ").bold = True
    info2.add_run("{{kode_paket}}")
    
    info3 = doc.add_paragraph()
    info3.add_run("Tahun Anggaran: ").bold = True
    info3.add_run("{{tahun_anggaran}}")
    
    doc.add_paragraph()
    
    # Survey Sources info
    doc.add_paragraph("Sumber Survey Harga:")
    
    survey_info = doc.add_table(rows=4, cols=4)
    survey_info.style = 'Table Grid'
    
    # Header
    headers = ["No", "Jenis", "Nama Sumber", "Alamat/Link"]
    for i, h in enumerate(headers):
        cell = survey_info.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "2E86AB")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = None  # Will be white from Word
    
    # Survey 1-3
    for i in range(1, 4):
        row = survey_info.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = "{{" + f"survey{i}_jenis" + "}}"
        row.cells[2].text = "{{" + f"survey{i}_nama" + "}}"
        row.cells[3].text = "{{" + f"survey{i}_alamat_lengkap" + "}}"
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Item Table with survey prices
    doc.add_paragraph("Daftar Item dan Hasil Survey:")
    
    # Create main item table with loop placeholders
    table = doc.add_table(rows=2, cols=9)
    table.style = 'Table Grid'
    
    # Header row
    headers = [
        "No", "Uraian Barang/Jasa", "Spesifikasi", "Satuan", "Volume",
        "Harga Survey 1", "Harga Survey 2", "Harga Survey 3", "Harga Rata-rata"
    ]
    
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        set_cell_shading(cell, "2E86AB")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    # Template row with placeholders - THIS IS THE KEY!
    template_row = table.rows[1]
    template_row.cells[0].text = "{{no}}"
    template_row.cells[1].text = "{{item.uraian}}"
    template_row.cells[2].text = "{{item.spesifikasi}}"
    template_row.cells[3].text = "{{item.satuan}}"
    template_row.cells[4].text = "{{item.volume}}"
    template_row.cells[5].text = "{{item.harga_survey1}}"
    template_row.cells[6].text = "{{item.harga_survey2}}"
    template_row.cells[7].text = "{{item.harga_survey3}}"
    template_row.cells[8].text = "{{item.harga_rata}}"
    
    # Set column widths
    widths = [Cm(1), Cm(4), Cm(3), Cm(1.5), Cm(1.5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width
    
    doc.add_paragraph()
    
    # Summary
    summary = doc.add_paragraph()
    summary.add_run("Keterangan:").bold = True
    doc.add_paragraph("1. Harga survey merupakan harga per satuan sebelum pajak")
    doc.add_paragraph("2. Minimal 3 sumber harga untuk setiap item")
    doc.add_paragraph("3. Sumber harga dapat berupa: Toko fisik, Marketplace, Kontrak sebelumnya, Brosur/Katalog")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    sign_table = doc.add_table(rows=5, cols=2)
    sign_table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_hari_ini_fmt}}"
    sign_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen"
    sign_table.rows[3].cells[1].text = "{{ppk_nama}}"
    sign_table.rows[4].cells[1].text = "NIP. {{ppk_nip}}"
    
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    filepath = os.path.join(TEMPLATES_DIR, "survey_harga.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


def create_hps_template():
    """Create HPS template with item loop"""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("HARGA PERKIRAAN SENDIRI (HPS)")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Paket info
    info1 = doc.add_paragraph()
    info1.add_run("Nama Paket\t\t: ").bold = True
    info1.add_run("{{nama_paket}}")
    
    info2 = doc.add_paragraph()
    info2.add_run("Kode Paket\t\t: ").bold = True
    info2.add_run("{{kode_paket}}")
    
    info3 = doc.add_paragraph()
    info3.add_run("Tahun Anggaran\t: ").bold = True
    info3.add_run("{{tahun_anggaran}}")
    
    info4 = doc.add_paragraph()
    info4.add_run("Pagu Anggaran\t: ").bold = True
    info4.add_run("{{nilai_pagu_fmt}}")
    
    info5 = doc.add_paragraph()
    info5.add_run("Metode HPS\t\t: ").bold = True
    info5.add_run("{{metode_hps}}")
    
    doc.add_paragraph()
    
    # Item Table with HPS
    doc.add_paragraph("A. DAFTAR BARANG/JASA DAN PERHITUNGAN HPS")
    
    table = doc.add_table(rows=2, cols=8)
    table.style = 'Table Grid'
    
    # Header
    headers = [
        "No", "Uraian Barang/Jasa", "Spesifikasi", "Satuan", "Volume",
        "Harga Satuan Survey", "Harga HPS/Satuan", "Total HPS"
    ]
    
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        set_cell_shading(cell, "27AE60")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    # Template row with placeholders
    template_row = table.rows[1]
    template_row.cells[0].text = "{{no}}"
    template_row.cells[1].text = "{{item.uraian}}"
    template_row.cells[2].text = "{{item.spesifikasi}}"
    template_row.cells[3].text = "{{item.satuan}}"
    template_row.cells[4].text = "{{item.volume}}"
    template_row.cells[5].text = "{{item.harga_rata}}"
    template_row.cells[6].text = "{{item.harga_hps_satuan}}"
    template_row.cells[7].text = "{{item.total_hps}}"
    
    # Column widths
    widths = [Cm(1), Cm(4), Cm(3), Cm(1.5), Cm(1.5), Cm(2.5), Cm(2.5), Cm(3)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width
    
    doc.add_paragraph()
    
    # Summary section
    doc.add_paragraph("B. REKAPITULASI")
    
    recap_table = doc.add_table(rows=4, cols=2)
    recap_table.style = 'Table Grid'
    
    recap_table.rows[0].cells[0].text = "Subtotal"
    recap_table.rows[0].cells[1].text = "{{subtotal_item_fmt}}"
    
    recap_table.rows[1].cells[0].text = "PPN 11%"
    recap_table.rows[1].cells[1].text = "{{ppn_item_fmt}}"
    
    recap_table.rows[2].cells[0].text = "TOTAL HPS"
    recap_table.rows[2].cells[1].text = "{{grand_total_item_fmt}}"
    set_cell_shading(recap_table.rows[2].cells[0], "D5F5E3")
    set_cell_shading(recap_table.rows[2].cells[1], "D5F5E3")
    for cell in recap_table.rows[2].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    recap_table.rows[3].cells[0].text = "Terbilang"
    recap_table.rows[3].cells[1].text = "{{grand_total_item_terbilang}}"
    
    # Set width
    recap_table.columns[0].width = Cm(4)
    recap_table.columns[1].width = Cm(8)
    
    doc.add_paragraph()
    
    # Notes
    doc.add_paragraph("C. KETERANGAN")
    doc.add_paragraph("1. HPS disusun berdasarkan data/informasi yang dapat dipertanggungjawabkan")
    doc.add_paragraph("2. Harga satuan HPS sudah termasuk biaya umum dan keuntungan yang wajar")
    doc.add_paragraph("3. HPS bersifat rahasia dan tidak boleh diumumkan")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    sign_table = doc.add_table(rows=5, cols=2)
    sign_table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_hari_ini_fmt}}"
    sign_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen"
    sign_table.rows[3].cells[1].text = "{{ppk_nama}}"
    sign_table.rows[4].cells[1].text = "NIP. {{ppk_nip}}"
    
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    filepath = os.path.join(TEMPLATES_DIR, "hps.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


def create_ba_survey_harga_template():
    """Create BA Survey Harga template"""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("BERITA ACARA SURVEY HARGA")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Number and date
    info1 = doc.add_paragraph()
    info1.add_run("Nomor\t: ").bold = True
    info1.add_run("{{nomor_ba_survey}}")
    
    info2 = doc.add_paragraph()
    info2.add_run("Tanggal\t: ").bold = True
    info2.add_run("{{tanggal_hari_ini_fmt}}")
    
    doc.add_paragraph()
    
    # Body
    body = doc.add_paragraph()
    body.add_run(
        "Pada hari ini {{hari}} tanggal {{tanggal_hari_ini_fmt}}, kami yang bertanda tangan di bawah ini, "
        "Pejabat Pembuat Komitmen pada {{satker_nama}}, telah melakukan survey harga untuk kegiatan pengadaan:\n\n"
    )
    
    info_paket = doc.add_paragraph()
    info_paket.add_run("Nama Paket\t\t: ").bold = True
    info_paket.add_run("{{nama_paket}}")
    
    info_kode = doc.add_paragraph()
    info_kode.add_run("Kode Paket\t\t: ").bold = True
    info_kode.add_run("{{kode_paket}}")
    
    info_nilai = doc.add_paragraph()
    info_nilai.add_run("Pagu Anggaran\t: ").bold = True
    info_nilai.add_run("{{nilai_pagu_fmt}}")
    
    doc.add_paragraph()
    
    # Survey sources
    doc.add_paragraph("Survey dilakukan pada sumber-sumber berikut:")
    
    survey_table = doc.add_table(rows=4, cols=4)
    survey_table.style = 'Table Grid'
    
    headers = ["No", "Jenis Sumber", "Nama Sumber", "Alamat/Kota"]
    for i, h in enumerate(headers):
        cell = survey_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "3498DB")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    for i in range(1, 4):
        row = survey_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = "{{" + f"survey{i}_jenis" + "}}"
        row.cells[2].text = "{{" + f"survey{i}_nama" + "}}"
        row.cells[3].text = "{{" + f"survey{i}_alamat_lengkap" + "}}"
    
    doc.add_paragraph()
    
    # Result table
    doc.add_paragraph("Dengan hasil survey sebagai berikut:")
    
    result_table = doc.add_table(rows=2, cols=8)
    result_table.style = 'Table Grid'
    
    headers = ["No", "Uraian", "Satuan", "Vol", "Survey 1", "Survey 2", "Survey 3", "Rata-rata"]
    for i, h in enumerate(headers):
        cell = result_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "3498DB")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    # Template row
    template_row = result_table.rows[1]
    template_row.cells[0].text = "{{no}}"
    template_row.cells[1].text = "{{item.uraian}}"
    template_row.cells[2].text = "{{item.satuan}}"
    template_row.cells[3].text = "{{item.volume}}"
    template_row.cells[4].text = "{{item.harga_survey1}}"
    template_row.cells[5].text = "{{item.harga_survey2}}"
    template_row.cells[6].text = "{{item.harga_survey3}}"
    template_row.cells[7].text = "{{item.harga_rata}}"
    
    doc.add_paragraph()
    
    # Closing
    closing = doc.add_paragraph()
    closing.add_run(
        "Demikian Berita Acara ini dibuat dengan sebenar-benarnya untuk dapat dipergunakan sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    sign_table = doc.add_table(rows=5, cols=2)
    sign_table.rows[0].cells[1].text = "{{satker_kota}}, {{tanggal_hari_ini_fmt}}"
    sign_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen"
    sign_table.rows[3].cells[1].text = "{{ppk_nama}}"
    sign_table.rows[4].cells[1].text = "NIP. {{ppk_nip}}"
    
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    filepath = os.path.join(TEMPLATES_DIR, "ba_survey_harga.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


def update_items_formatted_with_hps():
    """
    Update template engine to include HPS fields in items_formatted
    """
    # This is the code snippet that should be added to template engine's get_data_for_merge
    code = '''
    # Add HPS fields to items_formatted
    for i, item in enumerate(items, 1):
        data['items_formatted'].append({
            'no': i,
            'nomor_urut': item['nomor_urut'],
            'kategori': item.get('kategori', ''),
            'kelompok': item.get('kelompok', ''),
            'uraian': item.get('uraian', ''),
            'spesifikasi': item.get('spesifikasi', ''),
            'satuan': item.get('satuan', ''),
            'volume': item.get('volume', 0),
            'volume_fmt': f"{item.get('volume', 0):,.2f}".replace(",", "."),
            'harga_dasar': item.get('harga_dasar', 0),
            'harga_dasar_fmt': format_rupiah(item.get('harga_dasar', 0)),
            'harga_survey1': item.get('harga_survey1', 0) or 0,
            'harga_survey1_fmt': format_rupiah(item.get('harga_survey1', 0)),
            'harga_survey2': item.get('harga_survey2', 0) or 0,
            'harga_survey2_fmt': format_rupiah(item.get('harga_survey2', 0)),
            'harga_survey3': item.get('harga_survey3', 0) or 0,
            'harga_survey3_fmt': format_rupiah(item.get('harga_survey3', 0)),
            'harga_rata': item.get('harga_rata', 0) or 0,
            'harga_rata_fmt': format_rupiah(item.get('harga_rata', 0)),
            # HPS fields
            'harga_hps_satuan': item.get('harga_hps_satuan', 0) or 0,
            'harga_hps_satuan_fmt': format_rupiah(item.get('harga_hps_satuan', 0)),
            'total_hps': item.get('total_hps', 0) or 0,
            'total_hps_fmt': format_rupiah(item.get('total_hps', 0)),
            # Kontrak fields
            'harga_kontrak_satuan': item.get('harga_kontrak_satuan', 0) or 0,
            'harga_kontrak_satuan_fmt': format_rupiah(item.get('harga_kontrak_satuan', 0)),
            'total_kontrak': item.get('total_kontrak', 0) or 0,
            'total_kontrak_fmt': format_rupiah(item.get('total_kontrak', 0)),
            # Other
            'total': item.get('total', 0),
            'total_fmt': format_rupiah(item.get('total', 0)),
            'keterangan': item.get('keterangan', '')
        })
    '''
    return code


if __name__ == "__main__":
    import sys
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    
    print("Creating Survey Harga and HPS templates...")
    print()
    
    create_survey_harga_template()
    create_hps_template()
    create_ba_survey_harga_template()
    
    print()
    print("=" * 50)
    print("Templates created successfully!")
    print()
    print("Templates now include proper {{item.xxx}} placeholders")
    print("that will be replaced with actual item data.")
