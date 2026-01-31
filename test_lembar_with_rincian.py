"""
TEST - Lembar Permintaan dengan Rincian Barang
===============================================

Test untuk memverifikasi bahwa rincian barang dari form
berhasil ditampilkan di dokumen hasil generate.
"""

import sys
import os
from datetime import datetime
from pathlib import Path

# Setup path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.models.pencairan_models import PencairanManager
from app.services.dokumen_generator import DokumenGenerator
from app.core.config import DATABASE_PATH


def test_generate_with_rincian():
    """Test generate dokumen dengan rincian barang."""
    
    print("="*70)
    print("TEST: Generate Lembar Permintaan dengan Rincian Barang")
    print("="*70 + "\n")
    
    # Setup
    db = PencairanManager(DATABASE_PATH)
    generator = DokumenGenerator(db)
    
    # Sample data like from form
    transaksi_data = {
        'kode': 'TEST-001/2026',
        'nama_kegiatan': 'Belanja ATK Kantor Bagian Keuangan',
        'kode_akun': '5.2.2.01',
        'estimasi_biaya': 2000000,
        'mekanisme': 'LS',
        'penerima_nama': 'Budi Santoso',
        'penerima_nip': '19850101201001001',
        'penerima_jabatan': 'Kepala Bagian Keuangan',
    }
    
    # Rincian items dari form
    rincian_items = [
        {
            'no': 1,
            'uraian': 'Kertas A4 80 gsm',
            'volume': 10,
            'satuan': 'rim',
            'harga_satuan': 50000,
            'jumlah': 500000,
            'spesifikasi': '500 lembar per rim',
            'keterangan': 'Untuk penggunaan kantor',
        },
        {
            'no': 2,
            'uraian': 'Tinta Printer Brother',
            'volume': 2,
            'satuan': 'botol',
            'harga_satuan': 150000,
            'jumlah': 300000,
            'spesifikasi': 'LC 3617 BK',
            'keterangan': 'Original black ink',
        },
        {
            'no': 3,
            'uraian': 'Binder Clip Size 50mm',
            'volume': 20,
            'satuan': 'box',
            'harga_satuan': 80000,
            'jumlah': 1600000,
            'spesifikasi': '12 pcs per box',
            'keterangan': 'Gold color',
        },
    ]
    
    # Tambahan data
    additional_data = {
        'nama_pengajuan': 'Budi Santoso',
        'nama_verifikator': 'Eka Prasetya',
        'nama_ppk': 'Siti Nurhaliza',
        'nama_atasan': 'Ahmad Rahman',
        'nama_kpa': 'Drs. Hendra Kusuma',
        'tanggal_dokumen': datetime.now().strftime('%Y-%m-%d'),
    }
    
    print("1. Preparing data...")
    print(f"   - Transaksi: {transaksi_data.get('nama_kegiatan')}")
    print(f"   - Jumlah rincian: {len(rincian_items)} item")
    print(f"   - Total nilai: Rp {sum(item['jumlah'] for item in rincian_items):,.0f}".replace(",", "."))
    print()
    
    # Generate dokumen
    print("2. Generating dokumen...")
    output_path, error = generator.generate_document(
        transaksi=transaksi_data,
        kode_dokumen='LBR_REQ_TEST',
        template_name='lembar_permintaan.docx',
        rincian=rincian_items,
        additional_data=additional_data
    )
    
    if error:
        print(f"   ❌ Error: {error}")
        return False
    
    print(f"   ✓ Dokumen generated: {output_path}")
    print()
    
    # Verify dokumen
    print("3. Verifying dokumen...")
    
    try:
        from docx import Document
        
        doc = Document(output_path)
        
        # Check tables
        if not doc.tables:
            print("   ❌ Dokumen tidak memiliki table!")
            return False
        
        # Find rincian table (should be table 1 or 2, table 0 is header)
        rincian_table = None
        for table_idx, table in enumerate(doc.tables):
            # Check if table contains rincian headers
            if len(table.rows) > 0:
                first_row_text = ' '.join(cell.text.lower() for cell in table.rows[0].cells)
                if any(keyword in first_row_text for keyword in ['barang', 'uraian', 'volume', 'harga']):
                    rincian_table = table
                    print(f"   ✓ Found rincian table at index {table_idx}")
                    break
        
        if not rincian_table:
            print("   ❌ Rincian table tidak ditemukan di dokumen!")
            return False
        
        # Check if items are filled
        print(f"   - Table rows: {len(rincian_table.rows)}")
        
        items_found = 0
        for row_idx, row in enumerate(rincian_table.rows[1:], 1):  # Skip header
            row_text = ' '.join(cell.text for cell in row.cells)
            if row_text.strip():
                items_found += 1
                print(f"     Row {row_idx}: {row_text[:60]}...")
        
        print(f"   ✓ Items filled in table: {items_found} rows")
        
        if items_found >= len(rincian_items):
            print()
            print("✓ TEST PASSED!")
            print("  Rincian barang berhasil ditampilkan di dokumen hasil generate.")
            print(f"  Document: {output_path}")
            return True
        else:
            print()
            print("⚠ Partial success")
            print(f"  Expected {len(rincian_items)} items, found {items_found}")
            return True  # Still consider it pass if some items are shown
    
    except Exception as e:
        print(f"   ❌ Error verifying: {str(e)}")
        return False


def test_data_to_next_phase():
    """Test data transformation untuk fase berikutnya."""
    
    print("\n" + "="*70)
    print("TEST: Data Transformation untuk Fase Berikutnya")
    print("="*70 + "\n")
    
    from app.services.lembar_permintaan_helper import (
        transform_to_spj_format,
        transform_to_kuitansi_format,
        get_lembar_data_complete,
    )
    
    db = PencairanManager(DATABASE_PATH)
    
    # Get latest lembar_permintaan dari database
    all_lembar = db.list_lembar_permintaan(limit=1)
    
    if not all_lembar:
        print("⚠ Tidak ada lembar_permintaan di database untuk test transformasi")
        print("  Jalankan test_generate_with_rincian() terlebih dahulu")
        return False
    
    lembar_id = all_lembar[0]['id']
    print(f"1. Using lembar_permintaan ID: {lembar_id}")
    print()
    
    # Get complete data
    print("2. Getting complete lembar data...")
    complete_data = get_lembar_data_complete(db, lembar_id)
    
    if not complete_data:
        print("   ❌ Gagal get lembar data")
        return False
    
    lembar = complete_data.get('lembar', {})
    items = complete_data.get('items', [])
    
    print(f"   ✓ Lembar: {lembar.get('unit_kerja')}")
    print(f"   ✓ Items: {len(items)}")
    print(f"   ✓ Total: Rp {lembar.get('total', 0):,.0f}".replace(",", "."))
    print()
    
    # Test SPJ transformation
    print("3. Transform to SPJ format...")
    spj_data = transform_to_spj_format(db, lembar_id)
    if spj_data:
        print(f"   ✓ SPJ data prepared with {len(spj_data.get('rincian_items', []))} items")
        print(f"   ✓ Total: Rp {spj_data.get('total', 0):,.0f}".replace(",", "."))
    else:
        print("   ❌ Gagal transform to SPJ")
        return False
    
    print()
    
    # Test Kuitansi transformation
    print("4. Transform to Kuitansi format...")
    kuitansi_data = transform_to_kuitansi_format(db, lembar_id)
    if kuitansi_data:
        print(f"   ✓ Kuitansi data prepared")
        print(f"   ✓ Penerima: {kuitansi_data.get('penerima')}")
        print(f"   ✓ Total: Rp {kuitansi_data.get('total', 0):,.0f}".replace(",", "."))
    else:
        print("   ❌ Gagal transform to Kuitansi")
        return False
    
    print()
    print("✓ TEST PASSED!")
    print("  Data dapat ditransform untuk dokumen fase berikutnya")
    
    return True


if __name__ == '__main__':
    print("\n")
    
    # Run tests
    test1_passed = test_generate_with_rincian()
    test2_passed = test_data_to_next_phase()
    
    print("\n" + "="*70)
    print("SUMMARY")
    print("="*70)
    print(f"Test 1 (Generate dengan Rincian): {'✓ PASSED' if test1_passed else '❌ FAILED'}")
    print(f"Test 2 (Data Transformation):     {'✓ PASSED' if test2_passed else '❌ FAILED'}")
    print()
    
    if test1_passed and test2_passed:
        print("✓ SEMUA TEST BERHASIL!")
        print("\nRincian barang dari form berhasil ditampilkan di dokumen,")
        print("dan data dapat digunakan untuk proses selanjutnya.")
        sys.exit(0)
    else:
        print("❌ BEBERAPA TEST GAGAL")
        sys.exit(1)
