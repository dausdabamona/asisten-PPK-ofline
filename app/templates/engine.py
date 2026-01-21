"""
PPK DOCUMENT FACTORY v3.0 - Template Engine
============================================
Template engine for Word (.docx) and Excel (.xlsx) merge
"""

import os
import re
import shutil
from datetime import datetime, date
from typing import Dict, List, Optional, Any, Tuple
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet

from app.core.config import (
    WORD_TEMPLATES_DIR, EXCEL_TEMPLATES_DIR, BACKUP_TEMPLATES_DIR,
    OUTPUT_DIR, TAHUN_ANGGARAN, ALL_PLACEHOLDERS, BULAN_INDONESIA
)
from app.core.database import get_db_manager


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def format_rupiah(nilai: float, with_prefix: bool = True) -> str:
    """Format angka ke format Rupiah"""
    if nilai is None:
        nilai = 0
    formatted = f"{int(nilai):,}".replace(",", ".")
    return f"Rp {formatted}" if with_prefix else formatted


def format_angka(nilai: float) -> str:
    """Format angka dengan pemisah ribuan"""
    if nilai is None:
        nilai = 0
    return f"{int(nilai):,}".replace(",", ".")


def terbilang(n: float) -> str:
    """Convert number to Indonesian words"""
    if n == 0:
        return "nol rupiah"
    
    satuan = ["", "satu", "dua", "tiga", "empat", "lima", 
              "enam", "tujuh", "delapan", "sembilan", "sepuluh", "sebelas"]
    
    def convert(num: int) -> str:
        num = int(num)
        if num < 0:
            return "minus " + convert(-num)
        elif num < 12:
            return satuan[num]
        elif num < 20:
            return satuan[num - 10] + " belas"
        elif num < 100:
            return satuan[num // 10] + " puluh " + satuan[num % 10]
        elif num < 200:
            return "seratus " + convert(num - 100)
        elif num < 1000:
            return satuan[num // 100] + " ratus " + convert(num % 100)
        elif num < 2000:
            return "seribu " + convert(num - 1000)
        elif num < 1000000:
            return convert(num // 1000) + " ribu " + convert(num % 1000)
        elif num < 1000000000:
            return convert(num // 1000000) + " juta " + convert(num % 1000000)
        elif num < 1000000000000:
            return convert(num // 1000000000) + " miliar " + convert(num % 1000000000)
        else:
            return convert(num // 1000000000000) + " triliun " + convert(num % 1000000000000)
    
    result = convert(n).strip()
    while "  " in result:
        result = result.replace("  ", " ")
    return result.strip() + " rupiah"


def format_tanggal(tgl: Any, fmt: str = "long") -> str:
    """Format date to Indonesian format"""
    if tgl is None:
        return ""
    
    if isinstance(tgl, str):
        try:
            if "-" in tgl:
                tgl = datetime.strptime(tgl, "%Y-%m-%d").date()
            elif "/" in tgl:
                tgl = datetime.strptime(tgl, "%d/%m/%Y").date()
            else:
                return tgl
        except ValueError:
            return tgl
    
    if isinstance(tgl, datetime):
        tgl = tgl.date()
    
    if fmt == "long":
        return f"{tgl.day} {BULAN_INDONESIA[tgl.month]} {tgl.year}"
    elif fmt == "short":
        return f"{tgl.day:02d}/{tgl.month:02d}/{tgl.year}"
    elif fmt == "full":
        hari = ['Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu', 'Minggu']
        return f"{hari[tgl.weekday()]}, {tgl.day} {BULAN_INDONESIA[tgl.month]} {tgl.year}"
    return f"{tgl.day} {BULAN_INDONESIA[tgl.month]} {tgl.year}"


def format_nip(nip: str) -> str:
    """Format NIP dengan spasi"""
    if not nip:
        return ""
    nip = nip.replace(" ", "")
    if len(nip) == 18:
        return f"{nip[:8]} {nip[8:14]} {nip[14]} {nip[15:]}"
    return nip


def format_npwp(npwp: str) -> str:
    """Format NPWP"""
    if not npwp:
        return ""
    npwp = npwp.replace(".", "").replace("-", "")
    if len(npwp) == 15:
        return f"{npwp[:2]}.{npwp[2:5]}.{npwp[5:8]}.{npwp[8]}-{npwp[9:12]}.{npwp[12:]}"
    return npwp


# ============================================================================
# PLACEHOLDER PATTERN
# ============================================================================

# Match {{placeholder}} or {{placeholder:format}}
PLACEHOLDER_PATTERN = re.compile(r'\{\{([a-zA-Z_][a-zA-Z0-9_]*)(:[a-zA-Z_]+)?\}\}')


class TemplateEngine:
    """
    Template Engine for merging data with Word and Excel templates
    
    Supports:
    - Simple placeholders: {{nama_paket}}
    - Formatted placeholders: {{nilai_kontrak:rupiah}}, {{tanggal_spk:long}}
    - Auto-calculated fields
    - Table row duplication for items
    """
    
    def __init__(self):
        self.db = get_db_manager()
        self.formatters = {
            'rupiah': format_rupiah,
            'angka': format_angka,
            'terbilang': terbilang,
            'tanggal': format_tanggal,
            'tanggal_long': lambda x: format_tanggal(x, 'long'),
            'tanggal_short': lambda x: format_tanggal(x, 'short'),
            'tanggal_full': lambda x: format_tanggal(x, 'full'),
            'nip': format_nip,
            'npwp': format_npwp,
            'upper': lambda x: str(x).upper() if x else '',
            'lower': lambda x: str(x).lower() if x else '',
            'title': lambda x: str(x).title() if x else '',
        }
    
    # =========================================================================
    # DATA PREPARATION
    # =========================================================================
    
    def prepare_data(self, paket_id: int, doc_type: str) -> Dict:
        """
        Prepare complete data dictionary from paket for template merge
        
        Args:
            paket_id: ID of paket
            doc_type: Document type for specific data preparation
        
        Returns:
            Dictionary with all placeholder values
        """
        paket = self.db.get_paket(paket_id)
        if not paket:
            raise ValueError(f"Paket {paket_id} tidak ditemukan")
        
        satker = self.db.get_satker()
        
        # Build complete data dictionary
        data = {}
        
        # Paket data
        data['nama_paket'] = paket.get('nama', '')
        data['kode_paket'] = paket.get('kode', '')
        data['lokasi_pekerjaan'] = paket.get('lokasi', '')
        data['jenis_pengadaan'] = paket.get('jenis_pengadaan', '')
        data['metode_pengadaan'] = paket.get('metode_pengadaan', '')
        data['sumber_dana'] = paket.get('sumber_dana', '')
        data['kode_akun'] = paket.get('kode_akun', '')
        data['tahun_anggaran'] = paket.get('tahun_anggaran', TAHUN_ANGGARAN)
        
        # Nilai
        data['nilai_pagu'] = paket.get('nilai_pagu', 0)
        data['nilai_hps'] = paket.get('nilai_hps', 0)
        data['nilai_hps_terbilang'] = terbilang(paket.get('nilai_hps', 0))
        data['nilai_kontrak'] = paket.get('nilai_kontrak', 0)
        data['nilai_kontrak_terbilang'] = terbilang(paket.get('nilai_kontrak', 0))
        
        # Calculate taxes
        nilai_kontrak = paket.get('nilai_kontrak', 0)
        is_pkp = paket.get('penyedia_is_pkp', True)
        
        data['nilai_ppn'] = round(nilai_kontrak * 0.11) if is_pkp else 0
        data['nilai_bruto'] = nilai_kontrak + data['nilai_ppn']
        
        tarif_pph = paket.get('tarif_pph', 0.015)
        data['nilai_pph'] = round(nilai_kontrak * tarif_pph)
        data['jenis_pph'] = paket.get('jenis_pph', 'PPh 22')
        data['tarif_pph'] = tarif_pph * 100  # As percentage
        
        data['nilai_bersih'] = data['nilai_bruto'] - data['nilai_pph']
        data['nilai_bersih_terbilang'] = terbilang(data['nilai_bersih'])
        
        # Waktu
        data['tanggal_mulai'] = paket.get('tanggal_mulai', '')
        data['tanggal_selesai'] = paket.get('tanggal_selesai', '')
        data['jangka_waktu'] = paket.get('jangka_waktu', 30)
        
        # Get document dates/numbers from existing documents
        docs = self.db.get_documents(paket_id)
        for doc in docs:
            dtype = doc['doc_type'].lower()
            data[f'tanggal_{dtype}'] = doc.get('tanggal', '')
            data[f'nomor_{dtype}'] = doc.get('nomor', '')
        
        # Satker
        data['satker_kode'] = satker.get('kode', '')
        data['satker_nama'] = satker.get('nama', '')
        data['satker_alamat'] = satker.get('alamat', '')
        data['satker_kota'] = satker.get('kota', '')
        data['satker_provinsi'] = satker.get('provinsi', '')
        data['satker_telepon'] = satker.get('telepon', '')
        data['satker_email'] = satker.get('email', '')
        data['kementerian'] = satker.get('kementerian', '')
        data['eselon1'] = satker.get('eselon1', '')
        
        # PPK
        data['ppk_nama'] = paket.get('ppk_nama', '')
        data['ppk_nip'] = paket.get('ppk_nip', '')
        data['ppk_jabatan'] = paket.get('ppk_jabatan', 'Pejabat Pembuat Komitmen')
        
        # PPSPM
        data['ppspm_nama'] = paket.get('ppspm_nama', '')
        data['ppspm_nip'] = paket.get('ppspm_nip', '')
        
        # Bendahara
        data['bendahara_nama'] = paket.get('bendahara_nama', '')
        data['bendahara_nip'] = paket.get('bendahara_nip', '')
        
        # Penyedia
        data['penyedia_nama'] = paket.get('penyedia_nama', '')
        data['penyedia_alamat'] = paket.get('penyedia_alamat', '')
        data['penyedia_npwp'] = paket.get('penyedia_npwp', '')
        data['penyedia_rekening'] = paket.get('penyedia_rekening', '')
        data['penyedia_bank'] = paket.get('penyedia_bank', '')
        data['penyedia_is_pkp'] = paket.get('penyedia_is_pkp', True)
        data['direktur_nama'] = paket.get('direktur_nama', '')
        data['direktur_jabatan'] = 'Direktur'
        
        # Additional data from JSON
        if paket.get('additional_data'):
            data.update(paket['additional_data'])
        
        # =====================================================================
        # ITEM BARANG / Bill of Quantity
        # =====================================================================
        item_summary = self.db.get_item_barang_summary(paket_id)
        items = item_summary['items']
        
        data['item_barang'] = items
        data['item_count'] = item_summary['count']
        data['subtotal_item'] = item_summary['subtotal']
        data['ppn_item'] = item_summary['ppn']
        data['pph_item'] = item_summary['pph']
        data['grand_total_item'] = item_summary['grand_total']
        data['nilai_bersih_item'] = item_summary['nilai_bersih']
        
        # Pre-formatted items for loops
        data['items_formatted'] = []
        for i, item in enumerate(items, 1):
            data['items_formatted'].append({
                'no': i,
                'nomor_urut': item['nomor_urut'],
                'kategori': item.get('kategori', ''),
                'kelompok': item.get('kelompok', ''),
                'uraian': item.get('uraian', ''),
                'nama_item': item.get('uraian', ''),  # Alias for uraian
                'nama_barang': item.get('uraian', ''),  # Alias for uraian
                'spesifikasi': item.get('spesifikasi', ''),
                'satuan': item.get('satuan', ''),
                'volume': item.get('volume', 0),
                'volume_fmt': f"{item.get('volume', 0):,.2f}".replace(",", "."),
                # Harga dasar
                'harga_dasar': item.get('harga_dasar', 0),
                'harga_dasar_fmt': format_rupiah(item.get('harga_dasar', 0)),
                # Survey prices
                'harga_survey1': item.get('harga_survey1', 0) or 0,
                'harga_survey1_fmt': format_rupiah(item.get('harga_survey1', 0) or 0),
                'harga_survey2': item.get('harga_survey2', 0) or 0,
                'harga_survey2_fmt': format_rupiah(item.get('harga_survey2', 0) or 0),
                'harga_survey3': item.get('harga_survey3', 0) or 0,
                'harga_survey3_fmt': format_rupiah(item.get('harga_survey3', 0) or 0),
                'harga_rata': item.get('harga_rata', 0) or 0,
                'harga_rata_fmt': format_rupiah(item.get('harga_rata', 0) or 0),
                # HPS prices
                'harga_hps_satuan': item.get('harga_hps_satuan', 0) or 0,
                'harga_hps_satuan_fmt': format_rupiah(item.get('harga_hps_satuan', 0) or 0),
                'total_hps': item.get('total_hps', 0) or 0,
                'total_hps_fmt': format_rupiah(item.get('total_hps', 0) or 0),
                # Kontrak prices
                'harga_kontrak_satuan': item.get('harga_kontrak_satuan', 0) or 0,
                'harga_kontrak_satuan_fmt': format_rupiah(item.get('harga_kontrak_satuan', 0) or 0),
                'total_kontrak': item.get('total_kontrak', 0) or 0,
                'total_kontrak_fmt': format_rupiah(item.get('total_kontrak', 0) or 0),
                # Selisih
                'selisih_harga': item.get('selisih_harga', 0) or 0,
                'selisih_harga_fmt': format_rupiah(item.get('selisih_harga', 0) or 0),
                'selisih_total': item.get('selisih_total', 0) or 0,
                'selisih_total_fmt': format_rupiah(item.get('selisih_total', 0) or 0),
                # Jumlah (alias for total calculations)
                'jumlah': item.get('total', 0) or (item.get('volume', 0) * item.get('harga_rata', 0)),
                'jumlah_fmt': format_rupiah(item.get('total', 0) or (item.get('volume', 0) * item.get('harga_rata', 0))),
                # Other
                'total': item.get('total', 0),
                'total_fmt': format_rupiah(item.get('total', 0)),
                'keterangan': item.get('keterangan', '')
            })
        
        # Summary formatted
        data['subtotal_item_fmt'] = format_rupiah(item_summary['subtotal'])
        data['ppn_item_fmt'] = format_rupiah(item_summary['ppn'])
        data['pph_item_fmt'] = format_rupiah(item_summary['pph'])
        data['grand_total_item_fmt'] = format_rupiah(item_summary['grand_total'])
        data['nilai_bersih_item_fmt'] = format_rupiah(item_summary['nilai_bersih'])
        data['grand_total_item_terbilang'] = terbilang(item_summary['grand_total'])
        data['nilai_bersih_item_terbilang'] = terbilang(item_summary['nilai_bersih'])
        
        # =====================================================================
        # SURVEY TOKO / SUMBER HARGA
        # =====================================================================
        survey_tokos = self.db.get_survey_toko(paket_id)
        data['survey_toko'] = survey_tokos
        
        # Individual toko for direct access (survey1, survey2, survey3)
        for i in range(1, 4):
            prefix = f'survey{i}'
            if i <= len(survey_tokos):
                toko = survey_tokos[i-1]
                data[f'{prefix}_toko'] = toko.get('nama_toko', '')
                data[f'{prefix}_nama'] = toko.get('nama_toko', '')
                data[f'{prefix}_alamat'] = toko.get('alamat', '')
                data[f'{prefix}_kota'] = toko.get('kota', '')
                data[f'{prefix}_telepon'] = toko.get('telepon', '')
                data[f'{prefix}_email'] = toko.get('email', '')
                data[f'{prefix}_website'] = toko.get('website', '')
                data[f'{prefix}_jenis'] = toko.get('jenis_survey', '')
                data[f'{prefix}_tanggal'] = toko.get('tanggal_survey', '')
                data[f'{prefix}_surveyor'] = toko.get('nama_surveyor', '')
                data[f'{prefix}_nip_surveyor'] = toko.get('nip_surveyor', '')
                data[f'{prefix}_platform'] = toko.get('platform_online', '')
                data[f'{prefix}_link'] = toko.get('link_produk', '')
                data[f'{prefix}_kontrak_nomor'] = toko.get('nomor_kontrak_referensi', '')
                data[f'{prefix}_kontrak_tahun'] = toko.get('tahun_kontrak_referensi', '')
                data[f'{prefix}_instansi'] = toko.get('instansi_referensi', '')
                data[f'{prefix}_keterangan'] = toko.get('keterangan', '')
                
                # Full address format
                alamat_parts = [toko.get('alamat', ''), toko.get('kota', '')]
                data[f'{prefix}_alamat_lengkap'] = ', '.join([p for p in alamat_parts if p])
            else:
                # Empty placeholders
                data[f'{prefix}_toko'] = ''
                data[f'{prefix}_nama'] = ''
                data[f'{prefix}_alamat'] = ''
                data[f'{prefix}_kota'] = ''
                data[f'{prefix}_telepon'] = ''
                data[f'{prefix}_email'] = ''
                data[f'{prefix}_website'] = ''
                data[f'{prefix}_jenis'] = ''
                data[f'{prefix}_tanggal'] = ''
                data[f'{prefix}_surveyor'] = ''
                data[f'{prefix}_nip_surveyor'] = ''
                data[f'{prefix}_platform'] = ''
                data[f'{prefix}_link'] = ''
                data[f'{prefix}_kontrak_nomor'] = ''
                data[f'{prefix}_kontrak_tahun'] = ''
                data[f'{prefix}_instansi'] = ''
                data[f'{prefix}_keterangan'] = ''
                data[f'{prefix}_alamat_lengkap'] = ''
        
        # =====================================================================
        # TIM PEMERIKSA
        # =====================================================================
        tim_pemeriksa = self.db.get_tim_pemeriksa(paket_id)
        data['tim_pemeriksa'] = tim_pemeriksa
        
        # Individual pemeriksa for direct access
        for i, member in enumerate(tim_pemeriksa, 1):
            data[f'pemeriksa{i}_nama'] = member.get('nama', '')
            data[f'pemeriksa{i}_nip'] = member.get('nip', '')
            data[f'pemeriksa{i}_jabatan'] = member.get('jabatan_tim', '')
        
        # =====================================================================
        # PEJABAT PENGADAAN & DAFTAR PEJABAT (from paket_pejabat table)
        # =====================================================================
        try:
            from app.core.database_v4 import get_db_manager_v4
            db_v4 = get_db_manager_v4()
            paket_pejabat = db_v4.get_paket_pejabat(paket_id)
            
            # Build daftar_pejabat list for loops
            daftar_pejabat = []
            
            for pp in paket_pejabat:
                peran = pp.get('peran', '')
                nama = pp.get('nama', '')
                nip = pp.get('nip', '')
                jabatan = pp.get('jabatan', '')
                pangkat = pp.get('pangkat', '')
                golongan = pp.get('golongan', '')
                
                # Format nama dengan gelar
                nama_lengkap = nama
                if pp.get('gelar_depan'):
                    nama_lengkap = f"{pp['gelar_depan']} {nama_lengkap}"
                if pp.get('gelar_belakang'):
                    nama_lengkap = f"{nama_lengkap}, {pp['gelar_belakang']}"
                
                # Map to specific role placeholders
                if peran == 'PPK':
                    data['ppk_nama'] = nama_lengkap
                    data['ppk_nip'] = nip
                    data['ppk_jabatan'] = jabatan or 'Pejabat Pembuat Komitmen'
                    data['ppk_pangkat'] = pangkat
                    data['ppk_golongan'] = golongan
                elif peran == 'PEJABAT_PENGADAAN':
                    data['pejabat_pengadaan_nama'] = nama_lengkap
                    data['pejabat_pengadaan_nip'] = nip
                    data['pejabat_pengadaan_jabatan'] = jabatan or 'Pejabat Pengadaan'
                    data['pejabat_pengadaan_pangkat'] = pangkat
                    data['pejabat_pengadaan_golongan'] = golongan
                    data['pp_nama'] = nama_lengkap  # Alias
                    data['pp_nip'] = nip
                elif peran == 'PPSPM':
                    data['ppspm_nama'] = nama_lengkap
                    data['ppspm_nip'] = nip
                    data['ppspm_jabatan'] = jabatan or 'Pejabat Penandatangan SPM'
                    data['ppspm_pangkat'] = pangkat
                    data['ppspm_golongan'] = golongan
                elif peran == 'BENDAHARA':
                    data['bendahara_nama'] = nama_lengkap
                    data['bendahara_nip'] = nip
                    data['bendahara_jabatan'] = jabatan or 'Bendahara Pengeluaran'
                    data['bendahara_pangkat'] = pangkat
                    data['bendahara_golongan'] = golongan
                elif peran == 'KETUA_PPHP':
                    data['ketua_pphp_nama'] = nama_lengkap
                    data['ketua_pphp_nip'] = nip
                    data['ketua_pphp_jabatan'] = jabatan
                    data['ketua_pphp_pangkat'] = pangkat
                    data['ketua_pphp_golongan'] = golongan
                    # Also set as pemeriksa1 (Ketua)
                    data['pemeriksa1_nama'] = nama_lengkap
                    data['pemeriksa1_nip'] = nip
                    data['pemeriksa1_jabatan'] = 'Ketua'
                
                # Add to daftar_pejabat
                daftar_pejabat.append({
                    'peran': peran,
                    'nama': nama_lengkap,
                    'nip': nip,
                    'jabatan': jabatan,
                    'pangkat': pangkat,
                    'golongan': golongan,
                })
            
            # Get ANGGOTA_PPHP separately for tim_pemeriksa loop
            anggota_pphp = db_v4.get_paket_pejabat_by_role(paket_id, 'ANGGOTA_PPHP')
            for i, anggota in enumerate(anggota_pphp, 2):  # Start from 2 (Ketua is 1)
                nama_lengkap = anggota.get('nama', '')
                if anggota.get('gelar_depan'):
                    nama_lengkap = f"{anggota['gelar_depan']} {nama_lengkap}"
                if anggota.get('gelar_belakang'):
                    nama_lengkap = f"{nama_lengkap}, {anggota['gelar_belakang']}"
                
                data[f'pemeriksa{i}_nama'] = nama_lengkap
                data[f'pemeriksa{i}_nip'] = anggota.get('nip', '')
                data[f'pemeriksa{i}_jabatan'] = 'Anggota'
            
            data['daftar_pejabat'] = daftar_pejabat
            
        except Exception as e:
            # Fallback if v4 module not available
            data['pejabat_pengadaan_nama'] = ''
            data['pejabat_pengadaan_nip'] = ''
            data['pejabat_pengadaan_jabatan'] = ''
            data['pp_nama'] = ''
            data['pp_nip'] = ''
            data['daftar_pejabat'] = []
        
        # =====================================================================
        # HARGA KONTRAK FINAL (from item_barang)
        # =====================================================================
        harga_kontrak_items = []
        selisih_total = 0
        for item in items:
            harga_hps = item.get('harga_hps_satuan') or item.get('harga_dasar', 0)
            harga_kontrak = item.get('harga_kontrak_satuan', 0) or harga_hps
            selisih = harga_hps - harga_kontrak
            
            harga_kontrak_items.append({
                'no': item.get('nomor_urut', 0),
                'uraian': item.get('uraian', ''),
                'volume': item.get('volume', 0),
                'satuan': item.get('satuan', ''),
                'harga_hps': harga_hps,
                'harga_hps_fmt': format_rupiah(harga_hps),
                'harga_kontrak': harga_kontrak,
                'harga_kontrak_fmt': format_rupiah(harga_kontrak),
                'selisih': selisih,
                'selisih_fmt': format_rupiah(selisih),
                'total_hps': item.get('total_hps', 0) or item.get('total', 0),
                'total_hps_fmt': format_rupiah(item.get('total_hps', 0) or item.get('total', 0)),
                'total_kontrak': item.get('total_kontrak', 0) or item.get('total', 0),
                'total_kontrak_fmt': format_rupiah(item.get('total_kontrak', 0) or item.get('total', 0)),
            })
            selisih_total += selisih * item.get('volume', 0)
        
        data['harga_kontrak_items'] = harga_kontrak_items
        data['selisih_total'] = selisih_total
        data['selisih_total_fmt'] = format_rupiah(selisih_total)
        
        # =====================================================================
        # DOCUMENT TIMELINE (Nomor & Tanggal Dokumen)
        # =====================================================================
        from app.ui.timeline_manager import format_tanggal_indonesia, DOKUMEN_TIMELINE
        
        timeline_entries = self.db.get_dokumen_timeline(paket_id)
        timeline_dict = {t['doc_type']: t for t in timeline_entries}
        
        # Create placeholders for each document type
        for doc_config in DOKUMEN_TIMELINE:
            doc_code = doc_config['code']
            doc_code_lower = doc_code.lower()
            
            entry = timeline_dict.get(doc_code, {})
            
            # Nomor dokumen
            nomor = entry.get('nomor_dokumen', '')
            data[f'nomor_{doc_code_lower}'] = nomor
            data[f'{doc_code_lower}_nomor'] = nomor
            
            # Tanggal dokumen (various formats)
            tanggal_raw = entry.get('tanggal_dokumen', '')
            data[f'tanggal_{doc_code_lower}'] = format_tanggal_indonesia(tanggal_raw, 'long') if tanggal_raw else ''
            data[f'tanggal_{doc_code_lower}_full'] = format_tanggal_indonesia(tanggal_raw, 'full') if tanggal_raw else ''
            data[f'tanggal_{doc_code_lower}_short'] = format_tanggal_indonesia(tanggal_raw, 'short') if tanggal_raw else ''
            data[f'{doc_code_lower}_tanggal'] = format_tanggal_indonesia(tanggal_raw, 'long') if tanggal_raw else ''
        
        # Also format existing date fields in Indonesian
        if paket.get('tanggal_mulai'):
            data['tanggal_mulai_indo'] = format_tanggal_indonesia(paket['tanggal_mulai'], 'long')
            data['tanggal_mulai_full'] = format_tanggal_indonesia(paket['tanggal_mulai'], 'full')
        
        if paket.get('tanggal_selesai'):
            data['tanggal_selesai_indo'] = format_tanggal_indonesia(paket['tanggal_selesai'], 'long')
            data['tanggal_selesai_full'] = format_tanggal_indonesia(paket['tanggal_selesai'], 'full')
        
        # Current date in Indonesian
        from datetime import date
        today = date.today()
        data['tanggal_hari_ini'] = format_tanggal_indonesia(today, 'long')
        data['tanggal_hari_ini_full'] = format_tanggal_indonesia(today, 'full')
        
        return data
    
    def _process_table_rows(self, doc: Document, data: Dict):
        """
        Process table rows with item loops
        Looks for rows containing {{item.xxx}} placeholders and duplicates them
        Optimized for up to 500 items
        
        FIXED: Handle split runs in cells (Word often splits placeholders across runs)
        """
        from copy import deepcopy
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        for table in doc.tables:
            rows_to_process = []
            
            # Find rows with item placeholders
            for row_idx, row in enumerate(table.rows):
                row_text = ''
                for cell in row.cells:
                    # Get full cell text by joining all runs
                    cell_text = self._get_cell_full_text(cell)
                    row_text += cell_text
                
                # Check if row contains item placeholder
                if '{{item.' in row_text or '{{no}}' in row_text:
                    rows_to_process.append(row_idx)
            
            # Process template rows (only process first one found to avoid duplicates)
            if not rows_to_process:
                continue
            
            template_row_idx = rows_to_process[0]
            items = data.get('items_formatted', [])
            if not items:
                continue
            
            template_row = table.rows[template_row_idx]
            
            # Get cell templates (merge all runs to get complete text)
            cell_templates = []
            for cell in template_row.cells:
                cell_text = self._get_cell_full_text(cell)
                cell_templates.append(cell_text)
            
            # Store template row XML
            template_tr = template_row._tr
            tbl = template_tr.getparent()
            
            # Fill first item into existing template row
            first_item = items[0]
            for cell_idx, cell in enumerate(template_row.cells):
                if cell_idx < len(cell_templates):
                    template_text = cell_templates[cell_idx]
                    result = self._replace_item_placeholders(template_text, first_item)
                    # Set cell text safely
                    self._set_cell_text(cell, result)
            
            # Insert new rows for remaining items (from index 1 onwards)
            # Insert after the template row
            insert_position = list(tbl).index(template_tr) + 1
            
            for i, item in enumerate(items[1:], start=1):
                # Deep copy the template row
                new_tr = deepcopy(template_tr)
                
                # Insert at position
                tbl.insert(insert_position + i - 1, new_tr)
            
            # Now fill all the newly created rows
            # Re-get all rows after modification
            all_rows = table.rows
            
            for i, item in enumerate(items[1:], start=1):
                target_row_idx = template_row_idx + i
                if target_row_idx < len(all_rows):
                    target_row = all_rows[target_row_idx]
                    for cell_idx, cell in enumerate(target_row.cells):
                        if cell_idx < len(cell_templates):
                            template_text = cell_templates[cell_idx]
                            result = self._replace_item_placeholders(template_text, item)
                            # Set cell text safely
                            self._set_cell_text(cell, result)
    
    def _get_cell_full_text(self, cell) -> str:
        """
        Get full cell text by joining all paragraphs and runs.
        This handles cases where Word splits placeholders across multiple runs.
        """
        full_text = ''
        for para in cell.paragraphs:
            for run in para.runs:
                full_text += run.text
            full_text += '\n'
        return full_text.strip()
    
    def _set_cell_text(self, cell, text: str):
        """
        Safely set cell text, clearing all existing runs first.
        Preserves first paragraph/run formatting if available.
        """
        # Clear all existing content
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ''
        
        # Set new text
        if cell.paragraphs:
            first_para = cell.paragraphs[0]
            if first_para.runs:
                first_para.runs[0].text = text
            else:
                # No runs exist, add the text directly
                first_para.add_run(text)
        else:
            # No paragraphs, use cell.text
            cell.text = text
    
    def _replace_item_placeholders(self, template_text: str, item: dict) -> str:
        """Replace item placeholders in template text"""
        result = template_text
        
        # Basic info
        result = result.replace('{{no}}', str(item.get('no', '')))
        result = result.replace('{{item.no}}', str(item.get('no', '')))
        result = result.replace('{{item.nomor_urut}}', str(item.get('nomor_urut', '')))
        result = result.replace('{{item.kategori}}', str(item.get('kategori', '')))
        result = result.replace('{{item.kelompok}}', str(item.get('kelompok', '')))
        result = result.replace('{{item.uraian}}', str(item.get('uraian', '')))
        result = result.replace('{{item.nama_item}}', str(item.get('uraian', '')))  # Alias
        result = result.replace('{{item.nama}}', str(item.get('uraian', '')))  # Alias
        result = result.replace('{{item.spesifikasi}}', str(item.get('spesifikasi', '')))
        result = result.replace('{{item.satuan}}', str(item.get('satuan', '')))
        result = result.replace('{{item.volume}}', str(item.get('volume_fmt', '')))
        result = result.replace('{{item.volume_angka}}', str(item.get('volume', '')))
        
        # Harga dasar
        result = result.replace('{{item.harga_dasar}}', str(item.get('harga_dasar_fmt', '')))
        result = result.replace('{{item.harga_satuan}}', str(item.get('harga_dasar_fmt', '')))
        
        # Survey prices
        result = result.replace('{{item.harga_survey1}}', str(item.get('harga_survey1_fmt', '')))
        result = result.replace('{{item.harga_survey2}}', str(item.get('harga_survey2_fmt', '')))
        result = result.replace('{{item.harga_survey3}}', str(item.get('harga_survey3_fmt', '')))
        result = result.replace('{{item.harga_rata}}', str(item.get('harga_rata_fmt', '')))
        
        # HPS prices
        result = result.replace('{{item.harga_hps_satuan}}', str(item.get('harga_hps_satuan_fmt', '')))
        result = result.replace('{{item.harga_hps}}', str(item.get('harga_hps_satuan_fmt', '')))  # Alias
        result = result.replace('{{item.total_hps}}', str(item.get('total_hps_fmt', '')))
        
        # Kontrak prices
        result = result.replace('{{item.harga_kontrak_satuan}}', str(item.get('harga_kontrak_satuan_fmt', '')))
        result = result.replace('{{item.harga_kontrak}}', str(item.get('harga_kontrak_satuan_fmt', '')))  # Alias
        result = result.replace('{{item.total_kontrak}}', str(item.get('total_kontrak_fmt', '')))
        
        # Selisih
        result = result.replace('{{item.selisih_harga}}', str(item.get('selisih_harga_fmt', '')))
        result = result.replace('{{item.selisih_total}}', str(item.get('selisih_total_fmt', '')))
        
        # Total & keterangan
        result = result.replace('{{item.total}}', str(item.get('total_fmt', '')))
        result = result.replace('{{item.jumlah}}', str(item.get('total_fmt', '')))
        result = result.replace('{{item.keterangan}}', str(item.get('keterangan', '')))
        
        return result
    
    def _process_excel_table_rows(self, ws, data: Dict, start_row: int = None):
        """
        Process Excel rows with item loops
        Finds rows containing {{item.xxx}} and duplicates them for each item
        Optimized for up to 500 items
        """
        items = data.get('items_formatted', [])
        if not items:
            return
        
        # Find template row (row with {{item. placeholders)
        template_row = None
        for row in ws.iter_rows(min_row=1, max_row=50):  # Search only first 50 rows
            for cell in row:
                if cell.value and '{{item.' in str(cell.value):
                    template_row = cell.row
                    break
            if template_row:
                break
        
        if not template_row:
            return
        
        # Get template row content and styles
        max_col = ws.max_column
        cell_templates = []
        cell_styles = []
        
        for col in range(1, max_col + 1):
            cell = ws.cell(row=template_row, column=col)
            cell_templates.append(cell.value)
            # Store style info
            cell_styles.append({
                'font': cell.font.copy() if cell.font else None,
                'border': cell.border.copy() if cell.border else None,
                'fill': cell.fill.copy() if cell.fill else None,
                'alignment': cell.alignment.copy() if cell.alignment else None,
                'number_format': cell.number_format
            })
        
        # Find rows after template that need to be moved down
        # (e.g., summary rows, totals)
        rows_after = []
        last_row = ws.max_row
        
        # Store data from rows after template row
        for row_idx in range(template_row + 1, min(last_row + 1, template_row + 20)):
            row_data = []
            row_styles = []
            has_data = False
            for col in range(1, max_col + 1):
                cell = ws.cell(row=row_idx, column=col)
                if cell.value:
                    has_data = True
                row_data.append(cell.value)
                row_styles.append({
                    'font': cell.font.copy() if cell.font else None,
                    'border': cell.border.copy() if cell.border else None,
                    'fill': cell.fill.copy() if cell.fill else None,
                    'alignment': cell.alignment.copy() if cell.alignment else None,
                    'number_format': cell.number_format
                })
            if has_data:
                rows_after.append((row_data, row_styles))
        
        # Calculate how many rows to insert
        num_items = len(items)
        rows_to_add = num_items - 1  # First item uses template row
        
        # Insert blank rows if needed (batch insert is faster)
        if rows_to_add > 0:
            ws.insert_rows(template_row + 1, amount=rows_to_add)
        
        # Fill in item data
        for i, item in enumerate(items):
            target_row = template_row + i
            
            for col, template_text in enumerate(cell_templates, 1):
                cell = ws.cell(row=target_row, column=col)
                
                # Apply style
                style = cell_styles[col - 1]
                if style['font']:
                    cell.font = style['font']
                if style['border']:
                    cell.border = style['border']
                if style['fill']:
                    cell.fill = style['fill']
                if style['alignment']:
                    cell.alignment = style['alignment']
                if style['number_format']:
                    cell.number_format = style['number_format']
                
                if template_text is None:
                    continue
                
                result = str(template_text)
                
                # Replace placeholders
                result = result.replace('{{no}}', str(item['no']))
                result = result.replace('{{item.nomor_urut}}', str(item['nomor_urut']))
                result = result.replace('{{item.kategori}}', item['kategori'])
                result = result.replace('{{item.kelompok}}', item['kelompok'])
                result = result.replace('{{item.uraian}}', item['uraian'])
                result = result.replace('{{item.spesifikasi}}', item['spesifikasi'])
                result = result.replace('{{item.satuan}}', item['satuan'])
                result = result.replace('{{item.volume}}', str(item['volume']))
                result = result.replace('{{item.harga_dasar}}', str(item['harga_dasar']))
                result = result.replace('{{item.harga_satuan}}', str(item['harga_dasar']))
                result = result.replace('{{item.harga_survey1}}', str(item.get('harga_survey1', 0) or 0))
                result = result.replace('{{item.harga_survey2}}', str(item.get('harga_survey2', 0) or 0))
                result = result.replace('{{item.harga_survey3}}', str(item.get('harga_survey3', 0) or 0))
                result = result.replace('{{item.harga_rata}}', str(item.get('harga_rata', 0) or 0))
                result = result.replace('{{item.total}}', str(item['total']))
                result = result.replace('{{item.jumlah}}', str(item['total']))
                result = result.replace('{{item.keterangan}}', item['keterangan'])
                
                # Try to convert to number if it looks numeric
                try:
                    # Check if result is purely numeric (no text mixed in)
                    clean_result = result.replace('.', '').replace(',', '').replace('-', '')
                    if clean_result.isdigit() and result.replace(',', '.').replace('.', '', 1).lstrip('-').isdigit():
                        result = float(result.replace(',', '.'))
                except:
                    pass
                
                cell.value = result
        
        # Restore rows after items (summary/total rows)
        if rows_after:
            summary_start_row = template_row + num_items
            for row_offset, (row_data, row_styles) in enumerate(rows_after):
                target_row = summary_start_row + row_offset
                for col, value in enumerate(row_data, 1):
                    cell = ws.cell(row=target_row, column=col)
                    cell.value = value
                    
                    # Apply style
                    style = row_styles[col - 1]
                    if style['font']:
                        cell.font = style['font']
                    if style['border']:
                        cell.border = style['border']
                    if style['fill']:
                        cell.fill = style['fill']
                    if style['alignment']:
                        cell.alignment = style['alignment']
                    if style['number_format']:
                        cell.number_format = style['number_format']
    
    def format_value(self, value: Any, format_type: str = None) -> str:
        """Format a value according to type"""
        if value is None:
            return ''
        
        if format_type and format_type in self.formatters:
            return self.formatters[format_type](value)
        
        # Auto-detect formatting based on value type
        if isinstance(value, (int, float)) and value > 1000:
            return format_angka(value)
        elif isinstance(value, (date, datetime)):
            return format_tanggal(value)
        
        return str(value)
    
    # =========================================================================
    # WORD TEMPLATE MERGE
    # =========================================================================
    
    def merge_word(self, template_path: str, data: Dict, output_path: str) -> str:
        """
        Merge data into Word template
        
        Args:
            template_path: Path to .docx template
            data: Dictionary of placeholder values
            output_path: Path to save output
        
        Returns:
            Path to generated document
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template tidak ditemukan: {template_path}")
        
        doc = Document(template_path)
        
        # Process table rows with item loops FIRST (before placeholder replacement)
        self._process_table_rows(doc, data)
        
        # Process all paragraphs
        for para in doc.paragraphs:
            self._replace_placeholders_in_paragraph(para, data)
        
        # Process all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(para, data)
        
        # Process headers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    self._replace_placeholders_in_paragraph(para, data)
                # Process tables in header
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                self._replace_placeholders_in_paragraph(para, data)
            if section.footer:
                for para in section.footer.paragraphs:
                    self._replace_placeholders_in_paragraph(para, data)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc.save(output_path)
        return output_path
    
    def _replace_placeholders_in_paragraph(self, para, data: Dict):
        """Replace placeholders in a paragraph while preserving formatting"""
        # Get full text by joining all runs (handles split placeholders)
        full_text = ''.join(run.text for run in para.runs)
        
        # Also check para.text in case runs is empty
        if not full_text:
            full_text = para.text
        
        if '{{' not in full_text:
            return
        
        # Find all placeholders
        matches = PLACEHOLDER_PATTERN.findall(full_text)
        
        for match in matches:
            placeholder = match[0]
            format_type = match[1][1:] if match[1] else None  # Remove leading ':'
            
            # Get value from data
            value = data.get(placeholder, '')
            
            # Format value
            formatted = self.format_value(value, format_type)
            
            # Build placeholder string to replace
            if format_type:
                placeholder_str = f'{{{{{placeholder}:{format_type}}}}}'
            else:
                placeholder_str = f'{{{{{placeholder}}}}}'
            
            # Replace in paragraph
            if placeholder_str in full_text:
                self._replace_in_runs(para, placeholder_str, formatted)
                full_text = full_text.replace(placeholder_str, formatted)
    
    def _replace_in_runs(self, para, old_text: str, new_text: str):
        """Replace text in runs while trying to preserve formatting"""
        # Simple approach: concatenate all runs, replace, then redistribute
        # This handles cases where placeholder is split across runs
        
        full_text = ''.join(run.text for run in para.runs)
        
        if old_text not in full_text:
            return
        
        new_full = full_text.replace(old_text, new_text)
        
        # Clear existing runs and set new text to first run
        if para.runs:
            # Keep formatting from first run
            first_run = para.runs[0]
            
            # Clear all runs
            for run in para.runs:
                run.text = ''
            
            # Set new text to first run
            first_run.text = new_full
    
    # =========================================================================
    # EXCEL TEMPLATE MERGE
    # =========================================================================
    
    def merge_excel(self, template_path: str, data: Dict, output_path: str,
                    sheet_name: str = None, items: List[Dict] = None) -> str:
        """
        Merge data into Excel template
        
        Args:
            template_path: Path to .xlsx template
            data: Dictionary of placeholder values
            output_path: Path to save output
            sheet_name: Specific sheet to process (None = all sheets)
            items: List of items for table rows (optional)
        
        Returns:
            Path to generated document
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template tidak ditemukan: {template_path}")
        
        wb = load_workbook(template_path)
        
        # Process sheets
        sheets_to_process = [wb[sheet_name]] if sheet_name else wb.worksheets
        
        for ws in sheets_to_process:
            self._process_excel_sheet(ws, data, items)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        wb.save(output_path)
        return output_path
    
    def _process_excel_sheet(self, ws: Worksheet, data: Dict, items: List[Dict] = None):
        """Process a single Excel sheet with proper item row insertion"""
        from openpyxl.utils import get_column_letter
        from copy import copy
        
        # First pass: find item template row
        item_start_row = None
        item_row_template = []
        
        for row_idx, row in enumerate(ws.iter_rows(), 1):
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    if '{{item.' in cell.value or '{{no}}' in cell.value:
                        if item_start_row is None:
                            item_start_row = row_idx
                            # Store entire row template with styles
                            item_row_template = []
                            for c in row:
                                item_row_template.append({
                                    'column': c.column,
                                    'value': c.value,
                                    'font': copy(c.font) if c.font else None,
                                    'border': copy(c.border) if c.border else None,
                                    'fill': copy(c.fill) if c.fill else None,
                                    'alignment': copy(c.alignment) if c.alignment else None,
                                    'number_format': c.number_format,
                                })
                        break
            if item_start_row:
                break
        
        # Insert item rows if we have items and found template
        items_to_use = items or data.get('items_formatted', [])
        
        if items_to_use and item_start_row and item_row_template:
            # Insert additional rows for items (keep first row as template was there)
            num_items = len(items_to_use)
            if num_items > 1:
                ws.insert_rows(item_start_row + 1, num_items - 1)
            
            # Fill all item rows
            for idx, item in enumerate(items_to_use):
                row_num = item_start_row + idx
                
                for cell_template in item_row_template:
                    col = cell_template['column']
                    value_template = cell_template['value'] or ''
                    
                    # Replace item placeholders
                    value = str(value_template) if value_template else ''
                    
                    # Replace {{no}} with item index (1-based)
                    if '{{no}}' in value:
                        value = value.replace('{{no}}', str(idx + 1))
                    
                    # Replace {{item.xxx}} placeholders
                    if '{{item.' in value:
                        import re
                        # Find all item placeholders
                        item_placeholders = re.findall(r'\{\{item\.([a-zA-Z0-9_]+)\}\}', value)
                        for key in item_placeholders:
                            placeholder = f'{{{{item.{key}}}}}'
                            if key in item:
                                val = item[key]
                                # Format value appropriately
                                if isinstance(val, (int, float)) and val != 0:
                                    # Check if it's a price field
                                    if 'harga' in key.lower() or 'total' in key.lower() or 'jumlah' in key.lower():
                                        formatted = f"{int(val):,}".replace(",", ".")
                                    else:
                                        formatted = str(val) if not isinstance(val, float) else f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                                else:
                                    formatted = str(val) if val else ''
                                value = value.replace(placeholder, formatted)
                            else:
                                # Key not found in item - replace with empty string
                                value = value.replace(placeholder, '')
                    
                    # Convert numeric strings to numbers for proper Excel handling
                    try:
                        if value and value.replace('.', '').replace(',', '').isdigit():
                            # It's a number
                            value = float(value.replace('.', '').replace(',', '.'))
                    except:
                        pass
                    
                    # Create cell with value
                    cell = ws.cell(row=row_num, column=col, value=value)
                    
                    # Copy styles from template
                    if cell_template['font']:
                        cell.font = copy(cell_template['font'])
                    if cell_template['border']:
                        cell.border = copy(cell_template['border'])
                    if cell_template['fill']:
                        cell.fill = copy(cell_template['fill'])
                    if cell_template['alignment']:
                        cell.alignment = copy(cell_template['alignment'])
                    if cell_template['number_format']:
                        cell.number_format = cell_template['number_format']
        
        # Second pass: replace regular placeholders (skip item rows)
        skip_rows = set(range(item_start_row, item_start_row + len(items_to_use))) if item_start_row and items_to_use else set()
        
        for row_idx, row in enumerate(ws.iter_rows(), 1):
            if row_idx in skip_rows:
                continue
            for cell in row:
                if cell.value and isinstance(cell.value, str) and '{{' in cell.value:
                    if '{{item.' not in cell.value and '{{no}}' not in cell.value:
                        cell.value = self._replace_excel_placeholders(cell.value, data)
    
    def _replace_excel_placeholders(self, text: str, data: Dict) -> Any:
        """Replace placeholders in Excel cell value"""
        if not text or '{{' not in text:
            return text
        
        matches = PLACEHOLDER_PATTERN.findall(text)
        result = text
        
        for match in matches:
            placeholder = match[0]
            format_type = match[1][1:] if match[1] else None
            
            value = data.get(placeholder, '')
            formatted = self.format_value(value, format_type)
            
            if format_type:
                placeholder_str = f'{{{{{placeholder}:{format_type}}}}}'
            else:
                placeholder_str = f'{{{{{placeholder}}}}}'
            
            result = result.replace(placeholder_str, formatted)
        
        # If entire cell is a single placeholder and value is numeric, return as number
        if text.startswith('{{') and text.endswith('}}') and result.replace('.', '').isdigit():
            return float(result.replace('.', '')) if '.' in result else int(result)
        
        return result
    
    def _insert_item_rows(self, ws: Worksheet, start_row: int, 
                          template: List[Tuple], items: List[Dict]):
        """Insert multiple rows for items"""
        # This is a simplified version - full implementation would handle
        # row insertion and formula copying
        
        for idx, item in enumerate(items):
            row_num = start_row + idx
            
            for col, value_template in template:
                if '{{item.' in str(value_template):
                    # Replace item placeholders
                    value = value_template
                    for key, val in item.items():
                        placeholder = f'{{{{item.{key}}}}}'
                        if placeholder in value:
                            formatted = self.format_value(val)
                            value = value.replace(placeholder, str(formatted))
                    
                    # Also replace {{item.no}} with row number
                    value = value.replace('{{item.no}}', str(idx + 1))
                    
                    ws.cell(row=row_num, column=col, value=value)
    
    # =========================================================================
    # DOCUMENT GENERATION
    # =========================================================================
    
    def generate_document(self, paket_id: int, doc_type: str,
                          additional_data: Dict = None) -> Tuple[str, str]:
        """
        Generate document from template
        
        Args:
            paket_id: ID of paket
            doc_type: Document type (SPK, SPMK, etc.)
            additional_data: Additional data to merge
        
        Returns:
            Tuple of (filepath, nomor_dokumen)
        """
        from app.core.config import DOCUMENT_TEMPLATES, NUMBERING_PREFIXES
        
        # Get template config
        template_config = DOCUMENT_TEMPLATES.get(doc_type)
        if not template_config:
            raise ValueError(f"Document type {doc_type} tidak dikenal")
        
        # Check if custom template exists in database
        db_template = self.db.get_template(doc_type)
        
        if db_template:
            template_path = db_template['filepath']
        else:
            # Use default template
            template_dir = (WORD_TEMPLATES_DIR if template_config['type'] == 'word' 
                          else EXCEL_TEMPLATES_DIR)
            template_path = os.path.join(template_dir, template_config['template'])
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(
                f"Template untuk {doc_type} tidak ditemukan: {template_path}\n"
                f"Silakan upload template melalui Template Manager."
            )
        
        # Prepare data
        data = self.prepare_data(paket_id, doc_type)
        
        # Add additional data
        if additional_data:
            data.update(additional_data)
        
        # Generate document number
        nomor = self.db.get_next_number(doc_type)
        data[f'nomor_{doc_type.lower()}'] = nomor
        
        # Set current date if not provided
        if f'tanggal_{doc_type.lower()}' not in data or not data[f'tanggal_{doc_type.lower()}']:
            data[f'tanggal_{doc_type.lower()}'] = date.today()
        
        # Build output path
        paket = self.db.get_paket(paket_id)
        paket_folder = f"{paket['kode']}_{paket['nama'][:30]}"
        paket_folder = re.sub(r'[<>:"/\\|?*]', '_', paket_folder)
        
        output_dir = os.path.join(
            OUTPUT_DIR, 
            str(TAHUN_ANGGARAN),
            paket_folder,
            doc_type
        )
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate filename
        nomor_safe = nomor.replace('/', '_')
        ext = 'docx' if template_config['type'] == 'word' else 'xlsx'
        filename = f"{doc_type}_{nomor_safe}.{ext}"
        output_path = os.path.join(output_dir, filename)
        
        # Merge template
        if template_config['type'] == 'word':
            self.merge_word(template_path, data, output_path)
        else:
            sheet_name = template_config.get('sheet')
            # Pass items_formatted for Excel item loops
            items = data.get('items_formatted', [])
            self.merge_excel(template_path, data, output_path, sheet_name, items)
        
        # Save document record
        self.db.save_document(paket_id, doc_type, {
            'nomor': nomor,
            'tanggal': data.get(f'tanggal_{doc_type.lower()}', date.today()),
            'filename': filename,
            'filepath': output_path,
            'template_used': template_path,
            'data': data
        })
        
        return output_path, nomor
    
    def generate_batch(self, paket_id: int, doc_types: List[str],
                      additional_data: Dict = None) -> List[Tuple[str, str, str]]:
        """
        Generate multiple documents at once
        
        Args:
            paket_id: ID of paket
            doc_types: List of document types to generate
            additional_data: Additional data for all documents
        
        Returns:
            List of tuples (doc_type, filepath, nomor)
        """
        results = []
        
        for doc_type in doc_types:
            try:
                filepath, nomor = self.generate_document(paket_id, doc_type, additional_data)
                results.append((doc_type, filepath, nomor))
            except Exception as e:
                results.append((doc_type, None, f"Error: {str(e)}"))
        
        return results


# ============================================================================
# TEMPLATE MANAGER
# ============================================================================

class TemplateManager:
    """
    Manager for template upload, versioning, and variable mapping
    """
    
    def __init__(self):
        self.db = get_db_manager()
    
    def upload_template(self, doc_type: str, source_path: str,
                       description: str = None, uploaded_by: str = None) -> Dict:
        """
        Upload new template with version control
        
        Args:
            doc_type: Document type (SPK, HPS, etc.)
            source_path: Path to template file
            description: Optional description
            uploaded_by: Name of uploader
        
        Returns:
            Dictionary with upload result
        """
        from app.core.config import DOCUMENT_TEMPLATES
        
        if not os.path.exists(source_path):
            raise FileNotFoundError(f"File tidak ditemukan: {source_path}")
        
        # Determine template type
        ext = os.path.splitext(source_path)[1].lower()
        if ext not in ['.docx', '.xlsx']:
            raise ValueError("Hanya file .docx dan .xlsx yang diizinkan")
        
        template_type = 'word' if ext == '.docx' else 'excel'
        
        # Get template config
        template_config = DOCUMENT_TEMPLATES.get(doc_type)
        if not template_config:
            raise ValueError(f"Document type {doc_type} tidak dikenal")
        
        # Check existing template for backup
        existing = self.db.get_template(doc_type)
        if existing and os.path.exists(existing['filepath']):
            # Backup existing
            backup_name = f"{doc_type}_v{existing['version']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
            backup_path = os.path.join(BACKUP_TEMPLATES_DIR, backup_name)
            shutil.copy2(existing['filepath'], backup_path)
        
        # Copy new template
        target_dir = WORD_TEMPLATES_DIR if template_type == 'word' else EXCEL_TEMPLATES_DIR
        filename = f"{doc_type.lower()}{ext}"
        target_path = os.path.join(target_dir, filename)
        
        shutil.copy2(source_path, target_path)
        
        # Extract placeholders from template
        placeholders = self.extract_placeholders(target_path, template_type)
        
        # Save to database
        template_id = self.db.save_template({
            'code': doc_type,
            'name': template_config['name'],
            'type': template_type,
            'filename': filename,
            'filepath': target_path,
            'placeholders': placeholders,
            'description': description,
            'uploaded_by': uploaded_by
        })
        
        return {
            'success': True,
            'template_id': template_id,
            'filepath': target_path,
            'placeholders': placeholders,
            'message': f"Template {doc_type} berhasil di-upload"
        }
    
    def extract_placeholders(self, filepath: str, template_type: str) -> List[str]:
        """Extract placeholders from template file"""
        placeholders = set()
        
        try:
            if template_type == 'word':
                doc = Document(filepath)
                
                # Check paragraphs
                for para in doc.paragraphs:
                    matches = PLACEHOLDER_PATTERN.findall(para.text)
                    for match in matches:
                        placeholders.add(match[0])
                
                # Check tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                matches = PLACEHOLDER_PATTERN.findall(para.text)
                                for match in matches:
                                    placeholders.add(match[0])
            
            else:  # Excel
                wb = load_workbook(filepath)
                for ws in wb.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str):
                                matches = PLACEHOLDER_PATTERN.findall(cell.value)
                                for match in matches:
                                    placeholders.add(match[0])
        
        except Exception as e:
            print(f"Warning: Could not extract placeholders: {e}")
        
        return sorted(list(placeholders))
    
    def get_template_info(self, doc_type: str) -> Optional[Dict]:
        """Get template information with placeholder details"""
        template = self.db.get_template(doc_type)
        
        if template:
            # Add placeholder details
            placeholder_details = []
            for ph in template.get('placeholders', []):
                if ph in ALL_PLACEHOLDERS:
                    placeholder_details.append({
                        'name': ph,
                        **ALL_PLACEHOLDERS[ph]
                    })
                else:
                    placeholder_details.append({
                        'name': ph,
                        'label': ph,
                        'type': 'unknown',
                        'required': False
                    })
            
            template['placeholder_details'] = placeholder_details
        
        return template
    
    def get_all_templates_status(self) -> List[Dict]:
        """Get status of all templates"""
        from app.core.config import DOCUMENT_TEMPLATES
        
        results = []
        existing = {t['code']: t for t in self.db.get_all_templates()}
        
        for code, config in DOCUMENT_TEMPLATES.items():
            if code in existing:
                results.append({
                    'code': code,
                    'name': config['name'],
                    'type': config['type'],
                    'status': 'uploaded',
                    'version': existing[code]['version'],
                    'filepath': existing[code]['filepath'],
                    'uploaded_at': existing[code]['uploaded_at']
                })
            else:
                # Check if default exists
                template_dir = (WORD_TEMPLATES_DIR if config['type'] == 'word' 
                              else EXCEL_TEMPLATES_DIR)
                default_path = os.path.join(template_dir, config['template'])
                
                results.append({
                    'code': code,
                    'name': config['name'],
                    'type': config['type'],
                    'status': 'exists' if os.path.exists(default_path) else 'missing',
                    'version': 0,
                    'filepath': default_path if os.path.exists(default_path) else None
                })
        
        return results


# ============================================================================
# SINGLETON
# ============================================================================

_template_engine = None
_template_manager = None

def get_template_engine() -> TemplateEngine:
    global _template_engine
    if _template_engine is None:
        _template_engine = TemplateEngine()
    return _template_engine

def get_template_manager() -> TemplateManager:
    global _template_manager
    if _template_manager is None:
        _template_manager = TemplateManager()
    return _template_manager


__all__ = [
    'TemplateEngine', 'get_template_engine',
    'TemplateManager', 'get_template_manager',
    'format_rupiah', 'format_angka', 'terbilang', 'format_tanggal'
]
