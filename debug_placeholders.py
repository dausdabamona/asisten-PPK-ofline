"""
Debug script untuk cek placeholder dalam template lembar permintaan
"""
from pathlib import Path
from docx import Document
import re

BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"

# Find lembar permintaan template
lembar_template = TEMPLATES_DIR / "word" / "lembar_permintaan.docx"

if not lembar_template.exists():
    print(f"Template tidak ditemukan: {lembar_template}")
    # List available templates
    print("\nTemplate yang tersedia:")
    for f in TEMPLATES_DIR.rglob("*.docx"):
        print(f"  - {f.relative_to(TEMPLATES_DIR)}")
else:
    print(f"Opening template: {lembar_template}")
    doc = Document(str(lembar_template))
    
    # Find all placeholders
    placeholders = set()
    pattern = r'\{\{(\w+)(?::(\w+))?\}\}'
    
    print("\n=== PLACEHOLDERS IN PARAGRAPHS ===")
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if '{{' in paragraph.text:
            print(f"\nParagraph {para_idx}: {paragraph.text[:100]}")
            matches = re.findall(pattern, paragraph.text)
            for key, fmt in matches:
                placeholders.add(key)
                print(f"  - {key} (format: {fmt or 'none'})")
    
    print("\n=== PLACEHOLDERS IN TABLES ===")
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if '{{' in cell.text:
                    print(f"  Cell [{row_idx},{cell_idx}]: {cell.text[:80]}")
                    matches = re.findall(pattern, cell.text)
                    for key, fmt in matches:
                        placeholders.add(key)
                        print(f"    - {key}")
    
    print("\n=== ALL UNIQUE PLACEHOLDERS ===")
    for ph in sorted(placeholders):
        print(f"  - {ph}")
    
    print("\n=== CHECKING IF DATA KEYS MATCH ===")
    # Simulate what data will look like
    test_data = {
        'hari_tanggal': '15 Januari 2024',
        'unit_kerja': 'Dinas Pendidikan',
        'sumber_dana': 'APBD',
        'nama_pengajuan': 'John Doe',
        'nama_verifikator': 'Jane Smith',
        'nama_ppk': 'PPK Officer',
        'nama_atasan': 'Director',
        'nama_kpa': 'KPA Officer',
    }
    
    missing = []
    for ph in sorted(placeholders):
        if ph in test_data:
            print(f"  ✓ {ph}: {test_data[ph]}")
        else:
            print(f"  ✗ {ph}: MISSING FROM DATA")
            missing.append(ph)
    
    if missing:
        print(f"\nMISSING {len(missing)} placeholders from data dict:")
        for ph in missing:
            print(f"  - {ph}")
