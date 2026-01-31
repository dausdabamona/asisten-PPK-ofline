"""
PPK DOCUMENT FACTORY - Dokumen Generator Service
=================================================
Service untuk generate dokumen dari template untuk workflow pencairan.

Features:
- Generate dokumen Word/Excel dari template
- Replace placeholder dengan data transaksi
- Simpan ke folder terstruktur per transaksi
- Track status dokumen di database
"""

import os
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, Tuple, List
from docx import Document
from docx.shared import Pt
import openpyxl

# Base paths
BASE_DIR = Path(__file__).parent.parent.parent
TEMPLATES_DIR = BASE_DIR / "templates"
OUTPUT_DIR = BASE_DIR / "output" / "dokumen"


def format_rupiah(value: float) -> str:
    """Format angka ke format Rupiah."""
    if value is None:
        return "Rp 0"
    return f"Rp {value:,.0f}".replace(",", ".")


def terbilang(n: float) -> str:
    """Konversi angka ke terbilang dalam Bahasa Indonesia."""
    if n == 0:
        return "nol rupiah"

    satuan = ["", "satu", "dua", "tiga", "empat", "lima",
              "enam", "tujuh", "delapan", "sembilan", "sepuluh", "sebelas"]

    n = int(n)

    if n < 0:
        return "minus " + terbilang(-n)
    elif n < 12:
        return satuan[n]
    elif n < 20:
        return satuan[n - 10] + " belas"
    elif n < 100:
        return satuan[n // 10] + " puluh " + satuan[n % 10]
    elif n < 200:
        return "seratus " + terbilang(n - 100)
    elif n < 1000:
        return satuan[n // 100] + " ratus " + terbilang(n % 100)
    elif n < 2000:
        return "seribu " + terbilang(n - 1000)
    elif n < 1000000:
        return terbilang(n // 1000) + " ribu " + terbilang(n % 1000)
    elif n < 1000000000:
        return terbilang(n // 1000000) + " juta " + terbilang(n % 1000000)
    elif n < 1000000000000:
        return terbilang(n // 1000000000) + " miliar " + terbilang(n % 1000000000)
    else:
        return terbilang(n // 1000000000000) + " triliun " + terbilang(n % 1000000000000)


def format_tanggal(date_str: str) -> str:
    """Format tanggal ke format Indonesia."""
    if not date_str:
        return ""
    try:
        if isinstance(date_str, datetime):
            dt = date_str
        else:
            dt = datetime.strptime(str(date_str)[:10], "%Y-%m-%d")

        bulan = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
                 "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
        return f"{dt.day} {bulan[dt.month]} {dt.year}"
    except:
        return str(date_str)


class DokumenGenerator:
    """Service untuk generate dokumen dari template."""

    # Mapping format functions
    FORMATTERS = {
        'rupiah': format_rupiah,
        'terbilang': lambda x: terbilang(x).title() + " Rupiah",
        'tanggal': format_tanggal,
        'upper': lambda x: str(x).upper() if x else "",
        'lower': lambda x: str(x).lower() if x else "",
        'title': lambda x: str(x).title() if x else "",
    }

    def __init__(self, db=None):
        """Initialize generator with optional database connection."""
        self.db = db
        self._ensure_directories()

    def _ensure_directories(self):
        """Pastikan folder output ada."""
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    def _sanitize_folder_name(self, name: str) -> str:
        """Sanitize folder name - remove invalid characters."""
        if not name:
            return "unnamed"
        # Replace invalid characters with underscore
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            name = name.replace(char, '_')
        # Remove leading/trailing spaces and dots
        name = name.strip('. ')
        # Limit length
        if len(name) > 50:
            name = name[:50]
        return name or "unnamed"

    def get_output_folder(self, transaksi: Dict[str, Any] = None,
                          mekanisme: str = None, nama_kegiatan: str = None) -> Path:
        """
        Get atau buat folder output untuk transaksi.

        Format folder: output/dokumen/{MEKANISME}/{BULAN_TAHUN}/{NAMA_KEGIATAN}/
        Contoh: output/dokumen/UP/Januari_2026/Belanja_ATK_Kantor/

        Args:
            transaksi: Dictionary data transaksi (opsional)
            mekanisme: Jenis pembayaran (UP, TUP, LS)
            nama_kegiatan: Nama kegiatan
        """
        # Get values from transaksi if provided
        if transaksi:
            mekanisme = mekanisme or transaksi.get('mekanisme', 'LAINNYA')
            nama_kegiatan = nama_kegiatan or transaksi.get('nama_kegiatan', 'Kegiatan')

        # Default values
        mekanisme = mekanisme or 'LAINNYA'
        nama_kegiatan = nama_kegiatan or 'Kegiatan'

        # Get current month and year in Indonesian
        now = datetime.now()
        bulan_names = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
                       "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
        bulan_tahun = f"{bulan_names[now.month]}_{now.year}"

        # Sanitize nama kegiatan for folder name
        nama_folder = self._sanitize_folder_name(nama_kegiatan)

        # Build folder path: MEKANISME/BULAN_TAHUN/NAMA_KEGIATAN
        folder = OUTPUT_DIR / mekanisme.upper() / bulan_tahun / nama_folder
        folder.mkdir(parents=True, exist_ok=True)

        return folder

    def get_output_folder_legacy(self, transaksi_kode: str) -> Path:
        """Get folder output berdasarkan kode transaksi (legacy method)."""
        folder = OUTPUT_DIR / transaksi_kode
        folder.mkdir(parents=True, exist_ok=True)
        return folder

    def get_template_path(self, template_name: str) -> Optional[Path]:
        """Get path template file."""
        if not template_name:
            return None

        # Check in word folder first
        word_path = TEMPLATES_DIR / "word" / template_name
        if word_path.exists():
            return word_path

        # Check in excel folder
        excel_path = TEMPLATES_DIR / "excel" / template_name
        if excel_path.exists():
            return excel_path

        # Check root templates folder
        root_path = TEMPLATES_DIR / template_name
        if root_path.exists():
            return root_path

        return None

    def prepare_data(self, transaksi: Dict[str, Any], satker: Dict[str, Any] = None,
                     pegawai: Dict[str, Any] = None, rincian: List[Dict] = None) -> Dict[str, Any]:
        """Prepare data dictionary for template merging."""
        data = {}

        # Data transaksi
        data['kode_transaksi'] = transaksi.get('kode', '')
        data['nama_kegiatan'] = transaksi.get('nama_kegiatan', '')
        data['kode_akun'] = transaksi.get('kode_akun', '')
        data['estimasi_biaya'] = transaksi.get('estimasi_biaya', 0)
        data['uang_muka'] = transaksi.get('uang_muka', 0)
        data['realisasi'] = transaksi.get('realisasi', 0)
        data['selisih'] = data['realisasi'] - data['uang_muka']
        data['mekanisme'] = transaksi.get('mekanisme', '')
        data['jenis_kegiatan'] = transaksi.get('jenis_kegiatan', '')

        # Tanggal
        data['tanggal_pengajuan'] = transaksi.get('tanggal_pengajuan', '')
        data['tanggal_pencairan'] = transaksi.get('tanggal_pencairan', '')
        data['tanggal_selesai'] = transaksi.get('tanggal_selesai', '')
        data['tanggal_hari_ini'] = datetime.now().strftime("%Y-%m-%d")

        # Data satker
        if satker:
            data['satker_nama'] = satker.get('nama', '')
            data['satker_kode'] = satker.get('kode', '')
            data['satker_alamat'] = satker.get('alamat', '')
            data['satker_kota'] = satker.get('kota', '')
            data['kementerian'] = satker.get('kementerian', '')
            data['unit_organisasi'] = satker.get('unit_organisasi', '')
            data['lokasi'] = satker.get('kota', '')

        # Data penerima/pegawai
        if pegawai:
            data['penerima_nama'] = pegawai.get('nama', transaksi.get('penerima_nama', ''))
            data['penerima_nip'] = pegawai.get('nip', transaksi.get('penerima_nip', ''))
            data['penerima_jabatan'] = pegawai.get('jabatan', transaksi.get('penerima_jabatan', ''))
        else:
            data['penerima_nama'] = transaksi.get('penerima_nama', '')
            data['penerima_nip'] = transaksi.get('penerima_nip', '')
            data['penerima_jabatan'] = transaksi.get('penerima_jabatan', '')

        # Data PPK dari satker
        if satker:
            data['ppk_nama'] = satker.get('ppk_nama', '')
            data['ppk_nip'] = satker.get('ppk_nip', '')
            data['kpa_nama'] = satker.get('kpa_nama', '')
            data['kpa_nip'] = satker.get('kpa_nip', '')
            data['bendahara_nama'] = satker.get('bendahara_nama', '')
            data['bendahara_nip'] = satker.get('bendahara_nip', '')

        # Rincian items
        if rincian:
            data['rincian_items'] = rincian
            data['total_rincian'] = sum(item.get('jumlah', 0) for item in rincian)
            data['jumlah_item'] = len(rincian)

        return data

    def _apply_format(self, value: Any, format_type: str) -> str:
        """Apply format to value."""
        if format_type in self.FORMATTERS:
            try:
                return self.FORMATTERS[format_type](value)
            except:
                return str(value) if value else ""
        return str(value) if value else ""

    def _replace_placeholder(self, text: str, data: Dict[str, Any]) -> str:
        """Replace all {{placeholder}} or {{placeholder:format}} in text."""
        if not text or '{{' not in text:
            return text

        import re
        pattern = r'\{\{(\w+)(?::(\w+))?\}\}'

        def replacer(match):
            key = match.group(1)
            fmt = match.group(2)
            value = data.get(key, '')

            if fmt:
                return self._apply_format(value, fmt)
            return str(value) if value else ""

        return re.sub(pattern, replacer, text)

    def _flatten_rincian_items(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Flatten rincian_items list to indexed placeholders.
        
        Converts list like:
            rincian_items: [
                {nama_barang: 'Item 1', harga_satuan: 10000, ...},
                {nama_barang: 'Item 2', harga_satuan: 20000, ...}
            ]
        
        To flattened dict like:
            item_1_nama: 'Item 1'
            item_1_harga: '10000'
            item_2_nama: 'Item 2'
            item_2_harga: '20000'
            ...
        
        This matches the template placeholders like {{item_1_nama}}, {{item_2_harga}}, etc.
        Empty items are filled with blank values to clear template placeholders.
        """
        data = dict(data)  # Copy to avoid modifying original
        
        rincian_items = data.get('rincian_items', [])
        if not rincian_items:
            return data
        
        # Maximum 5 items based on template (item_1 to item_5)
        max_items = 5
        
        # Field mapping from rincian_items keys to template placeholder names
        field_mapping = {
            'no': 'no',
            'item_no': 'no',
            'uraian': 'nama',  # template uses 'item_N_nama'
            'nama_barang': 'nama',
            'spesifikasi': 'spek',
            'volume': 'volume',
            'satuan': 'satuan',
            'harga_satuan': 'harga',
            'jumlah': 'total',  # template uses 'item_N_total' for total price
            'total_item': 'total',
            'keterangan': 'ket',
        }
        
        # Flatten all items (including empty ones)
        for item_num in range(1, max_items + 1):
            # Get item or use empty dict for extra slots
            if item_num - 1 < len(rincian_items):
                item = rincian_items[item_num - 1]
            else:
                item = {}
            
            # Map each field to indexed placeholder
            for source_key, target_suffix in field_mapping.items():
                placeholder_key = f'item_{item_num}_{target_suffix}'
                value = item.get(source_key, '')
                
                # Format numeric values
                if target_suffix in ['harga', 'total'] and isinstance(value, (int, float)) and value > 0:
                    value = f"Rp {value:,.0f}".replace(",", ".")
                
                data[placeholder_key] = str(value) if value else ""
        
        return data


    def _process_paragraph(self, paragraph, data: Dict[str, Any]):
        """Process a paragraph to replace placeholders."""
        if not paragraph.text or '{{' not in paragraph.text:
            return

        # Get full text and replace
        full_text = paragraph.text
        new_text = self._replace_placeholder(full_text, data)

        if full_text != new_text:
            # Clear runs and set new text
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_text

    def _process_table(self, table, data: Dict[str, Any]):
        """Process a table to replace placeholders."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._process_paragraph(paragraph, data)

    def merge_word_document(self, template_path: Path, data: Dict[str, Any],
                           output_path: Path) -> bool:
        """Merge data into Word document template."""
        try:
            doc = Document(str(template_path))

            # IMPORTANT: Flatten rincian_items to item_N_* placeholders BEFORE processing
            # This allows the template's indexed placeholders (item_1_no, item_2_nama, etc.) to work
            if data.get('rincian_items'):
                data = self._flatten_rincian_items(data)

            # Process all paragraphs
            for paragraph in doc.paragraphs:
                self._process_paragraph(paragraph, data)

            # Process all tables
            for table_idx, table in enumerate(doc.tables):
                self._process_table(table, data)

            # Process headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        self._process_paragraph(paragraph, data)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        self._process_paragraph(paragraph, data)

            # Save
            doc.save(str(output_path))
            return True

        except Exception as e:
            print(f"Error merging Word document: {e}")
            return False

    def _is_rincian_table(self, table) -> bool:
        """Check if table adalah tabel rincian barang (dari template lembar permintaan)."""
        # Cek header row untuk keywords rincian
        if len(table.rows) > 0:
            header_text = ' '.join(cell.text.lower() for cell in table.rows[0].cells)
            rincian_keywords = ['barang', 'uraian', 'volume', 'satuan', 'harga', 'jumlah', 'spesifikasi']
            if any(keyword in header_text for keyword in rincian_keywords):
                return True
        return False

    def _fill_rincian_table(self, table, rincian_items: List[Dict], data: Dict[str, Any]):
        """
        Fill rincian items ke dalam tabel.
        Asumsi: Row pertama adalah header, row 1-5 adalah data items dengan placeholders.
        """
        if not rincian_items or len(table.rows) < 2:
            return

        # Mapping kolom berdasarkan header
        header_row = table.rows[0]
        column_mapping = {}
        
        headers = [cell.text.lower().strip() for cell in header_row.cells]
        
        # Map header ke column index
        for col_idx, header in enumerate(headers):
            if any(x in header for x in ['no', 'nomor']):
                column_mapping['no'] = col_idx
            elif any(x in header for x in ['uraian', 'barang', 'nama']):
                column_mapping['uraian'] = col_idx
            elif any(x in header for x in ['spesifikasi', 'spec']):
                column_mapping['spesifikasi'] = col_idx
            elif any(x in header for x in ['volume', 'vol']):
                column_mapping['volume'] = col_idx
            elif any(x in header for x in ['satuan']):
                column_mapping['satuan'] = col_idx
            elif any(x in header for x in ['harga', 'price']):
                column_mapping['harga_satuan'] = col_idx
            elif any(x in header for x in ['jumlah', 'total', 'total']):
                column_mapping['jumlah'] = col_idx
            elif any(x in header for x in ['ket', 'keterangan']):
                column_mapping['keterangan'] = col_idx

        # Fill items ke rows yang tersedia (skip header row)
        start_row = 1
        for item_idx, item in enumerate(rincian_items):
            row_idx = start_row + item_idx
            
            # Jika row tidak ada, create new row
            while len(table.rows) <= row_idx:
                table.add_row()
            
            row = table.rows[row_idx]

            # Fill columns
            for col_key, col_idx in column_mapping.items():
                if col_idx < len(row.cells):
                    value = item.get(col_key, '')
                    
                    # Format berdasarkan tipe data
                    if col_key in ['harga_satuan', 'jumlah']:
                        if isinstance(value, (int, float)):
                            value = f"Rp {value:,.0f}".replace(",", ".")
                    
                    # Set cell text
                    row.cells[col_idx].text = str(value or '')

        # Remove empty rows di atas placeholder rows
        while len(table.rows) > start_row + len(rincian_items) + 3:  # Keep 3 extra rows untuk subtotal
            # Delete last row
            tbl = table._element
            tbl.remove(table.rows[-1]._element)

    def merge_excel_document(self, template_path: Path, data: Dict[str, Any],
                            output_path: Path) -> bool:
        """Merge data into Excel document template."""
        try:
            wb = openpyxl.load_workbook(str(template_path))

            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str) and '{{' in cell.value:
                            cell.value = self._replace_placeholder(cell.value, data)

            wb.save(str(output_path))
            return True

        except Exception as e:
            print(f"Error merging Excel document: {e}")
            return False

    def generate_document(self, transaksi: Dict[str, Any], kode_dokumen: str,
                         template_name: str, satker: Dict[str, Any] = None,
                         pegawai: Dict[str, Any] = None,
                         rincian: List[Dict] = None,
                         additional_data: Dict[str, Any] = None,
                         is_draft: bool = True) -> Tuple[Optional[str], Optional[str]]:
        """
        Generate dokumen dari template.
        File akan menimpa file lama jika sudah ada (berdasarkan kode_dokumen).

        Args:
            transaksi: Data transaksi
            kode_dokumen: Kode dokumen (e.g., LBR_REQ, KUIT_UM)
            template_name: Nama file template
            satker: Data satuan kerja
            pegawai: Data pegawai/penerima
            rincian: List rincian barang/jasa
            additional_data: Data tambahan untuk merge
            is_draft: True jika dokumen adalah draft (fase 1), False jika final

        Returns:
            Tuple of (output_path, error_message)
        """
        # Get template
        template_path = self.get_template_path(template_name)
        if not template_path:
            return None, f"Template '{template_name}' tidak ditemukan"

        # Prepare data
        data = self.prepare_data(transaksi, satker, pegawai, rincian)

        # Add additional data
        if additional_data:
            data.update(additional_data)

        # Create output folder with new naming: MEKANISME/BULAN_TAHUN/NAMA_KEGIATAN
        output_folder = self.get_output_folder(transaksi=transaksi)

        # Generate output filename - tanpa timestamp agar menimpa file lama
        ext = template_path.suffix
        status_suffix = "_DRAFT" if is_draft else ""
        output_filename = f"{kode_dokumen}{status_suffix}{ext}"
        output_path = output_folder / output_filename

        # Remove existing file if exists (overwrite)
        if output_path.exists():
            try:
                os.remove(output_path)
            except Exception as e:
                print(f"Warning: Could not remove existing file: {e}")

        # Merge document
        if ext.lower() == '.docx':
            success = self.merge_word_document(template_path, data, output_path)
        elif ext.lower() == '.xlsx':
            success = self.merge_excel_document(template_path, data, output_path)
        else:
            return None, f"Format template '{ext}' tidak didukung"

        if success:
            return str(output_path), None
        else:
            return None, "Gagal generate dokumen"

    def open_document(self, filepath: str) -> bool:
        """Open document with default application."""
        import subprocess
        import platform

        try:
            if platform.system() == 'Windows':
                os.startfile(filepath)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', filepath])
            else:  # Linux
                subprocess.run(['xdg-open', filepath])
            return True
        except Exception as e:
            print(f"Error opening document: {e}")
            return False

    def open_folder(self, transaksi: Dict[str, Any] = None,
                    mekanisme: str = None, nama_kegiatan: str = None) -> bool:
        """Open output folder for transaksi."""
        folder = self.get_output_folder(transaksi, mekanisme, nama_kegiatan)
        return self.open_document(str(folder))

    def save_lembar_permintaan_to_db(self, lembar_data: Dict, items: List[Dict] = None) -> Tuple[int, Optional[str]]:
        """
        Save lembar_permintaan ke database beserta items.
        
        Args:
            lembar_data: Dictionary dengan data lembar permintaan:
                - hari_tanggal, unit_kerja, sumber_dana, kode_transaksi
                - subtotal, ppn, total
                - nama_pengajuan, nama_verifikator, nama_ppk, nama_atasan, nama_kpa
                - file_path, status, tahun_anggaran, created_by
            items: List dari item/rincian barang
        
        Returns:
            Tuple of (lembar_permintaan_id, error_message)
            - Jika berhasil: (id, None)
            - Jika error: (None, error_message)
        """
        try:
            if not self.db:
                return None, "Database connection tidak tersedia"
            
            # Create lembar_permintaan
            lembar_id = self.db.create_lembar_permintaan(lembar_data)
            
            # Add items
            if items:
                for item in items:
                    self.db.add_lembar_permintaan_item(lembar_id, item)
            
            return lembar_id, None
            
        except Exception as e:
            error_msg = f"Error saving lembar_permintaan to database: {str(e)}"
            print(f"❌ {error_msg}")
            return None, error_msg

    def generate_and_save_lembar_permintaan(
        self,
        lembar_data: Dict[str, Any],
        items: List[Dict] = None,
        template_name: str = 'lembar_permintaan.docx',
        output_dir: str = None,
        created_by: str = 'system'
    ) -> Tuple[int, Optional[str], Optional[str]]:
        """
        Generate lembar_permintaan document dan save data ke database secara bersamaan.
        
        Args:
            lembar_data: Data lembar permintaan
            items: List item/rincian barang
            template_name: Nama template file
            output_dir: Output directory untuk dokumen
            created_by: User yang membuat
        
        Returns:
            Tuple of (lembar_id, output_path, error_message)
            - Jika berhasil: (lembar_id, output_path, None)
            - Jika error: (None, None, error_message)
        """
        try:
            # Prepare placeholder mapping dari item
            placeholders = {
                'hari_tanggal': lembar_data.get('hari_tanggal', ''),
                'unit_kerja': lembar_data.get('unit_kerja', ''),
                'sumber_dana': lembar_data.get('sumber_dana', ''),
            }
            
            # Add item placeholders
            if items:
                for idx, item in enumerate(items, 1):
                    placeholders[f'item_{idx}_no'] = item.get('item_no', idx)
                    placeholders[f'item_{idx}_nama'] = item.get('nama_barang', '')
                    placeholders[f'item_{idx}_spek'] = item.get('spesifikasi', '')
                    placeholders[f'item_{idx}_volume'] = item.get('volume', '')
                    placeholders[f'item_{idx}_satuan'] = item.get('satuan', '')
                    placeholders[f'item_{idx}_harga'] = item.get('harga_satuan', '')
                    placeholders[f'item_{idx}_total'] = item.get('total_item', '')
                    placeholders[f'item_{idx}_ket'] = item.get('keterangan', '')
            
            # Add totals
            placeholders['subtotal'] = lembar_data.get('subtotal', 0)
            placeholders['ppn'] = lembar_data.get('ppn', 0)
            placeholders['total'] = lembar_data.get('total', 0)
            
            # Add signatures
            placeholders['nama_pengajuan'] = lembar_data.get('nama_pengajuan', '')
            placeholders['nama_verifikator'] = lembar_data.get('nama_verifikator', '')
            placeholders['nama_ppk'] = lembar_data.get('nama_ppk', '')
            placeholders['nama_atasan'] = lembar_data.get('nama_atasan', '')
            placeholders['nama_kpa'] = lembar_data.get('nama_kpa', '')
            
            # Generate dokumen
            doc_path, doc_error = self.generate_document_from_template(
                template_name,
                placeholders,
                output_dir
            )
            
            if not doc_path:
                return None, None, f"Gagal generate dokumen: {doc_error}"
            
            # Save ke database
            lembar_data['file_path'] = doc_path
            lembar_data['status'] = lembar_data.get('status', 'draft')
            lembar_data['created_by'] = created_by
            
            lembar_id, db_error = self.save_lembar_permintaan_to_db(lembar_data, items)
            
            if db_error:
                return None, doc_path, db_error
            
            print(f"✓ Lembar Permintaan tersimpan:")
            print(f"  ID: {lembar_id}")
            print(f"  Document: {doc_path}")
            
            return lembar_id, doc_path, None
            
        except Exception as e:
            error_msg = f"Error dalam generate_and_save_lembar_permintaan: {str(e)}"
            print(f"❌ {error_msg}")
            return None, None, error_msg

    def generate_document_from_template(
        self,
        template_name: str,
        placeholders: Dict[str, Any],
        output_dir: str = None
    ) -> Tuple[Optional[str], Optional[str]]:
        """
        Generate dokumen dari template dengan placeholder replacement.
        
        Returns:
            Tuple of (output_path, error_message)
        """
        try:
            template_path = self.get_template_path(template_name)
            if not template_path:
                return None, f"Template '{template_name}' tidak ditemukan"
            
            # Load dan process dokumen
            doc = Document(str(template_path))
            
            # Replace placeholders dalam paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, value in placeholders.items():
                    placeholder_text = f'{{{{{placeholder}}}}}'
                    if placeholder_text in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder_text in run.text:
                                run.text = run.text.replace(placeholder_text, str(value or ''))
            
            # Replace placeholders dalam tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in placeholders.items():
                                placeholder_text = f'{{{{{placeholder}}}}}'
                                if placeholder_text in paragraph.text:
                                    for run in paragraph.runs:
                                        if placeholder_text in run.text:
                                            run.text = run.text.replace(placeholder_text, str(value or ''))
            
            # Create output directory
            if not output_dir:
                output_dir = Path(BASE_DIR) / "output" / "dokumen"
            else:
                output_dir = Path(output_dir)
            
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate output filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"lembar_permintaan_{timestamp}.docx"
            output_path = output_dir / output_filename
            
            # Save dokumen
            doc.save(str(output_path))
            
            return str(output_path), None
            
        except Exception as e:
            error_msg = f"Error generate_document_from_template: {str(e)}"
            print(f"❌ {error_msg}")
            return None, error_msg


# Singleton instance
_generator_instance = None

def get_dokumen_generator(db=None) -> DokumenGenerator:
    """Get singleton instance of DokumenGenerator."""
    global _generator_instance
    if _generator_instance is None:
        _generator_instance = DokumenGenerator(db)
    return _generator_instance


# ============================================================================
# HELPER FUNCTIONS untuk Lembar Permintaan Workflow
# ============================================================================

def get_lembar_permintaan_data_for_next_phase(
    db,
    lembar_id: int,
    target_dokumen: str = None
) -> Optional[Dict[str, Any]]:
    """
    Helper function untuk get lembar_permintaan data yang sudah diformat.
    Digunakan di fase berikutnya (SPJ, Kuitansi, dst).
    
    Args:
        db: Database manager instance
        lembar_id: ID lembar_permintaan
        target_dokumen: Jenis dokumen target (opsional)
    
    Returns:
        Formatted data siap digunakan di next phase
    """
    if not db:
        return None
    
    return db.get_lembar_permintaan_for_next_phase(lembar_id, target_dokumen)


def generate_dokumen_from_lembar(
    db,
    lembar_id: int,
    dokumen_type: str,
    template_name: str,
    output_dir: str = None,
    created_by: str = 'system'
) -> Tuple[Optional[str], Optional[str]]:
    """
    Generate dokumen fase selanjutnya berdasarkan data lembar_permintaan.
    
    Args:
        db: Database manager instance
        lembar_id: ID lembar_permintaan
        dokumen_type: Tipe dokumen (SPJ, KUITANSI, REALISASI, dst)
        template_name: Nama template file
        output_dir: Output directory
        created_by: User yang membuat
    
    Returns:
        Tuple of (output_path, error_message)
    """
    try:
        # Get data lembar dengan formatting untuk next phase
        data = get_lembar_permintaan_data_for_next_phase(db, lembar_id)
        if not data:
            return None, f"Lembar Permintaan {lembar_id} tidak ditemukan"
        
        # Get generator
        generator = get_dokumen_generator(db)
        
        # Prepare placeholders dari data lembar
        placeholders = {
            'tanggal': data.get('tanggal', ''),
            'unit_kerja': data.get('unit_kerja', ''),
            'sumber_dana': data.get('sumber_dana', ''),
            'subtotal': data.get('subtotal', 0),
            'ppn': data.get('ppn', 0),
            'total': data.get('total', 0),
            'pengajuan_oleh': data.get('pengajuan_oleh', ''),
            'ppk': data.get('ppk', ''),
            'kpa': data.get('kpa', ''),
        }
        
        # Add rincian items
        items = data.get('rincian_items', [])
        for idx, item in enumerate(items, 1):
            placeholders[f'item_{idx}_nama'] = item.get('nama_barang', '')
            placeholders[f'item_{idx}_spek'] = item.get('spesifikasi', '')
            placeholders[f'item_{idx}_volume'] = item.get('volume', '')
            placeholders[f'item_{idx}_satuan'] = item.get('satuan', '')
            placeholders[f'item_{idx}_harga'] = item.get('harga_satuan', '')
            placeholders[f'item_{idx}_total'] = item.get('jumlah', '')
        
        # Generate dokumen
        doc_path, doc_error = generator.generate_document_from_template(
            template_name,
            placeholders,
            output_dir
        )
        
        if doc_error:
            return None, doc_error
        
        # Update lembar status
        if db:
            db.update_lembar_permintaan(lembar_id, {'status': 'used'})
        
        print(f"✓ Dokumen {dokumen_type} generated dari Lembar Permintaan {lembar_id}")
        return doc_path, None
        
    except Exception as e:
        error_msg = f"Error generate_dokumen_from_lembar: {str(e)}"
        print(f"❌ {error_msg}")
        return None, error_msg
