#!/usr/bin/env python3
"""
Create Lembar Permintaan Template
=================================
Template untuk Lembar Permintaan Pencairan Dana dengan:
- Tabel rincian barang/jasa
- Kolom validator dan atasan langsung
- Spasi 1, tanpa jarak setelah paragraf
- Lebar kolom menyesuaikan isi
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates", "word")


def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            edge_element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            edge_element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(edge_element)
    tcPr.append(tcBorders)


def set_no_border(cell):
    """Remove cell border."""
    set_cell_border(cell,
        top={'val': 'nil'},
        left={'val': 'nil'},
        bottom={'val': 'nil'},
        right={'val': 'nil'}
    )


def set_paragraph_format(paragraph, font_size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Set paragraph format with single spacing and no spacing after."""
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')


def add_formatted_paragraph(doc, text, font_size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add paragraph with proper formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    return p


def set_table_column_widths(table, widths):
    """Set column widths for table."""
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width


def autofit_table(table):
    """Set table to autofit content."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)


def create_lembar_permintaan():
    """Create Lembar Permintaan template."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Header - KOP SURAT
    add_formatted_paragraph(doc, "{{kop_satker}}", font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "{{alamat_satker}}", font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "", font_size=6)  # Small spacer

    # Garis pembatas
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p_border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '000000')
    p_border.append(bottom)
    p._p.get_or_add_pPr().append(p_border)

    # Title
    add_formatted_paragraph(doc, "", font_size=6)
    add_formatted_paragraph(doc, "LEMBAR PERMINTAAN PENCAIRAN DANA", font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "Nomor: {{nomor_dokumen}}", font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "", font_size=6)

    # Info Table (without border)
    info_table = doc.add_table(rows=6, cols=3)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    info_data = [
        ("Tanggal", ":", "{{tanggal_dokumen}}"),
        ("Nama Kegiatan", ":", "{{nama_kegiatan}}"),
        ("Kode Akun", ":", "{{kode_akun}}"),
        ("Penerima", ":", "{{penerima_nama}}"),
        ("NIP", ":", "{{penerima_nip}}"),
        ("Jabatan", ":", "{{penerima_jabatan}}"),
    ]

    for i, (label, sep, value) in enumerate(info_data):
        row = info_table.rows[i]
        cells = row.cells

        cells[0].text = label
        cells[1].text = sep
        cells[2].text = value

        for cell in cells:
            set_no_border(cell)
            for p in cell.paragraphs:
                set_paragraph_format(p, font_size=11)

    # Set column widths for info table
    set_table_column_widths(info_table, [Cm(3.5), Cm(0.5), Cm(10)])

    add_formatted_paragraph(doc, "", font_size=6)

    # Rincian Barang/Jasa Section
    add_formatted_paragraph(doc, "Rincian Barang/Jasa:", font_size=11, bold=True)
    add_formatted_paragraph(doc, "", font_size=3)

    # Rincian Table (with border)
    rincian_table = doc.add_table(rows=2, cols=6)
    rincian_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rincian_table.style = 'Table Grid'

    # Header row
    headers = ["No", "Uraian", "Volume", "Satuan", "Harga Satuan", "Jumlah"]
    header_row = rincian_table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            set_paragraph_format(p, font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Template row with placeholders - MUST match generator placeholders
    template_row = rincian_table.rows[1]
    template_data = ["{{rincian_no}}", "{{rincian_uraian}}", "{{rincian_volume}}", "{{rincian_satuan}}", "{{rincian_harga}}", "{{rincian_jumlah}}"]
    for i, data in enumerate(template_data):
        cell = template_row.cells[i]
        cell.text = data
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        align = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3] else (WD_ALIGN_PARAGRAPH.RIGHT if i in [4, 5] else WD_ALIGN_PARAGRAPH.LEFT)
        for p in cell.paragraphs:
            set_paragraph_format(p, font_size=10, align=align)

    # Set column widths - No kecil, Uraian lebar, sisanya sedang
    set_table_column_widths(rincian_table, [Cm(1), Cm(6), Cm(1.5), Cm(1.5), Cm(2.5), Cm(2.5)])

    add_formatted_paragraph(doc, "", font_size=6)

    # Total Section
    total_table = doc.add_table(rows=4, cols=2)
    total_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    total_data = [
        ("Sub Total", "{{subtotal}}"),
        ("PPN ({{ppn_persen}}%)", "{{ppn_nilai}}"),
        ("Total", "{{total_dengan_ppn}}"),
        ("Uang Muka Diterima ({{uang_muka_persen}}%)", "{{uang_muka_nilai}}"),
    ]

    for i, (label, value) in enumerate(total_data):
        row = total_table.rows[i]
        cells = row.cells
        cells[0].text = label
        cells[1].text = value

        for j, cell in enumerate(cells):
            set_no_border(cell)
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            bold = True if i == 3 else False
            for p in cell.paragraphs:
                set_paragraph_format(p, font_size=11, bold=bold, align=align)

    set_table_column_widths(total_table, [Cm(5), Cm(3)])

    add_formatted_paragraph(doc, "", font_size=12)

    # Terbilang
    add_formatted_paragraph(doc, "Terbilang: {{terbilang}}", font_size=11)
    add_formatted_paragraph(doc, "", font_size=12)

    # Signature Section - 3 columns (Pemohon, Validator, Atasan Langsung)
    add_formatted_paragraph(doc, "", font_size=6)

    sig_table = doc.add_table(rows=6, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Headers
    sig_headers = ["Pemohon/Penerima,", "Validator,", "Mengetahui,\nAtasan Langsung"]
    for i, header in enumerate(sig_headers):
        cell = sig_table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            set_paragraph_format(p, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_no_border(cell)

    # Rows 1-4: Empty for signature space
    for row_idx in range(1, 5):
        for cell in sig_table.rows[row_idx].cells:
            cell.text = ""
            set_no_border(cell)

    # Row 5: Names and NIP
    sig_names = [
        "{{penerima_nama}}\nNIP. {{penerima_nip}}",
        "{{validator_nama}}\nNIP. {{validator_nip}}",
        "{{atasan_nama}}\nNIP. {{atasan_nip}}"
    ]
    for i, name in enumerate(sig_names):
        cell = sig_table.rows[5].cells[i]
        cell.text = name
        for p in cell.paragraphs:
            set_paragraph_format(p, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_no_border(cell)

    set_table_column_widths(sig_table, [Cm(5), Cm(5), Cm(5)])

    add_formatted_paragraph(doc, "", font_size=12)

    # Persetujuan PPK Section
    add_formatted_paragraph(doc, "", font_size=6)

    ppk_table = doc.add_table(rows=6, cols=1)
    ppk_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    ppk_cell = ppk_table.rows[0].cells[0]
    ppk_cell.text = "Disetujui oleh:"
    set_no_border(ppk_cell)
    for p in ppk_cell.paragraphs:
        set_paragraph_format(p, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    ppk_cell = ppk_table.rows[1].cells[0]
    ppk_cell.text = "Pejabat Pembuat Komitmen,"
    set_no_border(ppk_cell)
    for p in ppk_cell.paragraphs:
        set_paragraph_format(p, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Empty rows for signature
    for row_idx in range(2, 5):
        cell = ppk_table.rows[row_idx].cells[0]
        cell.text = ""
        set_no_border(cell)

    ppk_cell = ppk_table.rows[5].cells[0]
    ppk_cell.text = "{{ppk_nama}}\nNIP. {{ppk_nip}}"
    set_no_border(ppk_cell)
    for p in ppk_cell.paragraphs:
        set_paragraph_format(p, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_table_column_widths(ppk_table, [Cm(6)])

    # Save document
    output_path = os.path.join(TEMPLATES_DIR, "lembar_permintaan.docx")
    doc.save(output_path)
    print(f"Template saved: {output_path}")

    return output_path


if __name__ == "__main__":
    create_lembar_permintaan()
    print("Done!")
