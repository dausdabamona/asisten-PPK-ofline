#!/usr/bin/env python3
"""
Create All Missing Template Drafts (Comprehensive)
===================================================
Script untuk membuat template draft lengkap untuk semua dokumen
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os

WORD_TEMPLATES_DIR = Path("templates/word")
os.makedirs(WORD_TEMPLATES_DIR, exist_ok=True)

def add_header(doc, title, subtitle=None):
    """Add standard header"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.bold = True
    
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(11)
    
    doc.add_paragraph()

def add_signature_block(doc, title, name_placeholder, nip_placeholder=""):
    """Add signature block"""
    doc.add_paragraph()
    doc.add_paragraph(title)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(name_placeholder)
    if nip_placeholder:
        doc.add_paragraph(f"NIP: {nip_placeholder}")

# ============================================================================
# JAMUAN TAMU TEMPLATES
# ============================================================================

def create_daftar_hadir_jamuan_tamu():
    """Create Daftar Hadir Jamuan Tamu"""
    doc = Document()
    
    add_header(doc, "DAFTAR HADIR JAMUAN TAMU", "{{satker_nama}}")
    
    # Info
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Nama Kegiatan"
    cells[1].text = "{{nama_kegiatan}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Tanggal"
    cells[1].text = "{{tanggal_kegiatan:tanggal_long}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Tempat"
    cells[1].text = "{{tempat_kegiatan}}"
    
    doc.add_paragraph()
    
    # Peserta
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "No"
    hdr_cells[1].text = "Nama"
    hdr_cells[2].text = "Jabatan/Instansi"
    hdr_cells[3].text = "Tanda Tangan"
    
    for i in range(1, 11):
        cells = table.rows[i].cells
        cells[0].text = str(i)
        cells[1].text = ""
        cells[2].text = ""
        cells[3].text = ""
    
    doc.add_paragraph()
    
    add_signature_block(doc, "Panitia", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "daftar_hadir_jamuan_tamu.docx")
    print("✅ Dibuat: daftar_hadir_jamuan_tamu.docx")

def create_nota_dinas_jamuan_tamu():
    """Create Nota Dinas Jamuan Tamu"""
    doc = Document()
    
    add_header(doc, "NOTA DINAS", "{{satker_nama}}")
    
    doc.add_paragraph("Kepada        : Yth. Pejabat Pengadaan {{satker_nama}}")
    doc.add_paragraph("Dari          : Pejabat Pembuat Komitmen")
    doc.add_paragraph("Tanggal       : {{tanggal_nota:tanggal_long}}")
    doc.add_paragraph("Perihal       : Persetujuan Pelaksanaan Jamuan Tamu")
    doc.add_paragraph()
    
    doc.add_paragraph("Dengan hormat, kami mengajukan permohonan untuk melaksanakan kegiatan jamuan tamu dengan rincian:")
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Nama Kegiatan"
    cells[1].text = "{{nama_kegiatan}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Tanggal Pelaksanaan"
    cells[1].text = "{{tanggal_kegiatan:tanggal_long}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Jumlah Peserta"
    cells[1].text = "{{jumlah_peserta}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Estimasi Biaya"
    cells[1].text = "{{estimasi_biaya:rupiah}}"
    
    doc.add_paragraph()
    doc.add_paragraph("Demikian permohonan ini kami sampaikan. Atas perhatian dan persetujuannya kami ucapkan terima kasih.")
    
    doc.add_paragraph()
    add_signature_block(doc, "PPK", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "nota_dinas_jamuan_tamu.docx")
    print("✅ Dibuat: nota_dinas_jamuan_tamu.docx")

def create_kuitansi_jamuan_tamu():
    """Create Kuitansi Jamuan Tamu"""
    doc = Document()
    
    add_header(doc, "KUITANSI JAMUAN TAMU", "{{satker_nama}}")
    
    doc.add_paragraph("Nomor Kuitansi: {{nomor_kuitansi}}")
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Tanggal"
    cells[1].text = "{{tanggal_kuitansi:tanggal_long}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Kegiatan"
    cells[1].text = "{{nama_kegiatan}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Jumlah Peserta"
    cells[1].text = "{{jumlah_peserta}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Biaya per Peserta"
    cells[1].text = "{{biaya_per_peserta:rupiah}}"
    
    cells = table.rows[4].cells
    cells[0].text = "Total Biaya"
    cells[1].text = "{{total_biaya:rupiah}}"
    
    doc.add_paragraph()
    doc.add_paragraph("Telah diterima pembayaran jamuan tamu sebesar {{total_biaya:terbilang}} ({{total_biaya:rupiah}})")
    doc.add_paragraph()
    
    add_signature_block(doc, "Bendahara", "{{bendahara_nama}}")
    
    doc.save(WORD_TEMPLATES_DIR / "kuitansi_jamuan_tamu.docx")
    print("✅ Dibuat: kuitansi_jamuan_tamu.docx")

def create_laporan_jamuan_tamu():
    """Create Laporan Jamuan Tamu"""
    doc = Document()
    
    add_header(doc, "LAPORAN PELAKSANAAN JAMUAN TAMU", "{{satker_nama}}")
    
    p = doc.add_paragraph("I. PENDAHULUAN")
    p.style = 'Heading 2'
    doc.add_paragraph("Kegiatan jamuan tamu dilaksanakan dengan tujuan untuk mempererat hubungan dan meningkatkan kerjasama dengan stakeholder terkait.")
    
    doc.add_paragraph()
    p = doc.add_paragraph("II. PELAKSANAAN")
    p.style = 'Heading 2'
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Tanggal"
    cells[1].text = "{{tanggal_kegiatan:tanggal_long}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Waktu"
    cells[1].text = "{{waktu_kegiatan}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Tempat"
    cells[1].text = "{{tempat_kegiatan}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Peserta"
    cells[1].text = "{{daftar_peserta}}"
    
    doc.add_paragraph()
    p = doc.add_paragraph("III. HASIL DAN PEMBAHASAN")
    p.style = 'Heading 2'
    doc.add_paragraph("{{hasil_kegiatan}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("IV. KESIMPULAN")
    p.style = 'Heading 2'
    doc.add_paragraph("Kegiatan jamuan tamu berjalan lancar dan mencapai tujuan yang diharapkan.")
    
    doc.add_paragraph()
    add_signature_block(doc, "PPK", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "laporan_jamuan_tamu.docx")
    print("✅ Dibuat: laporan_jamuan_tamu.docx")

def create_sk_kpa_jamuan_tamu():
    """Create SK KPA Jamuan Tamu"""
    doc = Document()
    
    add_header(doc, "SURAT KEPUTUSAN KPA", "{{satker_nama}}")
    
    doc.add_paragraph("NOMOR: {{nomor_sk:text}}")
    doc.add_paragraph()
    doc.add_paragraph("TENTANG:")
    doc.add_paragraph("PERSETUJUAN PELAKSANAAN JAMUAN TAMU")
    doc.add_paragraph()
    
    p = doc.add_paragraph("MENIMBANG :")
    p.style = 'Heading 2'
    
    doc.add_paragraph("a. Bahwa dalam rangka meningkatkan hubungan kerjasama dengan stakeholder;")
    doc.add_paragraph("b. Bahwa perlu diterbitkan surat keputusan tentang persetujuan jamuan tamu.")
    
    doc.add_paragraph()
    p = doc.add_paragraph("MEMUTUSKAN :")
    p.style = 'Heading 2'
    
    doc.add_paragraph("1. Menyetujui pelaksanaan jamuan tamu sebagai berikut:")
    doc.add_paragraph(f"   • Nama Kegiatan    : {{{{nama_kegiatan}}}}")
    doc.add_paragraph(f"   • Tanggal          : {{{{tanggal_kegiatan:tanggal_long}}}}")
    doc.add_paragraph(f"   • Tempat           : {{{{tempat_kegiatan}}}}")
    doc.add_paragraph(f"   • Estimasi Biaya   : {{{{estimasi_biaya:rupiah}}}}")
    
    doc.add_paragraph("2. Biaya jamuan tamu dibebankan pada APBN {{satker_nama}} tahun {{tahun_anggaran}}.")
    
    doc.add_paragraph()
    add_signature_block(doc, "KPA", "{{kpa_nama}}", "{{kpa_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "sk_kpa_jamuan_tamu.docx")
    print("✅ Dibuat: sk_kpa_jamuan_tamu.docx")

# ============================================================================
# DAFTAR HADIR SWAKELOLA
# ============================================================================

def create_daftar_hadir_swakelola():
    """Create Daftar Hadir Swakelola"""
    doc = Document()
    
    add_header(doc, "DAFTAR HADIR KEGIATAN SWAKELOLA", "{{satker_nama}}")
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Nama Kegiatan"
    cells[1].text = "{{nama_kegiatan}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Tanggal"
    cells[1].text = "{{tanggal_kegiatan:tanggal_long}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Tempat"
    cells[1].text = "{{tempat_kegiatan}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Tipe Swakelola"
    cells[1].text = "{{tipe_swakelola}}"
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=16, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "No"
    hdr_cells[1].text = "Nama"
    hdr_cells[2].text = "Jabatan"
    hdr_cells[3].text = "Tanda Tangan"
    
    for i in range(1, 16):
        cells = table.rows[i].cells
        cells[0].text = str(i)
        cells[1].text = ""
        cells[2].text = ""
        cells[3].text = ""
    
    doc.add_paragraph()
    add_signature_block(doc, "Ketua Tim", "{{ketua_tim_nama}}", "{{ketua_tim_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "daftar_hadir_swakelola.docx")
    print("✅ Dibuat: daftar_hadir_swakelola.docx")

# ============================================================================
# KUITANSI TAMBAHAN
# ============================================================================

def create_kuitansi_honorarium():
    """Create Kuitansi Honorarium"""
    doc = Document()
    
    add_header(doc, "KUITANSI HONORARIUM", "{{satker_nama}}")
    
    doc.add_paragraph("Nomor: {{nomor_kuitansi}}")
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Tanggal"
    cells[1].text = "{{tanggal_kuitansi:tanggal_long}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Penerima"
    cells[1].text = "{{penerima_nama}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Jabatan/Keahlian"
    cells[1].text = "{{penerima_jabatan}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Kegiatan"
    cells[1].text = "{{nama_kegiatan}}"
    
    cells = table.rows[4].cells
    cells[0].text = "Honorarium"
    cells[1].text = "{{nominal_honorarium:rupiah}}"
    
    doc.add_paragraph()
    doc.add_paragraph("Telah diterima honorarium sebesar {{nominal_honorarium:terbilang}} ({{nominal_honorarium:rupiah}})")
    doc.add_paragraph()
    
    add_signature_block(doc, "Bendahara", "{{bendahara_nama}}")
    
    doc.save(WORD_TEMPLATES_DIR / "kuitansi_honorarium.docx")
    print("✅ Dibuat: kuitansi_honorarium.docx")

def create_kuitansi_honor_pengelola():
    """Create Kuitansi Honor Pengelola"""
    doc = Document()
    
    add_header(doc, "KUITANSI HONOR PENGELOLA", "{{satker_nama}}")
    
    doc.add_paragraph("Nomor: {{nomor_kuitansi}}")
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Tanggal"
    cells[1].text = "{{tanggal_kuitansi:tanggal_long}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Pengelola"
    cells[1].text = "{{pengelola_nama}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Periode"
    cells[1].text = "{{periode}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Kegiatan"
    cells[1].text = "{{nama_kegiatan}}"
    
    cells = table.rows[4].cells
    cells[0].text = "Honor"
    cells[1].text = "{{nominal_honor:rupiah}}"
    
    doc.add_paragraph()
    doc.add_paragraph("Telah diterima honor sebesar {{nominal_honor:terbilang}} ({{nominal_honor:rupiah}})")
    doc.add_paragraph()
    
    add_signature_block(doc, "Bendahara", "{{bendahara_nama}}")
    
    doc.save(WORD_TEMPLATES_DIR / "kuitansi_honor_pengelola.docx")
    print("✅ Dibuat: kuitansi_honor_pengelola.docx")

# ============================================================================
# DOKUMEN LAINNYA
# ============================================================================

def create_bukti_serah_terima_um():
    """Create Bukti Serah Terima Uang Muka"""
    doc = Document()
    
    add_header(doc, "BUKTI SERAH TERIMA UANG MUKA", "{{satker_nama}}")
    
    doc.add_paragraph("Nomor: {{nomor_bukti}}")
    doc.add_paragraph("Tanggal: {{tanggal_bukti:tanggal_long}}")
    doc.add_paragraph()
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Paket Pekerjaan"
    cells[1].text = "{{nama_paket}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Penyedia"
    cells[1].text = "{{penyedia_nama}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Jumlah Uang Muka"
    cells[1].text = "{{jumlah_um:rupiah}}"
    
    doc.add_paragraph()
    doc.add_paragraph("Bahwa telah diterima uang muka sebesar {{jumlah_um:terbilang}} ({{jumlah_um:rupiah}}) untuk paket {{nama_paket}}.")
    doc.add_paragraph()
    
    doc.add_paragraph("Penyedia:")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("({{penyedia_nama}})")
    
    doc.save(WORD_TEMPLATES_DIR / "bukti_serah_terima_um.docx")
    print("✅ Dibuat: bukti_serah_terima_um.docx")

def create_notulen():
    """Create Notulen Rapat"""
    doc = Document()
    
    add_header(doc, "NOTULEN RAPAT", "{{satker_nama}}")
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Tanggal"
    cells[1].text = "{{tanggal_rapat:tanggal_long}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Waktu"
    cells[1].text = "{{waktu_rapat}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Tempat"
    cells[1].text = "{{tempat_rapat}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Peserta"
    cells[1].text = "{{peserta_rapat}}"
    
    cells = table.rows[4].cells
    cells[0].text = "Agenda"
    cells[1].text = "{{agenda_rapat}}"
    
    doc.add_paragraph()
    
    p = doc.add_paragraph("RINGKASAN PEMBAHASAN")
    p.style = 'Heading 2'
    doc.add_paragraph("{{ringkasan_pembahasan}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("KEPUTUSAN/KESIMPULAN")
    p.style = 'Heading 2'
    doc.add_paragraph("{{keputusan}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("TINDAK LANJUT")
    p.style = 'Heading 2'
    doc.add_paragraph("{{tindak_lanjut}}")
    
    doc.add_paragraph()
    add_signature_block(doc, "Pemimpin Rapat", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "notulen.docx")
    print("✅ Dibuat: notulen.docx")

def create_lembar_permintaan():
    """Create Lembar Permintaan"""
    doc = Document()
    
    add_header(doc, "LEMBAR PERMINTAAN/ORDER", "{{satker_nama}}")
    
    doc.add_paragraph("Nomor: {{nomor_permintaan}}")
    doc.add_paragraph("Tanggal: {{tanggal_permintaan:tanggal_long}}")
    doc.add_paragraph()
    
    p = doc.add_paragraph("DIKIRIM KEPADA:")
    p.style = 'Heading 2'
    
    doc.add_paragraph("Nama: {{penyedia_nama}}")
    doc.add_paragraph("Alamat: {{penyedia_alamat}}")
    doc.add_paragraph("Telepon: {{penyedia_telepon}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("RINCIAN PERMINTAAN")
    p.style = 'Heading 2'
    
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "No"
    hdr_cells[1].text = "Deskripsi Barang"
    hdr_cells[2].text = "Qty"
    hdr_cells[3].text = "Harga Satuan"
    hdr_cells[4].text = "Jumlah"
    
    for i in range(1, 6):
        cells = table.rows[i].cells
        cells[0].text = str(i)
        cells[1].text = ""
        cells[2].text = ""
        cells[3].text = ""
        cells[4].text = ""
    
    doc.add_paragraph()
    
    doc.add_paragraph("Catatan: {{catatan}}")
    doc.add_paragraph()
    
    add_signature_block(doc, "Pemesan", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "lembar_permintaan.docx")
    print("✅ Dibuat: lembar_permintaan.docx")

def create_lpj():
    """Create Laporan Pertanggungjawaban"""
    doc = Document()
    
    add_header(doc, "LAPORAN PERTANGGUNGJAWABAN", "{{satker_nama}}")
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Periode"
    cells[1].text = "{{periode_laporan}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Penanggung Jawab"
    cells[1].text = "{{ppk_nama}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Paket Kerja"
    cells[1].text = "{{nama_paket}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Nilai"
    cells[1].text = "{{nilai_paket:rupiah}}"
    
    doc.add_paragraph()
    
    p = doc.add_paragraph("RINGKASAN KEGIATAN")
    p.style = 'Heading 2'
    doc.add_paragraph("{{ringkasan_kegiatan}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("HASIL YANG DICAPAI")
    p.style = 'Heading 2'
    doc.add_paragraph("{{hasil_dicapai}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("KENDALA DAN SOLUSI")
    p.style = 'Heading 2'
    doc.add_paragraph("{{kendala_solusi}}")
    
    doc.add_paragraph()
    add_signature_block(doc, "PPK", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "lpj.docx")
    print("✅ Dibuat: lpj.docx")

def create_ba_mc():
    """Create BA Monitoring dan Evaluasi"""
    doc = Document()
    
    add_header(doc, "BERITA ACARA MONITORING DAN EVALUASI", "{{satker_nama}}")
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Tanggal"
    cells[1].text = "{{tanggal_monev:tanggal_long}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Paket"
    cells[1].text = "{{nama_paket}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Lokasi"
    cells[1].text = "{{lokasi_pekerjaan}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Progress"
    cells[1].text = "{{progres_persen}}%"
    
    cells = table.rows[4].cells
    cells[0].text = "Kondisi"
    cells[1].text = "{{kondisi_pekerjaan}}"
    
    doc.add_paragraph()
    
    p = doc.add_paragraph("HASIL MONITORING")
    p.style = 'Heading 2'
    doc.add_paragraph("{{hasil_monitoring}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("REKOMENDASI")
    p.style = 'Heading 2'
    doc.add_paragraph("{{rekomendasi}}")
    
    doc.add_paragraph()
    add_signature_block(doc, "Monitor", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "ba_mc.docx")
    print("✅ Dibuat: ba_mc.docx")

def create_monev():
    """Create Laporan Monitoring dan Evaluasi"""
    doc = Document()
    
    add_header(doc, "LAPORAN MONITORING DAN EVALUASI", "{{satker_nama}}")
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Periode"
    cells[1].text = "{{periode_monev}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Paket"
    cells[1].text = "{{nama_paket}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Progress"
    cells[1].text = "{{progres_persen}}%"
    
    cells = table.rows[3].cells
    cells[0].text = "Status"
    cells[1].text = "{{status_paket}}"
    
    doc.add_paragraph()
    
    p = doc.add_paragraph("ANALISIS")
    p.style = 'Heading 2'
    doc.add_paragraph("{{analisis_monev}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("REKOMENDASI")
    p.style = 'Heading 2'
    doc.add_paragraph("{{rekomendasi_monev}}")
    
    doc.add_paragraph()
    add_signature_block(doc, "Monitor", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "monev.docx")
    print("✅ Dibuat: monev.docx")

def create_laporan_kegiatan():
    """Create Laporan Kegiatan"""
    doc = Document()
    
    add_header(doc, "LAPORAN KEGIATAN", "{{satker_nama}}")
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Periode"
    cells[1].text = "{{periode_laporan}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Kegiatan"
    cells[1].text = "{{nama_kegiatan}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Penanggungjawab"
    cells[1].text = "{{ppk_nama}}"
    
    doc.add_paragraph()
    
    p = doc.add_paragraph("DESKRIPSI KEGIATAN")
    p.style = 'Heading 2'
    doc.add_paragraph("{{deskripsi_kegiatan}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("HASIL DAN MANFAAT")
    p.style = 'Heading 2'
    doc.add_paragraph("{{hasil_manfaat}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("KESIMPULAN")
    p.style = 'Heading 2'
    doc.add_paragraph("{{kesimpulan}}")
    
    doc.add_paragraph()
    add_signature_block(doc, "PPK", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "laporan_kegiatan.docx")
    print("✅ Dibuat: laporan_kegiatan.docx")

def create_sk_kpa():
    """Create Surat Keputusan KPA"""
    doc = Document()
    
    add_header(doc, "SURAT KEPUTUSAN KPA", "{{satker_nama}}")
    
    doc.add_paragraph("NOMOR: {{nomor_sk}}")
    doc.add_paragraph()
    doc.add_paragraph("TENTANG:")
    doc.add_paragraph("{{tentang_sk}}")
    doc.add_paragraph()
    
    p = doc.add_paragraph("MENIMBANG:")
    p.style = 'Heading 2'
    doc.add_paragraph("{{menimbang}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("MEMUTUSKAN:")
    p.style = 'Heading 2'
    doc.add_paragraph("{{memutuskan}}")
    
    doc.add_paragraph()
    add_signature_block(doc, "KPA", "{{kpa_nama}}", "{{kpa_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "sk_kpa.docx")
    print("✅ Dibuat: sk_kpa.docx")

def create_sskk():
    """Create SSKK"""
    doc = Document()
    
    add_header(doc, "SURAT SETORAN KEMASAN/KARUNG", "{{satker_nama}}")
    
    doc.add_paragraph("Nomor: {{nomor_sskk}}")
    doc.add_paragraph("Tanggal: {{tanggal_sskk:tanggal_long}}")
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Jenis Kemasan"
    cells[1].text = "{{jenis_kemasan}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Jumlah"
    cells[1].text = "{{jumlah_kemasan}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Kondisi"
    cells[1].text = "{{kondisi_kemasan}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Catatan"
    cells[1].text = "{{catatan_kemasan}}"
    
    doc.add_paragraph()
    
    add_signature_block(doc, "Pengguna Barang", "{{ppk_nama}}", "{{ppk_nip}}")
    
    doc.save(WORD_TEMPLATES_DIR / "sskk.docx")
    print("✅ Dibuat: sskk.docx")

def create_ssuk():
    """Create SSUK"""
    doc = Document()
    
    add_header(doc, "SURAT SETORAN UANG KEMBALIAN", "{{satker_nama}}")
    
    doc.add_paragraph("Nomor: {{nomor_ssuk}}")
    doc.add_paragraph("Tanggal: {{tanggal_ssuk:tanggal_long}}")
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = "Paket"
    cells[1].text = "{{nama_paket}}"
    
    cells = table.rows[1].cells
    cells[0].text = "Nominal Kembalian"
    cells[1].text = "{{nominal_kembalian:rupiah}}"
    
    cells = table.rows[2].cells
    cells[0].text = "Terbilang"
    cells[1].text = "{{nominal_kembalian:terbilang}}"
    
    cells = table.rows[3].cells
    cells[0].text = "Keterangan"
    cells[1].text = "{{keterangan_kembalian}}"
    
    doc.add_paragraph()
    doc.add_paragraph("Telah disetor uang sebesar {{nominal_kembalian:rupiah}} ke rekening Bendahara")
    doc.add_paragraph()
    
    add_signature_block(doc, "Bendahara", "{{bendahara_nama}}")
    
    doc.save(WORD_TEMPLATES_DIR / "ssuk.docx")
    print("✅ Dibuat: ssuk.docx")

def create_spby():
    """Create SPBY"""
    doc = Document()
    
    add_header(doc, "SURAT PENAWARAN BARANG/JASA", "{{satker_nama}}")
    
    doc.add_paragraph("Nomor: {{nomor_spby}}")
    doc.add_paragraph("Tanggal: {{tanggal_spby:tanggal_long}}")
    doc.add_paragraph()
    
    p = doc.add_paragraph("DIKIRIM KEPADA:")
    p.style = 'Heading 2'
    
    doc.add_paragraph("Nama: {{penyedia_nama}}")
    doc.add_paragraph("Alamat: {{penyedia_alamat}}")
    
    doc.add_paragraph()
    p = doc.add_paragraph("PENAWARAN")
    p.style = 'Heading 2'
    
    doc.add_paragraph("Kami menawarkan barang/jasa sebagai berikut:")
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "No"
    hdr_cells[1].text = "Deskripsi"
    hdr_cells[2].text = "Harga"
    hdr_cells[3].text = "Jumlah"
    
    for i in range(1, 5):
        cells = table.rows[i].cells
        cells[0].text = str(i)
        cells[1].text = ""
        cells[2].text = ""
        cells[3].text = ""
    
    doc.add_paragraph()
    
    doc.add_paragraph("Penawaran berlaku hingga: {{tanggal_berlaku_penawaran:tanggal_long}}")
    doc.add_paragraph()
    
    add_signature_block(doc, "Penyedia", "{{penyedia_nama}}")
    
    doc.save(WORD_TEMPLATES_DIR / "spby.docx")
    print("✅ Dibuat: spby.docx")

if __name__ == "__main__":
    print("=" * 70)
    print("MEMBUAT TEMPLATE DRAFT LENGKAP")
    print("=" * 70)
    print()
    
    try:
        # Jamuan Tamu
        create_daftar_hadir_jamuan_tamu()
        create_nota_dinas_jamuan_tamu()
        create_kuitansi_jamuan_tamu()
        create_laporan_jamuan_tamu()
        create_sk_kpa_jamuan_tamu()
        
        # Swakelola
        create_daftar_hadir_swakelola()
        
        # Kuitansi
        create_kuitansi_honorarium()
        create_kuitansi_honor_pengelola()
        
        # Dokumen Lainnya
        create_bukti_serah_terima_um()
        create_notulen()
        create_lembar_permintaan()
        create_lpj()
        create_ba_mc()
        create_monev()
        create_laporan_kegiatan()
        create_sk_kpa()
        create_sskk()
        create_ssuk()
        create_spby()
        
        print()
        print("=" * 70)
        print("✅ SEMUA TEMPLATE BERHASIL DIBUAT!")
        print("=" * 70)
        print()
        print("Total: 22 template draft baru + 4 sebelumnya = 26 template")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
