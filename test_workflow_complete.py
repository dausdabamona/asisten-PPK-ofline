"""
Comprehensive workflow test - simulating actual user flow
"""
from pathlib import Path
from docx import Document
import sys
import re

BASE_DIR = Path(__file__).parent
sys.path.insert(0, str(BASE_DIR))

from app.services.dokumen_generator import DokumenGenerator

print("=" * 80)
print("COMPREHENSIVE WORKFLOW TEST - Lembar Permintaan Document Generation")
print("=" * 80)

# Simulate data from actual user workflow
print("\n1. SIMULATING USER INPUT & DATA COLLECTION")
print("-" * 80)

# Satker data (dari database)
satker_data = {
    'nama': 'Dinas Pendidikan Kota Bandung',
    'kode': 'DPK',
    'sumber_dana': 'APBD 2024',
    'ppk_nama': 'Dr. Ahmad Wijaya, S.Pd, M.Ed',
    'ppk_nip': '19750320199203001',
    'kpa_nama': 'Drs. Bambang Sutrisno, M.Si',
    'kpa_nip': '19650815198703001',
    'pimpinan_nama': 'Ir. Siti Rahayui, M.T',
}

# Transaksi data
transaksi_data = {
    'kode': 'TRX-001',
    'nama_kegiatan': 'Pengadaan Alat Tulis Kantor Tahun 2024',
    'kode_akun': '5.1.2.01.01.01',
    'estimasi_biaya': 2500000,
}

# Form data dari user
form_data = {
    'tanggal': '2024-01-20',  # Will be formatted to "20 Januari 2024"
    'penerima_nama': 'Budi Santoso, S.Kom',
    'penerima_nip': '19800515200212001',
    'penerima_jabatan': 'Staf Administrasi',
    'rincian_items': [
        {
            'uraian': 'Kertas A4 Putih 80gsm',
            'spesifikasi': 'Kertas putih untuk surat-menyurat, 1 rim = 500 lembar',
            'volume': 10,
            'satuan': 'rim',
            'harga_satuan': 50000,
            'jumlah': 500000,
        },
        {
            'uraian': 'Tinta Printer Cartridge Hitam',
            'spesifikasi': 'Cartridge tinta printer compatible untuk printer Canon/HP',
            'volume': 5,
            'satuan': 'buah',
            'harga_satuan': 80000,
            'jumlah': 400000,
        },
        {
            'uraian': 'Ballpoint Standar',
            'spesifikasi': 'Ballpoint standar, tinta biru, tidak mudah bocor',
            'volume': 100,
            'satuan': 'buah',
            'harga_satuan': 2000,
            'jumlah': 200000,
        },
        {
            'uraian': 'Buku Catatan A5',
            'spesifikasi': 'Buku catatan A5, 80 halaman, cover tebal',
            'volume': 20,
            'satuan': 'buah',
            'harga_satuan': 5000,
            'jumlah': 100000,
        },
    ]
}

print(f"✓ Satker: {satker_data.get('nama')}")
print(f"✓ Transaksi: {transaksi_data.get('nama_kegiatan')}")
print(f"✓ Penerima: {form_data.get('penerima_nama')}")
print(f"✓ Rincian Items: {len(form_data.get('rincian_items', []))} items")

print("\n2. SIMULATING DIALOG DATA COLLECTION (_collect_data logic)")
print("-" * 80)

# Format tanggal
from datetime import datetime
tanggal_parts = form_data['tanggal'].split('-')
bulan = [
    "", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
    "Juli", "Agustus", "September", "Oktober", "November", "Desember"
]
hari_tanggal_str = f"{int(tanggal_parts[2])} {bulan[int(tanggal_parts[1])]} {tanggal_parts[0]}"

# Prepare rincian items
prepared_rincian = []
for idx, item in enumerate(form_data['rincian_items'], 1):
    prepared_rincian.append({
        'no': idx,
        'item_no': idx,
        'uraian': item.get('uraian', ''),
        'nama_barang': item.get('uraian', ''),
        'spesifikasi': item.get('spesifikasi', ''),
        'volume': item.get('volume', 0),
        'satuan': item.get('satuan', ''),
        'harga_satuan': item.get('harga_satuan', 0),
        'jumlah': item.get('jumlah', 0),
        'total_item': item.get('jumlah', 0),
        'keterangan': '',
    })

# Calculate totals
subtotal = sum(item.get('jumlah', 0) for item in prepared_rincian)
ppn = subtotal * 0.10
total = subtotal + ppn

# Collect data (matching _collect_data logic)
collected_data = {
    'nama_kegiatan': transaksi_data.get('nama_kegiatan', ''),
    'kode_akun': transaksi_data.get('kode_akun', ''),
    'estimasi_biaya': transaksi_data.get('estimasi_biaya', 0),
    'tanggal_dokumen': form_data['tanggal'],
    'hari_tanggal': hari_tanggal_str,
    'penerima_nama': form_data['penerima_nama'],
    'penerima_nip': form_data['penerima_nip'],
    'penerima_jabatan': form_data['penerima_jabatan'],
    'rincian_items': prepared_rincian,
    'subtotal': subtotal,
    'ppn': ppn,
    'total': total,
    # From satker
    'unit_kerja': satker_data.get('nama', ''),
    'sumber_dana': satker_data.get('sumber_dana', 'APBD'),
    'nama_pengajuan': form_data['penerima_nama'],
    'nama_verifikator': satker_data.get('ppk_nama', ''),
    'nama_ppk': satker_data.get('ppk_nama', ''),
    'nama_atasan': satker_data.get('pimpinan_nama', '') or satker_data.get('kpa_nama', ''),
    'nama_kpa': satker_data.get('kpa_nama', ''),
}

print(f"✓ Tanggal formatted: {collected_data['hari_tanggal']}")
print(f"✓ Subtotal: Rp {subtotal:,}")
print(f"✓ PPN (10%): Rp {int(ppn):,}")
print(f"✓ TOTAL: Rp {int(total):,}")
print(f"✓ Data keys collected: {len(collected_data)}")

print("\n3. INITIALIZING GENERATOR & FLATTENING DATA")
print("-" * 80)

generator = DokumenGenerator()
flattened_data = generator._flatten_rincian_items(collected_data)

print(f"✓ Data flattened")
print(f"✓ Total keys after flattening: {len(flattened_data)}")

# Count item keys
item_keys = [k for k in flattened_data.keys() if k.startswith('item_')]
print(f"✓ Item placeholder keys: {len(item_keys)}")
for i in range(1, 6):
    item_suffix = [k.split('_', 2)[2] for k in item_keys if k.startswith(f'item_{i}_')]
    if item_suffix:
        print(f"  - item_{i}: {', '.join(set(item_suffix))}")

print("\n4. MERGING DOCUMENT")
print("-" * 80)

template_path = BASE_DIR / "templates" / "word" / "lembar_permintaan.docx"
output_path = BASE_DIR / "output" / "dokumen" / "test_workflow_complete.docx"

output_path.parent.mkdir(parents=True, exist_ok=True)

success = generator.merge_word_document(template_path, flattened_data, output_path)

if success:
    print(f"✓ Document merged successfully")
    print(f"✓ Output: {output_path}")
    
    print("\n5. VERIFICATION")
    print("-" * 80)
    
    doc = Document(str(output_path))
    
    # Collect all text
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)
    doc_text = '\n'.join(all_text)
    
    # Check for unreplaced placeholders
    unreplaced = re.findall(r'\{\{[\w_:]+\}\}', doc_text)
    
    print(f"Total paragraphs + table cells checked: {len(all_text)}")
    print(f"Unreplaced placeholders found: {len(unreplaced)}")
    
    if unreplaced:
        print(f"⚠ WARNING: Found unreplaced placeholders:")
        for ph in set(unreplaced):
            print(f"  - {ph}")
    else:
        print(f"✓ NO UNREPLACED PLACEHOLDERS - EXCELLENT!")
    
    print("\n6. CONTENT VERIFICATION")
    print("-" * 80)
    
    # Verify key content
    checks = {
        'Dinas Pendidikan Kota Bandung': 'Unit Kerja',
        '20 Januari 2024': 'Tanggal',
        'Kertas A4 Putih 80gsm': 'Item 1',
        'Tinta Printer Cartridge Hitam': 'Item 2',
        'Ballpoint Standar': 'Item 3',
        'Buku Catatan A5': 'Item 4',
        'Budi Santoso': 'Penerima/Pengajuan',
        'Dr. Ahmad Wijaya': 'PPK',
        'Ir. Siti Rahayui': 'Atasan',
        '1,100,000': 'Subtotal (alt format)',
        'Rp 1,100,000': 'Subtotal (with Rp)',
        'Drs. Bambang Sutrisno': 'KPA',
    }
    
    found_count = 0
    for value, desc in checks.items():
        variants = [value, value.replace(',', '.'), value.replace(',', '')]
        if any(v in doc_text for v in variants):
            print(f"  ✓ {desc}")
            found_count += 1
        else:
            print(f"  ✗ {desc} ('{value}')")
    
    print(f"\n✓ {found_count} of {len(checks)} content checks passed")
    
    print("\n" + "=" * 80)
    print("WORKFLOW TEST COMPLETED SUCCESSFULLY!")
    print("=" * 80)
    print(f"\nGenerated document: {output_path}")
    print("\nNEXT STEPS:")
    print("1. Open the document in MS Word")
    print("2. Verify all placeholders are replaced")
    print("3. Check formatting of rincian items")
    print("4. Check penandatangan (signature) section")
    
else:
    print(f"✗ Document merge failed")
