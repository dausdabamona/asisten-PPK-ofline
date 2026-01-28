"""
Script untuk membuat template Kuitansi untuk SPM Lainnya:
1. Kuitansi Uang Muka (SPM Lainnya)
2. Kuitansi Rampung (SPM Lainnya)

Digunakan untuk pengajuan SPM Lainnya seperti:
- Honorarium
- Jamuan Tamu
- PJLP
- Lainnya
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

TEMPLATES_DIR = "templates/word"


def set_cell_shading(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_kop_surat(doc):
    """Add standard letter header"""
    kop1 = doc.add_paragraph()
    run = kop1.add_run("{{kementerian}}")
    run.bold = True
    run.font.size = Pt(12)
    kop1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    kop2 = doc.add_paragraph()
    run = kop2.add_run("{{eselon1}}")
    run.bold = True
    run.font.size = Pt(11)
    kop2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    kop3 = doc.add_paragraph()
    run = kop3.add_run("{{satker_nama}}")
    run.bold = True
    run.font.size = Pt(12)
    kop3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    alamat = doc.add_paragraph()
    run = alamat.add_run("{{satker_alamat}}, {{satker_kota}}")
    run.font.size = Pt(10)
    alamat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    border = doc.add_paragraph()
    run = border.add_run("═" * 75)
    run.font.size = Pt(10)
    border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()


def create_kuitansi_uang_muka_spm_lainnya():
    """Kuitansi Uang Muka untuk SPM Lainnya"""
    doc = Document()
    
    # Header
    add_kop_surat(doc)
    
    # Judul
    judul = doc.add_paragraph()
    run = judul.add_run("KUITANSI/TANDA TERIMA")
    run.bold = True
    run.font.size = Pt(12)
    judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_judul = doc.add_paragraph()
    run = sub_judul.add_run("UANG MUKA")
    run.bold = True
    run.font.size = Pt(11)
    sub_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    jenis = doc.add_paragraph()
    run = jenis.add_run("{{jenis_pembayaran_nama}}")
    run.bold = True
    run.font.size = Pt(10)
    jenis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # A. DATA DIPA
    p_dipa = doc.add_paragraph()
    run = p_dipa.add_run("A. DATA DIPA")
    run.bold = True
    run.font.size = Pt(11)
    
    table_dipa = doc.add_table(rows=6, cols=2)
    table_dipa.style = 'Table Grid'
    
    # Set column widths
    for row in table_dipa.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)
    
    dipa_data = [
        ("Kode Satker", "{{satker_kode}}"),
        ("Nama Satker", "{{satker_nama}}"),
        ("Tahun Anggaran", "{{tahun_anggaran}}"),
        ("Sumber Dana", "{{sumber_dana}}"),
        ("Kode Akun/MAK", "{{kode_akun}}"),
        ("Kegiatan/Keperluan", "{{perihal}}"),
    ]
    
    for i, (label, value) in enumerate(dipa_data):
        row = table_dipa.rows[i]
        row.cells[0].text = label
        row.cells[1].text = ": " + value
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # B. DATA KUITANSI
    p_kuitansi = doc.add_paragraph()
    run = p_kuitansi.add_run("B. DATA KUITANSI/PENERIMAAN UANG MUKA")
    run.bold = True
    run.font.size = Pt(11)
    
    table_kuitansi = doc.add_table(rows=5, cols=2)
    table_kuitansi.style = 'Table Grid'
    
    for row in table_kuitansi.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)
    
    kuitansi_data = [
        ("Nomor Kuitansi", "{{nomor_kuitansi}}"),
        ("Sudah Terima Dari", "Bendahara Pengeluaran {{satker_nama}}"),
        ("Uang Sebanyak", "Rp {{uang_muka_format}}"),
        ("Terbilang", "{{uang_muka_terbilang}}"),
        ("Untuk Keperluan", "{{perihal}}"),
    ]
    
    for i, (label, value) in enumerate(kuitansi_data):
        row = table_kuitansi.rows[i]
        row.cells[0].text = label
        row.cells[1].text = ": " + value
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # C. DATA PENERIMA
    p_penerima = doc.add_paragraph()
    run = p_penerima.add_run("C. DATA PENERIMA UANG MUKA")
    run.bold = True
    run.font.size = Pt(11)
    
    table_penerima = doc.add_table(rows=4, cols=2)
    table_penerima.style = 'Table Grid'
    
    for row in table_penerima.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)
    
    penerima_data = [
        ("Nama Penerima", "{{penerima_nama}}"),
        ("NIP/Identitas", "{{penerima_nip}}"),
        ("Jabatan", "{{penerima_jabatan}}"),
        ("Tanda Tangan", ""),
    ]
    
    for i, (label, value) in enumerate(penerima_data):
        row = table_penerima.rows[i]
        row.cells[0].text = label
        row.cells[1].text = ": " + value
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # D. DASAR PEMBAYARAN
    p_dasar = doc.add_paragraph()
    run = p_dasar.add_run("D. DASAR PEMBAYARAN")
    run.bold = True
    run.font.size = Pt(11)
    
    table_dasar = doc.add_table(rows=2, cols=2)
    table_dasar.style = 'Table Grid'
    
    for row in table_dasar.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)
    
    dasar_data = [
        ("SK KPA Nomor", "{{nomor_sk_kpa}} tanggal {{tanggal_sk_kpa}}"),
        ("Dasar Pembayaran", "{{dasar_pembayaran}}"),
    ]
    
    for i, (label, value) in enumerate(dasar_data):
        row = table_dasar.rows[i]
        row.cells[0].text = label
        row.cells[1].text = ": " + value
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # E. TANDA TANGAN
    p_ttd = doc.add_paragraph()
    run = p_ttd.add_run("E. PENGESAHAN PEMBAYARAN")
    run.bold = True
    run.font.size = Pt(11)
    
    table_ttd = doc.add_table(rows=4, cols=3)
    table_ttd.style = 'Table Grid'
    
    # Headers
    header_cells = table_ttd.rows[0].cells
    header_labels = [
        "Setuju Dibayar,\nPejabat Pembuat Komitmen",
        "Lunas Dibayar,\nBendahara Pengeluaran",
        f"{{{{satker_kota}}}}, {{{{tanggal_kuitansi}}}}\nYang Menerima,"
    ]
    
    for i, label in enumerate(header_labels):
        header_cells[i].text = label
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Space for signatures (row 1-2 are blank)
    for row_idx in range(1, 3):
        for col_idx in range(3):
            table_ttd.rows[row_idx].cells[col_idx].text = ""
    
    # Names and NIPs (row 3)
    name_cells = table_ttd.rows[3].cells
    names = [
        "{{ppk_nama}}\nNIP. {{ppk_nip}}",
        "{{bendahara_nama}}\nNIP. {{bendahara_nip}}",
        "{{penerima_nama}}\nNIP. {{penerima_nip}}"
    ]
    
    for i, name in enumerate(names):
        name_cells[i].text = name
        for para in name_cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(os.path.join(TEMPLATES_DIR, "kuitansi_uang_muka_spm_lainnya.docx"))
    print("✓ Created: kuitansi_uang_muka_spm_lainnya.docx")


def create_kuitansi_rampung_spm_lainnya():
    """Kuitansi Rampung untuk SPM Lainnya"""
    doc = Document()
    
    # Header
    add_kop_surat(doc)
    
    # Judul
    judul = doc.add_paragraph()
    run = judul.add_run("KUITANSI/BUKTI PEMBAYARAN")
    run.bold = True
    run.font.size = Pt(12)
    judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_judul = doc.add_paragraph()
    run = sub_judul.add_run("RAMPUNG (PEMBAYARAN FINAL)")
    run.bold = True
    run.font.size = Pt(11)
    sub_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    jenis = doc.add_paragraph()
    run = jenis.add_run("{{jenis_pembayaran_nama}}")
    run.bold = True
    run.font.size = Pt(10)
    jenis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # A. DATA DIPA
    p_dipa = doc.add_paragraph()
    run = p_dipa.add_run("A. DATA DIPA")
    run.bold = True
    run.font.size = Pt(11)
    
    table_dipa = doc.add_table(rows=6, cols=2)
    table_dipa.style = 'Table Grid'
    
    for row in table_dipa.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)
    
    dipa_data = [
        ("Kode Satker", "{{satker_kode}}"),
        ("Nama Satker", "{{satker_nama}}"),
        ("Tahun Anggaran", "{{tahun_anggaran}}"),
        ("Sumber Dana", "{{sumber_dana}}"),
        ("Kode Akun/MAK", "{{kode_akun}}"),
        ("Kegiatan/Keperluan", "{{perihal}}"),
    ]
    
    for i, (label, value) in enumerate(dipa_data):
        row = table_dipa.rows[i]
        row.cells[0].text = label
        row.cells[1].text = ": " + value
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # B. DATA PENERIMA
    p_penerima = doc.add_paragraph()
    run = p_penerima.add_run("B. DATA PENERIMA PEMBAYARAN")
    run.bold = True
    run.font.size = Pt(11)
    
    table_penerima = doc.add_table(rows=4, cols=2)
    table_penerima.style = 'Table Grid'
    
    for row in table_penerima.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)
    
    penerima_data = [
        ("Nama Penerima", "{{penerima_nama}}"),
        ("NIP/Identitas", "{{penerima_nip}}"),
        ("Jabatan", "{{penerima_jabatan}}"),
        ("Tanda Tangan", ""),
    ]
    
    for i, (label, value) in enumerate(penerima_data):
        row = table_penerima.rows[i]
        row.cells[0].text = label
        row.cells[1].text = ": " + value
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # C. RINCIAN PEMBAYARAN
    p_rincian = doc.add_paragraph()
    run = p_rincian.add_run("C. RINCIAN PEMBAYARAN")
    run.bold = True
    run.font.size = Pt(11)
    
    table_rincian = doc.add_table(rows=5, cols=3)
    table_rincian.style = 'Table Grid'
    
    # Header
    header = table_rincian.rows[0]
    header.cells[0].text = "Uraian"
    header.cells[1].text = "Jumlah"
    header.cells[2].text = "Keterangan"
    
    for cell in header.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    rincian_items = [
        ("Nilai Uang Muka", "{{uang_muka_format}}", ""),
        ("Biaya Tambahan/Pengganti", "{{biaya_tambahan_format}}", ""),
        ("TOTAL BIAYA", "{{total_biaya_format}}", ""),
        ("Uang Muka Sebelumnya", "{{uang_muka_format}}", ""),
    ]
    
    for i, (uraian, jumlah, ket) in enumerate(rincian_items):
        row = table_rincian.rows[i + 1]
        row.cells[0].text = uraian
        row.cells[1].text = jumlah
        row.cells[2].text = ket
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # D. RINGKASAN PEMBAYARAN
    p_ringkas = doc.add_paragraph()
    run = p_ringkas.add_run("D. RINGKASAN PEMBAYARAN")
    run.bold = True
    run.font.size = Pt(11)
    
    table_ringkas = doc.add_table(rows=3, cols=2)
    table_ringkas.style = 'Table Grid'
    
    for row in table_ringkas.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)
    
    ringkas_data = [
        ("Total Biaya", "{{total_biaya_format}}"),
        ("Dikurangi Uang Muka", "{{uang_muka_format}}"),
        ("PEMBAYARAN RAMPUNG", "{{pembayaran_rampung_format}}"),
    ]
    
    for i, (label, value) in enumerate(ringkas_data):
        row = table_ringkas.rows[i]
        row.cells[0].text = label
        row.cells[1].text = ": " + value
        
        # Bold untuk row terakhir
        is_last = (i == len(ringkas_data) - 1)
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].bold = is_last
                para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Terbilang
    p_terbilang = doc.add_paragraph()
    p_terbilang.add_run("Terbilang: ").bold = True
    p_terbilang.add_run("{{pembayaran_rampung_terbilang}}")
    for run in p_terbilang.runs:
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # E. DASAR PEMBAYARAN
    p_dasar = doc.add_paragraph()
    run = p_dasar.add_run("E. DASAR PEMBAYARAN")
    run.bold = True
    run.font.size = Pt(11)
    
    table_dasar = doc.add_table(rows=2, cols=2)
    table_dasar.style = 'Table Grid'
    
    for row in table_dasar.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)
    
    dasar_data = [
        ("SK KPA Nomor", "{{nomor_sk_kpa}} tanggal {{tanggal_sk_kpa}}"),
        ("Dasar Pembayaran", "{{dasar_pembayaran}}"),
    ]
    
    for i, (label, value) in enumerate(dasar_data):
        row = table_dasar.rows[i]
        row.cells[0].text = label
        row.cells[1].text = ": " + value
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # F. PENGESAHAN PEMBAYARAN
    p_ttd = doc.add_paragraph()
    run = p_ttd.add_run("F. PENGESAHAN PEMBAYARAN")
    run.bold = True
    run.font.size = Pt(11)
    
    table_ttd = doc.add_table(rows=4, cols=3)
    table_ttd.style = 'Table Grid'
    
    # Headers
    header_cells = table_ttd.rows[0].cells
    header_labels = [
        "Setuju Dibayar,\nPejabat Pembuat Komitmen",
        "Lunas Dibayar,\nBendahara Pengeluaran",
        f"{{{{satker_kota}}}}, {{{{tanggal_kuitansi}}}}\nYang Menerima,"
    ]
    
    for i, label in enumerate(header_labels):
        header_cells[i].text = label
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Space for signatures (row 1-2 are blank)
    for row_idx in range(1, 3):
        for col_idx in range(3):
            table_ttd.rows[row_idx].cells[col_idx].text = ""
    
    # Names and NIPs (row 3)
    name_cells = table_ttd.rows[3].cells
    names = [
        "{{ppk_nama}}\nNIP. {{ppk_nip}}",
        "{{bendahara_nama}}\nNIP. {{bendahara_nip}}",
        "{{penerima_nama}}\nNIP. {{penerima_nip}}"
    ]
    
    for i, name in enumerate(names):
        name_cells[i].text = name
        for para in name_cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(os.path.join(TEMPLATES_DIR, "kuitansi_rampung_spm_lainnya.docx"))
    print("✓ Created: kuitansi_rampung_spm_lainnya.docx")


def main():
    """Main function"""
    print("=" * 70)
    print("MEMBUAT TEMPLATE KUITANSI UNTUK SPM LAINNYA")
    print("=" * 70)
    
    try:
        create_kuitansi_uang_muka_spm_lainnya()
        create_kuitansi_rampung_spm_lainnya()
        
        print("\n" + "=" * 70)
        print("✓ BERHASIL! Template kuitansi SPM Lainnya telah dibuat:")
        print("  1. kuitansi_uang_muka_spm_lainnya.docx")
        print("  2. kuitansi_rampung_spm_lainnya.docx")
        print("=" * 70)
        
    except Exception as e:
        print(f"\n✗ ERROR: {str(e)}")
        raise


if __name__ == "__main__":
    main()
