#!/usr/bin/env python3
"""
Fix SPR and SPPR templates:
1. Fix kop surat (use proper format)
2. Fix date format to Indonesian format (tanggal-Bulan-tahun)
3. Fix pejabat to only use PPK atas nama KPA
4. Remove tata cara pengisian section
"""

import sys
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_DIR = Path(__file__).parent.parent / "templates" / "word"


def create_spr_template():
    """Create new SPR template with proper format."""
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # === KOP SURAT ===
    # Will be replaced by satker data
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kop.add_run("{{kop_surat}}")
    run.bold = True
    run.font.size = Pt(14)

    # Separator line
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === TITLE ===
    doc.add_paragraph()  # Spacing

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SURAT PENDEBITAN REKENING (SPR)")
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)

    # Nomor dan Tanggal
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Nomor: {{nomor_dokumen}}    Tanggal: {{tanggal_dokumen:tanggal}}")

    doc.add_paragraph()  # Spacing

    # === KEPADA ===
    kepada = doc.add_paragraph()
    kepada.add_run("Kepada Yth.").bold = True
    doc.add_paragraph("Pimpinan {{nama_bank}}")
    doc.add_paragraph("di tempat")

    doc.add_paragraph()  # Spacing

    # === ISI SURAT ===
    isi = doc.add_paragraph()
    isi.add_run("\tSaya yang bertanda tangan di bawah ini selaku Pejabat Pembuat Komitmen atas nama Kuasa Pengguna Anggaran, memerintahkan Bendahara Pengeluaran agar melakukan pendebitan tunai melalui teller bank yang Saudara pimpin dengan keterangan sebagai berikut:")

    doc.add_paragraph()  # Spacing

    # === TABEL DATA ===
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    data_rows = [
        ("Nomor Rekening", "{{nomor_rekening}}"),
        ("Nama Rekening", "{{nama_rekening}}"),
        ("Sejumlah", "Rp {{nilai:rupiah}}"),
        ("Terbilang", "{{nilai:terbilang}}"),
        ("Hari/Tanggal", "{{tanggal_pencairan:tanggal}}"),
    ]

    for i, (label, value) in enumerate(data_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"
        # Set column widths
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()  # Spacing

    # === PENUTUP ===
    penutup1 = doc.add_paragraph()
    penutup1.add_run("Berkenaan dengan hal tersebut, mohon bantuan Saudara untuk membantu kelancaran transaksi dimaksud.")

    penutup2 = doc.add_paragraph()
    penutup2.add_run("Demikian disampaikan, atas bantuan dan kerja sama yang baik diucapkan terima kasih.")

    doc.add_paragraph()  # Spacing

    # === TANDA TANGAN ===
    ttd_table = doc.add_table(rows=4, cols=1)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # Tempat dan tanggal
    cell0 = ttd_table.rows[0].cells[0]
    cell0.text = "{{kota}}, {{tanggal_dokumen:tanggal}}"
    cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Jabatan
    cell1 = ttd_table.rows[1].cells[0]
    cell1.text = "Pejabat Pembuat Komitmen"
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Spacing untuk tanda tangan
    cell2 = ttd_table.rows[2].cells[0]
    cell2.text = "\n\n\n"

    # Nama dan NIP
    cell3 = ttd_table.rows[3].cells[0]
    p = cell3.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{nama_ppk}}")
    run.bold = True
    cell3.add_paragraph("NIP. {{nip_ppk}}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = TEMPLATE_DIR / "spr.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


def create_sppr_template():
    """Create new SPPR template with proper format."""
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # === KOP SURAT ===
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kop.add_run("{{kop_surat}}")
    run.bold = True
    run.font.size = Pt(14)

    # Separator line
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === TITLE ===
    doc.add_paragraph()  # Spacing

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SURAT PERINTAH PENDEBITAN REKENING (SPPR)")
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)

    # Nomor dan Tanggal
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Nomor: {{nomor_dokumen}}    Tanggal: {{tanggal_dokumen:tanggal}}")

    doc.add_paragraph()  # Spacing

    # === KEPADA ===
    kepada = doc.add_paragraph()
    kepada.add_run("Kepada Yth.").bold = True
    doc.add_paragraph("Pimpinan {{nama_bank}}")
    doc.add_paragraph("di tempat")

    doc.add_paragraph()  # Spacing

    # === ISI SURAT ===
    isi = doc.add_paragraph()
    isi.add_run("\tSaya yang bertanda tangan di bawah ini selaku Pejabat Pembuat Komitmen atas nama Kuasa Pengguna Anggaran, memerintahkan agar dilakukan pendebitan rekening menggunakan kartu debit dengan keterangan sebagai berikut:")

    doc.add_paragraph()  # Spacing

    # === TABEL DATA ===
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    data_rows = [
        ("Nomor Rekening", "{{nomor_rekening}}"),
        ("Nama Rekening", "{{nama_rekening}}"),
        ("Sejumlah", "Rp {{nilai:rupiah}}"),
        ("Terbilang", "{{nilai:terbilang}}"),
        ("Hari/Tanggal", "{{tanggal_pencairan:tanggal}}"),
    ]

    for i, (label, value) in enumerate(data_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()  # Spacing

    # === PENUTUP ===
    penutup1 = doc.add_paragraph()
    penutup1.add_run("Berkenaan dengan hal tersebut, mohon bantuan Saudara untuk membantu kelancaran transaksi dimaksud.")

    penutup2 = doc.add_paragraph()
    penutup2.add_run("Demikian disampaikan, atas bantuan dan kerja sama yang baik diucapkan terima kasih.")

    doc.add_paragraph()  # Spacing

    # === TANDA TANGAN ===
    ttd_table = doc.add_table(rows=4, cols=1)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # Tempat dan tanggal
    cell0 = ttd_table.rows[0].cells[0]
    cell0.text = "{{kota}}, {{tanggal_dokumen:tanggal}}"
    cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Jabatan
    cell1 = ttd_table.rows[1].cells[0]
    cell1.text = "Pejabat Pembuat Komitmen"
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Spacing untuk tanda tangan
    cell2 = ttd_table.rows[2].cells[0]
    cell2.text = "\n\n\n"

    # Nama dan NIP
    cell3 = ttd_table.rows[3].cells[0]
    p = cell3.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{nama_ppk}}")
    run.bold = True
    cell3.add_paragraph("NIP. {{nip_ppk}}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = TEMPLATE_DIR / "sppr.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Fixing SPR and SPPR templates...")
    print("-" * 50)

    create_spr_template()
    create_sppr_template()

    print("-" * 50)
    print("Done! Templates updated successfully.")
