"""
Script untuk update template Kuitansi Uang Muka dan Kuitansi Rampung
dengan data DIPA yang lengkap
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates", "word")


def add_heading(doc, text, level=1, center=False):
    """Add heading with custom style"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_paragraph(doc, text, bold=False, center=False):
    """Add a paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def create_kuitansi_uang_muka():
    """Kuitansi Uang Muka Perjalanan Dinas dengan Data DIPA"""
    doc = Document()

    # Header Kop Surat
    add_heading(doc, "{{kementerian}}", 1, True)
    add_heading(doc, "{{eselon1}}", 2, True)
    add_heading(doc, "{{satker_nama}}", 2, True)

    doc.add_paragraph()
    add_heading(doc, "KUITANSI/TANDA TERIMA", 1, True)
    add_heading(doc, "UANG MUKA PERJALANAN DINAS", 2, True)

    doc.add_paragraph()

    # ===== DATA DIPA =====
    add_paragraph(doc, "A. DATA DIPA", bold=True)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    dipa_items = [
        ("Kode Satker", ": {{satker_kode}}"),
        ("Nama Satker", ": {{satker_nama}}"),
        ("Tahun Anggaran", ": {{tahun_anggaran}}"),
        ("Sumber Dana", ": {{sumber_dana}}"),
        ("Kode Akun/MAK", ": {{kode_akun}}"),
        ("Kegiatan", ": {{nama_paket}}"),
    ]
    for i, (label, value) in enumerate(dipa_items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()

    # ===== DATA KUITANSI =====
    add_paragraph(doc, "B. DATA KUITANSI", bold=True)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    items = [
        ("Nomor Bukti", ": {{nomor_sppd}}"),
        ("Sudah Terima Dari", ": Bendahara Pengeluaran {{satker_nama}}"),
        ("Uang Sebanyak", ": Rp {{uang_muka}}"),
        ("Terbilang", ": {{uang_muka_terbilang}}"),
        ("Untuk Pembayaran", ": Uang Muka Perjalanan Dinas"),
    ]
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()

    # ===== DATA PERJALANAN =====
    add_paragraph(doc, "C. DATA PERJALANAN DINAS", bold=True)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    perjalanan_items = [
        ("Nama Pelaksana", ": {{pelaksana_nama}}"),
        ("NIP", ": {{pelaksana_nip}}"),
        ("Pangkat/Golongan", ": {{pelaksana_pangkat}}"),
        ("Jabatan", ": {{pelaksana_jabatan}}"),
        ("Tujuan", ": {{kota_tujuan}}, {{provinsi_tujuan}}"),
        ("Maksud Perjalanan", ": {{maksud_perjalanan}}"),
    ]
    for i, (label, value) in enumerate(perjalanan_items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    # Dasar
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    dasar_items = [
        ("Surat Tugas Nomor", ": {{nomor_surat_tugas}} tanggal {{tanggal_surat_tugas}}"),
        ("SPPD Nomor", ": {{nomor_sppd}}"),
    ]
    for i, (label, value) in enumerate(dasar_items):
        row = table2.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()
    doc.add_paragraph()

    # ===== TANDA TANGAN =====
    table = doc.add_table(rows=2, cols=3)

    # Baris 1: Header
    table.rows[0].cells[0].text = "Setuju Dibayar,\nPejabat Pembuat Komitmen"
    table.rows[0].cells[1].text = "Lunas Dibayar,\nBendahara Pengeluaran"
    table.rows[0].cells[2].text = f"{{{{satker_kota}}}}, {{{{tanggal_surat_tugas}}}}\nYang Menerima,"

    # Baris 2: Nama & NIP
    table.rows[1].cells[0].text = "\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"
    table.rows[1].cells[1].text = "\n\n\n\n{{bendahara_nama}}\nNIP. {{bendahara_nip}}"
    table.rows[1].cells[2].text = "\n\n\n\n{{pelaksana_nama}}\nNIP. {{pelaksana_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "kuitansi_uang_muka.docx"))
    print("Created: kuitansi_uang_muka.docx")


def create_kuitansi_rampung():
    """Kuitansi Rampung Perjalanan Dinas dengan Data DIPA"""
    doc = Document()

    # Header Kop Surat
    add_heading(doc, "{{kementerian}}", 1, True)
    add_heading(doc, "{{eselon1}}", 2, True)
    add_heading(doc, "{{satker_nama}}", 2, True)

    doc.add_paragraph()
    add_heading(doc, "KUITANSI/BUKTI PEMBAYARAN", 1, True)
    add_heading(doc, "PERJALANAN DINAS (RAMPUNG)", 2, True)

    doc.add_paragraph()

    # ===== DATA DIPA =====
    add_paragraph(doc, "A. DATA DIPA", bold=True)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    dipa_items = [
        ("Kode Satker", ": {{satker_kode}}"),
        ("Nama Satker", ": {{satker_nama}}"),
        ("Tahun Anggaran", ": {{tahun_anggaran}}"),
        ("Sumber Dana", ": {{sumber_dana}}"),
        ("Kode Akun/MAK", ": {{kode_akun}}"),
        ("Kegiatan", ": {{nama_paket}}"),
    ]
    for i, (label, value) in enumerate(dipa_items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()

    # ===== DATA PERJALANAN =====
    add_paragraph(doc, "B. DATA PERJALANAN DINAS", bold=True)
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    perjalanan_items = [
        ("Nama Pelaksana", ": {{pelaksana_nama}}"),
        ("NIP", ": {{pelaksana_nip}}"),
        ("Pangkat/Golongan", ": {{pelaksana_pangkat}}"),
        ("Jabatan", ": {{pelaksana_jabatan}}"),
        ("Tujuan", ": {{kota_tujuan}}, {{provinsi_tujuan}}"),
        ("Tanggal Berangkat", ": {{tanggal_berangkat}}"),
        ("Tanggal Kembali", ": {{tanggal_kembali}}"),
        ("Lama Perjalanan", ": {{lama_perjalanan}} hari"),
    ]
    for i, (label, value) in enumerate(perjalanan_items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    # Dasar
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    dasar_items = [
        ("Surat Tugas Nomor", ": {{nomor_surat_tugas}} tanggal {{tanggal_surat_tugas}}"),
        ("SPPD Nomor", ": {{nomor_sppd}}"),
    ]
    for i, (label, value) in enumerate(dasar_items):
        row = table2.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph()

    # ===== RINCIAN BIAYA =====
    add_paragraph(doc, "C. RINCIAN BIAYA PERJALANAN DINAS", bold=True)
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'

    # Header
    hdr = table.rows[0]
    hdr.cells[0].text = "No"
    hdr.cells[1].text = "Uraian Biaya"
    hdr.cells[2].text = "Jumlah (Rp)"

    biaya_items = [
        ("1", "Biaya Transportasi (Pergi-Pulang)", "{{biaya_transport}}"),
        ("2", "Uang Harian ({{lama_perjalanan}} hari)", "{{biaya_uang_harian}}"),
        ("3", "Biaya Penginapan", "{{biaya_penginapan}}"),
        ("4", "Uang Representasi", "{{biaya_representasi}}"),
        ("5", "Biaya Lain-lain", "{{biaya_lain_lain}}"),
        ("", "TOTAL BIAYA PERJALANAN DINAS", "{{total_biaya}}"),
        ("", "Uang Muka yang Telah Diterima", "{{uang_muka}}"),
        ("", "SISA KURANG / (LEBIH) BAYAR", "{{kekurangan_bayar}} / ({{kelebihan_bayar}})"),
    ]

    for i, (no, uraian, jumlah) in enumerate(biaya_items):
        row = table.rows[i + 1]
        row.cells[0].text = no
        row.cells[1].text = uraian
        row.cells[2].text = jumlah

    doc.add_paragraph()
    add_paragraph(doc, "Terbilang: {{sisa_terbilang}}")

    doc.add_paragraph()
    doc.add_paragraph()

    # ===== TANDA TANGAN =====
    table = doc.add_table(rows=2, cols=3)

    # Baris 1: Header
    table.rows[0].cells[0].text = "Setuju Dibayar,\nPejabat Pembuat Komitmen"
    table.rows[0].cells[1].text = "Lunas Dibayar,\nBendahara Pengeluaran"
    table.rows[0].cells[2].text = f"{{{{satker_kota}}}}, {{{{tanggal_kembali}}}}\nYang Menerima,"

    # Baris 2: Nama & NIP
    table.rows[1].cells[0].text = "\n\n\n\n{{ppk_nama}}\nNIP. {{ppk_nip}}"
    table.rows[1].cells[1].text = "\n\n\n\n{{bendahara_nama}}\nNIP. {{bendahara_nip}}"
    table.rows[1].cells[2].text = "\n\n\n\n{{pelaksana_nama}}\nNIP. {{pelaksana_nip}}"

    doc.save(os.path.join(TEMPLATES_DIR, "kuitansi_rampung.docx"))
    print("Created: kuitansi_rampung.docx")


def main():
    print("Updating Kuitansi templates with DIPA data...")
    create_kuitansi_uang_muka()
    create_kuitansi_rampung()
    print("\nTemplates updated successfully!")


if __name__ == "__main__":
    main()
