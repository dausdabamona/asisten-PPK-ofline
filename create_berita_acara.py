"""
Script untuk membuat template:
1. Berita Acara Pemeriksaan (BAHP) - Barang, Jasa Lainnya, Konstruksi
2. Berita Acara Serah Terima (BAST) - Barang, Jasa Lainnya, Konstruksi
3. Berita Acara Pembayaran (BAP) - Termin, Sekaligus

Karakteristik berbeda sesuai jenis pengadaan dan cara pembayaran
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = "templates/word"


def set_cell_shading(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_kop_surat(doc):
    """Add standard letter header"""
    kop1 = doc.add_paragraph()
    kop1.add_run("{{kementerian}}").bold = True
    kop1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    kop2 = doc.add_paragraph()
    kop2.add_run("{{eselon1}}").bold = True
    kop2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    kop3 = doc.add_paragraph()
    run = kop3.add_run("{{satker_nama}}")
    run.bold = True
    run.font.size = Pt(14)
    kop3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    alamat = doc.add_paragraph()
    alamat.add_run("{{satker_alamat}}, {{satker_kota}}")
    alamat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    border = doc.add_paragraph()
    border.add_run("═" * 75)
    border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()


def add_ttd_pphp(doc, include_ppk=True):
    """Add signature block for PPHP team"""
    doc.add_paragraph()
    
    if include_ppk:
        # TTD dengan PPK
        sign_table = doc.add_table(rows=8, cols=2)
        
        sign_table.rows[0].cells[0].text = "Pejabat/Panitia Penerima Hasil Pekerjaan"
        sign_table.rows[0].cells[1].text = "Mengetahui,"
        
        sign_table.rows[1].cells[0].text = ""
        sign_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen"
        
        # Space
        sign_table.rows[4].cells[0].text = "{{pphp_ketua_nama}}"
        sign_table.rows[4].cells[1].text = "{{ppk_nama}}"
        
        sign_table.rows[5].cells[0].text = "{{pphp_ketua_jabatan}}"
        sign_table.rows[5].cells[1].text = "NIP. {{ppk_nip}}"
        
        # Bold names
        for cell in [sign_table.rows[4].cells[0], sign_table.rows[4].cells[1]]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.underline = True
    else:
        # TTD PPHP saja
        sign_table = doc.add_table(rows=6, cols=1)
        
        sign_table.rows[0].cells[0].text = "Pejabat/Panitia Penerima Hasil Pekerjaan"
        sign_table.rows[4].cells[0].text = "{{pphp_ketua_nama}}"
        sign_table.rows[5].cells[0].text = "{{pphp_ketua_jabatan}}"
        
        for para in sign_table.rows[4].cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.underline = True
    
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_ttd_serah_terima(doc):
    """Add signature block for BAST (PPK, Penyedia, PPHP)"""
    doc.add_paragraph()
    
    sign_table = doc.add_table(rows=8, cols=3)
    
    # Headers
    sign_table.rows[0].cells[0].text = "PIHAK PERTAMA"
    sign_table.rows[0].cells[1].text = "PIHAK KEDUA"
    sign_table.rows[0].cells[2].text = "Mengetahui,"
    
    sign_table.rows[1].cells[0].text = "Pejabat Pembuat Komitmen,"
    sign_table.rows[1].cells[1].text = "{{penyedia_nama}},"
    sign_table.rows[1].cells[2].text = "PPHP,"
    
    # Names
    sign_table.rows[5].cells[0].text = "{{ppk_nama}}"
    sign_table.rows[5].cells[1].text = "{{direktur_nama}}"
    sign_table.rows[5].cells[2].text = "{{pphp_ketua_nama}}"
    
    sign_table.rows[6].cells[0].text = "NIP. {{ppk_nip}}"
    sign_table.rows[6].cells[1].text = "Direktur"
    sign_table.rows[6].cells[2].text = "{{pphp_ketua_jabatan}}"
    
    # Bold names
    for cell in [sign_table.rows[5].cells[0], sign_table.rows[5].cells[1], sign_table.rows[5].cells[2]]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.underline = True
    
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================================================================
# BERITA ACARA PEMERIKSAAN - BARANG
# ============================================================================

def create_bahp_barang():
    """Create BAHP template for Pengadaan Barang"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA PEMERIKSAAN HASIL PEKERJAAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("PENGADAAN BARANG")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bahp}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_bahp}} tanggal {{tanggal_bahp_terbilang}} bulan {{bulan_bahp}} "
        "tahun {{tahun_bahp_terbilang}} ({{tanggal_bahp_fmt}}), kami yang bertanda tangan di bawah ini:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    p1.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Info PPHP
    p2 = doc.add_paragraph()
    p2.add_run("Pejabat/Panitia Penerima Hasil Pekerjaan (PPHP):").bold = True
    
    pphp_table = doc.add_table(rows=3, cols=4)
    pphp_table.style = 'Table Grid'
    
    # Header
    headers = ["No", "Nama", "Jabatan", "Kedudukan dalam Tim"]
    for i, h in enumerate(headers):
        pphp_table.rows[0].cells[i].text = h
        for para in pphp_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    # Data PPHP
    pphp_data = [
        ("1.", "{{pphp_ketua_nama}}", "{{pphp_ketua_jabatan}}", "Ketua"),
        ("2.", "{{pphp_anggota1_nama}}", "{{pphp_anggota1_jabatan}}", "Anggota"),
    ]
    for i, (no, nama, jabatan, kedudukan) in enumerate(pphp_data, 1):
        pphp_table.rows[i].cells[0].text = no
        pphp_table.rows[i].cells[1].text = nama
        pphp_table.rows[i].cells[2].text = jabatan
        pphp_table.rows[i].cells[3].text = kedudukan
    
    doc.add_paragraph()
    
    p3 = doc.add_paragraph()
    p3.add_run(
        "berdasarkan Surat Keputusan Kuasa Pengguna Anggaran Nomor {{nomor_sk_pphp}} tanggal "
        "{{tanggal_sk_pphp}}, telah melakukan pemeriksaan terhadap barang yang diserahkan oleh:"
    )
    p3.paragraph_format.line_spacing = 1.5
    
    # Info Penyedia
    penyedia_table = doc.add_table(rows=4, cols=3)
    penyedia_info = [
        ("Nama Perusahaan", ":", "{{penyedia_nama}}"),
        ("Alamat", ":", "{{penyedia_alamat}}"),
        ("Nama Direktur", ":", "{{direktur_nama}}"),
        ("Nomor Kontrak/SPK", ":", "{{nomor_spk}} tanggal {{tanggal_spk_fmt}}"),
    ]
    for i, (label, sep, value) in enumerate(penyedia_info):
        penyedia_table.rows[i].cells[0].text = label
        penyedia_table.rows[i].cells[1].text = sep
        penyedia_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Info Paket
    p4 = doc.add_paragraph()
    p4.add_run("A. DATA PAKET PENGADAAN").bold = True
    
    paket_table = doc.add_table(rows=5, cols=3)
    paket_table.style = 'Table Grid'
    paket_info = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("3.", "Waktu Pelaksanaan", "{{jangka_waktu}} hari kalender"),
        ("4.", "Tanggal Kontrak", "{{tanggal_spk_fmt}}"),
        ("5.", "Tanggal Selesai Kontrak", "{{tanggal_selesai_fmt}}"),
    ]
    for i, (no, label, value) in enumerate(paket_info):
        paket_table.rows[i].cells[0].text = no
        paket_table.rows[i].cells[1].text = label
        paket_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Hasil Pemeriksaan Barang
    p5 = doc.add_paragraph()
    p5.add_run("B. HASIL PEMERIKSAAN BARANG").bold = True
    
    # Tabel pemeriksaan barang
    pemeriksaan_table = doc.add_table(rows=5, cols=7)
    pemeriksaan_table.style = 'Table Grid'
    
    headers = ["No", "Uraian Barang", "Spesifikasi", "Satuan", "Volume Kontrak", "Volume Diterima", "Keterangan"]
    for i, h in enumerate(headers):
        pemeriksaan_table.rows[0].cells[i].text = h
        for para in pemeriksaan_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    # Template rows
    for row_idx in range(1, 5):
        pemeriksaan_table.rows[row_idx].cells[0].text = "{{item_no_" + str(row_idx) + "}}"
        pemeriksaan_table.rows[row_idx].cells[1].text = "{{item_uraian_" + str(row_idx) + "}}"
        pemeriksaan_table.rows[row_idx].cells[2].text = "{{item_spek_" + str(row_idx) + "}}"
        pemeriksaan_table.rows[row_idx].cells[3].text = "{{item_satuan_" + str(row_idx) + "}}"
        pemeriksaan_table.rows[row_idx].cells[4].text = "{{item_vol_kontrak_" + str(row_idx) + "}}"
        pemeriksaan_table.rows[row_idx].cells[5].text = "{{item_vol_terima_" + str(row_idx) + "}}"
        pemeriksaan_table.rows[row_idx].cells[6].text = "{{item_ket_" + str(row_idx) + "}}"
    
    doc.add_paragraph()
    
    # Checklist Pemeriksaan Barang
    p6 = doc.add_paragraph()
    p6.add_run("C. CHECKLIST PEMERIKSAAN").bold = True
    
    checklist_table = doc.add_table(rows=8, cols=4)
    checklist_table.style = 'Table Grid'
    
    checklist_headers = ["No", "Item Pemeriksaan", "Sesuai", "Tidak Sesuai"]
    for i, h in enumerate(checklist_headers):
        checklist_table.rows[0].cells[i].text = h
        for para in checklist_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    checklist_items = [
        "Jumlah/kuantitas barang sesuai kontrak",
        "Spesifikasi teknis sesuai kontrak",
        "Merk/tipe sesuai penawaran",
        "Kondisi barang baru dan tidak cacat",
        "Kelengkapan (manual, aksesoris, dll)",
        "Sertifikat garansi tersedia",
        "Kemasan/packaging dalam kondisi baik",
    ]
    
    for i, item in enumerate(checklist_items, 1):
        checklist_table.rows[i].cells[0].text = str(i)
        checklist_table.rows[i].cells[1].text = item
        checklist_table.rows[i].cells[2].text = "☐"
        checklist_table.rows[i].cells[3].text = "☐"
        
        for j in [0, 2, 3]:
            for para in checklist_table.rows[i].cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Kesimpulan
    p7 = doc.add_paragraph()
    p7.add_run("D. KESIMPULAN").bold = True
    
    kesimpulan = doc.add_paragraph()
    kesimpulan.add_run(
        "Berdasarkan hasil pemeriksaan tersebut di atas, kami menyatakan bahwa barang yang "
        "diserahkan oleh {{penyedia_nama}}:"
    )
    
    opsi_table = doc.add_table(rows=2, cols=2)
    opsi_table.rows[0].cells[0].text = "☐"
    opsi_table.rows[0].cells[1].text = "DITERIMA - sesuai dengan spesifikasi dan kuantitas dalam kontrak"
    opsi_table.rows[1].cells[0].text = "☐"
    opsi_table.rows[1].cells[1].text = "DITOLAK - tidak sesuai spesifikasi/kuantitas (lihat keterangan)"
    
    doc.add_paragraph()
    
    # Catatan
    p8 = doc.add_paragraph()
    p8.add_run("E. CATATAN/KETERANGAN").bold = True
    
    catatan = doc.add_paragraph()
    catatan.add_run("{{catatan_pemeriksaan}}")
    catatan.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # Penutup
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara Pemeriksaan ini dibuat dengan sebenarnya untuk dipergunakan "
        "sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    
    add_ttd_pphp(doc, include_ppk=True)
    
    filepath = os.path.join(TEMPLATES_DIR, "bahp_barang.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# BERITA ACARA PEMERIKSAAN - JASA LAINNYA
# ============================================================================

def create_bahp_jasa_lainnya():
    """Create BAHP template for Jasa Lainnya"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA PEMERIKSAAN HASIL PEKERJAAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("PENGADAAN JASA LAINNYA")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bahp}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_bahp}} tanggal {{tanggal_bahp_fmt}}, kami yang bertanda tangan di bawah ini "
        "selaku Pejabat/Panitia Penerima Hasil Pekerjaan berdasarkan SK KPA Nomor {{nomor_sk_pphp}}, "
        "telah melakukan pemeriksaan terhadap hasil pekerjaan:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    p1.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Info Paket
    p2 = doc.add_paragraph()
    p2.add_run("A. INFORMASI PAKET PEKERJAAN").bold = True
    
    paket_table = doc.add_table(rows=7, cols=3)
    paket_table.style = 'Table Grid'
    paket_info = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Nomor Kontrak/SPK", "{{nomor_spk}}"),
        ("3.", "Tanggal Kontrak", "{{tanggal_spk_fmt}}"),
        ("4.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("5.", "Penyedia Jasa", "{{penyedia_nama}}"),
        ("6.", "Waktu Pelaksanaan", "{{jangka_waktu}} hari kalender"),
        ("7.", "Tanggal Selesai", "{{tanggal_selesai_fmt}}"),
    ]
    for i, (no, label, value) in enumerate(paket_info):
        paket_table.rows[i].cells[0].text = no
        paket_table.rows[i].cells[1].text = label
        paket_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Pemeriksaan Output/Deliverables
    p3 = doc.add_paragraph()
    p3.add_run("B. PEMERIKSAAN OUTPUT/DELIVERABLES").bold = True
    
    output_table = doc.add_table(rows=6, cols=5)
    output_table.style = 'Table Grid'
    
    headers = ["No", "Output/Deliverable", "Kriteria Penerimaan", "Status", "Keterangan"]
    for i, h in enumerate(headers):
        output_table.rows[0].cells[i].text = h
        for para in output_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    # Template rows for deliverables
    for row_idx in range(1, 6):
        output_table.rows[row_idx].cells[0].text = str(row_idx)
        output_table.rows[row_idx].cells[1].text = "{{output_" + str(row_idx) + "}}"
        output_table.rows[row_idx].cells[2].text = "{{kriteria_" + str(row_idx) + "}}"
        output_table.rows[row_idx].cells[3].text = "{{status_" + str(row_idx) + "}}"
        output_table.rows[row_idx].cells[4].text = "{{ket_" + str(row_idx) + "}}"
    
    doc.add_paragraph()
    
    # Pemeriksaan Kualitas Jasa
    p4 = doc.add_paragraph()
    p4.add_run("C. PEMERIKSAAN KUALITAS PELAKSANAAN JASA").bold = True
    
    kualitas_table = doc.add_table(rows=7, cols=4)
    kualitas_table.style = 'Table Grid'
    
    kualitas_headers = ["No", "Aspek Pemeriksaan", "Sesuai", "Tidak Sesuai"]
    for i, h in enumerate(kualitas_headers):
        kualitas_table.rows[0].cells[i].text = h
        for para in kualitas_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    kualitas_items = [
        "Ruang lingkup pekerjaan sesuai KAK",
        "Metodologi pelaksanaan sesuai proposal",
        "Kualitas output sesuai standar yang ditetapkan",
        "Timeline/jadwal pelaksanaan terpenuhi",
        "Personil yang ditugaskan sesuai kontrak",
        "Laporan pelaksanaan lengkap",
    ]
    
    for i, item in enumerate(kualitas_items, 1):
        kualitas_table.rows[i].cells[0].text = str(i)
        kualitas_table.rows[i].cells[1].text = item
        kualitas_table.rows[i].cells[2].text = "☐"
        kualitas_table.rows[i].cells[3].text = "☐"
        
        for j in [0, 2, 3]:
            for para in kualitas_table.rows[i].cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Verifikasi Dokumen
    p5 = doc.add_paragraph()
    p5.add_run("D. VERIFIKASI DOKUMEN PENDUKUNG").bold = True
    
    dokumen_items = [
        "Laporan Akhir Pelaksanaan Pekerjaan",
        "Dokumentasi Kegiatan (foto, video, dll)",
        "Daftar Hadir/Absensi (jika relevan)",
        "Sertifikat/Output Pekerjaan",
        "Softcopy seluruh dokumen hasil pekerjaan",
    ]
    
    for i, item in enumerate(dokumen_items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ☐ {item}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # Kesimpulan
    p6 = doc.add_paragraph()
    p6.add_run("E. KESIMPULAN").bold = True
    
    kesimpulan = doc.add_paragraph()
    kesimpulan.add_run(
        "Berdasarkan hasil pemeriksaan, kami menyatakan bahwa pekerjaan jasa yang dilaksanakan oleh "
        "{{penyedia_nama}}:"
    )
    
    opsi_table = doc.add_table(rows=3, cols=2)
    opsi_table.rows[0].cells[0].text = "☐"
    opsi_table.rows[0].cells[1].text = "DITERIMA SELURUHNYA - semua output sesuai KAK dan memenuhi kriteria"
    opsi_table.rows[1].cells[0].text = "☐"
    opsi_table.rows[1].cells[1].text = "DITERIMA DENGAN CATATAN - perlu perbaikan minor (lihat catatan)"
    opsi_table.rows[2].cells[0].text = "☐"
    opsi_table.rows[2].cells[1].text = "DITOLAK - tidak memenuhi kriteria penerimaan"
    
    doc.add_paragraph()
    
    # Catatan
    p7 = doc.add_paragraph()
    p7.add_run("F. CATATAN/REKOMENDASI").bold = True
    
    catatan = doc.add_paragraph()
    catatan.add_run("{{catatan_pemeriksaan}}")
    catatan.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara ini dibuat dengan sebenarnya untuk dipergunakan sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    
    add_ttd_pphp(doc, include_ppk=True)
    
    filepath = os.path.join(TEMPLATES_DIR, "bahp_jasa_lainnya.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# BERITA ACARA PEMERIKSAAN - KONSTRUKSI
# ============================================================================

def create_bahp_konstruksi():
    """Create BAHP template for Konstruksi"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA PEMERIKSAAN HASIL PEKERJAAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("PEKERJAAN KONSTRUKSI")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bahp}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_bahp}} tanggal {{tanggal_bahp_fmt}}, bertempat di lokasi pekerjaan "
        "{{lokasi_pekerjaan}}, kami yang bertanda tangan di bawah ini:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    p1.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Tim Pemeriksa
    p2 = doc.add_paragraph()
    p2.add_run("TIM PEMERIKSA:").bold = True
    
    tim_table = doc.add_table(rows=4, cols=4)
    tim_table.style = 'Table Grid'
    
    headers = ["No", "Nama", "Jabatan", "Kedudukan"]
    for i, h in enumerate(headers):
        tim_table.rows[0].cells[i].text = h
        for para in tim_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    tim_data = [
        ("1.", "{{pphp_ketua_nama}}", "{{pphp_ketua_jabatan}}", "Ketua PPHP"),
        ("2.", "{{konsultan_pengawas_nama}}", "Konsultan Pengawas", "Pengawas Lapangan"),
        ("3.", "{{pphp_anggota1_nama}}", "{{pphp_anggota1_jabatan}}", "Anggota PPHP"),
    ]
    for i, (no, nama, jabatan, kedudukan) in enumerate(tim_data, 1):
        tim_table.rows[i].cells[0].text = no
        tim_table.rows[i].cells[1].text = nama
        tim_table.rows[i].cells[2].text = jabatan
        tim_table.rows[i].cells[3].text = kedudukan
    
    doc.add_paragraph()
    
    # Info Paket
    p3 = doc.add_paragraph()
    p3.add_run("A. INFORMASI KONTRAK").bold = True
    
    paket_table = doc.add_table(rows=8, cols=3)
    paket_table.style = 'Table Grid'
    paket_info = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Nomor Kontrak", "{{nomor_spk}}"),
        ("3.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("4.", "Penyedia Jasa", "{{penyedia_nama}}"),
        ("5.", "Waktu Pelaksanaan", "{{jangka_waktu}} hari kalender"),
        ("6.", "Tanggal SPMK", "{{tanggal_spmk_fmt}}"),
        ("7.", "Tanggal Selesai Kontrak", "{{tanggal_selesai_fmt}}"),
        ("8.", "Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
    ]
    for i, (no, label, value) in enumerate(paket_info):
        paket_table.rows[i].cells[0].text = no
        paket_table.rows[i].cells[1].text = label
        paket_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Pemeriksaan Volume/Kuantitas
    p4 = doc.add_paragraph()
    p4.add_run("B. PEMERIKSAAN VOLUME PEKERJAAN").bold = True
    
    volume_table = doc.add_table(rows=6, cols=7)
    volume_table.style = 'Table Grid'
    
    vol_headers = ["No", "Uraian Pekerjaan", "Satuan", "Vol. Kontrak", "Vol. Terpasang", "Deviasi", "Ket"]
    for i, h in enumerate(vol_headers):
        volume_table.rows[0].cells[i].text = h
        for para in volume_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    for row_idx in range(1, 6):
        volume_table.rows[row_idx].cells[0].text = str(row_idx)
        volume_table.rows[row_idx].cells[1].text = "{{pekerjaan_" + str(row_idx) + "}}"
        volume_table.rows[row_idx].cells[2].text = "{{satuan_" + str(row_idx) + "}}"
        volume_table.rows[row_idx].cells[3].text = "{{vol_kontrak_" + str(row_idx) + "}}"
        volume_table.rows[row_idx].cells[4].text = "{{vol_terpasang_" + str(row_idx) + "}}"
        volume_table.rows[row_idx].cells[5].text = "{{deviasi_" + str(row_idx) + "}}"
        volume_table.rows[row_idx].cells[6].text = "{{ket_" + str(row_idx) + "}}"
    
    doc.add_paragraph()
    
    # Pemeriksaan Mutu
    p5 = doc.add_paragraph()
    p5.add_run("C. PEMERIKSAAN MUTU PEKERJAAN").bold = True
    
    mutu_table = doc.add_table(rows=8, cols=4)
    mutu_table.style = 'Table Grid'
    
    mutu_headers = ["No", "Item Pemeriksaan Mutu", "Sesuai", "Tidak Sesuai"]
    for i, h in enumerate(mutu_headers):
        mutu_table.rows[0].cells[i].text = h
        for para in mutu_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    mutu_items = [
        "Pekerjaan sesuai gambar rencana (shop drawing)",
        "Material yang digunakan sesuai spesifikasi",
        "Hasil uji mutu material (beton, baja, dll) memenuhi standar",
        "Dimensi/elevasi sesuai toleransi yang diizinkan",
        "Finishing/tampilan visual memenuhi standar",
        "Fungsi bangunan/konstruksi berjalan baik",
        "Tidak ada kerusakan/cacat yang terlihat",
    ]
    
    for i, item in enumerate(mutu_items, 1):
        mutu_table.rows[i].cells[0].text = str(i)
        mutu_table.rows[i].cells[1].text = item
        mutu_table.rows[i].cells[2].text = "☐"
        mutu_table.rows[i].cells[3].text = "☐"
        
        for j in [0, 2, 3]:
            for para in mutu_table.rows[i].cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Verifikasi Dokumen Konstruksi
    p6 = doc.add_paragraph()
    p6.add_run("D. VERIFIKASI DOKUMEN TEKNIS").bold = True
    
    dokumen_items = [
        "As Built Drawing",
        "Laporan Hasil Uji Material/Mutu",
        "Foto Dokumentasi Progress 0%, 50%, 100%",
        "Laporan Mingguan/Bulanan",
        "Berita Acara Perubahan (jika ada CCO)",
        "Manual Operasional dan Pemeliharaan",
        "Sertifikat/Izin yang diperlukan",
    ]
    
    for i, item in enumerate(dokumen_items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ☐ {item}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # Progres Fisik
    p7 = doc.add_paragraph()
    p7.add_run("E. REKAPITULASI PROGRES FISIK").bold = True
    
    progres_table = doc.add_table(rows=4, cols=3)
    progres_table.style = 'Table Grid'
    progres_data = [
        ("Progres Rencana", ":", "{{progres_rencana}}%"),
        ("Progres Realisasi", ":", "{{progres_realisasi}}%"),
        ("Deviasi", ":", "{{deviasi_progres}}%"),
    ]
    for i, (label, sep, value) in enumerate(progres_data):
        progres_table.rows[i].cells[0].text = label
        progres_table.rows[i].cells[1].text = sep
        progres_table.rows[i].cells[2].text = value
    
    # Status
    status = doc.add_paragraph()
    status.add_run("Status: ")
    status.add_run("☐ On Schedule / ☐ Ahead / ☐ Behind Schedule")
    
    doc.add_paragraph()
    
    # Kesimpulan
    p8 = doc.add_paragraph()
    p8.add_run("F. KESIMPULAN").bold = True
    
    kesimpulan = doc.add_paragraph()
    kesimpulan.add_run("Berdasarkan hasil pemeriksaan, pekerjaan konstruksi:")
    
    opsi_table = doc.add_table(rows=3, cols=2)
    opsi_table.rows[0].cells[0].text = "☐"
    opsi_table.rows[0].cells[1].text = "DITERIMA - volume dan mutu sesuai kontrak, siap untuk PHO"
    opsi_table.rows[1].cells[0].text = "☐"
    opsi_table.rows[1].cells[1].text = "DITERIMA DENGAN CATATAN - perlu perbaikan (lihat daftar defect)"
    opsi_table.rows[2].cells[0].text = "☐"
    opsi_table.rows[2].cells[1].text = "DITOLAK - tidak memenuhi spesifikasi teknis"
    
    doc.add_paragraph()
    
    # Daftar Defect
    p9 = doc.add_paragraph()
    p9.add_run("G. DAFTAR PEKERJAAN YANG PERLU DIPERBAIKI (DEFECT LIST)").bold = True
    
    defect_table = doc.add_table(rows=4, cols=4)
    defect_table.style = 'Table Grid'
    defect_headers = ["No", "Item Pekerjaan", "Jenis Cacat/Kekurangan", "Batas Waktu Perbaikan"]
    for i, h in enumerate(defect_headers):
        defect_table.rows[0].cells[i].text = h
        for para in defect_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    for row_idx in range(1, 4):
        defect_table.rows[row_idx].cells[0].text = str(row_idx)
    
    doc.add_paragraph()
    
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara Pemeriksaan ini dibuat dengan sebenarnya untuk dipergunakan "
        "sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    
    add_ttd_pphp(doc, include_ppk=True)
    
    filepath = os.path.join(TEMPLATES_DIR, "bahp_konstruksi.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# BERITA ACARA SERAH TERIMA - BARANG
# ============================================================================

def create_bast_barang():
    """Create BAST template for Pengadaan Barang"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA SERAH TERIMA").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("PENGADAAN BARANG")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bast}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_bast}} tanggal {{tanggal_bast_fmt}}, kami yang bertanda tangan di bawah ini:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    
    doc.add_paragraph()
    
    # PIHAK PERTAMA
    pihak1 = doc.add_paragraph()
    pihak1.add_run("1. Nama\t\t: {{ppk_nama}}\n")
    pihak1.add_run("   Jabatan\t\t: Pejabat Pembuat Komitmen\n")
    pihak1.add_run("   Alamat\t\t: {{satker_nama}}\n")
    pihak1.add_run("   Selanjutnya disebut ").italic = True
    pihak1.add_run("PIHAK PERTAMA").bold = True
    
    doc.add_paragraph()
    
    # PIHAK KEDUA
    pihak2 = doc.add_paragraph()
    pihak2.add_run("2. Nama\t\t: {{direktur_nama}}\n")
    pihak2.add_run("   Jabatan\t\t: Direktur {{penyedia_nama}}\n")
    pihak2.add_run("   Alamat\t\t: {{penyedia_alamat}}\n")
    pihak2.add_run("   Selanjutnya disebut ").italic = True
    pihak2.add_run("PIHAK KEDUA").bold = True
    
    doc.add_paragraph()
    
    # Pernyataan
    p2 = doc.add_paragraph()
    p2.add_run("Dengan ini menyatakan bahwa:")
    
    # Info Kontrak
    info_table = doc.add_table(rows=5, cols=3)
    info_table.style = 'Table Grid'
    info_data = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Nomor Kontrak/SPK", "{{nomor_spk}}"),
        ("3.", "Tanggal Kontrak", "{{tanggal_spk_fmt}}"),
        ("4.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("5.", "Jangka Waktu", "{{jangka_waktu}} hari kalender"),
    ]
    for i, (no, label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = no
        info_table.rows[i].cells[1].text = label
        info_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Daftar Barang
    p3 = doc.add_paragraph()
    p3.add_run("Daftar Barang yang Diserahterimakan:").bold = True
    
    barang_table = doc.add_table(rows=5, cols=6)
    barang_table.style = 'Table Grid'
    
    headers = ["No", "Uraian Barang", "Spesifikasi/Merk", "Satuan", "Volume", "Keterangan"]
    for i, h in enumerate(headers):
        barang_table.rows[0].cells[i].text = h
        for para in barang_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    for row_idx in range(1, 5):
        barang_table.rows[row_idx].cells[0].text = "{{no_" + str(row_idx) + "}}"
        barang_table.rows[row_idx].cells[1].text = "{{uraian_" + str(row_idx) + "}}"
        barang_table.rows[row_idx].cells[2].text = "{{spek_" + str(row_idx) + "}}"
        barang_table.rows[row_idx].cells[3].text = "{{satuan_" + str(row_idx) + "}}"
        barang_table.rows[row_idx].cells[4].text = "{{volume_" + str(row_idx) + "}}"
        barang_table.rows[row_idx].cells[5].text = "{{ket_" + str(row_idx) + "}}"
    
    doc.add_paragraph()
    
    # Pernyataan Serah Terima
    pernyataan = [
        "PIHAK KEDUA telah menyerahkan barang sebagaimana tersebut di atas kepada PIHAK PERTAMA dalam kondisi baik dan lengkap.",
        "PIHAK PERTAMA telah menerima barang tersebut setelah dilakukan pemeriksaan oleh Pejabat/Panitia Penerima Hasil Pekerjaan (PPHP) berdasarkan Berita Acara Pemeriksaan Nomor {{nomor_bahp}} tanggal {{tanggal_bahp_fmt}}.",
        "Barang yang diserahkan telah sesuai dengan spesifikasi teknis dan kuantitas yang tercantum dalam kontrak.",
        "PIHAK KEDUA memberikan garansi atas barang selama {{masa_garansi}} bulan terhitung sejak tanggal Berita Acara ini.",
        "Dengan ditandatanganinya Berita Acara ini, maka hak kepemilikan atas barang beralih kepada PIHAK PERTAMA.",
    ]
    
    for i, item in enumerate(pernyataan, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara Serah Terima ini dibuat dalam rangkap 2 (dua) untuk dipergunakan "
        "sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    
    add_ttd_serah_terima(doc)
    
    filepath = os.path.join(TEMPLATES_DIR, "bast_barang.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# BERITA ACARA SERAH TERIMA - JASA LAINNYA
# ============================================================================

def create_bast_jasa_lainnya():
    """Create BAST template for Jasa Lainnya"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA SERAH TERIMA HASIL PEKERJAAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("PENGADAAN JASA LAINNYA")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bast}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_bast}} tanggal {{tanggal_bast_fmt}}, kami yang bertanda tangan di bawah ini:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    
    doc.add_paragraph()
    
    # Para Pihak
    pihak1 = doc.add_paragraph()
    pihak1.add_run("1. {{ppk_nama}}, selaku Pejabat Pembuat Komitmen pada {{satker_nama}}, "
                   "selanjutnya disebut ")
    pihak1.add_run("PIHAK PERTAMA").bold = True
    
    doc.add_paragraph()
    
    pihak2 = doc.add_paragraph()
    pihak2.add_run("2. {{direktur_nama}}, selaku Direktur {{penyedia_nama}}, "
                   "selanjutnya disebut ")
    pihak2.add_run("PIHAK KEDUA").bold = True
    
    doc.add_paragraph()
    
    # Menyatakan
    p2 = doc.add_paragraph()
    p2.add_run("Dengan ini menyatakan bahwa berdasarkan:")
    
    dasar = [
        "Kontrak/SPK Nomor {{nomor_spk}} tanggal {{tanggal_spk_fmt}};",
        "Berita Acara Pemeriksaan Hasil Pekerjaan Nomor {{nomor_bahp}} tanggal {{tanggal_bahp_fmt}};",
    ]
    
    for i, item in enumerate(dasar, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # Info Paket
    p3 = doc.add_paragraph()
    p3.add_run("Untuk pekerjaan:").bold = True
    
    info_table = doc.add_table(rows=5, cols=3)
    info_table.style = 'Table Grid'
    info_data = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("3.", "Jangka Waktu", "{{jangka_waktu}} hari kalender"),
        ("4.", "Tanggal Mulai", "{{tanggal_mulai_fmt}}"),
        ("5.", "Tanggal Selesai", "{{tanggal_selesai_fmt}}"),
    ]
    for i, (no, label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = no
        info_table.rows[i].cells[1].text = label
        info_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Hasil Pekerjaan yang Diserahkan
    p4 = doc.add_paragraph()
    p4.add_run("Hasil Pekerjaan/Output yang Diserahterimakan:").bold = True
    
    output_table = doc.add_table(rows=5, cols=4)
    output_table.style = 'Table Grid'
    
    headers = ["No", "Output/Deliverable", "Jumlah", "Keterangan"]
    for i, h in enumerate(headers):
        output_table.rows[0].cells[i].text = h
        for para in output_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    
    for row_idx in range(1, 5):
        output_table.rows[row_idx].cells[0].text = str(row_idx)
        output_table.rows[row_idx].cells[1].text = "{{output_" + str(row_idx) + "}}"
        output_table.rows[row_idx].cells[2].text = "{{jumlah_" + str(row_idx) + "}}"
        output_table.rows[row_idx].cells[3].text = "{{ket_" + str(row_idx) + "}}"
    
    doc.add_paragraph()
    
    # Pernyataan
    pernyataan = [
        "PIHAK KEDUA telah menyelesaikan seluruh pekerjaan dan menyerahkan hasil pekerjaan kepada PIHAK PERTAMA sesuai dengan ruang lingkup dalam Kerangka Acuan Kerja (KAK).",
        "PIHAK PERTAMA telah menerima hasil pekerjaan tersebut setelah diperiksa dan dinyatakan sesuai oleh PPHP.",
        "Seluruh hak kekayaan intelektual atas hasil pekerjaan beralih kepada PIHAK PERTAMA.",
        "PIHAK KEDUA telah menyerahkan seluruh softcopy dan hardcopy dokumen hasil pekerjaan.",
    ]
    
    for i, item in enumerate(pernyataan, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara Serah Terima ini dibuat untuk dipergunakan sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    
    add_ttd_serah_terima(doc)
    
    filepath = os.path.join(TEMPLATES_DIR, "bast_jasa_lainnya.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# BERITA ACARA SERAH TERIMA - KONSTRUKSI (PHO & FHO)
# ============================================================================

def create_bast_konstruksi_pho():
    """Create BAST PHO (Provisional Hand Over) template for Konstruksi"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA SERAH TERIMA PERTAMA").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("(PROVISIONAL HAND OVER / PHO)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle2 = doc.add_paragraph()
    subtitle2.add_run("PEKERJAAN KONSTRUKSI")
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bast_pho}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_bast}} tanggal {{tanggal_bast_fmt}}, bertempat di {{lokasi_pekerjaan}}, "
        "kami yang bertanda tangan di bawah ini:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    
    doc.add_paragraph()
    
    # Para Pihak
    pihak1 = doc.add_paragraph()
    pihak1.add_run("1. {{ppk_nama}}, selaku Pejabat Pembuat Komitmen, selanjutnya disebut ")
    pihak1.add_run("PENGGUNA JASA").bold = True
    
    doc.add_paragraph()
    
    pihak2 = doc.add_paragraph()
    pihak2.add_run("2. {{direktur_nama}}, selaku Direktur {{penyedia_nama}}, selanjutnya disebut ")
    pihak2.add_run("PENYEDIA JASA").bold = True
    
    doc.add_paragraph()
    
    # Info Kontrak
    p2 = doc.add_paragraph()
    p2.add_run("Berdasarkan Kontrak:").bold = True
    
    kontrak_table = doc.add_table(rows=8, cols=3)
    kontrak_table.style = 'Table Grid'
    kontrak_data = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Nomor Kontrak", "{{nomor_spk}}"),
        ("3.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("4.", "Tanggal SPMK", "{{tanggal_spmk_fmt}}"),
        ("5.", "Waktu Pelaksanaan", "{{jangka_waktu}} hari kalender"),
        ("6.", "Tanggal Selesai Kontrak", "{{tanggal_selesai_fmt}}"),
        ("7.", "Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
        ("8.", "Masa Pemeliharaan", "{{masa_pemeliharaan}} hari kalender"),
    ]
    for i, (no, label, value) in enumerate(kontrak_data):
        kontrak_table.rows[i].cells[0].text = no
        kontrak_table.rows[i].cells[1].text = label
        kontrak_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Rekapitulasi Pekerjaan
    p3 = doc.add_paragraph()
    p3.add_run("Rekapitulasi Progres Pekerjaan:").bold = True
    
    progres_table = doc.add_table(rows=5, cols=3)
    progres_table.style = 'Table Grid'
    progres_data = [
        ("a.", "Progres Fisik", "{{progres_fisik}}%"),
        ("b.", "Progres Keuangan", "{{progres_keuangan}}%"),
        ("c.", "Pembayaran yang telah diterima", "{{pembayaran_diterima_fmt}}"),
        ("d.", "Retensi yang ditahan", "{{retensi_fmt}}"),
    ]
    for i, (no, label, value) in enumerate(progres_data):
        progres_table.rows[i].cells[0].text = no
        progres_table.rows[i].cells[1].text = label
        progres_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Pernyataan PHO
    p4 = doc.add_paragraph()
    p4.add_run("Dengan ini menyatakan bahwa:").bold = True
    
    pernyataan = [
        "PENYEDIA JASA telah menyelesaikan pekerjaan konstruksi dengan progres fisik 100% (seratus persen) sesuai kontrak.",
        "Hasil pekerjaan telah diperiksa oleh Konsultan Pengawas dan PPHP berdasarkan Berita Acara Pemeriksaan Nomor {{nomor_bahp}} tanggal {{tanggal_bahp_fmt}}.",
        "PENGGUNA JASA menerima penyerahan pertama (PHO) atas pekerjaan konstruksi tersebut.",
        "Daftar pekerjaan yang masih perlu diperbaiki (defect list) terlampir dan harus diselesaikan selama masa pemeliharaan.",
        "Masa Pemeliharaan dimulai sejak tanggal {{tanggal_bast_fmt}} sampai dengan {{tanggal_fho_fmt}} ({{masa_pemeliharaan}} hari kalender).",
        "Selama masa pemeliharaan, PENYEDIA JASA wajib memperbaiki kerusakan yang terjadi bukan karena kesalahan PENGGUNA JASA tanpa biaya tambahan.",
        "Retensi sebesar {{retensi_fmt}} (5% dari nilai kontrak) akan ditahan sampai dengan Serah Terima Akhir (FHO).",
    ]
    
    for i, item in enumerate(pernyataan, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara Serah Terima Pertama (PHO) ini dibuat untuk dipergunakan sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    
    add_ttd_serah_terima(doc)
    
    filepath = os.path.join(TEMPLATES_DIR, "bast_konstruksi_pho.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


def create_bast_konstruksi_fho():
    """Create BAST FHO (Final Hand Over) template for Konstruksi"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA SERAH TERIMA AKHIR").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("(FINAL HAND OVER / FHO)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle2 = doc.add_paragraph()
    subtitle2.add_run("PEKERJAAN KONSTRUKSI")
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bast_fho}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_fho}} tanggal {{tanggal_fho_fmt}}, bertempat di {{lokasi_pekerjaan}}, "
        "kami yang bertanda tangan di bawah ini:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    
    doc.add_paragraph()
    
    # Para Pihak
    pihak1 = doc.add_paragraph()
    pihak1.add_run("1. {{ppk_nama}}, selaku Pejabat Pembuat Komitmen, selanjutnya disebut ")
    pihak1.add_run("PENGGUNA JASA").bold = True
    
    doc.add_paragraph()
    
    pihak2 = doc.add_paragraph()
    pihak2.add_run("2. {{direktur_nama}}, selaku Direktur {{penyedia_nama}}, selanjutnya disebut ")
    pihak2.add_run("PENYEDIA JASA").bold = True
    
    doc.add_paragraph()
    
    # Mengacu PHO
    p2 = doc.add_paragraph()
    p2.add_run("Menunjuk pada:").bold = True
    
    referensi = [
        "Kontrak/SPK Nomor {{nomor_spk}} tanggal {{tanggal_spk_fmt}};",
        "Berita Acara Serah Terima Pertama (PHO) Nomor {{nomor_bast_pho}} tanggal {{tanggal_bast_fmt}};",
        "Masa Pemeliharaan yang telah berakhir pada tanggal {{tanggal_fho_fmt}};",
    ]
    
    for i, item in enumerate(referensi, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # Info Kontrak
    p3 = doc.add_paragraph()
    p3.add_run("Untuk pekerjaan:").bold = True
    
    kontrak_table = doc.add_table(rows=5, cols=3)
    kontrak_table.style = 'Table Grid'
    kontrak_data = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("3.", "Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
        ("4.", "Masa Pemeliharaan", "{{masa_pemeliharaan}} hari kalender"),
        ("5.", "Retensi", "{{retensi_fmt}}"),
    ]
    for i, (no, label, value) in enumerate(kontrak_data):
        kontrak_table.rows[i].cells[0].text = no
        kontrak_table.rows[i].cells[1].text = label
        kontrak_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Pernyataan FHO
    p4 = doc.add_paragraph()
    p4.add_run("Dengan ini menyatakan bahwa:").bold = True
    
    pernyataan = [
        "Masa Pemeliharaan telah berakhir pada tanggal {{tanggal_fho_fmt}}.",
        "Selama masa pemeliharaan, PENYEDIA JASA telah memperbaiki seluruh kerusakan/cacat yang tercantum dalam defect list.",
        "Hasil pemeriksaan akhir menyatakan bahwa pekerjaan dalam kondisi baik dan berfungsi sebagaimana mestinya.",
        "PENGGUNA JASA menerima penyerahan akhir (FHO) atas pekerjaan konstruksi tersebut.",
        "Dengan ditandatanganinya Berita Acara ini, maka retensi sebesar {{retensi_fmt}} dapat dibayarkan kepada PENYEDIA JASA.",
        "Tanggung jawab PENYEDIA JASA atas pekerjaan berakhir dengan ditandatanganinya Berita Acara ini, kecuali untuk tanggung jawab kegagalan bangunan sesuai ketentuan perundang-undangan.",
    ]
    
    for i, item in enumerate(pernyataan, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara Serah Terima Akhir (FHO) ini dibuat untuk dipergunakan sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    
    add_ttd_serah_terima(doc)
    
    filepath = os.path.join(TEMPLATES_DIR, "bast_konstruksi_fho.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# BERITA ACARA PEMBAYARAN - SEKALIGUS
# ============================================================================

def create_bap_sekaligus():
    """Create BAP template for Pembayaran Sekaligus (100%)"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA PEMBAYARAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("(PEMBAYARAN SEKALIGUS 100%)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bap}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_bap}} tanggal {{tanggal_bap_fmt}}, kami yang bertanda tangan di bawah ini:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    
    doc.add_paragraph()
    
    # Para Pihak
    pihak1 = doc.add_paragraph()
    pihak1.add_run("1. Nama\t\t: {{ppk_nama}}\n")
    pihak1.add_run("   Jabatan\t\t: Pejabat Pembuat Komitmen\n")
    pihak1.add_run("   Alamat\t\t: {{satker_nama}}\n")
    pihak1.add_run("   Selanjutnya disebut ").italic = True
    pihak1.add_run("PIHAK PERTAMA").bold = True
    
    doc.add_paragraph()
    
    pihak2 = doc.add_paragraph()
    pihak2.add_run("2. Nama\t\t: {{direktur_nama}}\n")
    pihak2.add_run("   Jabatan\t\t: Direktur {{penyedia_nama}}\n")
    pihak2.add_run("   Alamat\t\t: {{penyedia_alamat}}\n")
    pihak2.add_run("   Selanjutnya disebut ").italic = True
    pihak2.add_run("PIHAK KEDUA").bold = True
    
    doc.add_paragraph()
    
    p2 = doc.add_paragraph()
    p2.add_run("Berdasarkan:").bold = True
    
    dasar = [
        "Kontrak/SPK Nomor {{nomor_spk}} tanggal {{tanggal_spk_fmt}};",
        "Berita Acara Pemeriksaan Nomor {{nomor_bahp}} tanggal {{tanggal_bahp_fmt}};",
        "Berita Acara Serah Terima Nomor {{nomor_bast}} tanggal {{tanggal_bast_fmt}};",
    ]
    
    for i, item in enumerate(dasar, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # Info Paket
    p3 = doc.add_paragraph()
    p3.add_run("Untuk pekerjaan:").bold = True
    
    paket_table = doc.add_table(rows=4, cols=3)
    paket_table.style = 'Table Grid'
    paket_data = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Jenis Pengadaan", "{{jenis_pengadaan}}"),
        ("3.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("4.", "Sumber Dana", "{{sumber_dana}} - MAK {{kode_akun}}"),
    ]
    for i, (no, label, value) in enumerate(paket_data):
        paket_table.rows[i].cells[0].text = no
        paket_table.rows[i].cells[1].text = label
        paket_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Rincian Pembayaran
    p4 = doc.add_paragraph()
    p4.add_run("Rincian Pembayaran:").bold = True
    
    pembayaran_table = doc.add_table(rows=7, cols=3)
    pembayaran_table.style = 'Table Grid'
    
    pembayaran_data = [
        ("A.", "Nilai Kontrak (sebelum pajak)", "{{nilai_sebelum_pajak_fmt}}"),
        ("B.", "PPN 11%", "{{ppn_fmt}}"),
        ("C.", "Nilai Kontrak (termasuk PPN)", "{{nilai_kontrak_fmt}}"),
        ("D.", "PPh Pasal {{jenis_pph}}", "{{pph_fmt}}"),
        ("E.", "Potongan Lainnya", "{{potongan_lain_fmt}}"),
        ("F.", "Jumlah yang Dibayarkan (C - D - E)", "{{jumlah_dibayar_fmt}}"),
    ]
    
    for i, (kode, label, value) in enumerate(pembayaran_data):
        pembayaran_table.rows[i].cells[0].text = kode
        pembayaran_table.rows[i].cells[1].text = label
        pembayaran_table.rows[i].cells[2].text = value
    
    # Terbilang
    pembayaran_table.rows[6].cells[0].text = ""
    pembayaran_table.rows[6].cells[1].text = "Terbilang"
    pembayaran_table.rows[6].cells[2].text = "{{jumlah_dibayar_terbilang}}"
    
    doc.add_paragraph()
    
    # Rekening
    p5 = doc.add_paragraph()
    p5.add_run("Pembayaran ditransfer ke rekening:").bold = True
    
    rek_table = doc.add_table(rows=4, cols=3)
    rek_data = [
        ("Nama Rekening", ":", "{{penyedia_nama}}"),
        ("Nomor Rekening", ":", "{{penyedia_rekening}}"),
        ("Nama Bank", ":", "{{penyedia_bank}}"),
        ("Cabang", ":", "{{penyedia_bank_cabang}}"),
    ]
    for i, (label, sep, value) in enumerate(rek_data):
        rek_table.rows[i].cells[0].text = label
        rek_table.rows[i].cells[1].text = sep
        rek_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara Pembayaran ini dibuat untuk dipergunakan sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    
    # TTD
    sign_table = doc.add_table(rows=7, cols=2)
    sign_table.rows[0].cells[0].text = "PIHAK KEDUA"
    sign_table.rows[0].cells[1].text = "PIHAK PERTAMA"
    sign_table.rows[1].cells[0].text = "{{penyedia_nama}},"
    sign_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen,"
    sign_table.rows[5].cells[0].text = "{{direktur_nama}}"
    sign_table.rows[5].cells[1].text = "{{ppk_nama}}"
    sign_table.rows[6].cells[0].text = "Direktur"
    sign_table.rows[6].cells[1].text = "NIP. {{ppk_nip}}"
    
    for cell in [sign_table.rows[5].cells[0], sign_table.rows[5].cells[1]]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.underline = True
    
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(TEMPLATES_DIR, "bap_sekaligus.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# BERITA ACARA PEMBAYARAN - TERMIN
# ============================================================================

def create_bap_termin():
    """Create BAP template for Pembayaran Termin (Progress-based)"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_kop_surat(doc)
    
    # Title
    title = doc.add_paragraph()
    title.add_run("BERITA ACARA PEMBAYARAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("TERMIN KE-{{nomor_termin}}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_bap}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run(
        "Pada hari ini {{hari_bap}} tanggal {{tanggal_bap_fmt}}, kami yang bertanda tangan di bawah ini:"
    )
    p1.paragraph_format.first_line_indent = Cm(1.25)
    
    doc.add_paragraph()
    
    # Para Pihak
    pihak1 = doc.add_paragraph()
    pihak1.add_run("1. {{ppk_nama}}, selaku Pejabat Pembuat Komitmen, selanjutnya disebut ")
    pihak1.add_run("PIHAK PERTAMA").bold = True
    
    doc.add_paragraph()
    
    pihak2 = doc.add_paragraph()
    pihak2.add_run("2. {{direktur_nama}}, selaku Direktur {{penyedia_nama}}, selanjutnya disebut ")
    pihak2.add_run("PIHAK KEDUA").bold = True
    
    doc.add_paragraph()
    
    p2 = doc.add_paragraph()
    p2.add_run("Berdasarkan:").bold = True
    
    dasar = [
        "Kontrak/SPK Nomor {{nomor_spk}} tanggal {{tanggal_spk_fmt}};",
        "Berita Acara Kemajuan Pekerjaan Nomor {{nomor_ba_progres}} tanggal {{tanggal_ba_progres}};",
        "Laporan Kemajuan Pekerjaan yang disetujui Konsultan Pengawas/Direksi Teknis;",
    ]
    
    for i, item in enumerate(dasar, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # Info Paket
    p3 = doc.add_paragraph()
    p3.add_run("Untuk pekerjaan:").bold = True
    
    paket_table = doc.add_table(rows=5, cols=3)
    paket_table.style = 'Table Grid'
    paket_data = [
        ("1.", "Nama Paket", "{{nama_paket}}"),
        ("2.", "Jenis Pengadaan", "{{jenis_pengadaan}}"),
        ("3.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("4.", "Sumber Dana", "{{sumber_dana}} - MAK {{kode_akun}}"),
        ("5.", "Jangka Waktu", "{{jangka_waktu}} hari kalender"),
    ]
    for i, (no, label, value) in enumerate(paket_data):
        paket_table.rows[i].cells[0].text = no
        paket_table.rows[i].cells[1].text = label
        paket_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Rekapitulasi Progres
    p4 = doc.add_paragraph()
    p4.add_run("Rekapitulasi Kemajuan Pekerjaan:").bold = True
    
    progres_table = doc.add_table(rows=5, cols=3)
    progres_table.style = 'Table Grid'
    progres_data = [
        ("a.", "Progres Fisik s.d. Termin Ini", "{{progres_fisik_sd_termin}}%"),
        ("b.", "Progres Fisik Termin Sebelumnya", "{{progres_fisik_sebelumnya}}%"),
        ("c.", "Progres Fisik Termin Ini", "{{progres_fisik_termin}}%"),
        ("d.", "Progres Keuangan s.d. Termin Ini", "{{progres_keuangan_sd_termin}}%"),
    ]
    for i, (kode, label, value) in enumerate(progres_data):
        progres_table.rows[i].cells[0].text = kode
        progres_table.rows[i].cells[1].text = label
        progres_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Rincian Pembayaran Termin
    p5 = doc.add_paragraph()
    p5.add_run("Rincian Pembayaran Termin Ke-{{nomor_termin}}:").bold = True
    
    pembayaran_table = doc.add_table(rows=10, cols=3)
    pembayaran_table.style = 'Table Grid'
    
    pembayaran_data = [
        ("A.", "Nilai Kontrak", "{{nilai_kontrak_fmt}}"),
        ("B.", "Nilai Pekerjaan s.d. Termin Ini ({{progres_fisik_sd_termin}}%)", "{{nilai_sd_termin_fmt}}"),
        ("C.", "Pembayaran s.d. Termin Sebelumnya", "{{pembayaran_sebelumnya_fmt}}"),
        ("D.", "Nilai Pekerjaan Termin Ini (B - C)", "{{nilai_termin_fmt}}"),
        ("E.", "Retensi 5%", "{{retensi_termin_fmt}}"),
        ("F.", "Nilai Setelah Retensi (D - E)", "{{nilai_setelah_retensi_fmt}}"),
        ("G.", "PPN 11%", "{{ppn_termin_fmt}}"),
        ("H.", "PPh Pasal {{jenis_pph}}", "{{pph_termin_fmt}}"),
        ("I.", "Jumlah yang Dibayarkan (F + G - H)", "{{jumlah_dibayar_termin_fmt}}"),
    ]
    
    for i, (kode, label, value) in enumerate(pembayaran_data):
        pembayaran_table.rows[i].cells[0].text = kode
        pembayaran_table.rows[i].cells[1].text = label
        pembayaran_table.rows[i].cells[2].text = value
    
    pembayaran_table.rows[9].cells[0].text = ""
    pembayaran_table.rows[9].cells[1].text = "Terbilang"
    pembayaran_table.rows[9].cells[2].text = "{{jumlah_dibayar_termin_terbilang}}"
    
    doc.add_paragraph()
    
    # Akumulasi Pembayaran
    p6 = doc.add_paragraph()
    p6.add_run("Akumulasi Pembayaran:").bold = True
    
    akum_table = doc.add_table(rows=5, cols=3)
    akum_table.style = 'Table Grid'
    akum_data = [
        ("1.", "Total Pembayaran s.d. Termin Ini", "{{total_pembayaran_sd_fmt}}"),
        ("2.", "Total Retensi yang Ditahan", "{{total_retensi_fmt}}"),
        ("3.", "Sisa Nilai Kontrak", "{{sisa_kontrak_fmt}}"),
        ("4.", "Persentase Pembayaran", "{{persen_pembayaran}}%"),
    ]
    for i, (no, label, value) in enumerate(akum_data):
        akum_table.rows[i].cells[0].text = no
        akum_table.rows[i].cells[1].text = label
        akum_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    # Rekening
    p7 = doc.add_paragraph()
    p7.add_run("Pembayaran ditransfer ke rekening:").bold = True
    
    rek_table = doc.add_table(rows=4, cols=3)
    rek_data = [
        ("Nama Rekening", ":", "{{penyedia_nama}}"),
        ("Nomor Rekening", ":", "{{penyedia_rekening}}"),
        ("Nama Bank", ":", "{{penyedia_bank}}"),
        ("Cabang", ":", "{{penyedia_bank_cabang}}"),
    ]
    for i, (label, sep, value) in enumerate(rek_data):
        rek_table.rows[i].cells[0].text = label
        rek_table.rows[i].cells[1].text = sep
        rek_table.rows[i].cells[2].text = value
    
    doc.add_paragraph()
    
    penutup = doc.add_paragraph()
    penutup.add_run(
        "Demikian Berita Acara Pembayaran Termin ini dibuat untuk dipergunakan sebagaimana mestinya."
    )
    
    doc.add_paragraph()
    
    # TTD
    sign_table = doc.add_table(rows=7, cols=2)
    sign_table.rows[0].cells[0].text = "PIHAK KEDUA"
    sign_table.rows[0].cells[1].text = "PIHAK PERTAMA"
    sign_table.rows[1].cells[0].text = "{{penyedia_nama}},"
    sign_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen,"
    sign_table.rows[5].cells[0].text = "{{direktur_nama}}"
    sign_table.rows[5].cells[1].text = "{{ppk_nama}}"
    sign_table.rows[6].cells[0].text = "Direktur"
    sign_table.rows[6].cells[1].text = "NIP. {{ppk_nip}}"
    
    for cell in [sign_table.rows[5].cells[0], sign_table.rows[5].cells[1]]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.underline = True
    
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(TEMPLATES_DIR, "bap_termin.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    
    print("=" * 60)
    print("MEMBUAT TEMPLATE BERITA ACARA")
    print("=" * 60)
    print()
    
    print("📋 BERITA ACARA PEMERIKSAAN (BAHP):")
    create_bahp_barang()
    create_bahp_jasa_lainnya()
    create_bahp_konstruksi()
    
    print()
    print("📋 BERITA ACARA SERAH TERIMA (BAST):")
    create_bast_barang()
    create_bast_jasa_lainnya()
    create_bast_konstruksi_pho()
    create_bast_konstruksi_fho()
    
    print()
    print("📋 BERITA ACARA PEMBAYARAN (BAP):")
    create_bap_sekaligus()
    create_bap_termin()
    
    print()
    print("=" * 60)
    print("SELESAI!")
    print("=" * 60)
    print()
    print("Template yang dibuat:")
    print("┌─────────────────────────────────────────────────────────┐")
    print("│ BAHP (Pemeriksaan)                                      │")
    print("│ ├── bahp_barang.docx        - Checklist barang, garansi │")
    print("│ ├── bahp_jasa_lainnya.docx  - Output/deliverables, KAK  │")
    print("│ └── bahp_konstruksi.docx    - Volume, mutu, defect list │")
    print("├─────────────────────────────────────────────────────────┤")
    print("│ BAST (Serah Terima)                                     │")
    print("│ ├── bast_barang.docx        - Garansi, kepemilikan      │")
    print("│ ├── bast_jasa_lainnya.docx  - HKI, softcopy/hardcopy    │")
    print("│ ├── bast_konstruksi_pho.docx- PHO, masa pemeliharaan    │")
    print("│ └── bast_konstruksi_fho.docx- FHO, retensi              │")
    print("├─────────────────────────────────────────────────────────┤")
    print("│ BAP (Pembayaran)                                        │")
    print("│ ├── bap_sekaligus.docx      - 100%, pajak, rekening     │")
    print("│ └── bap_termin.docx         - Progres, retensi, akum.   │")
    print("└─────────────────────────────────────────────────────────┘")
