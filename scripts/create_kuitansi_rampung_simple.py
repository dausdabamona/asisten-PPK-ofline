#!/usr/bin/env python3
"""
Script untuk membuat template Kuitansi Rampung Sederhana (non-Perjalanan Dinas).
Format simpel untuk belanja barang/jasa selain perjalanan dinas.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_single_spacing(paragraph):
    """Set paragraph to single spacing with no space before/after."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    borders = []
    if top:
        borders.append('top')
    if bottom:
        borders.append('bottom')
    if left:
        borders.append('left')
    if right:
        borders.append('right')

    for border_name in borders:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color='D9D9D9'):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)


def create_kuitansi_rampung_simple():
    """Create simple Kuitansi Rampung template for non-travel expenses."""
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== HEADER: KUITANSI RAMPUNG ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(title)
    title_run = title.add_run("KUITANSI RAMPUNG")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(subtitle)
    sub_run = subtitle.add_run("PERTANGGUNGJAWABAN BELANJA")
    sub_run.font.size = Pt(12)

    doc.add_paragraph()

    # ========== INFO TABLE ==========
    info_table = doc.add_table(rows=5, cols=3)
    info_table.autofit = False

    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(11)

    info_data = [
        ("Nomor Kuitansi", ":", "{{nomor_kuitansi}}"),
        ("Tanggal", ":", "{{tanggal_dokumen:tanggal_panjang}}"),
        ("Nama Kegiatan", ":", "{{nama_kegiatan}}"),
        ("Sumber Dana/MAK", ":", "{{kode_akun}}"),
        ("Unit Kerja", ":", "{{unit_kerja}}"),
    ]

    for i, (label, sep, value) in enumerate(info_data):
        row = info_table.rows[i]

        # Label cell
        cell_label = row.cells[0]
        p = cell_label.paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(label)
        run.font.size = Pt(11)

        # Separator cell
        cell_sep = row.cells[1]
        p = cell_sep.paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(sep)
        run.font.size = Pt(11)

        # Value cell
        cell_value = row.cells[2]
        p = cell_value.paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(value)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # ========== RINCIAN BELANJA TABLE ==========
    rincian_title = doc.add_paragraph()
    set_single_spacing(rincian_title)
    run = rincian_title.add_run("RINCIAN BELANJA:")
    run.bold = True
    run.font.size = Pt(11)

    # Rincian table header
    rincian_table = doc.add_table(rows=1, cols=6)
    rincian_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["No", "Uraian Barang/Jasa", "Volume", "Satuan", "Harga Satuan", "Jumlah"]
    header_widths = [Cm(0.8), Cm(6), Cm(1.5), Cm(1.5), Cm(3), Cm(3)]

    for i, (header, width) in enumerate(zip(headers, header_widths)):
        cell = rincian_table.rows[0].cells[i]
        cell.width = width
        set_cell_borders(cell)
        set_cell_shading(cell, 'E8E8E8')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_single_spacing(p)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)

    # Data rows placeholder (using jinja2 loop syntax)
    # Add sample row with template variables
    data_row = rincian_table.add_row()
    sample_data = ["{{loop.index}}", "{{item.uraian}}", "{{item.volume}}", "{{item.satuan}}",
                   "{{item.harga_satuan:rupiah}}", "{{item.jumlah:rupiah}}"]

    for i, (value, width) in enumerate(zip(sample_data, header_widths)):
        cell = data_row.cells[i]
        cell.width = width
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        if i in [0, 2, 3]:  # No, Volume, Satuan - center
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i in [4, 5]:  # Harga, Jumlah - right
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_single_spacing(p)
        run = p.add_run(value)
        run.font.size = Pt(10)

    # Add comment for template engine
    comment_p = doc.add_paragraph()
    set_single_spacing(comment_p)
    run = comment_p.add_run("{% for item in rincian_items %} ... {% endfor %}")
    run.font.size = Pt(8)
    run.italic = True

    doc.add_paragraph()

    # ========== TOTAL SECTION ==========
    total_table = doc.add_table(rows=1, cols=2)
    total_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    total_table.rows[0].cells[0].width = Cm(10)
    total_table.rows[0].cells[1].width = Cm(4)

    # Total label
    total_label_cell = total_table.rows[0].cells[0]
    p = total_label_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_single_spacing(p)
    run = p.add_run("TOTAL BELANJA:")
    run.bold = True
    run.font.size = Pt(11)

    # Total value
    total_value_cell = total_table.rows[0].cells[1]
    set_cell_borders(total_value_cell)
    p = total_value_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_single_spacing(p)
    run = p.add_run("{{total_realisasi:rupiah}}")
    run.bold = True
    run.font.size = Pt(11)

    # Terbilang
    terbilang_p = doc.add_paragraph()
    terbilang_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(terbilang_p)
    run = terbilang_p.add_run("({{total_realisasi:terbilang}})")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    # ========== PERHITUNGAN SELISIH ==========
    selisih_title = doc.add_paragraph()
    set_single_spacing(selisih_title)
    run = selisih_title.add_run("PERHITUNGAN:")
    run.bold = True
    run.font.size = Pt(11)

    # Selisih table
    selisih_table = doc.add_table(rows=3, cols=3)
    selisih_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    selisih_data = [
        ("Uang Muka Diterima", ":", "{{uang_muka:rupiah}}"),
        ("Total Realisasi Belanja", ":", "{{total_realisasi:rupiah}}"),
        ("Selisih", ":", "{{selisih:rupiah}}"),
    ]

    for i, (label, sep, value) in enumerate(selisih_data):
        row = selisih_table.rows[i]
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(4)

        # Label
        p = row.cells[0].paragraphs[0]
        set_single_spacing(p)
        run = p.add_run(label)
        run.font.size = Pt(11)
        if i == 2:
            run.bold = True

        # Separator
        p = row.cells[1].paragraphs[0]
        set_single_spacing(p)
        p.add_run(sep).font.size = Pt(11)

        # Value
        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_single_spacing(p)
        run = p.add_run(value)
        run.font.size = Pt(11)
        if i == 2:
            run.bold = True

    # Keterangan selisih
    ket_p = doc.add_paragraph()
    set_single_spacing(ket_p)
    run = ket_p.add_run("Keterangan: {{keterangan_selisih}}")
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # ========== SIGNATURES (2 columns) ==========
    sig_table = doc.add_table(rows=5, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in sig_table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(7)

    # Row 0: Headers
    left_header = sig_table.rows[0].cells[0]
    p = left_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("Yang Mempertanggungjawabkan,")
    run.font.size = Pt(11)

    right_header = sig_table.rows[0].cells[1]
    p = right_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("Mengetahui,")
    run.font.size = Pt(11)

    # Row 1: Jabatan
    left_jabatan = sig_table.rows[1].cells[0]
    p = left_jabatan.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("{{penerima_jabatan}}")
    run.font.size = Pt(10)

    right_jabatan = sig_table.rows[1].cells[1]
    p = right_jabatan.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("Pejabat Pembuat Komitmen")
    run.font.size = Pt(10)

    # Row 2: Empty space for signature
    for cell in sig_table.rows[2].cells:
        p = cell.paragraphs[0]
        set_single_spacing(p)
        p.add_run("\n\n\n")

    # Row 3: Names (underlined)
    left_name = sig_table.rows[3].cells[0]
    p = left_name.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("{{penerima_nama}}")
    run.font.size = Pt(11)
    run.underline = True

    right_name = sig_table.rows[3].cells[1]
    p = right_name.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("{{ppk_nama}}")
    run.font.size = Pt(11)
    run.underline = True

    # Row 4: NIP
    left_nip = sig_table.rows[4].cells[0]
    p = left_nip.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("NIP. {{penerima_nip}}")
    run.font.size = Pt(10)

    right_nip = sig_table.rows[4].cells[1]
    p = right_nip.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p)
    run = p.add_run("NIP. {{ppk_nip}}")
    run.font.size = Pt(10)

    # ========== FOOTER NOTE ==========
    doc.add_paragraph()
    doc.add_paragraph()

    note_p = doc.add_paragraph()
    set_single_spacing(note_p)
    note_run = note_p.add_run(
        "Catatan: Kuitansi rampung ini merupakan bukti pertanggungjawaban penggunaan "
        "uang muka untuk kegiatan tersebut di atas."
    )
    note_run.font.size = Pt(9)
    note_run.italic = True

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "word" / "kuitansi_rampung_simple.docx"
    doc.save(str(output_path))
    print(f"Template saved to: {output_path}")

    return output_path


if __name__ == "__main__":
    create_kuitansi_rampung_simple()
