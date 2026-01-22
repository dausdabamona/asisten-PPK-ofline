"""
Script untuk membuat template Rancangan Kontrak:
1. SPK Barang
2. SPK Jasa Lainnya  
3. SPK Konstruksi
4. Surat Perjanjian Barang
5. Surat Perjanjian Jasa Lainnya
6. Surat Perjanjian Konstruksi

Karakteristik berbeda sesuai jenis pengadaan
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = "templates/word"


def add_heading(doc, text, level=1):
    """Add heading with proper formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        run.font.size = Pt(11)
    return p


def add_pasal(doc, nomor, judul):
    """Add pasal heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Pasal {nomor}")
    run.bold = True
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(judul)
    run2.bold = True
    
    return p, p2


def add_ayat(doc, nomor, text):
    """Add numbered paragraph (ayat)"""
    p = doc.add_paragraph()
    p.add_run(f"({nomor}) ").bold = True
    p.add_run(text)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.75)
    return p


def add_huruf(doc, huruf, text):
    """Add lettered sub-item"""
    p = doc.add_paragraph()
    p.add_run(f"{huruf}. ")
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(1.5)
    return p


# =============================================================================
# SPK BARANG
# =============================================================================
def create_spk_barang():
    """Create SPK template for Pengadaan Barang - Format Tabel"""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    run = title.add_run("RANCANGAN SURAT PERINTAH KERJA (SPK)")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Main table with border
    main_table = doc.add_table(rows=4, cols=2)
    main_table.style = 'Table Grid'

    # Set column widths
    for row in main_table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(9)

    # Row 1: Header SPK and Satuan Kerja info
    cell_0_0 = main_table.rows[0].cells[0]
    cell_0_0.text = ""
    p = cell_0_0.paragraphs[0]
    run = p.add_run("SURAT PERINTAH KERJA (SPK)")
    run.bold = True

    cell_0_1 = main_table.rows[0].cells[1]
    cell_0_1.text = ""
    p = cell_0_1.paragraphs[0]
    p.add_run("SATUAN KERJA").bold = True
    p.add_run(" : {{satker_nama}}")
    cell_0_1.add_paragraph()
    p2 = cell_0_1.add_paragraph()
    p2.add_run("NOMOR DAN TANGGAL SPK").bold = True
    p2.add_run(" : ")
    cell_0_1.add_paragraph("{{nomor_spk}} tanggal {{tanggal_spk}}")

    # Row 2: Nama Pejabat Penandatangan Kontrak
    cell_1_0 = main_table.rows[1].cells[0]
    cell_1_0.text = ""
    p = cell_1_0.paragraphs[0]
    p.add_run("Nama Pejabat Penandatangan Kontrak").bold = True
    p.add_run(":")

    cell_1_1 = main_table.rows[1].cells[1]
    cell_1_1.text = "{{ppk_nama}}"

    # Row 3: Nama Penyedia
    cell_2_0 = main_table.rows[2].cells[0]
    cell_2_0.text = ""
    p = cell_2_0.paragraphs[0]
    p.add_run("Nama Penyedia").bold = True
    p.add_run(":")

    cell_2_1 = main_table.rows[2].cells[1]
    cell_2_1.text = "{{penyedia_nama}}"

    # Row 4: Paket Pengadaan and Document Numbers
    cell_3_0 = main_table.rows[3].cells[0]
    cell_3_0.text = ""
    p = cell_3_0.paragraphs[0]
    p.add_run("PAKET PENGADAAN").bold = True
    p.add_run(" :")
    cell_3_0.add_paragraph("{{nama_paket}}")

    cell_3_1 = main_table.rows[3].cells[1]
    cell_3_1.text = ""
    # Add document info
    p = cell_3_1.paragraphs[0]
    p.add_run("NOMOR SURAT UNDANGAN PENGADAAN LANGSUNG").bold = True
    p.add_run(" :")
    cell_3_1.add_paragraph("{{nomor_undangan}}")
    cell_3_1.add_paragraph()
    p2 = cell_3_1.add_paragraph()
    p2.add_run("TANGGAL SURAT UNDANGAN PENGADAAN LANGSUNG").bold = True
    p2.add_run(" :")
    cell_3_1.add_paragraph("{{tanggal_undangan}}")
    cell_3_1.add_paragraph()
    p3 = cell_3_1.add_paragraph()
    p3.add_run("NOMOR BERITA ACARA HASIL PENGADAAN LANGSUNG").bold = True
    p3.add_run(" :")
    cell_3_1.add_paragraph("{{nomor_ba_pl}}")
    cell_3_1.add_paragraph()
    p4 = cell_3_1.add_paragraph()
    p4.add_run("TANGGAL BERITA ACARA HASIL PENGADAAN LANGSUNG").bold = True
    p4.add_run(" :")
    cell_3_1.add_paragraph("{{tanggal_ba_pl}}")

    doc.add_paragraph()

    # Sumber Dana
    p_sumber = doc.add_paragraph()
    p_sumber.add_run("SUMBER DANA: ").bold = True
    p_sumber.add_run("[sebagai contoh, cantumkan \"dibebankan atas DIPA/DPA ")
    p_sumber.add_run("{{satker_nama}}")
    p_sumber.add_run(" Tahun Anggaran ")
    p_sumber.add_run("{{tahun_anggaran}}")
    p_sumber.add_run(" untuk mata anggaran kegiatan ")
    p_sumber.add_run("{{kode_akun}}")
    p_sumber.add_run("\"]")

    doc.add_paragraph()

    # Nilai Kontrak
    p_nilai = doc.add_paragraph()
    p_nilai.add_run("Nilai Kontrak termasuk Pajak Pertambahan Nilai (PPN) adalah sebesar ")
    p_nilai.add_run("{{nilai_kontrak_fmt}}").bold = True
    p_nilai.add_run(" (")
    p_nilai.add_run("{{nilai_kontrak_terbilang}}")
    p_nilai.add_run(" rupiah).")

    doc.add_paragraph()

    # Waktu Pelaksanaan
    p_waktu = doc.add_paragraph()
    p_waktu.add_run("WAKTU PELAKSANAAN PEKERJAAN: ").bold = True
    p_waktu.add_run("{{jangka_waktu}}").bold = True
    p_waktu.add_run(" (")
    p_waktu.add_run("{{jangka_waktu_terbilang}}")
    p_waktu.add_run(") hari kalender")

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature table
    ttd_table = doc.add_table(rows=7, cols=2)

    # Headers
    cell_left = ttd_table.rows[0].cells[0]
    cell_left.text = ""
    p = cell_left.paragraphs[0]
    p.add_run("Untuk dan atas nama").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = cell_left.add_paragraph()
    p2.add_run("Pejabat Penandatangan Kontrak").bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_right = ttd_table.rows[0].cells[1]
    cell_right.text = ""
    p = cell_right.paragraphs[0]
    p.add_run("Untuk dan atas nama Penyedia").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Note for PPK side
    note_left = ttd_table.rows[1].cells[0]
    note_left.text = ""
    p = note_left.paragraphs[0]
    p.add_run("[tanda tangan dan cap (jika salinan asli ini untuk Penyedia maka rekatkan meterai Rp 10.000,-)]").font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Note for Penyedia side
    note_right = ttd_table.rows[1].cells[1]
    note_right.text = ""
    p = note_right.paragraphs[0]
    p.add_run("[tanda tangan dan cap (jika salinan asli ini untuk proyek/satuan kerja Pejabat Penandatangan Kontrak maka rekatkan meterai Rp 10.000,-)]").font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Empty rows for signature space
    for i in range(2, 5):
        ttd_table.rows[i].cells[0].text = ""
        ttd_table.rows[i].cells[1].text = ""

    # Names
    name_left = ttd_table.rows[5].cells[0]
    name_left.text = ""
    p = name_left.paragraphs[0]
    run = p.add_run("{{ppk_nama}}")
    run.bold = True
    run.underline = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_right = ttd_table.rows[5].cells[1]
    name_right.text = ""
    p = name_right.paragraphs[0]
    run = p.add_run("{{direktur_nama}}")
    run.bold = True
    run.underline = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Positions
    pos_left = ttd_table.rows[6].cells[0]
    pos_left.text = ""
    p = pos_left.paragraphs[0]
    p.add_run("{{ppk_jabatan}}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pos_right = ttd_table.rows[6].cells[1]
    pos_right.text = ""
    p = pos_right.paragraphs[0]
    p.add_run("{{direktur_jabatan}}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page break for Syarat Umum
    doc.add_page_break()

    # =========================================================================
    # SYARAT UMUM SPK
    # =========================================================================
    su_title = doc.add_paragraph()
    run = su_title.add_run("SYARAT UMUM")
    run.bold = True
    run.font.size = Pt(14)
    su_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    su_subtitle = doc.add_paragraph()
    run = su_subtitle.add_run("SURAT PERINTAH KERJA (SPK)")
    run.bold = True
    run.font.size = Pt(14)
    su_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1. LINGKUP PEKERJAAN
    add_su_item(doc, 1, "LINGKUP PEKERJAAN",
        "Penyedia yang ditunjuk berkewajiban untuk menyelesaikan pekerjaan dalam jangka waktu yang ditentukan sesuai dengan volume, spesifikasi teknis dan harga yang tercantum dalam SPK.")

    # 2. HUKUM YANG BERLAKU
    add_su_item(doc, 2, "HUKUM YANG BERLAKU",
        "Keabsahan, interpretasi, dan pelaksanaan SPK ini didasarkan kepada hukum Republik Indonesia.")

    # 3. HARGA SPK
    p = doc.add_paragraph()
    p.add_run("3.\tHARGA SPK").bold = True
    doc.add_paragraph("a.\tPejabat Penandatangan Kontrak membayar kepada penyedia sebesar harga SPK.")
    doc.add_paragraph("b.\tHarga SPK telah memperhitungkan keuntungan, beban pajak dan biaya overhead serta biaya asuransi (apabila dipersyaratkan).")
    doc.add_paragraph("c.\tRincian harga SPK sesuai dengan rincian yang tercantum dalam daftar kuantitas dan harga.")
    doc.add_paragraph()

    # 4. HAK KEPEMILIKAN
    p = doc.add_paragraph()
    p.add_run("4.\tHAK KEPEMILIKAN").bold = True
    doc.add_paragraph("a.\tPejabat Penandatangan Kontrak berhak atas kepemilikan semua barang/bahan yang terkait langsung atau disediakan sehubungan dengan jasa yang diberikan oleh penyedia kepada Pejabat Penandatangan Kontrak. Jika diminta oleh Pejabat Penandatangan Kontrak maka penyedia berkewajiban untuk membantu secara optimal pengalihan hak kepemilikan tersebut kepada Pejabat Penandatangan Kontrak sesuai dengan hukum yang berlaku.")
    doc.add_paragraph("b.\tHak kepemilikan atas peralatan dan barang/bahan yang disediakan oleh Pejabat Penandatangan Kontrak tetap pada Pejabat Penandatangan Kontrak, dan semua peralatan tersebut harus dikembalikan kepada Pejabat Penandatangan Kontrak pada saat SPK berakhir atau jika tidak diperlukan lagi oleh penyedia. Semua peralatan tersebut harus dikembalikan dalam kondisi yang sama pada saat diberikan kepada penyedia dengan pengecualian keausan akibat pemakaian yang wajar.")
    doc.add_paragraph()

    # 5. CACAT MUTU
    add_su_item(doc, 5, "CACAT MUTU",
        "Pejabat Penandatangan Kontrak akan memeriksa setiap hasil pekerjaan penyedia dan memberitahukan secara tertulis penyedia atas setiap cacat mutu yang ditemukan. Pejabat Penandatangan Kontrak dapat memerintahkan penyedia untuk menguji pekerjaan yang dianggap oleh Pejabat Penandatangan Kontrak mengandung cacat mutu. Penyedia bertanggung jawab atas cacat mutu selama masa garansi.")

    # 6. PERPAJAKAN
    add_su_item(doc, 6, "PERPAJAKAN",
        "Penyedia berkewajiban untuk membayar semua pajak, bea, retribusi, dan pungutan lain yang sah yang dibebankan oleh hukum yang berlaku atas pelaksanaan SPK. Semua pengeluaran perpajakan ini dianggap telah termasuk dalam harga SPK.")

    # 7. PENGALIHAN DAN/ATAU SUBKONTRAK
    add_su_item(doc, 7, "PENGALIHAN DAN/ATAU SUBKONTRAK",
        "Penyedia dilarang untuk mengalihkan dan/atau mensubkontrakkan sebagian atau seluruh pekerjaan. Pengalihan seluruh pekerjaan hanya diperbolehkan dalam hal pergantian nama penyedia, baik sebagai akibat peleburan (merger) atau akibat lainnya.")

    # 8. JADWAL
    p = doc.add_paragraph()
    p.add_run("8.\tJADWAL").bold = True
    doc.add_paragraph("a.\tSPK ini berlaku efektif pada tanggal penandatanganan oleh para pihak atau pada tanggal yang ditetapkan dalam Surat Perintah Pengiriman.")
    doc.add_paragraph("b.\tWaktu pelaksanaan SPK adalah sejak tanggal mulai kerja yang tercantum dalam Surat Perintah Pengiriman.")
    doc.add_paragraph("c.\tPenyedia harus menyelesaikan pekerjaan sesuai jadwal yang ditentukan.")
    doc.add_paragraph("d.\tApabila penyedia tidak dapat menyelesaikan pekerjaan sesuai jadwal karena keadaan diluar pengendaliannya dan penyedia telah melaporkan kejadian tersebut kepada Pejabat Penandatangan Kontrak, maka Pejabat Penandatangan Kontrak dapat melakukan penjadwalan kembali pelaksanaan tugas penyedia dengan adendum SPK.")
    doc.add_paragraph()

    # 9. ASURANSI
    p = doc.add_paragraph()
    p.add_run("9.\tASURANSI").bold = True
    doc.add_paragraph("a.\tApabila dipersyaratkan, penyedia wajib menyediakan asuransi sejak Surat Perintah Pengiriman sampai dengan tanggal selesainya pemeliharaan untuk:")
    doc.add_paragraph("\t1)\tsemua barang dan peralatan yang mempunyai risiko tinggi terjadinya kecelakaan, pelaksanaan pekerjaan, serta pekerja untuk pelaksanaan pekerjaan, atas segala risiko terhadap kecelakaan, kerusakan, kehilangan, serta risiko lain yang tidak dapat diduga;")
    doc.add_paragraph("\t2)\tpihak ketiga sebagai akibat kecelakaan di tempat kerjanya; dan")
    doc.add_paragraph("b.\tBesarnya asuransi sudah diperhitungkan dalam penawaran dan termasuk dalam harga SPK.")
    doc.add_paragraph()

    # 10. PENANGGUNGAN DAN RISIKO
    p = doc.add_paragraph()
    p.add_run("10.\tPENANGGUNGAN DAN RISIKO").bold = True
    doc.add_paragraph("a.\tPenyedia berkewajiban untuk melindungi, membebaskan, dan menanggung tanpa batas Pejabat Penandatangan Kontrak beserta instansinya terhadap semua bentuk tuntutan, tanggung jawab, kewajiban, kehilangan, kerugian, denda, gugatan atau tuntutan hukum, proses pemeriksaan hukum, dan biaya yang dikenakan terhadap Pejabat Penandatangan Kontrak beserta instansinya sehubungan dengan klaim yang timbul dari hal-hal berikut terhitung sejak tanggal mulai kerja sampai dengan tanggal penandatanganan berita acara penyerahan akhir.")
    doc.add_paragraph("b.\tTerhitung sejak tanggal mulai kerja sampai dengan tanggal penandatanganan berita acara serah terima, semua risiko kehilangan atau kerusakan hasil pekerjaan ini merupakan risiko penyedia, kecuali kerugian atau kerusakan tersebut diakibatkan oleh kesalahan atau kelalaian Pejabat Penandatangan Kontrak.")
    doc.add_paragraph("c.\tPertanggungan asuransi yang dimiliki oleh penyedia tidak membatasi kewajiban penanggungan dalam syarat ini.")
    doc.add_paragraph("d.\tKehilangan atau kerusakan terhadap hasil pekerjaan sejak tanggal mulai kerja sampai batas akhir garansi, harus diperbaiki, diganti atau dilengkapi oleh penyedia atas tanggungannya sendiri jika kehilangan atau kerusakan tersebut terjadi akibat tindakan atau kelalaian penyedia.")
    doc.add_paragraph()

    # 11. PENGAWASAN DAN PEMERIKSAAN
    add_su_item(doc, 11, "PENGAWASAN DAN PEMERIKSAAN",
        "Pejabat Penandatangan Kontrak berwenang melakukan pengawasan dan pemeriksaan terhadap pelaksanaan pekerjaan yang dilaksanakan oleh penyedia. Pejabat Penandatangan Kontrak dapat memerintahkan kepada pihak lain untuk melakukan pengawasan dan pemeriksaan atas semua pelaksanaan pekerjaan yang dilaksanakan oleh penyedia.")

    # 12. PENGUJIAN
    add_su_item(doc, 12, "PENGUJIAN",
        "Jika Pejabat Penandatangan Kontrak atau Pengawas Pekerjaan memerintahkan penyedia untuk melakukan pengujian Cacat Mutu yang tidak tercantum dalam Spesifikasi Teknis dan Gambar, dan hasil uji coba menunjukkan adanya Cacat Mutu maka penyedia berkewajiban untuk menanggung biaya pengujian tersebut. Jika tidak ditemukan adanya Cacat Mutu maka uji coba tersebut dianggap sebagai Peristiwa Kompensasi.")

    # 13. LAPORAN HASIL PEKERJAAN
    p = doc.add_paragraph()
    p.add_run("13.\tLAPORAN HASIL PEKERJAAN").bold = True
    doc.add_paragraph("a.\tPemeriksaan pekerjaan dilakukan selama pelaksanaan Kontrak terhadap kemajuan pekerjaan dalam rangka pengawasan kualitas dan waktu pelaksanaan pekerjaan. Hasil pemeriksaan pekerjaan dituangkan dalam laporan kemajuan hasil pekerjaan.")
    doc.add_paragraph("b.\tUntuk merekam pelaksanaan pekerjaan, Pejabat Penandatangan Kontrak dapat menugaskan Pengawas Pekerjaan dan/atau tim teknis membuat foto-foto dokumentasi pelaksanaan pekerjaan di lokasi pekerjaan.")
    doc.add_paragraph()

    # 14. WAKTU PENYELESAIAN PEKERJAAN
    p = doc.add_paragraph()
    p.add_run("14.\tWAKTU PENYELESAIAN PEKERJAAN").bold = True
    doc.add_paragraph("a.\tKecuali SPK diputuskan lebih awal, penyedia berkewajiban untuk memulai pelaksanaan pekerjaan pada tanggal mulai kerja, dan melaksanakan pekerjaan sesuai dengan program mutu, serta menyelesaikan pekerjaan selambat-lambatnya pada tanggal penyelesaian yang ditetapkan dalam surat perintah pengiriman.")
    doc.add_paragraph("b.\tJika pekerjaan tidak selesai pada tanggal penyelesaian disebabkan karena kesalahan atau kelalaian penyedia maka penyedia dikenakan sanksi berupa denda keterlambatan.")
    doc.add_paragraph("c.\tJika keterlambatan tersebut disebabkan oleh Peristiwa Kompensasi maka Pejabat Penandatangan Kontrak memberikan tambahan perpanjangan waktu penyelesaian pekerjaan.")
    doc.add_paragraph("d.\tTanggal penyelesaian yang dimaksud dalam ketentuan ini adalah tanggal penyelesaian semua pekerjaan.")
    doc.add_paragraph()

    # 15. SERAH TERIMA PEKERJAAN
    p = doc.add_paragraph()
    p.add_run("15.\tSERAH TERIMA PEKERJAAN").bold = True
    doc.add_paragraph("a.\tSetelah pekerjaan selesai 100% (seratus persen), penyedia mengajukan permintaan secara tertulis kepada Pejabat Penandatangan Kontrak untuk penyerahan pekerjaan.")
    doc.add_paragraph("b.\tSebelum dilakukan serah terima, Pejabat Penandatangan Kontrak melakukan pemeriksaan terhadap hasil pekerjaan.")
    doc.add_paragraph("c.\tPejabat Penandatangan Kontrak dalam melakukan pemeriksaan hasil pekerjaan dapat dibantu oleh pengawas pekerjaan dan/atau tim teknis.")
    doc.add_paragraph("d.\tApabila terdapat kekurangan-kekurangan dan/atau cacat hasil pekerjaan, penyedia wajib memperbaiki/menyelesaikannya, atas perintah Pejabat Penandatangan Kontrak.")
    doc.add_paragraph("e.\tPejabat Penandatangan Kontrak menerima hasil pekerjaan setelah seluruh hasil pekerjaan dilaksanakan sesuai dengan ketentuan SPK.")
    doc.add_paragraph("f.\tPembayaran dilakukan sebesar 100% (seratus persen) dari harga SPK dan penyedia harus menyerahkan Sertifikat Garansi.")
    doc.add_paragraph()

    # 16. JAMINAN BEBAS CACAT MUTU/GARANSI
    p = doc.add_paragraph()
    p.add_run("16.\tJAMINAN BEBAS CACAT MUTU/GARANSI").bold = True
    doc.add_paragraph("a.\tPenyedia dengan jaminan pabrikan dari produsen pabrikan (jika ada) berkewajiban untuk menjamin bahwa selama penggunaan secara wajar, Barang tidak mengandung cacat mutu yang disebabkan oleh tindakan atau kelalaian Penyedia, atau cacat mutu akibat desain, bahan, dan cara kerja.")
    doc.add_paragraph("b.\tJaminan bebas cacat mutu ini berlaku selama masa garansi berlaku.")
    doc.add_paragraph("c.\tPejabat Penandatangan Kontrak akan menyampaikan pemberitahuan cacat mutu kepada Penyedia segera setelah ditemukan cacat mutu tersebut selama masa garansi berlaku.")
    doc.add_paragraph("d.\tTerhadap pemberitahuan cacat mutu oleh Pejabat Penandatangan Kontrak, Penyedia berkewajiban untuk memperbaiki, mengganti, dan/atau melengkapi Barang dalam jangka waktu sesuai dengan syarat dan ketentuan dalam Sertifikat Garansi.")
    doc.add_paragraph("e.\tJika Penyedia tidak memperbaiki, mengganti, atau melengkapi Barang akibat cacat mutu dalam jangka waktu sesuai dengan syarat dan ketentuan dalam Sertifikat Garansi, Pejabat Penandatangan Kontrak akan menghitung biaya perbaikan yang diperlukan, dan Pejabat Penandatangan Kontrak secara langsung atau melalui pihak ketiga yang ditunjuk oleh Pejabat Penandatangan Kontrak akan melakukan perbaikan tersebut.")
    doc.add_paragraph("f.\tSelain kewajiban penggantian biaya, Penyedia yang lalai memperbaiki cacat mutu dikenakan Sanksi Daftar Hitam.")
    doc.add_paragraph()

    # 17. PERUBAHAN SPK
    p = doc.add_paragraph()
    p.add_run("17.\tPERUBAHAN SPK").bold = True
    doc.add_paragraph("a.\tSPK hanya dapat diubah melalui adendum SPK.")
    doc.add_paragraph("b.\tPerubahan SPK dapat dilaksanakan dalam hal terdapat perbedaan antara kondisi lapangan pada saat pelaksanaan dengan SPK dan disetujui oleh para pihak, meliputi:")
    doc.add_paragraph("\t1)\tMenambah atau mengurangi volume yang tercantum dalam SPK;")
    doc.add_paragraph("\t2)\tMenambah dan/atau mengurangi jenis kegiatan;")
    doc.add_paragraph("\t3)\tMengubah spesifikasi teknis sesuai dengan kondisi lapangan; dan/atau")
    doc.add_paragraph("\t4)\tMengubah jadwal pelaksanaan pekerjaan.")
    doc.add_paragraph("c.\tUntuk kepentingan perubahan SPK, Pejabat Penandatangan Kontrak dapat dibantu Pejabat Peneliti Pelaksanaan Kontrak.")
    doc.add_paragraph()

    # 18. PERISTIWA KOMPENSASI
    p = doc.add_paragraph()
    p.add_run("18.\tPERISTIWA KOMPENSASI").bold = True
    doc.add_paragraph("a.\tPeristiwa Kompensasi dapat diberikan kepada penyedia dalam hal sebagai berikut:")
    doc.add_paragraph("\t1)\tPejabat Penandatangan Kontrak mengubah jadwal yang dapat mempengaruhi pelaksanaan pekerjaan;")
    doc.add_paragraph("\t2)\tKeterlambatan pembayaran kepada penyedia;")
    doc.add_paragraph("\t3)\tPejabat Penandatangan Kontrak tidak memberikan gambar-gambar, spesifikasi dan/atau instruksi sesuai jadwal yang dibutuhkan;")
    doc.add_paragraph("\t4)\tPenyedia belum bisa masuk ke lokasi sesuai jadwal;")
    doc.add_paragraph("\t5)\tPejabat Penandatangan Kontrak menginstruksikan kepada pihak penyedia untuk melakukan pengujian tambahan yang setelah dilaksanakan pengujian ternyata tidak ditemukan kerusakan/kegagalan/penyimpangan;")
    doc.add_paragraph("\t6)\tPejabat Penandatangan Kontrak memerintahkan penundaan pelaksanaan pekerjaan;")
    doc.add_paragraph("\t7)\tPejabat Penandatangan Kontrak memerintahkan untuk mengatasi kondisi tertentu yang tidak dapat diduga sebelumnya dan disebabkan oleh Pejabat Penandatangan Kontrak;")
    doc.add_paragraph("\t8)\tKetentuan lain dalam SPK.")
    doc.add_paragraph("b.\tJika Peristiwa Kompensasi mengakibatkan pengeluaran tambahan dan/atau keterlambatan penyelesaian pekerjaan maka Pejabat Penandatangan Kontrak berkewajiban untuk membayar ganti rugi dan/atau memberikan perpanjangan waktu penyelesaian pekerjaan.")
    doc.add_paragraph("c.\tGanti rugi hanya dapat dibayarkan jika berdasarkan data penunjang dan perhitungan kompensasi yang diajukan oleh penyedia kepada Pejabat Penandatangan Kontrak, dapat dibuktikan kerugian nyata akibat Peristiwa Kompensasi.")
    doc.add_paragraph("d.\tPerpanjangan waktu penyelesaian pekerjaan hanya dapat diberikan jika berdasarkan data penunjang dan perhitungan kompensasi yang diajukan oleh penyedia kepada Pejabat Penandatangan Kontrak, dapat dibuktikan perlunya tambahan waktu akibat Peristiwa Kompensasi.")
    doc.add_paragraph("e.\tPenyedia tidak berhak atas ganti rugi dan/atau perpanjangan waktu penyelesaian pekerjaan jika penyedia gagal atau lalai untuk memberikan peringatan dini dalam mengantisipasi atau mengatasi dampak Peristiwa Kompensasi.")
    doc.add_paragraph()

    # 19. PERPANJANGAN WAKTU
    p = doc.add_paragraph()
    p.add_run("19.\tPERPANJANGAN WAKTU").bold = True
    doc.add_paragraph("a.\tJika terjadi Peristiwa Kompensasi sehingga penyelesaian pekerjaan akan melampaui tanggal penyelesaian maka penyedia berhak untuk meminta perpanjangan tanggal penyelesaian berdasarkan data penunjang. Pejabat Penandatangan Kontrak berdasarkan pertimbangan Pengawas Pekerjaan memperpanjang tanggal penyelesaian pekerjaan secara tertulis. Perpanjangan tanggal penyelesaian harus dilakukan melalui adendum SPK.")
    doc.add_paragraph("b.\tPejabat Penandatangan Kontrak dapat menyetujui perpanjangan waktu pelaksanaan setelah melakukan penelitian terhadap usulan tertulis yang diajukan oleh penyedia.")
    doc.add_paragraph()

    # 20. PENGHENTIAN DAN PEMUTUSAN SPK
    p = doc.add_paragraph()
    p.add_run("20.\tPENGHENTIAN DAN PEMUTUSAN SPK").bold = True
    doc.add_paragraph("a.\tPenghentian SPK dapat dilakukan karena terjadi Keadaan Kahar.")
    doc.add_paragraph("b.\tDalam hal SPK dihentikan, Pejabat Penandatangan Kontrak wajib membayar kepada penyedia sesuai dengan prestasi pekerjaan yang telah dicapai.")
    doc.add_paragraph("c.\tPemutusan SPK dapat dilakukan oleh pihak Pejabat Penandatangan Kontrak atau pihak penyedia.")
    doc.add_paragraph("d.\tMenyimpang dari Pasal 1266 dan 1267 Kitab Undang-Undang Hukum Perdata, pemutusan SPK melalui pemberitahuan tertulis dapat dilakukan apabila:")
    doc.add_paragraph("\t1)\tPenyedia terbukti melakukan KKN, kecurangan dan/atau pemalsuan dalam proses Pengadaan yang diputuskan oleh instansi yang berwenang;")
    doc.add_paragraph("\t2)\tPengaduan tentang penyimpangan prosedur, dugaan KKN dan/atau pelanggaran persaingan sehat dalam pelaksanaan pengadaan dinyatakan benar oleh instansi yang berwenang;")
    doc.add_paragraph("\t3)\tPenyedia lalai/cidera janji dalam melaksanakan kewajibannya dan tidak memperbaiki kelalaiannya dalam jangka waktu yang telah ditetapkan;")
    doc.add_paragraph("\t4)\tPenyedia tanpa persetujuan Pejabat Penandatangan Kontrak, tidak memulai pelaksanaan pekerjaan;")
    doc.add_paragraph("\t5)\tPenyedia menghentikan pekerjaan dan penghentian ini tidak tercantum dalam program mutu serta tanpa persetujuan Pejabat Penandatangan Kontrak;")
    doc.add_paragraph("\t6)\tPenyedia berada dalam keadaan pailit;")
    doc.add_paragraph("\t7)\tPenyedia gagal memperbaiki kinerja setelah mendapat Surat Peringatan sebanyak 3 (tiga) kali;")
    doc.add_paragraph("\t8)\tPenyedia selama Masa SPK gagal memperbaiki Cacat Mutu dalam jangka waktu yang ditetapkan oleh Pejabat Penandatangan Kontrak;")
    doc.add_paragraph("\t9)\tPejabat Penandatangan Kontrak memerintahkan penyedia untuk menunda pelaksanaan atau kelanjutan pekerjaan, dan perintah tersebut tidak ditarik selama 28 (dua puluh delapan) hari; dan/atau")
    doc.add_paragraph("\t10)\tPejabat Penandatangan Kontrak tidak menerbitkan surat perintah pembayaran untuk pembayaran tagihan angsuran sesuai dengan yang disepakati sebagaimana tercantum dalam SPK.")
    doc.add_paragraph("e.\tDalam hal pemutusan SPK dilakukan karena kesalahan penyedia:")
    doc.add_paragraph("\t1)\tSisa uang muka harus dilunasi oleh Penyedia atau Jaminan Uang Muka dicairkan (apabila diberikan);")
    doc.add_paragraph("\t2)\tPenyedia membayar denda keterlambatan (apabila ada); dan/atau")
    doc.add_paragraph("\t3)\tPenyedia dikenakan Sanksi Daftar Hitam.")
    doc.add_paragraph("f.\tDalam hal pemutusan SPK dilakukan karena Pejabat Penandatangan Kontrak terlibat penyimpangan prosedur, melakukan KKN dan/atau pelanggaran persaingan sehat dalam pelaksanaan pengadaan, maka Pejabat Penandatangan Kontrak dikenakan sanksi berdasarkan peraturan perundang-undangan.")
    doc.add_paragraph()

    # 21. PEMBAYARAN
    p = doc.add_paragraph()
    p.add_run("21.\tPEMBAYARAN").bold = True
    doc.add_paragraph("a.\tPembayaran prestasi hasil pekerjaan yang disepakati dilakukan oleh Pejabat Penandatangan Kontrak, dengan ketentuan:")
    doc.add_paragraph("\t1)\tPenyedia telah mengajukan tagihan disertai laporan kemajuan hasil pekerjaan;")
    doc.add_paragraph("\t2)\tPembayaran dilakukan dengan [sistem termin/pembayaran secara sekaligus];")
    doc.add_paragraph("\t3)\tPembayaran harus dipotong denda (apabila ada), dan pajak;")
    doc.add_paragraph("b.\tPembayaran terakhir hanya dilakukan setelah pekerjaan selesai 100% (seratus persen) dan Berita Acara Serah Terima ditandatangani.")
    doc.add_paragraph("c.\tPejabat Penandatangan Kontrak dalam kurun waktu 7 (tujuh) hari kerja setelah pengajuan permintaan pembayaran dari penyedia harus sudah mengajukan surat permintaan pembayaran kepada Pejabat Penandatangan Surat Perintah Membayar (PPSPM).")
    doc.add_paragraph("d.\tBila terdapat ketidaksesuaian dalam perhitungan angsuran, tidak akan menjadi alasan untuk menunda pembayaran. Pejabat Penandatangan Kontrak dapat meminta penyedia untuk menyampaikan perhitungan prestasi sementara dengan mengesampingkan hal-hal yang sedang menjadi perselisihan.")
    doc.add_paragraph()

    # 22. DENDA
    p = doc.add_paragraph()
    p.add_run("22.\tDENDA").bold = True
    doc.add_paragraph("a.\tJika pekerjaan tidak dapat diselesaikan dalam jangka waktu pelaksanaan pekerjaan karena kesalahan atau kelalaian Penyedia maka Penyedia berkewajiban untuk membayar denda kepada Pejabat Penandatangan Kontrak sebesar 1/1000 (satu permil) dari nilai SPK (tidak termasuk PPN) untuk setiap hari keterlambatan atau 1/1000 (satu permil) dari nilai bagian SPK yang tercantum dalam SPK (tidak termasuk PPN).")
    doc.add_paragraph("b.\tPejabat Penandatangan Kontrak mengenakan Denda dengan memotong pembayaran prestasi pekerjaan penyedia. Pembayaran Denda tidak mengurangi tanggung jawab kontraktual penyedia.")
    doc.add_paragraph()

    # 23. PENYELESAIAN PERSELISIHAN
    add_su_item(doc, 23, "PENYELESAIAN PERSELISIHAN",
        "Pejabat Penandatangan Kontrak dan penyedia berkewajiban untuk berupaya sungguh-sungguh menyelesaikan secara damai semua perselisihan yang timbul dari atau berhubungan dengan SPK ini atau interpretasinya selama atau setelah pelaksanaan pekerjaan. Jika perselisihan tidak dapat diselesaikan secara musyawarah maka perselisihan akan diselesaikan melalui Layanan Penyelesaian Sengketa, arbitrase atau Pengadilan Negeri.")

    # 24. LARANGAN PEMBERIAN KOMISI
    add_su_item(doc, 24, "LARANGAN PEMBERIAN KOMISI",
        "Penyedia menjamin bahwa tidak satu pun personel satuan kerja Pejabat Penandatangan Kontrak telah atau akan menerima komisi atau keuntungan tidak sah lainnya baik langsung maupun tidak langsung dari SPK ini. Penyedia menyetujui bahwa pelanggaran syarat ini merupakan pelanggaran yang mendasar terhadap SPK ini.")

    # Save
    filepath = os.path.join(TEMPLATES_DIR, "spk_barang.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


def add_su_item(doc, nomor, judul, isi):
    """Helper function to add Syarat Umum item"""
    p = doc.add_paragraph()
    p.add_run(f"{nomor}.\t{judul}").bold = True
    doc.add_paragraph(isi)
    doc.add_paragraph()


# =============================================================================
# SPK JASA LAINNYA
# =============================================================================
def create_spk_jasa_lainnya():
    """Create SPK template for Jasa Lainnya - Format Tabel"""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    run = title.add_run("RANCANGAN SURAT PERINTAH KERJA (SPK)")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Main table with border
    main_table = doc.add_table(rows=4, cols=2)
    main_table.style = 'Table Grid'

    # Set column widths
    for row in main_table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(9)

    # Row 1: Header SPK and Satuan Kerja info
    cell_0_0 = main_table.rows[0].cells[0]
    cell_0_0.text = ""
    p = cell_0_0.paragraphs[0]
    run = p.add_run("SURAT PERINTAH KERJA (SPK)")
    run.bold = True

    cell_0_1 = main_table.rows[0].cells[1]
    cell_0_1.text = ""
    p = cell_0_1.paragraphs[0]
    p.add_run("SATUAN KERJA").bold = True
    p.add_run(" : {{satker_nama}}")
    cell_0_1.add_paragraph()
    p2 = cell_0_1.add_paragraph()
    p2.add_run("NOMOR DAN TANGGAL SPK").bold = True
    p2.add_run(" : ")
    cell_0_1.add_paragraph("{{nomor_spk}} tanggal {{tanggal_spk}}")

    # Row 2: Nama Pejabat Penandatangan Kontrak
    cell_1_0 = main_table.rows[1].cells[0]
    cell_1_0.text = ""
    p = cell_1_0.paragraphs[0]
    p.add_run("Nama Pejabat Penandatangan Kontrak").bold = True
    p.add_run(":")

    cell_1_1 = main_table.rows[1].cells[1]
    cell_1_1.text = "{{ppk_nama}}"

    # Row 3: Nama Penyedia
    cell_2_0 = main_table.rows[2].cells[0]
    cell_2_0.text = ""
    p = cell_2_0.paragraphs[0]
    p.add_run("Nama Penyedia").bold = True
    p.add_run(":")

    cell_2_1 = main_table.rows[2].cells[1]
    cell_2_1.text = "{{penyedia_nama}}"

    # Row 4: Paket Pengadaan and Document Numbers
    cell_3_0 = main_table.rows[3].cells[0]
    cell_3_0.text = ""
    p = cell_3_0.paragraphs[0]
    p.add_run("PAKET PENGADAAN").bold = True
    p.add_run(" :")
    cell_3_0.add_paragraph("{{nama_paket}}")

    cell_3_1 = main_table.rows[3].cells[1]
    cell_3_1.text = ""
    # Add document info
    p = cell_3_1.paragraphs[0]
    p.add_run("NOMOR SURAT UNDANGAN PENGADAAN LANGSUNG").bold = True
    p.add_run(" :")
    cell_3_1.add_paragraph("{{nomor_undangan}}")
    cell_3_1.add_paragraph()
    p2 = cell_3_1.add_paragraph()
    p2.add_run("TANGGAL SURAT UNDANGAN PENGADAAN LANGSUNG").bold = True
    p2.add_run(" :")
    cell_3_1.add_paragraph("{{tanggal_undangan}}")
    cell_3_1.add_paragraph()
    p3 = cell_3_1.add_paragraph()
    p3.add_run("NOMOR BERITA ACARA HASIL PENGADAAN LANGSUNG").bold = True
    p3.add_run(" :")
    cell_3_1.add_paragraph("{{nomor_ba_pl}}")
    cell_3_1.add_paragraph()
    p4 = cell_3_1.add_paragraph()
    p4.add_run("TANGGAL BERITA ACARA HASIL PENGADAAN LANGSUNG").bold = True
    p4.add_run(" :")
    cell_3_1.add_paragraph("{{tanggal_ba_pl}}")

    doc.add_paragraph()

    # Sumber Dana
    p_sumber = doc.add_paragraph()
    p_sumber.add_run("SUMBER DANA: ").bold = True
    p_sumber.add_run("[sebagai contoh, cantumkan \"dibebankan atas DIPA/DPA ")
    p_sumber.add_run("{{satker_nama}}")
    p_sumber.add_run(" Tahun Anggaran ")
    p_sumber.add_run("{{tahun_anggaran}}")
    p_sumber.add_run(" untuk mata anggaran kegiatan ")
    p_sumber.add_run("{{kode_akun}}")
    p_sumber.add_run("\"]")

    doc.add_paragraph()

    # Nilai Kontrak
    p_nilai = doc.add_paragraph()
    p_nilai.add_run("Nilai Kontrak termasuk Pajak Pertambahan Nilai (PPN) adalah sebesar ")
    p_nilai.add_run("{{nilai_kontrak_fmt}}").bold = True
    p_nilai.add_run(" (")
    p_nilai.add_run("{{nilai_kontrak_terbilang}}")
    p_nilai.add_run(" rupiah).")

    doc.add_paragraph()

    # Jenis Kontrak (khusus jasa lainnya)
    p_jenis = doc.add_paragraph()
    p_jenis.add_run("Jenis Kontrak ").bold = True
    p_jenis.add_run("{{jenis_kontrak}}")

    doc.add_paragraph()

    # Waktu Pelaksanaan
    p_waktu = doc.add_paragraph()
    p_waktu.add_run("WAKTU PELAKSANAAN PEKERJAAN: ").bold = True
    p_waktu.add_run("{{jangka_waktu}}").bold = True
    p_waktu.add_run(" (")
    p_waktu.add_run("{{jangka_waktu_terbilang}}")
    p_waktu.add_run(") hari kalender")

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature table
    ttd_table = doc.add_table(rows=7, cols=2)

    # Headers
    cell_left = ttd_table.rows[0].cells[0]
    cell_left.text = ""
    p = cell_left.paragraphs[0]
    p.add_run("Untuk dan atas nama").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = cell_left.add_paragraph()
    p2.add_run("Pejabat Penandatangan Kontrak").bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_right = ttd_table.rows[0].cells[1]
    cell_right.text = ""
    p = cell_right.paragraphs[0]
    p.add_run("Untuk dan atas nama Penyedia").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Note for PPK side
    note_left = ttd_table.rows[1].cells[0]
    note_left.text = ""
    p = note_left.paragraphs[0]
    p.add_run("[tanda tangan dan cap (jika salinan asli ini untuk Penyedia maka rekatkan meterai Rp10.000,-)]").font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Note for Penyedia side
    note_right = ttd_table.rows[1].cells[1]
    note_right.text = ""
    p = note_right.paragraphs[0]
    p.add_run("[tanda tangan dan cap (jika salinan asli ini untuk proyek/satuan kerja Pejabat Penandatangan Kontrak maka rekatkan meterai Rp10.000,-)]").font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Empty rows for signature space
    for i in range(2, 5):
        ttd_table.rows[i].cells[0].text = ""
        ttd_table.rows[i].cells[1].text = ""

    # Names
    name_left = ttd_table.rows[5].cells[0]
    name_left.text = ""
    p = name_left.paragraphs[0]
    run = p.add_run("{{ppk_nama}}")
    run.bold = True
    run.underline = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_right = ttd_table.rows[5].cells[1]
    name_right.text = ""
    p = name_right.paragraphs[0]
    run = p.add_run("{{direktur_nama}}")
    run.bold = True
    run.underline = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Positions
    pos_left = ttd_table.rows[6].cells[0]
    pos_left.text = ""
    p = pos_left.paragraphs[0]
    p.add_run("{{ppk_jabatan}}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pos_right = ttd_table.rows[6].cells[1]
    pos_right.text = ""
    p = pos_right.paragraphs[0]
    p.add_run("{{direktur_jabatan}}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page break for Syarat Umum
    doc.add_page_break()

    # =========================================================================
    # SYARAT UMUM SPK JASA LAINNYA
    # =========================================================================
    su_title = doc.add_paragraph()
    run = su_title.add_run("SYARAT UMUM")
    run.bold = True
    run.font.size = Pt(14)
    su_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    su_subtitle = doc.add_paragraph()
    run = su_subtitle.add_run("SURAT PERINTAH KERJA (SPK)")
    run.bold = True
    run.font.size = Pt(14)
    su_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1. LINGKUP PEKERJAAN
    add_su_item(doc, 1, "LINGKUP PEKERJAAN",
        "Penyedia yang ditunjuk berkewajiban untuk menyelesaikan pekerjaan dalam jangka waktu yang ditentukan sesuai dengan volume, spesifikasi teknis dan harga yang tercantum dalam SPK.")

    # 2. HUKUM YANG BERLAKU
    add_su_item(doc, 2, "HUKUM YANG BERLAKU",
        "Keabsahan, interpretasi, dan pelaksanaan SPK ini didasarkan kepada hukum Republik Indonesia.")

    # 3. HARGA SPK
    p = doc.add_paragraph()
    p.add_run("3.\tHARGA SPK").bold = True
    doc.add_paragraph("a.\tPejabat Penandatangan Kontrak membayar kepada penyedia sebesar harga SPK.")
    doc.add_paragraph("b.\tHarga SPK telah memperhitungkan keuntungan, beban pajak dan biaya overhead serta biaya asuransi (apabila dipersyaratkan).")
    doc.add_paragraph("c.\tRincian harga SPK sesuai dengan rincian yang tercantum dalam daftar kuantitas dan harga.")
    doc.add_paragraph()

    # 4. HAK KEPEMILIKAN
    p = doc.add_paragraph()
    p.add_run("4.\tHAK KEPEMILIKAN").bold = True
    doc.add_paragraph("a.\tPejabat Penandatangan Kontrak berhak atas kepemilikan semua barang/bahan yang terkait langsung atau disediakan sehubungan dengan jasa yang diberikan oleh penyedia kepada Pejabat Penandatangan Kontrak. Jika diminta oleh Pejabat Penandatangan Kontrak maka penyedia berkewajiban untuk membantu secara optimal pengalihan hak kepemilikan tersebut kepada Pejabat Penandatangan Kontrak sesuai dengan hukum yang berlaku.")
    doc.add_paragraph("b.\tHak kepemilikan atas peralatan dan barang/bahan yang disediakan oleh Pejabat Penandatangan Kontrak tetap pada Pejabat Penandatangan Kontrak, dan semua peralatan tersebut harus dikembalikan kepada Pejabat Penandatangan Kontrak pada saat SPK berakhir atau jika tidak diperlukan lagi oleh penyedia. Semua peralatan tersebut harus dikembalikan dalam kondisi yang sama pada saat diberikan kepada penyedia dengan pengecualian keausan akibat pemakaian yang wajar.")
    doc.add_paragraph()

    # 5. CACAT MUTU
    add_su_item(doc, 5, "CACAT MUTU",
        "Pejabat Penandatangan Kontrak akan memeriksa setiap hasil pekerjaan penyedia dan memberitahukan secara tertulis penyedia atas setiap cacat mutu yang ditemukan. Pejabat Penandatangan Kontrak dapat memerintahkan penyedia untuk menguji pekerjaan yang dianggap oleh Pejabat Penandatangan Kontrak mengandung cacat mutu. Penyedia bertanggung jawab atas cacat mutu selama masa garansi.")

    # 6. PERPAJAKAN
    add_su_item(doc, 6, "PERPAJAKAN",
        "Penyedia berkewajiban untuk membayar semua pajak, bea, retribusi, dan pungutan lain yang sah yang dibebankan oleh hukum yang berlaku atas pelaksanaan SPK. Semua pengeluaran perpajakan ini dianggap telah termasuk dalam harga SPK.")

    # 7. PENGALIHAN DAN/ATAU SUBKONTRAK
    add_su_item(doc, 7, "PENGALIHAN DAN/ATAU SUBKONTRAK",
        "Penyedia dilarang untuk mengalihkan dan/atau mensubkontrakkan sebagian atau seluruh pekerjaan. Pengalihan seluruh pekerjaan hanya diperbolehkan dalam hal pergantian nama penyedia, baik sebagai akibat peleburan (merger) atau akibat lainnya.")

    # 8. JADWAL
    p = doc.add_paragraph()
    p.add_run("8.\tJADWAL").bold = True
    doc.add_paragraph("a.\tSPK ini berlaku efektif pada tanggal penandatanganan oleh para pihak atau pada tanggal yang ditetapkan dalam Surat Perintah Mulai Kerja.")
    doc.add_paragraph("b.\tWaktu pelaksanaan SPK adalah sejak tanggal mulai kerja yang tercantum dalam Surat Perintah Mulai Kerja.")
    doc.add_paragraph("c.\tPenyedia harus menyelesaikan pekerjaan sesuai jadwal yang ditentukan.")
    doc.add_paragraph("d.\tApabila penyedia tidak dapat menyelesaikan pekerjaan sesuai jadwal karena keadaan diluar pengendaliannya dan penyedia telah melaporkan kejadian tersebut kepada Pejabat Penandatangan Kontrak, maka Pejabat Penandatangan Kontrak dapat melakukan penjadwalan kembali pelaksanaan tugas penyedia dengan adendum SPK.")
    doc.add_paragraph()

    # 9. ASURANSI
    p = doc.add_paragraph()
    p.add_run("9.\tASURANSI").bold = True
    doc.add_paragraph("a.\tApabila dipersyaratkan, penyedia wajib menyediakan asuransi sejak Surat Perintah Mulai Kerja sampai dengan tanggal selesainya pemeliharaan untuk:")
    doc.add_paragraph("\t1)\tsemua barang dan peralatan yang mempunyai risiko tinggi terjadinya kecelakaan, pelaksanaan pekerjaan, serta pekerja untuk pelaksanaan pekerjaan, atas segala risiko terhadap kecelakaan, kerusakan, kehilangan, serta risiko lain yang tidak dapat diduga;")
    doc.add_paragraph("\t2)\tpihak ketiga sebagai akibat kecelakaan di tempat kerjanya; dan")
    doc.add_paragraph("b.\tBesarnya asuransi sudah diperhitungkan dalam penawaran dan termasuk dalam harga SPK.")
    doc.add_paragraph()

    # 10. PENANGGUNGAN DAN RISIKO
    p = doc.add_paragraph()
    p.add_run("10.\tPENANGGUNGAN DAN RISIKO").bold = True
    doc.add_paragraph("a.\tPenyedia berkewajiban untuk melindungi, membebaskan, dan menanggung tanpa batas Pejabat Penandatangan Kontrak beserta instansinya terhadap semua bentuk tuntutan, tanggung jawab, kewajiban, kehilangan, kerugian, denda, gugatan atau tuntutan hukum, proses pemeriksaan hukum, dan biaya yang dikenakan terhadap Pejabat Penandatangan Kontrak beserta instansinya sehubungan dengan klaim yang timbul dari hal-hal berikut terhitung sejak tanggal mulai kerja sampai dengan tanggal penandatanganan berita acara penyerahan akhir.")
    doc.add_paragraph("b.\tTerhitung sejak tanggal mulai kerja sampai dengan tanggal penandatanganan berita acara serah terima, semua risiko kehilangan atau kerusakan hasil pekerjaan ini merupakan risiko penyedia, kecuali kerugian atau kerusakan tersebut diakibatkan oleh kesalahan atau kelalaian Pejabat Penandatangan Kontrak.")
    doc.add_paragraph("c.\tPertanggungan asuransi yang dimiliki oleh penyedia tidak membatasi kewajiban penanggungan dalam syarat ini.")
    doc.add_paragraph("d.\tKehilangan atau kerusakan terhadap hasil pekerjaan sejak tanggal mulai kerja sampai batas akhir garansi, harus diperbaiki, diganti atau dilengkapi oleh penyedia atas tanggungannya sendiri jika kehilangan atau kerusakan tersebut terjadi akibat tindakan atau kelalaian penyedia.")
    doc.add_paragraph()

    # 11. PENGAWASAN DAN PEMERIKSAAN
    add_su_item(doc, 11, "PENGAWASAN DAN PEMERIKSAAN",
        "Pejabat Penandatangan Kontrak berwenang melakukan pengawasan dan pemeriksaan terhadap pelaksanaan pekerjaan yang dilaksanakan oleh penyedia. Pejabat Penandatangan Kontrak dapat memerintahkan kepada pihak lain untuk melakukan pengawasan dan pemeriksaan atas semua pelaksanaan pekerjaan yang dilaksanakan oleh penyedia.")

    # 12. PENGUJIAN
    add_su_item(doc, 12, "PENGUJIAN",
        "Jika Pejabat Penandatangan Kontrak atau Pengawas Pekerjaan memerintahkan penyedia untuk melakukan pengujian Cacat Mutu yang tidak tercantum dalam Spesifikasi Teknis dan Gambar, dan hasil uji coba menunjukkan adanya Cacat Mutu maka penyedia berkewajiban untuk menanggung biaya pengujian tersebut. Jika tidak ditemukan adanya Cacat Mutu maka uji coba tersebut dianggap sebagai Peristiwa Kompensasi.")

    # 13. LAPORAN HASIL PEKERJAAN
    p = doc.add_paragraph()
    p.add_run("13.\tLAPORAN HASIL PEKERJAAN").bold = True
    doc.add_paragraph("a.\tPemeriksaan pekerjaan dilakukan selama pelaksanaan Kontrak terhadap kemajuan pekerjaan dalam rangka pengawasan kualitas dan waktu pelaksanaan pekerjaan. Hasil pemeriksaan pekerjaan dituangkan dalam laporan kemajuan hasil pekerjaan.")
    doc.add_paragraph("b.\tUntuk merekam pelaksanaan pekerjaan, Pejabat Penandatangan Kontrak dapat menugaskan Pengawas Pekerjaan dan/atau tim teknis membuat foto-foto dokumentasi pelaksanaan pekerjaan di lokasi pekerjaan.")
    doc.add_paragraph()

    # 14. WAKTU PENYELESAIAN PEKERJAAN
    p = doc.add_paragraph()
    p.add_run("14.\tWAKTU PENYELESAIAN PEKERJAAN").bold = True
    doc.add_paragraph("a.\tKecuali SPK diputuskan lebih awal, penyedia berkewajiban untuk memulai pelaksanaan pekerjaan pada tanggal mulai kerja, dan melaksanakan pekerjaan sesuai dengan program mutu, serta menyelesaikan pekerjaan selambat-lambatnya pada tanggal penyelesaian yang ditetapkan dalam Surat Perintah Mulai Kerja.")
    doc.add_paragraph("b.\tJika pekerjaan tidak selesai pada tanggal penyelesaian disebabkan karena kesalahan atau kelalaian penyedia maka penyedia dikenakan sanksi berupa denda keterlambatan.")
    doc.add_paragraph("c.\tJika keterlambatan tersebut disebabkan oleh Peristiwa Kompensasi maka Pejabat Penandatangan Kontrak memberikan tambahan perpanjangan waktu penyelesaian pekerjaan.")
    doc.add_paragraph("d.\tTanggal penyelesaian yang dimaksud dalam ketentuan ini adalah tanggal penyelesaian semua pekerjaan.")
    doc.add_paragraph()

    # 15. SERAH TERIMA PEKERJAAN
    p = doc.add_paragraph()
    p.add_run("15.\tSERAH TERIMA PEKERJAAN").bold = True
    doc.add_paragraph("a.\tSetelah pekerjaan selesai 100% (seratus persen), penyedia mengajukan permintaan secara tertulis kepada Pejabat Penandatangan Kontrak untuk penyerahan pekerjaan.")
    doc.add_paragraph("b.\tSebelum dilakukan serah terima, Pejabat Penandatangan Kontrak melakukan pemeriksaan terhadap hasil pekerjaan.")
    doc.add_paragraph("c.\tPejabat Penandatangan Kontrak dalam melakukan pemeriksaan hasil pekerjaan dapat dibantu oleh pengawas pekerjaan dan/atau tim teknis.")
    doc.add_paragraph("d.\tApabila terdapat kekurangan-kekurangan dan/atau cacat hasil pekerjaan, penyedia wajib memperbaiki/menyelesaikannya, atas perintah Pejabat Penandatangan Kontrak.")
    doc.add_paragraph("e.\tPejabat Penandatangan Kontrak menerima hasil pekerjaan setelah seluruh hasil pekerjaan dilaksanakan sesuai dengan ketentuan SPK.")
    doc.add_paragraph("f.\tPembayaran dilakukan sebesar 100% (seratus persen) dari harga SPK dan penyedia harus menyerahkan Sertifikat Garansi.")
    doc.add_paragraph()

    # 16. JAMINAN BEBAS CACAT MUTU/GARANSI
    p = doc.add_paragraph()
    p.add_run("16.\tJAMINAN BEBAS CACAT MUTU/GARANSI").bold = True
    doc.add_paragraph("a.\tPenyedia dengan jaminan pabrikan dari produsen pabrikan (jika ada) berkewajiban untuk menjamin bahwa selama penggunaan secara wajar, Barang tidak mengandung cacat mutu yang disebabkan oleh tindakan atau kelalaian Penyedia, atau cacat mutu akibat desain, bahan, dan cara kerja.")
    doc.add_paragraph("b.\tJaminan bebas cacat mutu ini berlaku selama masa garansi berlaku.")
    doc.add_paragraph("c.\tPejabat Penandatangan Kontrak akan menyampaikan pemberitahuan cacat mutu kepada Penyedia segera setelah ditemukan cacat mutu tersebut selama masa garansi berlaku.")
    doc.add_paragraph("d.\tTerhadap pemberitahuan cacat mutu oleh Pejabat Penandatangan Kontrak, Penyedia berkewajiban untuk memperbaiki, mengganti, dan/atau melengkapi Barang dalam jangka waktu sesuai dengan syarat dan ketentuan dalam Sertifikat Garansi.")
    doc.add_paragraph("e.\tJika Penyedia tidak memperbaiki, mengganti, atau melengkapi Barang akibat cacat mutu dalam jangka waktu sesuai dengan syarat dan ketentuan dalam Sertifikat Garansi, Pejabat Penandatangan Kontrak akan menghitung biaya perbaikan yang diperlukan, dan Pejabat Penandatangan Kontrak secara langsung atau melalui pihak ketiga yang ditunjuk oleh Pejabat Penandatangan Kontrak akan melakukan perbaikan tersebut. Penyedia berkewajiban untuk membayar biaya perbaikan atau penggantian tersebut sesuai dengan klaim yang diajukan secara tertulis oleh Pejabat Penandatangan Kontrak.")
    doc.add_paragraph("f.\tSelain kewajiban penggantian biaya, Penyedia yang lalai memperbaiki cacat mutu dikenakan Sanksi Daftar Hitam.")
    doc.add_paragraph()

    # 17. PERUBAHAN SPK
    p = doc.add_paragraph()
    p.add_run("17.\tPERUBAHAN SPK").bold = True
    doc.add_paragraph("a.\tSPK hanya dapat diubah melalui adendum SPK.")
    doc.add_paragraph("b.\tPerubahan SPK dapat dilaksanakan dalam hal terdapat perbedaan antara kondisi lapangan pada saat pelaksanaan dengan SPK dan disetujui oleh para pihak, meliputi:")
    doc.add_paragraph("\t1)\tmenambah atau mengurangi volume yang tercantum dalam SPK;")
    doc.add_paragraph("\t2)\tmenambah dan/atau mengurangi jenis kegiatan;")
    doc.add_paragraph("\t3)\tmengubah spesifikasi teknis sesuai dengan kondisi lapangan; dan/atau")
    doc.add_paragraph("\t4)\tmengubah jadwal pelaksanaan pekerjaan.")
    doc.add_paragraph("c.\tUntuk kepentingan perubahan SPK, Pejabat Penandatangan Kontrak dapat dibantu Pejabat Peneliti Pelaksanaan Kontrak.")
    doc.add_paragraph()

    # 18. PERISTIWA KOMPENSASI
    p = doc.add_paragraph()
    p.add_run("18.\tPERISTIWA KOMPENSASI").bold = True
    doc.add_paragraph("a.\tPeristiwa Kompensasi dapat diberikan kepada penyedia dalam hal sebagai berikut:")
    doc.add_paragraph("\t1)\tPejabat Penandatangan Kontrak mengubah jadwal yang dapat mempengaruhi pelaksanaan pekerjaan;")
    doc.add_paragraph("\t2)\tketerlambatan pembayaran kepada penyedia;")
    doc.add_paragraph("\t3)\tPejabat Penandatangan Kontrak tidak memberikan gambar-gambar, spesifikasi dan/atau instruksi sesuai jadwal yang dibutuhkan;")
    doc.add_paragraph("\t4)\tpenyedia belum bisa masuk ke lokasi sesuai jadwal;")
    doc.add_paragraph("\t5)\tPejabat Penandatangan Kontrak menginstruksikan kepada pihak penyedia untuk melakukan pengujian tambahan yang setelah dilaksanakan pengujian ternyata tidak ditemukan kerusakan/kegagalan/penyimpangan;")
    doc.add_paragraph("\t6)\tPejabat Penandatangan Kontrak memerintahkan penundaan pelaksanaan pekerjaan;")
    doc.add_paragraph("\t7)\tPejabat Penandatangan Kontrak memerintahkan untuk mengatasi kondisi tertentu yang tidak dapat diduga sebelumnya dan disebabkan oleh Pejabat Penandatangan Kontrak;")
    doc.add_paragraph("\t8)\tketentuan lain dalam SPK.")
    doc.add_paragraph("b.\tJika Peristiwa Kompensasi mengakibatkan pengeluaran tambahan dan/atau keterlambatan penyelesaian pekerjaan maka Pejabat Penandatangan Kontrak berkewajiban untuk membayar ganti rugi dan/atau memberikan perpanjangan waktu penyelesaian pekerjaan.")
    doc.add_paragraph("c.\tGanti rugi hanya dapat dibayarkan jika berdasarkan data penunjang dan perhitungan kompensasi yang diajukan oleh penyedia kepada Pejabat Penandatangan Kontrak, dapat dibuktikan kerugian nyata akibat Peristiwa Kompensasi.")
    doc.add_paragraph("d.\tPerpanjangan waktu penyelesaian pekerjaan hanya dapat diberikan jika berdasarkan data penunjang dan perhitungan kompensasi yang diajukan oleh penyedia kepada Pejabat Penandatangan Kontrak, dapat dibuktikan perlunya tambahan waktu akibat Peristiwa Kompensasi.")
    doc.add_paragraph("e.\tPenyedia tidak berhak atas ganti rugi dan/atau perpanjangan waktu penyelesaian pekerjaan jika penyedia gagal atau lalai untuk memberikan peringatan dini dalam mengantisipasi atau mengatasi dampak Peristiwa Kompensasi.")
    doc.add_paragraph()

    # 19. PERPANJANGAN WAKTU
    p = doc.add_paragraph()
    p.add_run("19.\tPERPANJANGAN WAKTU").bold = True
    doc.add_paragraph("a.\tJika terjadi Peristiwa Kompensasi sehingga penyelesaian pekerjaan akan melampaui tanggal penyelesaian maka penyedia berhak untuk meminta perpanjangan tanggal penyelesaian berdasarkan data penunjang. Pejabat Penandatangan Kontrak berdasarkan pertimbangan Pengawas Pekerjaan memperpanjang tanggal penyelesaian pekerjaan secara tertulis. Perpanjangan tanggal penyelesaian harus dilakukan melalui adendum SPK.")
    doc.add_paragraph("b.\tPejabat Penandatangan Kontrak dapat menyetujui perpanjangan waktu pelaksanaan setelah melakukan penelitian terhadap usulan tertulis yang diajukan oleh penyedia.")
    doc.add_paragraph()

    # 20. PENGHENTIAN DAN PEMUTUSAN SPK
    p = doc.add_paragraph()
    p.add_run("20.\tPENGHENTIAN DAN PEMUTUSAN SPK").bold = True
    doc.add_paragraph("a.\tPenghentian SPK dapat dilakukan karena terjadi Keadaan Kahar.")
    doc.add_paragraph("b.\tDalam hal SPK dihentikan, Pejabat Penandatangan Kontrak wajib membayar kepada penyedia sesuai dengan prestasi pekerjaan yang telah dicapai.")
    doc.add_paragraph("c.\tPemutusan SPK dapat dilakukan oleh pihak Pejabat Penandatangan Kontrak atau pihak penyedia.")
    doc.add_paragraph("d.\tMenyimpang dari Pasal 1266 dan 1267 Kitab Undang-Undang Hukum Perdata, pemutusan SPK melalui pemberitahuan tertulis dapat dilakukan apabila:")
    doc.add_paragraph("\t1)\tpenyedia terbukti melakukan KKN, kecurangan dan/atau pemalsuan dalam proses Pengadaan yang diputuskan oleh instansi yang berwenang;")
    doc.add_paragraph("\t2)\tpengaduan tentang penyimpangan prosedur, dugaan KKN dan/atau pelanggaran persaingan sehat dalam pelaksanaan pengadaan dinyatakan benar oleh instansi yang berwenang;")
    doc.add_paragraph("\t3)\tpenyedia lalai/cidera janji dalam melaksanakan kewajibannya dan tidak memperbaiki kelalaiannya dalam jangka waktu yang telah ditetapkan;")
    doc.add_paragraph("\t4)\tpenyedia tanpa persetujuan Pejabat Penandatangan Kontrak, tidak memulai pelaksanaan pekerjaan;")
    doc.add_paragraph("\t5)\tpenyedia menghentikan pekerjaan dan penghentian ini tidak tercantum dalam program mutu serta tanpa persetujuan Pejabat Penandatangan Kontrak;")
    doc.add_paragraph("\t6)\tpenyedia berada dalam keadaan pailit;")
    doc.add_paragraph("\t7)\tPenyedia gagal memperbaiki kinerja setelah mendapat Surat Peringatan sebanyak 3 (tiga) kali;")
    doc.add_paragraph("\t8)\tpenyedia selama Masa SPK gagal memperbaiki Cacat Mutu dalam jangka waktu yang ditetapkan oleh Pejabat Penandatangan Kontrak;")
    doc.add_paragraph("\t9)\tPejabat Penandatangan Kontrak memerintahkan penyedia untuk menunda pelaksanaan atau kelanjutan pekerjaan, dan perintah tersebut tidak ditarik selama 28 (dua puluh delapan) hari; dan/atau")
    doc.add_paragraph("\t10)\tPejabat Penandatangan Kontrak tidak menerbitkan surat perintah pembayaran untuk pembayaran tagihan angsuran sesuai dengan yang disepakati sebagaimana tercantum dalam SPK.")
    doc.add_paragraph("e.\tDalam hal pemutusan SPK dilakukan karena kesalahan penyedia:")
    doc.add_paragraph("\t1)\tSisa uang muka harus dilunasi oleh Penyedia atau Jaminan Uang Muka dicairkan (apabila diberikan);")
    doc.add_paragraph("\t2)\tpenyedia membayar denda keterlambatan (apabila ada); dan/atau")
    doc.add_paragraph("\t3)\tpenyedia dikenakan Sanksi Daftar Hitam.")
    doc.add_paragraph("f.\tDalam hal pemutusan SPK dilakukan karena Pejabat Penandatangan Kontrak terlibat penyimpangan prosedur, melakukan KKN dan/atau pelanggaran persaingan sehat dalam pelaksanaan pengadaan, maka Pejabat Penandatangan Kontrak dikenakan sanksi berdasarkan peraturan perundang-undangan.")
    doc.add_paragraph()

    # 21. PEMBAYARAN
    p = doc.add_paragraph()
    p.add_run("21.\tPEMBAYARAN").bold = True
    doc.add_paragraph("a.\tpembayaran prestasi hasil pekerjaan yang disepakati dilakukan oleh Pejabat Penandatangan Kontrak, dengan ketentuan:")
    doc.add_paragraph("\t1)\tpenyedia telah mengajukan tagihan disertai laporan kemajuan hasil pekerjaan;")
    doc.add_paragraph("\t2)\tpembayaran dilakukan dengan [sistem bulanan/sistem termin/pembayaran secara sekaligus];")
    doc.add_paragraph("\t3)\tpembayaran harus dipotong denda (apabila ada), dan pajak;")
    doc.add_paragraph("b.\tpembayaran terakhir hanya dilakukan setelah pekerjaan selesai 100% (seratus persen) dan Berita Acara Serah Terima ditandatangani.")
    doc.add_paragraph("c.\tPejabat Penandatangan Kontrak dalam kurun waktu 7 (tujuh) hari kerja setelah pengajuan permintaan pembayaran dari penyedia harus sudah mengajukan surat permintaan pembayaran kepada Pejabat Penandatangan Surat Perintah Membayar (PPSPM).")
    doc.add_paragraph("d.\tbila terdapat ketidaksesuaian dalam perhitungan angsuran, tidak akan menjadi alasan untuk menunda pembayaran. Pejabat Penandatangan Kontrak dapat meminta penyedia untuk menyampaikan perhitungan prestasi sementara dengan mengesampingkan hal-hal yang sedang menjadi perselisihan.")
    doc.add_paragraph()

    # 22. DENDA
    p = doc.add_paragraph()
    p.add_run("22.\tDENDA").bold = True
    doc.add_paragraph("a.\tJika pekerjaan tidak dapat diselesaikan dalam jangka waktu pelaksanaan pekerjaan karena kesalahan atau kelalaian Penyedia maka Penyedia berkewajiban untuk membayar denda kepada Pejabat Penandatangan Kontrak sebesar 1/1000 (satu permil) dari nilai SPK (tidak termasuk PPN) untuk setiap hari keterlambatan atau 1/1000 (satu permil) dari nilai bagian SPK yang tercantum dalam SPK (tidak termasuk PPN).")
    doc.add_paragraph("b.\tPejabat Penandatangan Kontrak mengenakan Denda dengan memotong pembayaran prestasi pekerjaan penyedia. Pembayaran Denda tidak mengurangi tanggung jawab kontraktual penyedia.")
    doc.add_paragraph()

    # 23. PENYELESAIAN PERSELISIHAN
    add_su_item(doc, 23, "PENYELESAIAN PERSELISIHAN",
        "Pejabat Penandatangan Kontrak dan penyedia berkewajiban untuk berupaya sungguh-sungguh menyelesaikan secara damai semua perselisihan yang timbul dari atau berhubungan dengan SPK ini atau interpretasinya selama atau setelah pelaksanaan pekerjaan. Jika perselisihan tidak dapat diselesaikan secara musyawarah maka perselisihan akan diselesaikan melalui Layanan Penyelesaian Sengketa, arbitrase atau Pengadilan Negeri.")

    # 24. LARANGAN PEMBERIAN KOMISI
    add_su_item(doc, 24, "LARANGAN PEMBERIAN KOMISI",
        "Penyedia menjamin bahwa tidak satu pun personel satuan kerja Pejabat Penandatangan Kontrak telah atau akan menerima komisi atau keuntungan tidak sah lainnya baik langsung maupun tidak langsung dari SPK ini. Penyedia menyetujui bahwa pelanggaran syarat ini merupakan pelanggaran yang mendasar terhadap SPK ini.")

    # Save
    filepath = os.path.join(TEMPLATES_DIR, "spk_jasa_lainnya.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# SPK KONSTRUKSI
# =============================================================================
def create_spk_konstruksi():
    """Create SPK template for Konstruksi"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    # Header
    title = doc.add_paragraph()
    title.add_run("SURAT PERINTAH KERJA (SPK)").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("PEKERJAAN KONSTRUKSI").bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_spk}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    doc.add_paragraph("Pada hari ini {{hari_spk}} tanggal {{tanggal_spk_terbilang}} bulan {{bulan_spk}} tahun {{tahun_spk_terbilang}} ({{tanggal_spk_fmt}}), kami yang bertanda tangan di bawah ini:")
    
    doc.add_paragraph()
    
    # PPK
    ppk_table = doc.add_table(rows=4, cols=3)
    ppk_data = [
        ("Nama", ":", "{{ppk_nama}}"),
        ("NIP", ":", "{{ppk_nip}}"),
        ("Jabatan", ":", "Pejabat Pembuat Komitmen"),
        ("Alamat", ":", "{{satker_nama}}, {{satker_alamat}}"),
    ]
    for i, (label, sep, val) in enumerate(ppk_data):
        ppk_table.rows[i].cells[0].text = label
        ppk_table.rows[i].cells[1].text = sep
        ppk_table.rows[i].cells[2].text = val
    
    doc.add_paragraph("selanjutnya disebut sebagai PIHAK KESATU")
    doc.add_paragraph()
    
    # Penyedia
    penyedia_table = doc.add_table(rows=6, cols=3)
    penyedia_data = [
        ("Nama Perusahaan", ":", "{{penyedia_nama}}"),
        ("Nama Direktur", ":", "{{direktur_nama}}"),
        ("Alamat", ":", "{{penyedia_alamat}}"),
        ("NPWP", ":", "{{penyedia_npwp}}"),
        ("Rekening", ":", "{{penyedia_rekening}} - {{penyedia_bank}}"),
        ("Kualifikasi", ":", "{{kualifikasi_penyedia}}"),
    ]
    for i, (label, sep, val) in enumerate(penyedia_data):
        penyedia_table.rows[i].cells[0].text = label
        penyedia_table.rows[i].cells[1].text = sep
        penyedia_table.rows[i].cells[2].text = val
    
    doc.add_paragraph("selanjutnya disebut sebagai PIHAK KEDUA")
    doc.add_paragraph()
    
    doc.add_paragraph("PARA PIHAK sepakat mengikatkan diri dalam SPK Pekerjaan Konstruksi dengan ketentuan:")
    doc.add_paragraph()
    
    # Pasal 1 - Lingkup Pekerjaan
    add_pasal(doc, 1, "LINGKUP PEKERJAAN")
    add_ayat(doc, 1, "PIHAK KEDUA berkewajiban melaksanakan pekerjaan:")
    
    lingkup_table = doc.add_table(rows=4, cols=3)
    lingkup_table.style = 'Table Grid'
    lingkup_data = [
        ("a.", "Nama Paket", "{{nama_paket}}"),
        ("b.", "Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
        ("c.", "Jenis Pekerjaan", "{{jenis_pekerjaan_konstruksi}}"),
        ("d.", "Volume Pekerjaan", "Sesuai Bill of Quantity (BOQ) terlampir"),
    ]
    for i, (no, label, val) in enumerate(lingkup_data):
        lingkup_table.rows[i].cells[0].text = no
        lingkup_table.rows[i].cells[1].text = label
        lingkup_table.rows[i].cells[2].text = val
    
    add_ayat(doc, 2, "Lingkup pekerjaan meliputi penyediaan tenaga kerja, bahan/material, peralatan, dan metode pelaksanaan sesuai spesifikasi teknis.")
    
    doc.add_paragraph()
    
    # Pasal 2 - Nilai Kontrak
    add_pasal(doc, 2, "NILAI KONTRAK")
    add_ayat(doc, 1, "Nilai Kontrak adalah sebesar {{nilai_kontrak_fmt}} ({{nilai_kontrak_terbilang}}) termasuk PPN 11%.")
    add_ayat(doc, 2, "Jenis kontrak adalah {{jenis_kontrak_konstruksi}}:")
    add_huruf(doc, "a", "Lump Sum: harga tetap tidak berubah; ATAU")
    add_huruf(doc, "b", "Unit Price: berdasarkan harga satuan dan volume terpasang.")
    
    doc.add_paragraph()
    
    # Pasal 3 - Jangka Waktu
    add_pasal(doc, 3, "JANGKA WAKTU PELAKSANAAN")
    add_ayat(doc, 1, "Jangka waktu pelaksanaan adalah {{jangka_waktu}} ({{jangka_waktu_terbilang}}) hari kalender.")
    add_ayat(doc, 2, "Waktu pelaksanaan:")
    add_huruf(doc, "a", "Tanggal mulai kerja: {{tanggal_mulai_fmt}}")
    add_huruf(doc, "b", "Tanggal selesai: {{tanggal_selesai_fmt}}")
    add_ayat(doc, 3, "PIHAK KEDUA wajib membuat jadwal pelaksanaan (time schedule) dan kurva S.")
    add_ayat(doc, 4, "Keterlambatan yang bukan kesalahan PIHAK KEDUA dapat diberikan perpanjangan waktu (addendum).")
    
    doc.add_paragraph()
    
    # Pasal 4 - Cara Pembayaran (Konstruksi - Termin & Retensi)
    add_pasal(doc, 4, "CARA PEMBAYARAN")
    add_ayat(doc, 1, "Pembayaran dilakukan secara bertahap (termin) berdasarkan progres fisik pekerjaan:")
    add_huruf(doc, "a", "Termin I: ... % setelah progres fisik mencapai ... %")
    add_huruf(doc, "b", "Termin II: ... % setelah progres fisik mencapai ... %")
    add_huruf(doc, "c", "Termin terakhir: setelah pekerjaan selesai 100%")
    add_ayat(doc, 2, "Setiap pembayaran dipotong retensi sebesar 5% (lima persen) sebagai jaminan pemeliharaan.")
    add_ayat(doc, 3, "Retensi dikembalikan setelah masa pemeliharaan berakhir dan tidak ada cacat/kerusakan.")
    add_ayat(doc, 4, "Pembayaran setelah:")
    add_huruf(doc, "a", "Progres pekerjaan diverifikasi oleh Pengawas/Konsultan Pengawas;")
    add_huruf(doc, "b", "Berita Acara Kemajuan Pekerjaan ditandatangani;")
    add_huruf(doc, "c", "Dokumen tagihan lengkap.")
    
    doc.add_paragraph()
    
    # Pasal 5 - Masa Pemeliharaan (Khusus Konstruksi)
    add_pasal(doc, 5, "MASA PEMELIHARAAN")
    add_ayat(doc, 1, "Masa pemeliharaan adalah {{masa_pemeliharaan}} ({{masa_pemeliharaan_terbilang}}) bulan sejak Serah Terima Pertama (PHO).")
    add_ayat(doc, 2, "Selama masa pemeliharaan, PIHAK KEDUA wajib:")
    add_huruf(doc, "a", "Memperbaiki kerusakan yang terjadi bukan karena kesalahan PIHAK KESATU;")
    add_huruf(doc, "b", "Menanggung biaya perbaikan tersebut.")
    add_ayat(doc, 3, "Setelah masa pemeliharaan berakhir, dilakukan Serah Terima Akhir (FHO).")
    
    doc.add_paragraph()
    
    # Pasal 6 - Tenaga Kerja dan K3
    add_pasal(doc, 6, "TENAGA KERJA DAN K3")
    add_ayat(doc, 1, "PIHAK KEDUA wajib mempekerjakan tenaga kerja yang kompeten dan bersertifikat sesuai kebutuhan.")
    add_ayat(doc, 2, "PIHAK KEDUA wajib menerapkan Sistem Manajemen Keselamatan dan Kesehatan Kerja (SMK3).")
    add_ayat(doc, 3, "PIHAK KEDUA bertanggung jawab atas keselamatan seluruh pekerja di lokasi proyek.")
    add_ayat(doc, 4, "Pelanggaran K3 yang menyebabkan kecelakaan kerja menjadi tanggung jawab PIHAK KEDUA.")
    
    doc.add_paragraph()
    
    # Pasal 7 - Sanksi
    add_pasal(doc, 7, "SANKSI DAN DENDA")
    add_ayat(doc, 1, "Keterlambatan dikenakan denda 1/1000 per hari dari nilai kontrak.")
    add_ayat(doc, 2, "Denda maksimal 5% dari nilai kontrak.")
    add_ayat(doc, 3, "Kegagalan bangunan dalam masa pemeliharaan menjadi tanggung jawab PIHAK KEDUA.")
    add_ayat(doc, 4, "Pemutusan kontrak dapat dilakukan jika pekerjaan tidak dapat diselesaikan.")
    
    doc.add_paragraph()
    
    # Pasal 8 - Penyelesaian Perselisihan
    add_pasal(doc, 8, "PENYELESAIAN PERSELISIHAN")
    add_ayat(doc, 1, "Perselisihan diselesaikan secara musyawarah.")
    add_ayat(doc, 2, "Jika tidak tercapai, melalui Badan Arbitrase atau Pengadilan Negeri {{satker_kota}}.")
    
    doc.add_paragraph()
    
    # Pasal 9 - Penutup
    add_pasal(doc, 9, "PENUTUP")
    add_ayat(doc, 1, "SPK ini berlaku sejak tanggal ditandatangani sampai dengan berakhirnya masa pemeliharaan.")
    add_ayat(doc, 2, "Dibuat dalam rangkap 2 (dua) dengan kekuatan hukum yang sama.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    ttd_table = doc.add_table(rows=6, cols=2)
    ttd_table.rows[0].cells[0].text = "PIHAK KEDUA"
    ttd_table.rows[0].cells[1].text = "PIHAK KESATU"
    ttd_table.rows[1].cells[0].text = "{{penyedia_nama}}"
    ttd_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen"
    ttd_table.rows[4].cells[0].text = "{{direktur_nama}}"
    ttd_table.rows[4].cells[1].text = "{{ppk_nama}}"
    ttd_table.rows[5].cells[0].text = "Direktur"
    ttd_table.rows[5].cells[1].text = "NIP. {{ppk_nip}}"
    
    for row in ttd_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(TEMPLATES_DIR, "spk_konstruksi.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# SURAT PERJANJIAN BARANG (untuk nilai > 200 juta)
# =============================================================================
def create_surat_perjanjian_barang():
    """Create Surat Perjanjian template for Pengadaan Barang (nilai besar)"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    # Header
    title = doc.add_paragraph()
    title.add_run("SURAT PERJANJIAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("PENGADAAN BARANG").bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_kontrak}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    p1 = doc.add_paragraph()
    p1.add_run("Surat Perjanjian ini berikut semua lampirannya (selanjutnya disebut \"Kontrak\") dibuat dan ditandatangani di {{satker_kota}} pada hari {{hari_kontrak}} tanggal {{tanggal_kontrak_terbilang}} bulan {{bulan_kontrak}} tahun {{tahun_kontrak_terbilang}} ({{tanggal_kontrak_fmt}}), antara:")
    
    doc.add_paragraph()
    
    # PPK
    doc.add_paragraph("I. {{ppk_nama}}, {{ppk_jabatan}} {{satker_nama}}, yang bertindak untuk dan atas nama {{satker_nama}}, berkedudukan di {{satker_alamat}}, berdasarkan Surat Keputusan {{sk_ppk}}, selanjutnya disebut \"PIHAK KESATU\".")
    
    doc.add_paragraph()
    
    # Penyedia
    doc.add_paragraph("II. {{direktur_nama}}, Direktur {{penyedia_nama}}, yang bertindak untuk dan atas nama {{penyedia_nama}}, berkedudukan di {{penyedia_alamat}}, NPWP {{penyedia_npwp}}, selanjutnya disebut \"PIHAK KEDUA\".")
    
    doc.add_paragraph()
    
    # Mengingat
    doc.add_paragraph("MENGINGAT BAHWA:")
    
    mengingat = [
        "PIHAK KESATU telah meminta PIHAK KEDUA untuk menyediakan barang sebagaimana diterangkan dalam Kontrak ini;",
        "PIHAK KEDUA sebagaimana dinyatakan kepada PIHAK KESATU, memiliki keahlian profesional, personil, dan sumber daya teknis, serta telah menyetujui untuk menyediakan barang sesuai dengan persyaratan dan ketentuan dalam Kontrak ini;",
        "PIHAK KESATU dan PIHAK KEDUA menyatakan memiliki kewenangan untuk menandatangani Kontrak ini, dan mengikat pihak yang diwakili;",
        "PIHAK KESATU dan PIHAK KEDUA mengakui dan menyatakan bahwa sehubungan dengan penandatanganan Kontrak ini masing-masing pihak telah membaca dan memahami isi Kontrak, dan sepakat untuk terikat dengan ketentuan dalam Kontrak.",
    ]
    
    for i, m in enumerate(mengingat, 1):
        p = doc.add_paragraph()
        p.add_run(f"({i}) ")
        p.add_run(m)
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    doc.add_paragraph("MAKA OLEH KARENA ITU, PIHAK KESATU dan PIHAK KEDUA (selanjutnya disebut \"PARA PIHAK\") dengan ini bersepakat dan menyetujui hal-hal sebagai berikut:")
    
    doc.add_paragraph()
    
    # Pasal 1 - Istilah dan Ungkapan
    add_pasal(doc, 1, "ISTILAH DAN UNGKAPAN")
    doc.add_paragraph("Istilah dan ungkapan dalam Kontrak ini memiliki arti dan makna yang sama seperti yang tercantum dalam Syarat-Syarat Umum Kontrak (SSUK).")
    
    doc.add_paragraph()
    
    # Pasal 2 - Ruang Lingkup
    add_pasal(doc, 2, "RUANG LINGKUP PEKERJAAN")
    add_ayat(doc, 1, "PIHAK KEDUA menyediakan barang kepada PIHAK KESATU sesuai dengan ketentuan Kontrak.")
    add_ayat(doc, 2, "Uraian pekerjaan yang harus dilaksanakan:")
    
    lingkup_table = doc.add_table(rows=4, cols=3)
    lingkup_table.style = 'Table Grid'
    lingkup_data = [
        ("a.", "Nama Paket", "{{nama_paket}}"),
        ("b.", "Kode RUP", "{{kode_rup}}"),
        ("c.", "Lokasi Pengiriman", "{{lokasi_pekerjaan}}"),
        ("d.", "Rincian Barang", "Sebagaimana tercantum dalam Lampiran"),
    ]
    for i, (no, label, val) in enumerate(lingkup_data):
        lingkup_table.rows[i].cells[0].text = no
        lingkup_table.rows[i].cells[1].text = label
        lingkup_table.rows[i].cells[2].text = val
    
    doc.add_paragraph()
    
    # Pasal 3 - Nilai Kontrak
    add_pasal(doc, 3, "HARGA KONTRAK")
    add_ayat(doc, 1, "PIHAK KESATU membayar kepada PIHAK KEDUA atas pelaksanaan pekerjaan berdasarkan Kontrak sebesar {{nilai_kontrak_fmt}} ({{nilai_kontrak_terbilang}}).")
    add_ayat(doc, 2, "Harga Kontrak telah memperhitungkan keuntungan dan biaya overhead serta pajak.")
    add_ayat(doc, 3, "Rincian Harga Kontrak tercantum dalam Lampiran.")
    
    doc.add_paragraph()
    
    # Pasal 4 - Jangka Waktu
    add_pasal(doc, 4, "JANGKA WAKTU PELAKSANAAN")
    add_ayat(doc, 1, "PIHAK KEDUA wajib menyelesaikan pengiriman barang dalam jangka waktu {{jangka_waktu}} ({{jangka_waktu_terbilang}}) hari kalender.")
    add_ayat(doc, 2, "Kontrak ini berlaku efektif sejak {{tanggal_mulai_fmt}} sampai dengan {{tanggal_selesai_fmt}}.")
    add_ayat(doc, 3, "Apabila PIHAK KEDUA tidak dapat menyelesaikan pekerjaan sesuai jadwal, maka PIHAK KEDUA dikenakan sanksi sesuai ketentuan SSUK.")
    
    doc.add_paragraph()
    
    # Pasal 5 - Dokumen Kontrak
    add_pasal(doc, 5, "DOKUMEN KONTRAK")
    add_ayat(doc, 1, "Dokumen-dokumen berikut merupakan satu kesatuan dan bagian yang tidak terpisahkan dari Kontrak ini:")
    add_huruf(doc, "a", "Surat Perjanjian;")
    add_huruf(doc, "b", "Surat Penawaran beserta lampirannya;")
    add_huruf(doc, "c", "Syarat-Syarat Umum Kontrak (SSUK);")
    add_huruf(doc, "d", "Syarat-Syarat Khusus Kontrak (SSKK);")
    add_huruf(doc, "e", "Spesifikasi Teknis;")
    add_huruf(doc, "f", "Gambar/Brosur (jika ada);")
    add_huruf(doc, "g", "Daftar Kuantitas dan Harga;")
    add_huruf(doc, "h", "Dokumen lain yang menjadi bagian dari Kontrak.")
    add_ayat(doc, 2, "Jika terdapat pertentangan, urutan pemberlakuan mengikuti urutan dokumen di atas.")
    
    doc.add_paragraph()
    
    # Pasal 6 - Jaminan
    add_pasal(doc, 6, "JAMINAN")
    add_ayat(doc, 1, "PIHAK KEDUA menyerahkan Jaminan Pelaksanaan sebesar {{nilai_jaminan_pelaksanaan_fmt}} ({{persen_jaminan_pelaksanaan}}% dari nilai kontrak).")
    add_ayat(doc, 2, "Jaminan Pelaksanaan dikembalikan setelah pekerjaan selesai dan BAST ditandatangani.")
    add_ayat(doc, 3, "Masa berlaku jaminan sesuai dengan masa kontrak ditambah {{masa_klaim_jaminan}} hari.")
    
    doc.add_paragraph()
    
    # Pasal 7 - Korespondensi
    add_pasal(doc, 7, "KORESPONDENSI")
    doc.add_paragraph("Semua pemberitahuan, permohonan, persetujuan dan korespondensi disampaikan secara tertulis kepada:")
    
    korespondensi = doc.add_table(rows=4, cols=2)
    korespondensi.style = 'Table Grid'
    korespondensi.rows[0].cells[0].text = "PIHAK KESATU"
    korespondensi.rows[0].cells[1].text = "PIHAK KEDUA"
    korespondensi.rows[1].cells[0].text = "{{ppk_nama}}"
    korespondensi.rows[1].cells[1].text = "{{direktur_nama}}"
    korespondensi.rows[2].cells[0].text = "{{satker_alamat}}"
    korespondensi.rows[2].cells[1].text = "{{penyedia_alamat}}"
    korespondensi.rows[3].cells[0].text = "Telp: {{satker_telepon}}"
    korespondensi.rows[3].cells[1].text = "Telp: {{penyedia_telepon}}"
    
    doc.add_paragraph()
    
    # Pasal 8 - Penutup
    add_pasal(doc, 8, "PENUTUP")
    doc.add_paragraph("Kontrak ini mulai berlaku efektif terhitung sejak tanggal yang ditetapkan. Kontrak ditandatangani oleh PARA PIHAK.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    ttd_table = doc.add_table(rows=6, cols=2)
    ttd_table.rows[0].cells[0].text = "PIHAK KEDUA"
    ttd_table.rows[0].cells[1].text = "PIHAK KESATU"
    ttd_table.rows[1].cells[0].text = "{{penyedia_nama}}"
    ttd_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen"
    ttd_table.rows[4].cells[0].text = "{{direktur_nama}}"
    ttd_table.rows[4].cells[1].text = "{{ppk_nama}}"
    ttd_table.rows[5].cells[0].text = "Direktur"
    ttd_table.rows[5].cells[1].text = "NIP. {{ppk_nip}}"
    
    for row in ttd_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(TEMPLATES_DIR, "surat_perjanjian_barang.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# SURAT PERJANJIAN JASA LAINNYA
# =============================================================================
def create_surat_perjanjian_jasa():
    """Create Surat Perjanjian template for Jasa Lainnya"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    # Header
    title = doc.add_paragraph()
    title.add_run("SURAT PERJANJIAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("PENGADAAN JASA LAINNYA").bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_kontrak}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    doc.add_paragraph("Surat Perjanjian ini berikut semua lampirannya (selanjutnya disebut \"Kontrak\") dibuat dan ditandatangani di {{satker_kota}} pada hari {{hari_kontrak}} tanggal {{tanggal_kontrak_fmt}}, antara:")
    
    doc.add_paragraph()
    
    doc.add_paragraph("I. {{ppk_nama}}, {{ppk_jabatan}} {{satker_nama}}, berkedudukan di {{satker_alamat}}, selanjutnya disebut \"PIHAK KESATU\".")
    doc.add_paragraph()
    doc.add_paragraph("II. {{direktur_nama}}, Direktur {{penyedia_nama}}, berkedudukan di {{penyedia_alamat}}, NPWP {{penyedia_npwp}}, selanjutnya disebut \"PIHAK KEDUA\".")
    
    doc.add_paragraph()
    doc.add_paragraph("PARA PIHAK sepakat dan menyetujui hal-hal sebagai berikut:")
    doc.add_paragraph()
    
    # Pasal 1
    add_pasal(doc, 1, "RUANG LINGKUP PEKERJAAN")
    add_ayat(doc, 1, "PIHAK KEDUA berkewajiban melaksanakan pekerjaan jasa:")
    
    lingkup_table = doc.add_table(rows=3, cols=3)
    lingkup_table.style = 'Table Grid'
    lingkup_data = [
        ("a.", "Nama Paket", "{{nama_paket}}"),
        ("b.", "Ruang Lingkup", "Sesuai KAK/Spesifikasi Teknis"),
        ("c.", "Output/Deliverable", "Sesuai KAK terlampir"),
    ]
    for i, (no, label, val) in enumerate(lingkup_data):
        lingkup_table.rows[i].cells[0].text = no
        lingkup_table.rows[i].cells[1].text = label
        lingkup_table.rows[i].cells[2].text = val
    
    add_ayat(doc, 2, "Rincian ruang lingkup, metodologi, dan jadwal pelaksanaan tercantum dalam KAK yang merupakan bagian tidak terpisahkan dari Kontrak.")
    
    doc.add_paragraph()
    
    # Pasal 2
    add_pasal(doc, 2, "HARGA KONTRAK DAN CARA PEMBAYARAN")
    add_ayat(doc, 1, "Harga Kontrak adalah {{nilai_kontrak_fmt}} ({{nilai_kontrak_terbilang}}) termasuk pajak.")
    add_ayat(doc, 2, "Pembayaran dilakukan secara:")
    add_huruf(doc, "a", "Sekaligus 100% setelah pekerjaan selesai; ATAU")
    add_huruf(doc, "b", "Bertahap sesuai progres/milestone yang disepakati")
    add_ayat(doc, 3, "Pembayaran dipotong pajak sesuai ketentuan yang berlaku.")
    
    doc.add_paragraph()
    
    # Pasal 3
    add_pasal(doc, 3, "JANGKA WAKTU")
    add_ayat(doc, 1, "Jangka waktu pelaksanaan adalah {{jangka_waktu}} ({{jangka_waktu_terbilang}}) hari kalender.")
    add_ayat(doc, 2, "Tanggal mulai: {{tanggal_mulai_fmt}}")
    add_ayat(doc, 3, "Tanggal selesai: {{tanggal_selesai_fmt}}")
    
    doc.add_paragraph()
    
    # Pasal 4
    add_pasal(doc, 4, "HAK KEKAYAAN INTELEKTUAL")
    add_ayat(doc, 1, "Seluruh hasil pekerjaan menjadi milik PIHAK KESATU.")
    add_ayat(doc, 2, "PIHAK KEDUA tidak berhak menggunakan hasil pekerjaan untuk kepentingan pihak lain tanpa persetujuan tertulis.")
    add_ayat(doc, 3, "PIHAK KEDUA menjamin hasil pekerjaan tidak melanggar hak kekayaan intelektual pihak ketiga.")
    
    doc.add_paragraph()
    
    # Pasal 5
    add_pasal(doc, 5, "KERAHASIAAN")
    add_ayat(doc, 1, "PIHAK KEDUA wajib menjaga kerahasiaan seluruh informasi yang diperoleh dalam pelaksanaan pekerjaan.")
    add_ayat(doc, 2, "Kewajiban kerahasiaan tetap berlaku meskipun Kontrak telah berakhir.")
    
    doc.add_paragraph()
    
    # Pasal 6
    add_pasal(doc, 6, "DOKUMEN KONTRAK")
    add_ayat(doc, 1, "Dokumen Kontrak terdiri dari:")
    add_huruf(doc, "a", "Surat Perjanjian;")
    add_huruf(doc, "b", "SSUK dan SSKK;")
    add_huruf(doc, "c", "Kerangka Acuan Kerja (KAK);")
    add_huruf(doc, "d", "Surat Penawaran;")
    add_huruf(doc, "e", "Daftar Harga;")
    add_huruf(doc, "f", "Dokumen lain yang relevan.")
    
    doc.add_paragraph()
    
    # Pasal 7
    add_pasal(doc, 7, "SANKSI DAN DENDA")
    add_ayat(doc, 1, "Keterlambatan dikenakan denda 1/1000 per hari dari nilai kontrak, maksimal 5%.")
    add_ayat(doc, 2, "Hasil pekerjaan yang tidak sesuai spesifikasi wajib diperbaiki tanpa biaya tambahan.")
    
    doc.add_paragraph()
    
    # Pasal 8
    add_pasal(doc, 8, "PENYELESAIAN PERSELISIHAN")
    add_ayat(doc, 1, "Perselisihan diselesaikan secara musyawarah.")
    add_ayat(doc, 2, "Jika tidak tercapai, melalui Pengadilan Negeri {{satker_kota}}.")
    
    doc.add_paragraph()
    
    # Pasal 9
    add_pasal(doc, 9, "PENUTUP")
    doc.add_paragraph("Kontrak ini ditandatangani oleh PARA PIHAK pada tanggal tersebut di atas dalam rangkap 2 (dua) yang masing-masing mempunyai kekuatan hukum yang sama.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    ttd_table = doc.add_table(rows=6, cols=2)
    ttd_table.rows[0].cells[0].text = "PIHAK KEDUA"
    ttd_table.rows[0].cells[1].text = "PIHAK KESATU"
    ttd_table.rows[1].cells[0].text = "{{penyedia_nama}}"
    ttd_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen"
    ttd_table.rows[4].cells[0].text = "{{direktur_nama}}"
    ttd_table.rows[4].cells[1].text = "{{ppk_nama}}"
    ttd_table.rows[5].cells[0].text = "Direktur"
    ttd_table.rows[5].cells[1].text = "NIP. {{ppk_nip}}"
    
    for row in ttd_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(TEMPLATES_DIR, "surat_perjanjian_jasa_lainnya.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# SURAT PERJANJIAN KONSTRUKSI
# =============================================================================
def create_surat_perjanjian_konstruksi():
    """Create Surat Perjanjian template for Konstruksi"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    # Header
    title = doc.add_paragraph()
    title.add_run("SURAT PERJANJIAN").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("PEKERJAAN KONSTRUKSI").bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nomor = doc.add_paragraph()
    nomor.add_run("Nomor: {{nomor_kontrak}}")
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Pembukaan
    doc.add_paragraph("Surat Perjanjian untuk melaksanakan Pekerjaan Konstruksi ini berikut semua lampirannya (selanjutnya disebut \"Kontrak\") dibuat dan ditandatangani di {{satker_kota}} pada hari {{hari_kontrak}} tanggal {{tanggal_kontrak_fmt}}, antara:")
    
    doc.add_paragraph()
    doc.add_paragraph("I. {{ppk_nama}}, {{ppk_jabatan}} {{satker_nama}}, berkedudukan di {{satker_alamat}}, selanjutnya disebut \"PENGGUNA JASA\".")
    doc.add_paragraph()
    doc.add_paragraph("II. {{direktur_nama}}, Direktur {{penyedia_nama}}, berkedudukan di {{penyedia_alamat}}, NPWP {{penyedia_npwp}}, Kualifikasi {{kualifikasi_penyedia}}, selanjutnya disebut \"PENYEDIA JASA\".")
    
    doc.add_paragraph()
    doc.add_paragraph("PARA PIHAK sepakat dan menyetujui hal-hal sebagai berikut:")
    doc.add_paragraph()
    
    # Pasal 1
    add_pasal(doc, 1, "LINGKUP PEKERJAAN")
    add_ayat(doc, 1, "PENYEDIA JASA berkewajiban melaksanakan pekerjaan konstruksi:")
    
    lingkup_table = doc.add_table(rows=5, cols=3)
    lingkup_table.style = 'Table Grid'
    lingkup_data = [
        ("a.", "Nama Paket", "{{nama_paket}}"),
        ("b.", "Lokasi Pekerjaan", "{{lokasi_pekerjaan}}"),
        ("c.", "Jenis Pekerjaan", "{{jenis_pekerjaan_konstruksi}}"),
        ("d.", "Klasifikasi", "{{klasifikasi_konstruksi}}"),
        ("e.", "Volume", "Sesuai BOQ terlampir"),
    ]
    for i, (no, label, val) in enumerate(lingkup_data):
        lingkup_table.rows[i].cells[0].text = no
        lingkup_table.rows[i].cells[1].text = label
        lingkup_table.rows[i].cells[2].text = val
    
    add_ayat(doc, 2, "Pekerjaan meliputi penyediaan tenaga kerja, bahan, peralatan, dan semua yang diperlukan.")
    add_ayat(doc, 3, "Gambar teknis, spesifikasi, dan BOQ tercantum dalam Lampiran.")
    
    doc.add_paragraph()
    
    # Pasal 2
    add_pasal(doc, 2, "HARGA KONTRAK")
    add_ayat(doc, 1, "Harga Kontrak adalah {{nilai_kontrak_fmt}} ({{nilai_kontrak_terbilang}}) termasuk PPN.")
    add_ayat(doc, 2, "Jenis Kontrak: {{jenis_kontrak_konstruksi}}")
    add_huruf(doc, "a", "Lump Sum: Harga tetap dan pasti;")
    add_huruf(doc, "b", "Unit Price: Berdasarkan harga satuan dan volume terpasang;")
    add_huruf(doc, "c", "Gabungan Lump Sum dan Unit Price.")
    add_ayat(doc, 3, "Harga sudah memperhitungkan keuntungan, overhead, dan risiko.")
    
    doc.add_paragraph()
    
    # Pasal 3
    add_pasal(doc, 3, "JANGKA WAKTU PELAKSANAAN")
    add_ayat(doc, 1, "Jangka waktu pelaksanaan: {{jangka_waktu}} ({{jangka_waktu_terbilang}}) hari kalender.")
    add_ayat(doc, 2, "Tanggal mulai kerja: {{tanggal_mulai_fmt}}")
    add_ayat(doc, 3, "Tanggal Serah Terima I (PHO): {{tanggal_selesai_fmt}}")
    add_ayat(doc, 4, "Masa Pemeliharaan: {{masa_pemeliharaan}} bulan")
    add_ayat(doc, 5, "PENYEDIA JASA wajib membuat time schedule dan kurva S.")
    
    doc.add_paragraph()
    
    # Pasal 4
    add_pasal(doc, 4, "CARA PEMBAYARAN")
    add_ayat(doc, 1, "Pembayaran dilakukan secara termin berdasarkan progres fisik:")
    
    termin_table = doc.add_table(rows=5, cols=3)
    termin_table.style = 'Table Grid'
    termin_table.rows[0].cells[0].text = "Termin"
    termin_table.rows[0].cells[1].text = "Progres Fisik"
    termin_table.rows[0].cells[2].text = "Pembayaran"
    termin_table.rows[1].cells[0].text = "I"
    termin_table.rows[1].cells[1].text = "{{progres_termin1}}%"
    termin_table.rows[1].cells[2].text = "{{pembayaran_termin1}}%"
    termin_table.rows[2].cells[0].text = "II"
    termin_table.rows[2].cells[1].text = "{{progres_termin2}}%"
    termin_table.rows[2].cells[2].text = "{{pembayaran_termin2}}%"
    termin_table.rows[3].cells[0].text = "III"
    termin_table.rows[3].cells[1].text = "{{progres_termin3}}%"
    termin_table.rows[3].cells[2].text = "{{pembayaran_termin3}}%"
    termin_table.rows[4].cells[0].text = "IV (Final)"
    termin_table.rows[4].cells[1].text = "100%"
    termin_table.rows[4].cells[2].text = "Sisa"
    
    add_ayat(doc, 2, "Setiap pembayaran dipotong retensi 5% sebagai jaminan pemeliharaan.")
    add_ayat(doc, 3, "Retensi dikembalikan setelah FHO dan tidak ada cacat.")
    
    doc.add_paragraph()
    
    # Pasal 5
    add_pasal(doc, 5, "JAMINAN-JAMINAN")
    add_ayat(doc, 1, "PENYEDIA JASA wajib menyerahkan:")
    add_huruf(doc, "a", "Jaminan Pelaksanaan sebesar {{persen_jaminan_pelaksanaan}}% dari nilai kontrak;")
    add_huruf(doc, "b", "Jaminan Uang Muka (jika ada uang muka) sebesar 100% dari uang muka;")
    add_huruf(doc, "c", "Jaminan Pemeliharaan sebesar 5% dari nilai kontrak atau dapat diganti dengan retensi.")
    add_ayat(doc, 2, "Jaminan dari Bank Umum atau Perusahaan Asuransi yang memiliki izin.")
    
    doc.add_paragraph()
    
    # Pasal 6
    add_pasal(doc, 6, "MASA PEMELIHARAAN")
    add_ayat(doc, 1, "Masa Pemeliharaan adalah {{masa_pemeliharaan}} bulan sejak PHO.")
    add_ayat(doc, 2, "Selama masa pemeliharaan, PENYEDIA JASA wajib:")
    add_huruf(doc, "a", "Memperbaiki kerusakan yang bukan kesalahan PENGGUNA JASA;")
    add_huruf(doc, "b", "Melakukan perawatan rutin sesuai spesifikasi;")
    add_huruf(doc, "c", "Menanggung biaya perbaikan.")
    add_ayat(doc, 3, "Setelah masa pemeliharaan berakhir dan tidak ada cacat, dilakukan FHO.")
    
    doc.add_paragraph()
    
    # Pasal 7
    add_pasal(doc, 7, "TENAGA AHLI DAN PERALATAN")
    add_ayat(doc, 1, "PENYEDIA JASA wajib menyediakan tenaga ahli sesuai yang ditawarkan:")
    add_huruf(doc, "a", "{{tenaga_ahli_1}}")
    add_huruf(doc, "b", "{{tenaga_ahli_2}}")
    add_huruf(doc, "c", "Tenaga teknis lainnya sesuai kebutuhan")
    add_ayat(doc, 2, "Penggantian tenaga ahli harus mendapat persetujuan PENGGUNA JASA.")
    add_ayat(doc, 3, "PENYEDIA JASA wajib menyediakan peralatan yang memadai.")
    
    doc.add_paragraph()
    
    # Pasal 8
    add_pasal(doc, 8, "KESELAMATAN DAN KESEHATAN KERJA (K3)")
    add_ayat(doc, 1, "PENYEDIA JASA wajib menerapkan SMK3 Konstruksi.")
    add_ayat(doc, 2, "PENYEDIA JASA bertanggung jawab penuh atas keselamatan di lokasi proyek.")
    add_ayat(doc, 3, "Kecelakaan kerja menjadi tanggung jawab PENYEDIA JASA.")
    add_ayat(doc, 4, "PENYEDIA JASA wajib mengasuransikan tenaga kerja dan pihak ketiga.")
    
    doc.add_paragraph()
    
    # Pasal 9
    add_pasal(doc, 9, "KEGAGALAN BANGUNAN")
    add_ayat(doc, 1, "PENYEDIA JASA bertanggung jawab atas kegagalan bangunan selama umur konstruksi.")
    add_ayat(doc, 2, "Pertanggungjawaban meliputi:")
    add_huruf(doc, "a", "Perbaikan/penggantian;")
    add_huruf(doc, "b", "Ganti rugi;")
    add_huruf(doc, "c", "Tuntutan pidana sesuai peraturan perundang-undangan.")
    
    doc.add_paragraph()
    
    # Pasal 10
    add_pasal(doc, 10, "SANKSI DAN DENDA")
    add_ayat(doc, 1, "Keterlambatan dikenakan denda 1/1000 per hari, maksimal 5%.")
    add_ayat(doc, 2, "Keterlambatan >50 hari dapat dilakukan pemutusan kontrak.")
    add_ayat(doc, 3, "Pekerjaan tidak sesuai mutu wajib diperbaiki tanpa biaya tambahan.")
    
    doc.add_paragraph()
    
    # Pasal 11
    add_pasal(doc, 11, "PENYELESAIAN PERSELISIHAN")
    add_ayat(doc, 1, "Musyawarah mufakat;")
    add_ayat(doc, 2, "Mediasi/konsiliasi;")
    add_ayat(doc, 3, "Arbitrase melalui BANI; atau")
    add_ayat(doc, 4, "Pengadilan Negeri {{satker_kota}}.")
    
    doc.add_paragraph()
    
    # Pasal 12
    add_pasal(doc, 12, "DOKUMEN KONTRAK")
    add_ayat(doc, 1, "Dokumen Kontrak terdiri dari:")
    add_huruf(doc, "a", "Surat Perjanjian;")
    add_huruf(doc, "b", "SSUK dan SSKK;")
    add_huruf(doc, "c", "Spesifikasi Teknis;")
    add_huruf(doc, "d", "Gambar-gambar;")
    add_huruf(doc, "e", "Bill of Quantity (BOQ);")
    add_huruf(doc, "f", "Surat Penawaran;")
    add_huruf(doc, "g", "Daftar Harga Satuan;")
    add_huruf(doc, "h", "Dokumen lain yang relevan.")
    
    doc.add_paragraph()
    
    # Pasal 13
    add_pasal(doc, 13, "PENUTUP")
    doc.add_paragraph("Kontrak ini ditandatangani PARA PIHAK pada tanggal tersebut di atas, dibuat dalam rangkap 2 (dua) bermaterai cukup yang masing-masing mempunyai kekuatan hukum yang sama.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TTD
    ttd_table = doc.add_table(rows=6, cols=2)
    ttd_table.rows[0].cells[0].text = "PENYEDIA JASA"
    ttd_table.rows[0].cells[1].text = "PENGGUNA JASA"
    ttd_table.rows[1].cells[0].text = "{{penyedia_nama}}"
    ttd_table.rows[1].cells[1].text = "Pejabat Pembuat Komitmen"
    ttd_table.rows[4].cells[0].text = "{{direktur_nama}}"
    ttd_table.rows[4].cells[1].text = "{{ppk_nama}}"
    ttd_table.rows[5].cells[0].text = "Direktur"
    ttd_table.rows[5].cells[1].text = "NIP. {{ppk_nip}}"
    
    for row in ttd_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(TEMPLATES_DIR, "surat_perjanjian_konstruksi.docx")
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# MAIN
# =============================================================================
if __name__ == "__main__":
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    
    print("=" * 60)
    print("MEMBUAT TEMPLATE RANCANGAN KONTRAK")
    print("=" * 60)
    print()
    
    print("📋 SPK (Surat Perintah Kerja) - untuk nilai s.d. 200 juta:")
    create_spk_barang()
    create_spk_jasa_lainnya()
    create_spk_konstruksi()
    
    print()
    print("📋 Surat Perjanjian - untuk nilai > 200 juta:")
    create_surat_perjanjian_barang()
    create_surat_perjanjian_jasa()
    create_surat_perjanjian_konstruksi()
    
    print()
    print("=" * 60)
    print("SELESAI!")
    print("=" * 60)
    print()
    print("KARAKTERISTIK PERBEDAAN TEMPLATE:")
    print()
    print("BARANG:")
    print("  - Fokus: Spesifikasi, pengiriman, garansi")
    print("  - Pembayaran: Sekaligus setelah barang diterima")
    print("  - Garansi: Jaminan kualitas barang")
    print()
    print("JASA LAINNYA:")
    print("  - Fokus: Ruang lingkup, output/deliverable, timeline")
    print("  - Pembayaran: Sekaligus atau termin sesuai progres")
    print("  - Khusus: Hak kekayaan intelektual, kerahasiaan")
    print()
    print("KONSTRUKSI:")
    print("  - Fokus: Gambar teknis, BOQ, tenaga ahli, K3")
    print("  - Pembayaran: Termin berdasarkan progres fisik")
    print("  - Jaminan: Pelaksanaan, Uang Muka, Pemeliharaan")
    print("  - Khusus: Masa pemeliharaan, PHO/FHO, Retensi 5%")
    print("  - Tanggung jawab: Kegagalan bangunan")
