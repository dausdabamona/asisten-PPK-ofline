"""
Test Script: Generate Lembar Permintaan Dokumen
Test creating lembar_permintaan with sample data
"""

from docx import Document
from datetime import datetime
import os

def test_create_lembar_permintaan():
    """Create test lembar permintaan document."""
    
    template_path = 'templates/word/lembar_permintaan.docx'
    
    if not os.path.exists(template_path):
        print(f"❌ Template tidak ditemukan: {template_path}")
        return False
    
    try:
        # Load template
        doc = Document(template_path)
        print(f"✓ Template loaded: {template_path}")
        
        # Sample data dari attachment
        sample_data = {
            'hari_tanggal': 'Rabu, 01 September 2021',
            'unit_kerja': 'Pusbinka',
            'sumber_dana': 'DIPA POLTEK KP TA. 2021',
            
            # Item 1
            'item_1_no': '1',
            'item_1_nama': 'Siempang Perwira (Kuning Biru)',
            'item_1_spek': 'Sesuai spesifikasi',
            'item_1_volume': '2',
            'item_1_satuan': 'Buah',
            'item_1_harga': '100.000',
            'item_1_total': '200.000',
            'item_1_ket': '',
            
            # Item 2
            'item_2_no': '2',
            'item_2_nama': 'Siempang Bintara (Merah Biru)',
            'item_2_spek': 'Sesuai spesifikasi',
            'item_2_volume': '2',
            'item_2_satuan': 'Buah',
            'item_2_harga': '100.000',
            'item_2_total': '200.000',
            'item_2_ket': '',
            
            # Item 3
            'item_3_no': '3',
            'item_3_nama': 'Ban Piket',
            'item_3_spek': 'Ban standar',
            'item_3_volume': '21',
            'item_3_satuan': 'Buah',
            'item_3_harga': '130.000',
            'item_3_total': '2.730.000',
            'item_3_ket': '',
            
            # Item 4
            'item_4_no': '4',
            'item_4_nama': 'Tali Kat/Senat',
            'item_4_spek': 'Tali standar',
            'item_4_volume': '60',
            'item_4_satuan': 'Buah',
            'item_4_harga': '150.000',
            'item_4_total': '9.000.000',
            'item_4_ket': '',
            
            # Item 5 (empty)
            'item_5_no': '',
            'item_5_nama': '',
            'item_5_spek': '',
            'item_5_volume': '',
            'item_5_satuan': '',
            'item_5_harga': '',
            'item_5_total': '',
            'item_5_ket': '',
            
            # Summary
            'subtotal': 'Rp 12.130.000',
            'ppn': 'Rp 1.213.000',
            'total': 'Rp 13.343.000',
            
            # Signature
            'nama_pengajuan': 'Fataha Ilyas Hasan, A.Md., S.Pi',
            'nama_verifikator': 'Abdullah Sidiq, A.Md.',
            'nama_ppk': 'Muhammad Ali Ulat, S.Pi., M.Si',
            'nama_atasan': 'Dr. Handayani, S.Pi, M',
            'nama_kpa': 'Ir. Suyatno, M.T.'
        }
        
        # Replace placeholders
        placeholder_count = 0
        for placeholder, value in sample_data.items():
            for paragraph in doc.paragraphs:
                if f'{{{{{placeholder}}}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{{{placeholder}}}}}', value)
                    placeholder_count += 1
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if f'{{{{{placeholder}}}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{{{placeholder}}}}}', value)
                            placeholder_count += 1
        
        # Save test document
        output_path = 'output/test_lembar_permintaan.docx'
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        print(f"✓ Test document created: {output_path}")
        print(f"✓ Placeholders replaced: {placeholder_count}")
        print()
        print("📋 RINGKASAN DOKUMEN:")
        print("─" * 60)
        print(f"  Unit Kerja: Pusbinka")
        print(f"  Tanggal: Rabu, 01 September 2021")
        print(f"  DIPA: DIPA POLTEK KP TA. 2021")
        print()
        print("  Item yang diajukan:")
        print("    1. Siempang Perwira (2 Buah × Rp 100.000) = Rp 200.000")
        print("    2. Siempang Bintara (2 Buah × Rp 100.000) = Rp 200.000")
        print("    3. Ban Piket (21 Buah × Rp 130.000) = Rp 2.730.000")
        print("    4. Tali Kat/Senat (60 Buah × Rp 150.000) = Rp 9.000.000")
        print()
        print("  SUB TOTAL: Rp 12.130.000")
        print("  PPN 10%: Rp 1.213.000")
        print("  ─" * 60)
        print("  TOTAL: Rp 13.343.000")
        print()
        print("✅ Test BERHASIL - Dokumen siap dilihat di output/test_lembar_permintaan.docx")
        return True
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = test_create_lembar_permintaan()
    exit(0 if success else 1)
