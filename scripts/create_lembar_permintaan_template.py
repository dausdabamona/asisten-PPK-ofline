"""
Script untuk membuat template Lembar Permintaan sesuai format yang diminta.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


def set_single_spacing(paragraph):
    """Set paragraph to single spacing with no space before/after."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), edge_data.get('val', 'single'))
            tag.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            tag.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def create_template():
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== HEADER ==========
    # Title 1
    title1 = doc.add_paragraph()
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(title1)
    run1 = title1.add_run("FORMULIR PERMINTAAN REALISASI KEGIATAN")
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.color.rgb = RGBColor(139, 0, 0)  # Dark red

    # Title 2 - Institution name
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(title2)
    run2 = title2.add_run("{{satker_nama}}")
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(139, 0, 0)
    run2.underline = True

    # Small spacing
    spacer = doc.add_paragraph()
    set_single_spacing(spacer)

    # ========== INFO FIELDS ==========
    # Create table for info fields (2 columns)
    info_table = doc.add_table(rows=3, cols=4)
    info_table.autofit = True

    # Row 1: Hari/Tanggal and Nama Kegiatan
    info_table.cell(0, 0).text = "Hari/Tanggal"
    info_table.cell(0, 1).text = ": {{tanggal_dokumen:tanggal_panjang}}"
    info_table.cell(0, 2).text = "Nama Kegiatan"
    info_table.cell(0, 3).text = ": {{nama_kegiatan}}"

    # Row 2: Unit Kerja
    info_table.cell(1, 0).text = "Unit Kerja"
    info_table.cell(1, 1).text = ": {{unit_kerja}}"
    info_table.cell(1, 2).text = ""
    info_table.cell(1, 3).text = ""

    # Row 3: Sumber Dana
    info_table.cell(2, 0).text = "Sumber Dana"
    info_table.cell(2, 1).text = ": {{sumber_dana}}"
    info_table.cell(2, 2).text = ""
    info_table.cell(2, 3).text = ""

    # Style info table - no borders
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ========== RINCIAN TABLE ==========
    # Main table with 7 columns: No, Nama Barang, Spesifikasi, Volume, Satuan, Harga Satuan, Total
    rincian_table = doc.add_table(rows=10, cols=7)
    rincian_table.style = 'Table Grid'
    rincian_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    widths = [Cm(0.8), Cm(4), Cm(2.5), Cm(1.5), Cm(1.5), Cm(2.5), Cm(2.5)]
    for i, width in enumerate(widths):
        for cell in rincian_table.columns[i].cells:
            cell.width = width

    # Header row
    headers = ["No", "Nama Barang", "Spesifikasi", "Volume", "Satuan", "Harga Satuan", "Total"]
    header_row = rincian_table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        # Center and bold
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows (1-6) with placeholders
    for row_num in range(1, 7):
        row = rincian_table.rows[row_num]
        row.cells[0].text = str(row_num)
        row.cells[1].text = f"{{{{rincian_{row_num}_nama}}}}"
        row.cells[2].text = f"{{{{rincian_{row_num}_spek}}}}"
        row.cells[3].text = f"{{{{rincian_{row_num}_vol}}}}"
        row.cells[4].text = f"{{{{rincian_{row_num}_satuan}}}}"
        row.cells[5].text = f"{{{{rincian_{row_num}_harga:rupiah}}}}"
        row.cells[6].text = f"{{{{rincian_{row_num}_total:rupiah}}}}"

        # Align
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # SUB TOTAL row
    subtotal_row = rincian_table.rows[7]
    subtotal_row.cells[0].merge(subtotal_row.cells[5])
    subtotal_row.cells[0].text = "SUB TOTAL"
    subtotal_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtotal_row.cells[0].paragraphs[0].runs[0].italic = True
    subtotal_row.cells[6].text = "{{subtotal:rupiah}}"
    subtotal_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # PPn row
    ppn_row = rincian_table.rows[8]
    ppn_row.cells[0].merge(ppn_row.cells[5])
    ppn_row.cells[0].text = "PPn {{ppn_persen}}"
    ppn_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ppn_row.cells[0].paragraphs[0].runs[0].italic = True
    ppn_row.cells[6].text = "{{ppn_amount:rupiah}}"
    ppn_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # TOTAL row
    total_row = rincian_table.rows[9]
    total_row.cells[0].merge(total_row.cells[5])
    total_row.cells[0].text = "TOTAL"
    total_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row.cells[0].paragraphs[0].runs[0].italic = True
    total_row.cells[6].text = "{{grand_total:rupiah}}"
    total_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Set font size for all cells
    for row in rincian_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ========== LOCATION AND DATE ==========
    loc_date = doc.add_paragraph()
    loc_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_single_spacing(loc_date)
    loc_date.add_run("{{satker_kota}}, {{tanggal_dokumen:tanggal_panjang}}")

    doc.add_paragraph()

    # ========== SIGNATURES (4 people) ==========
    # Create signature table (2x2 for 4 signatures)
    sig_table = doc.add_table(rows=6, cols=2)
    sig_table.autofit = True

    # Row 0: Titles
    sig_table.cell(0, 0).text = "Yang mengajukan"
    sig_table.cell(0, 1).text = "Verifikator"
    sig_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Row 1-2: Space for signature
    sig_table.cell(1, 0).text = ""
    sig_table.cell(1, 1).text = ""
    sig_table.cell(2, 0).text = ""
    sig_table.cell(2, 1).text = ""

    # Row 3: Names (Yang mengajukan and Verifikator)
    sig_table.cell(3, 0).text = "{{penerima_nama}}"
    p1 = sig_table.cell(3, 0).paragraphs[0]
    p1.runs[0].underline = True

    sig_table.cell(3, 1).text = "{{verifikator_nama}}"
    p2 = sig_table.cell(3, 1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.runs[0].underline = True

    # Row 4: Second row titles
    sig_table.cell(4, 0).text = "Mengetahui dan Menyetujui\nKuasa Pengguna Anggaran"
    sig_table.cell(4, 1).text = "Mengetahui\n{{jabatan_mengetahui}}"
    sig_table.cell(4, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Row 5: Names (KPA and Mengetahui)
    # Add empty rows for signature space
    empty_row = sig_table.add_row()
    empty_row.cells[0].text = ""
    empty_row.cells[1].text = ""

    name_row = sig_table.add_row()
    name_row.cells[0].text = "{{kpa_nama}}"
    name_row.cells[0].paragraphs[0].runs[0].underline = True
    name_row.cells[1].text = "{{mengetahui_nama}}"
    name_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    name_row.cells[1].paragraphs[0].runs[0].underline = True

    # Set font for signature table
    for row in sig_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # ========== FOOTER NOTES ==========
    notes = doc.add_paragraph()
    set_single_spacing(notes)
    notes_run = notes.add_run("Catatan Bagian Keuangan :")
    notes_run.italic = True
    notes_run.font.size = Pt(9)
    notes_run.font.color.rgb = RGBColor(139, 0, 0)

    note_items = [
        "Jumlah dana sesuai MAK",
        "Realisasi  s/d",
        "Sisa Dana",
        "Catatan :  *Coret yang tidk perlu"
    ]

    for item in note_items:
        p = doc.add_paragraph()
        set_single_spacing(p)
        r = p.add_run(item)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(139, 0, 0)

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "word" / "lembar_permintaan.docx"
    doc.save(str(output_path))
    print(f"Template saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_template()
