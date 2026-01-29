#!/usr/bin/env python3
"""
Create Surat Tugas template for Perjalanan Dinas LS
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEMPLATE_DIR = Path(__file__).parent.parent / "templates" / "word"


def create_surat_tugas_template():
    """Create Surat Tugas template."""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # === KOP SURAT ===
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kop.add_run("{{kop_surat}}")
    run.bold = True
    run.font.size = Pt(14)

    # Separator line
    sep = doc.add_paragraph("_" * 70)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === TITLE ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SURAT TUGAS")
    run.bold = True
    run.underline = True
    run.font.size = Pt(14)

    # Nomor
    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor.add_run("Nomor: {{nomor_dokumen}}")

    doc.add_paragraph()

    # === DASAR ===
    dasar = doc.add_paragraph()
    dasar.add_run("Dasar").bold = True
    dasar.add_run("\t: ")
    dasar.add_run("{{dasar_surat_tugas}}")

    doc.add_paragraph()

    # === MENUGASKAN ===
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p1.add_run("MENUGASKAN")
    run.bold = True

    doc.add_paragraph()

    # === KEPADA ===
    kepada = doc.add_paragraph()
    kepada.add_run("Kepada").bold = True

    # Tabel data pegawai
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    data_rows = [
        ("Nama", "{{penerima_nama}}"),
        ("NIP", "{{penerima_nip}}"),
        ("Pangkat/Gol", "{{penerima_pangkat}}"),
        ("Jabatan", "{{penerima_jabatan}}"),
    ]

    for i, (label, value) in enumerate(data_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = f": {value}"
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()

    # === UNTUK ===
    untuk = doc.add_paragraph()
    untuk.add_run("Untuk").bold = True
    untuk.add_run("\t: Melaksanakan Perjalanan Dinas dalam rangka {{tujuan_kegiatan}}")

    # Tujuan
    tujuan = doc.add_paragraph()
    tujuan.add_run("Tujuan").bold = True
    tujuan.add_run("\t: {{kota_tujuan}}")

    # Waktu
    waktu = doc.add_paragraph()
    waktu.add_run("Waktu").bold = True
    waktu.add_run("\t: {{tanggal_berangkat:tanggal}} s.d. {{tanggal_kembali:tanggal}}")

    # Lama
    lama = doc.add_paragraph()
    lama.add_run("Lama").bold = True
    lama.add_run("\t: {{lama_hari}} ({{lama_hari_terbilang}}) hari")

    doc.add_paragraph()

    # === PENUTUP ===
    penutup = doc.add_paragraph()
    penutup.add_run("Demikian Surat Tugas ini dibuat untuk dilaksanakan dengan penuh tanggung jawab.")

    doc.add_paragraph()
    doc.add_paragraph()

    # === TANDA TANGAN ===
    ttd_table = doc.add_table(rows=5, cols=1)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # Tempat dan tanggal
    cell0 = ttd_table.rows[0].cells[0]
    cell0.text = "{{kota}}, {{tanggal_dokumen:tanggal}}"
    cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Jabatan penandatangan
    cell1 = ttd_table.rows[1].cells[0]
    cell1.text = "{{jabatan_penandatangan}}"
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Spacing untuk tanda tangan
    cell2 = ttd_table.rows[2].cells[0]
    cell2.text = "\n\n\n"

    # Nama penandatangan
    cell3 = ttd_table.rows[3].cells[0]
    p = cell3.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{nama_penandatangan}}")
    run.bold = True

    # NIP
    cell4 = ttd_table.rows[4].cells[0]
    cell4.text = "NIP. {{nip_penandatangan}}"
    cell4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = TEMPLATE_DIR / "surat_tugas.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Creating Surat Tugas template...")
    create_surat_tugas_template()
    print("Done!")
