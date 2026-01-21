"""
Module untuk generate dokumen Word (Daftar Hadir)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def set_cell_border(cell, **kwargs):
    """
    Set border untuk cell tabel
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def generate_daftar_hadir_jamuan(data: dict, output_path: str) -> str:
    """
    Generate dokumen Daftar Hadir untuk Jamuan Tamu

    Args:
        data: Dictionary berisi informasi jamuan
            - nama_kegiatan: str
            - tanggal: str
            - waktu: str
            - tempat: str
            - jumlah_peserta: int
            - nama_ppk: str
            - nip_ppk: str
        output_path: Path untuk menyimpan file

    Returns:
        Path file yang dihasilkan
    """
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Header - Kop Surat
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("DAFTAR HADIR")
    run.bold = True
    run.font.size = Pt(14)

    # Sub judul
    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_header.add_run(f"JAMUAN TAMU/RAPAT\n{data.get('nama_kegiatan', '').upper()}")
    run.bold = True
    run.font.size = Pt(12)

    # Informasi kegiatan
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run(f"Hari/Tanggal\t: {data.get('tanggal', '-')}\n")
    info.add_run(f"Waktu\t\t: {data.get('waktu', '-')}\n")
    info.add_run(f"Tempat\t\t: {data.get('tempat', '-')}")

    doc.add_paragraph()

    # Tabel Daftar Hadir
    jumlah_peserta = data.get('jumlah_peserta', 10)
    table = doc.add_table(rows=jumlah_peserta + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header tabel
    headers = ['No', 'Nama', 'Jabatan/Instansi', 'No. HP', 'Tanda Tangan']
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        # Set border
        border_style = {'val': 'single', 'sz': 4, 'color': '000000'}
        set_cell_border(cell, top=border_style, bottom=border_style,
                       left=border_style, right=border_style)

    # Set column widths
    widths = [Cm(1), Cm(4), Cm(4), Cm(3), Cm(3.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            border_style = {'val': 'single', 'sz': 4, 'color': '000000'}
            set_cell_border(cell, top=border_style, bottom=border_style,
                           left=border_style, right=border_style)

    # Isi nomor
    for i in range(1, jumlah_peserta + 1):
        table.rows[i].cells[0].text = str(i)
        table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tanda tangan PPK
    doc.add_paragraph()
    doc.add_paragraph()

    ttd = doc.add_paragraph()
    ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ttd.add_run(f"..................., {data.get('tanggal', '-')}\n")
    ttd.add_run("Pejabat Pembuat Komitmen,\n\n\n\n\n")
    run = ttd.add_run(f"{data.get('nama_ppk', '.........................')}\n")
    run.bold = True
    run.underline = True
    ttd.add_run(f"NIP. {data.get('nip_ppk', '.........................')}")

    # Simpan dokumen
    filename = f"Daftar_Hadir_Jamuan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    full_path = os.path.join(output_path, filename)
    doc.save(full_path)

    return full_path


def generate_daftar_hadir_swakelola(data: dict, output_path: str) -> str:
    """
    Generate dokumen Daftar Hadir untuk kegiatan Swakelola

    Args:
        data: Dictionary berisi informasi swakelola
            - nama_kegiatan: str
            - nama_paket: str
            - tanggal: str
            - waktu: str
            - tempat: str
            - jumlah_peserta: int
            - nama_ppk: str
            - nip_ppk: str
            - tahun_anggaran: str
        output_path: Path untuk menyimpan file

    Returns:
        Path file yang dihasilkan
    """
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("DAFTAR HADIR")
    run.bold = True
    run.font.size = Pt(14)

    # Sub judul
    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_header.add_run(f"PELAKSANAAN KEGIATAN SWAKELOLA\n{data.get('nama_kegiatan', '').upper()}")
    run.bold = True
    run.font.size = Pt(12)

    # Informasi paket
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run(f"Nama Paket\t\t: {data.get('nama_paket', '-')}\n")
    info.add_run(f"Tahun Anggaran\t: {data.get('tahun_anggaran', '-')}\n")
    info.add_run(f"Hari/Tanggal\t\t: {data.get('tanggal', '-')}\n")
    info.add_run(f"Waktu\t\t\t: {data.get('waktu', '-')}\n")
    info.add_run(f"Tempat\t\t\t: {data.get('tempat', '-')}")

    doc.add_paragraph()

    # Tabel Daftar Hadir
    jumlah_peserta = data.get('jumlah_peserta', 15)
    table = doc.add_table(rows=jumlah_peserta + 1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header tabel
    headers = ['No', 'Nama', 'Jabatan', 'Unit Kerja', 'No. HP', 'Tanda Tangan']
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        border_style = {'val': 'single', 'sz': 4, 'color': '000000'}
        set_cell_border(cell, top=border_style, bottom=border_style,
                       left=border_style, right=border_style)

    # Set column widths
    widths = [Cm(0.8), Cm(3.5), Cm(2.5), Cm(3), Cm(2.5), Cm(3)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            border_style = {'val': 'single', 'sz': 4, 'color': '000000'}
            set_cell_border(cell, top=border_style, bottom=border_style,
                           left=border_style, right=border_style)

    # Isi nomor
    for i in range(1, jumlah_peserta + 1):
        table.rows[i].cells[0].text = str(i)
        table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tanda tangan PPK
    doc.add_paragraph()
    doc.add_paragraph()

    ttd = doc.add_paragraph()
    ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ttd.add_run(f"..................., {data.get('tanggal', '-')}\n")
    ttd.add_run("Pejabat Pembuat Komitmen,\n\n\n\n\n")
    run = ttd.add_run(f"{data.get('nama_ppk', '.........................')}\n")
    run.bold = True
    run.underline = True
    ttd.add_run(f"NIP. {data.get('nip_ppk', '.........................')}")

    # Simpan dokumen
    filename = f"Daftar_Hadir_Swakelola_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    full_path = os.path.join(output_path, filename)
    doc.save(full_path)

    return full_path
