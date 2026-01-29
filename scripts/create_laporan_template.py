"""
Script untuk membuat template Laporan Kegiatan dengan format yang lebih baik.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def create_laporan_kegiatan_template():
    """Create improved Laporan Kegiatan template."""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- HEADER / KOP SURAT ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("{{satker_nama}}")
    run.bold = True
    run.font.size = Pt(14)

    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = header2.add_run("{{satker_alamat}}")
    run2.font.size = Pt(11)

    # Garis pemisah
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = separator.add_run("_" * 70)
    run_sep.font.size = Pt(11)

    # --- JUDUL ---
    doc.add_paragraph()  # Spacing

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("LAPORAN KEGIATAN")
    run_title.bold = True
    run_title.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Nomor: {{nomor_dokumen}}")
    run_sub.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # --- INFORMASI KEGIATAN ---
    info_title = doc.add_paragraph()
    run_info = info_title.add_run("A. INFORMASI KEGIATAN")
    run_info.bold = True
    run_info.font.size = Pt(11)

    # Table for activity info
    table_info = doc.add_table(rows=6, cols=3)
    table_info.autofit = True

    info_items = [
        ("1.", "Nama Kegiatan", "{{nama_kegiatan}}"),
        ("2.", "Tanggal Pelaksanaan", "{{tanggal_dokumen}}"),
        ("3.", "Lokasi/Tempat", "{{lokasi}}"),
        ("4.", "Dasar Pelaksanaan", "{{nomor_dasar}}"),
        ("5.", "Pelaksana", "{{penerima_nama}}"),
        ("6.", "Kode Akun/MAK", "{{kode_akun}}"),
    ]

    for i, (no, label, value) in enumerate(info_items):
        row = table_info.rows[i]
        row.cells[0].text = no
        row.cells[1].text = label
        row.cells[2].text = ": " + value

        # Set width
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(9)

        # Format cells
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # --- URAIAN KEGIATAN ---
    uraian_title = doc.add_paragraph()
    run_uraian = uraian_title.add_run("B. URAIAN KEGIATAN")
    run_uraian.bold = True
    run_uraian.font.size = Pt(11)

    uraian_content = doc.add_paragraph()
    uraian_content.add_run("{{uraian_kegiatan}}")
    uraian_content.paragraph_format.first_line_indent = Cm(1)
    uraian_content.paragraph_format.space_after = Pt(6)

    # Placeholder untuk uraian default
    default_uraian = doc.add_paragraph()
    default_uraian.add_run(
        "Kegiatan telah dilaksanakan sesuai dengan rencana yang telah ditetapkan. "
        "Pelaksanaan berjalan dengan baik dan lancar."
    )
    default_uraian.paragraph_format.first_line_indent = Cm(1)

    doc.add_paragraph()  # Spacing

    # --- PESERTA/HADIR ---
    peserta_title = doc.add_paragraph()
    run_peserta = peserta_title.add_run("C. PESERTA/YANG HADIR")
    run_peserta.bold = True
    run_peserta.font.size = Pt(11)

    peserta_content = doc.add_paragraph()
    peserta_content.add_run("{{jumlah_peserta}} orang peserta (terlampir daftar hadir)")
    peserta_content.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()  # Spacing

    # --- HASIL/KESIMPULAN ---
    hasil_title = doc.add_paragraph()
    run_hasil = hasil_title.add_run("D. HASIL/KESIMPULAN")
    run_hasil.bold = True
    run_hasil.font.size = Pt(11)

    hasil_content = doc.add_paragraph()
    hasil_content.add_run("{{hasil_kegiatan}}")
    hasil_content.paragraph_format.first_line_indent = Cm(1)

    # Default text
    default_hasil = doc.add_paragraph()
    default_hasil.add_run(
        "Kegiatan telah selesai dilaksanakan dengan hasil sesuai target yang ditetapkan."
    )
    default_hasil.paragraph_format.first_line_indent = Cm(1)

    doc.add_paragraph()  # Spacing

    # --- REKOMENDASI ---
    rekomendasi_title = doc.add_paragraph()
    run_rekom = rekomendasi_title.add_run("E. REKOMENDASI/TINDAK LANJUT")
    run_rekom.bold = True
    run_rekom.font.size = Pt(11)

    rekomendasi_content = doc.add_paragraph()
    rekomendasi_content.add_run("{{rekomendasi}}")
    rekomendasi_content.paragraph_format.first_line_indent = Cm(1)

    doc.add_paragraph()  # Spacing

    # --- LAMPIRAN ---
    lampiran_title = doc.add_paragraph()
    run_lamp = lampiran_title.add_run("F. LAMPIRAN")
    run_lamp.bold = True
    run_lamp.font.size = Pt(11)

    lampiran_list = [
        "1. Daftar Hadir",
        "2. Dokumentasi Foto",
        "3. Notulensi (jika ada)",
        "4. Bukti-bukti pendukung lainnya",
    ]

    for item in lampiran_list:
        p = doc.add_paragraph()
        p.add_run(item)
        p.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing

    # --- TANDA TANGAN ---
    ttd_intro = doc.add_paragraph()
    ttd_intro.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ttd_intro.add_run("{{lokasi}}, {{tanggal_dokumen}}")

    ttd_jabatan = doc.add_paragraph()
    ttd_jabatan.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ttd_jabatan.add_run("Pelaksana Kegiatan,")

    # Space for signature
    for _ in range(3):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ttd_nama = doc.add_paragraph()
    ttd_nama.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_nama = ttd_nama.add_run("{{penerima_nama}}")
    run_nama.bold = True
    run_nama.underline = True

    ttd_nip = doc.add_paragraph()
    ttd_nip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ttd_nip.add_run("NIP. {{penerima_nip}}")

    doc.add_paragraph()  # Spacing

    # --- MENGETAHUI ---
    mengetahui = doc.add_paragraph()
    mengetahui.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mengetahui.add_run("Mengetahui,")

    ppk_jabatan = doc.add_paragraph()
    ppk_jabatan.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ppk_jabatan.add_run("Pejabat Pembuat Komitmen,")

    # Space for signature
    for _ in range(3):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    ppk_nama = doc.add_paragraph()
    ppk_nama.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_ppk = ppk_nama.add_run("{{ppk_nama}}")
    run_ppk.bold = True
    run_ppk.underline = True

    ppk_nip = doc.add_paragraph()
    ppk_nip.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ppk_nip.add_run("NIP. {{ppk_nip}}")

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "word" / "laporan_kegiatan.docx"
    doc.save(str(output_path))
    print(f"Template saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_laporan_kegiatan_template()
